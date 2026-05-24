"""Generate Word fix document for export skipped_phases bug."""
import sys
sys.path.insert(0, r"D:\Desktop New\AI-QMS-gene for test\AI-QMS-test")
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = Path(r"D:\Desktop New\AI-QMS-gene for test\AI-QMS-test\docs\fix_export_skipped_phases_20260524.docx")

doc = Document()

section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(3)
section.right_margin = Cm(2.5)


def h(text, level=2):
    return doc.add_heading(text, level=level)


def para(text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(10)
    return p


def code_block(text):
    p = doc.add_paragraph()
    p.style = doc.styles["No Spacing"]
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p


def add_table(rows_data, headers):
    n = len(headers)
    tbl = doc.add_table(rows=1 + len(rows_data), cols=n)
    tbl.style = "Table Grid"
    for ci, hdr in enumerate(headers):
        cell = tbl.rows[0].cells[ci]
        cell.text = hdr
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(255, 255, 255)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F4E79")
        tcPr.append(shd)
    for ri, row_data in enumerate(rows_data, 1):
        for ci, val in enumerate(row_data):
            cell = tbl.rows[ri].cells[ci]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
    return tbl


# ── Title ──
title = doc.add_heading("Bug Fix: Report Export Ignores skipped_phases", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run(
    "Date: 2026-05-24  |  Commit: 36fee11  |  File: src/utils/crossexam_export.py"
)
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(100, 100, 100)
doc.add_paragraph()

# ── 1. Problem ──
h("1. Problem Description")
para(
    "When generating reports via the pipeline, both 一般報告 (P0-P3) and 風險分析報告 (P0-P4) "
    "produced Excel and Word files that looked identical to users. Despite being called with "
    "different skipped_phases values, the exported content was the same."
)
doc.add_paragraph()

h("1.1  Observed Symptoms", level=3)
add_table(
    [
        ["合規分析 cols 14-15 (P4 data)", "Empty in 一般報告", "Identical (same data or both empty)"],
        ["合規分析 cols 16-17 (P5 data)", "Empty in both", "Same"],
        ["合規分析 column headers", "Grey in 一般報告 for P4/P5", "Both blue"],
        ["LLM 互動記錄 sheet", "Absent in 一般報告", "✅ Already correct"],
        ["交叉詰問 sheet", "Skip notice in both (P5 skipped)", "Filled with 311 rows from flat_rows"],
        ["Word Section 0.5 remediation col", "Empty in 一般報告", "Both showed remediation data"],
    ],
    headers=["Sheet / Section", "Expected Difference", "Actual"],
)
doc.add_paragraph()

# ── 2. Root Cause ──
h("2. Root Cause")
para(
    "export_deep_report_excel() accepted the skipped_phases parameter but only used it at "
    "line 2166 (inside the Phase Progress sheet builder) — after the compliance table was "
    "already written. All other sheets unconditionally output all data from flat_rows."
)
doc.add_paragraph()

h("2.1  _skipped_xl defined too late", level=3)
code_block(
    "# Was at line 2166 — AFTER Sheet 2 was already written\n"
    "_skipped_xl = skipped_phases or []\n\n"
    "# Sheet 2 (compliance) was built at lines 2118-2160, before _skipped_xl existed"
)

h("2.2  P4/P5 data always written in compliance sheet rows", level=3)
code_block(
    "# Always executed, regardless of skipped_phases:\n"
    'ws_comp.cell(row=ri, column=14, value=_safe_str(row.get("remediation_suggestion"), 500))\n'
    'ws_comp.cell(row=ri, column=15, value=_safe_str(row.get("remediation_regulation_cite"), 500))\n'
    'ws_comp.cell(row=ri, column=16, value=_safe_str(row.get("analyzer_position"), 500))\n'
    'ws_comp.cell(row=ri, column=17, value=_safe_str(row.get("verifier_position"), 500))'
)

h("2.3  交叉詰問 sheet always populated when P5 skipped", level=3)
code_block(
    "# elif flat_rows: ran unconditionally — even when phase_5 was skipped\n"
    "elif flat_rows:\n"
    "    for ri, row in enumerate(flat_rows, 2):\n"
    "        ...  # writes 311 rows of verification data even if P5 was skipped"
)

h("2.4  LLM 互動記錄 sheet did not filter by phase", level=3)
code_block(
    "# All interactions written including P5 verification records\n"
    "for ri, interaction in enumerate(interactions, 2):\n"
    "    ..."
)

h("2.5  Word Section 0.5 remediation column always shown", level=3)
code_block(
    '# In risk summary table, col 7 always showed remediation suggestion\n'
    '(r.get("remediation_suggestion") or "")[:150],'
)
doc.add_paragraph()

# ── 3. Report Types ──
h("3. Three Report Types — Definition")
add_table(
    [
        ["📄 一般報告",     '["phase_4", "phase_5", "phase_6"]', "P0 P0.5 P1 P2 P3"],
        ["⚖️ 風險分析報告", '["phase_5", "phase_6"]',            "P0 P0.5 P1 P2 P3 P4"],
        ["📋 深度報告",     "[]",                                 "P0 P0.5 P1 P2 P3 P4 P5 P6"],
    ],
    headers=["Report", "skipped_phases", "Phases Run"],
)
doc.add_paragraph()
para("Phase → interaction phase name mapping:")
code_block('phase_4  →  "remediation"\nphase_5  →  "verification"')
doc.add_paragraph()

# ── 4. Fix ──
h("4. Fix Applied")
para("File: src/utils/crossexam_export.py — 6 targeted changes across Excel and Word export functions.")
doc.add_paragraph()

h("Change 1 — Move _skipped_xl before Sheet 2", level=3)
code_block(
    "# BEFORE (defined at line 2166, after compliance sheet)\n"
    "# AFTER (defined right after _label_key, before wb = Workbook())\n"
    "_skipped_xl = skipped_phases or []"
)

h("Change 2 — Grey out skipped phase column headers", level=3)
code_block(
    '_skip_hdr_fill = PatternFill(start_color="BFBFBF", end_color="BFBFBF", fill_type="solid")\n'
    '_skip_hdr_font = Font(bold=True, color="666666", size=10)\n'
    "for ci, h in enumerate(comp_headers, 1):\n"
    "    c = ws_comp.cell(row=1, column=ci, value=h)\n"
    '    if ci in (14, 15) and "phase_4" in _skipped_xl:\n'
    "        c.fill = _skip_hdr_fill\n"
    "        c.font = _skip_hdr_font\n"
    '    elif ci in (16, 17) and "phase_5" in _skipped_xl:\n'
    "        c.fill = _skip_hdr_fill\n"
    "        c.font = _skip_hdr_font\n"
    "    else:\n"
    "        c.fill = header_fill\n"
    "        c.font = header_font"
)

h("Change 3 — Conditional P4/P5 cell data in compliance sheet rows", level=3)
code_block(
    '# P4 columns: only write when phase_4 not skipped\n'
    'if "phase_4" not in _skipped_xl:\n'
    '    ws_comp.cell(row=ri, column=14, value=_safe_str(row.get("remediation_suggestion"), 500))\n'
    '    ws_comp.cell(row=ri, column=15, value=_safe_str(row.get("remediation_regulation_cite"), 500))\n'
    '# P5 columns: only write when phase_5 not skipped\n'
    'if "phase_5" not in _skipped_xl:\n'
    '    ws_comp.cell(row=ri, column=16, value=_safe_str(row.get("analyzer_position"), 500))\n'
    '    ws_comp.cell(row=ri, column=17, value=_safe_str(row.get("verifier_position"), 500))'
)

h("Change 4 — 交叉詰問 sheet: block flat_rows path when P5 skipped", level=3)
code_block(
    "# BEFORE\n"
    "elif flat_rows:\n"
    "    for ri, row in enumerate(flat_rows, 2):\n"
    "        ...\n\n"
    "# AFTER\n"
    'elif flat_rows and "phase_5" not in _skipped_xl:\n'
    "    for ri, row in enumerate(flat_rows, 2):\n"
    "        ...\n\n"
    "# Skip notice when no data written\n"
    "_skip_note = (\n"
    '    "（P5 交叉詰問已跳過 — 此報告類型不包含交叉詰問資料）"\n'
    '    if "phase_5" in _skipped_xl\n'
    '    else "（無交叉詰問資料）"\n'
    ")\n"
    "ws_xe.cell(row=2, column=1, value=_skip_note)"
)

h("Change 5 — LLM 互動記錄: filter by skipped phases", level=3)
code_block(
    '_SKIP_PHASE_MAP = {"phase_4": "remediation", "phase_5": "verification"}\n'
    "_filtered_interactions = [\n"
    "    i for i in (interactions or [])\n"
    "    if i.get(\"phase\") not in {_SKIP_PHASE_MAP[p] for p in _skipped_xl if p in _SKIP_PHASE_MAP}\n"
    "] if interactions else None\n"
    "if _filtered_interactions:\n"
    '    ws_llm = wb.create_sheet(dh["xl_sheet_llm"])\n'
    "    ...\n"
    "    for ri, interaction in enumerate(_filtered_interactions, 2):\n"
    "        ..."
)

h("Change 6 — Word Section 0.5: hide remediation column when P4 skipped", level=3)
code_block(
    "# BEFORE\n"
    '(r.get("remediation_suggestion") or "")[:150],\n\n'
    "# AFTER\n"
    '"" if "phase_4" in _skipped else (r.get("remediation_suggestion") or "")[:150],'
)
doc.add_paragraph()

# ── 5. Verified Output ──
h("5. Verified Output Differences")
para(
    "Reports verified from pipeline run run_8c3ee3a447c1 (311 rows, regulatory_list, ISO 13485):"
)
doc.add_paragraph()
add_table(
    [
        ["合規分析 col 14-15 header fill", "#BFBFBF (grey)", "#4472C4 (blue)"],
        ["合規分析 col 16-17 header fill", "#BFBFBF (grey)", "#BFBFBF (grey)"],
        ["合規分析 P4 non-empty rows", "0", "1"],
        ["LLM 互動記錄 sheet", "Absent", "Present (P1/P2/P4 only)"],
        ["LLM 互動記錄 phases", "—", "{P1 差距掃描, P2 查核表驗證, P4 改善建議}"],
        ["交叉詰問 row 2", "Skip notice", "Skip notice (P5 skipped in both)"],
        ["Word S0.5 remediation col", "Empty", "Shows data when available"],
    ],
    headers=["Area", "一般報告 P0-P3 (v2)", "風險分析報告 P0-P4 (v2)"],
)
doc.add_paragraph()

# ── 6. Files Changed ──
h("6. Files Changed")
add_table(
    [["src/utils/crossexam_export.py", "6 targeted changes across Excel and Word export functions"]],
    headers=["File", "Change"],
)
para(
    "Sample corrected reports saved to doc/: "
    "doc/一般報告_P0-P3_v2.docx / _v2.xlsx  |  doc/風險分析報告_P0-P4_v2.docx / _v2.xlsx"
)

doc.save(str(OUTPUT))
print("Word saved:", OUTPUT)
