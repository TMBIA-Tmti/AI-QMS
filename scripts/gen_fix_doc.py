"""Generate Word fix document for report-type timeout bug."""
import sys
sys.path.insert(0, r"D:\Desktop New\AI-QMS-gene for test\AI-QMS-test")
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = Path(r"D:\Desktop New\AI-QMS-gene for test\AI-QMS-test\docs\fix_report_type_timeout_20260524.docx")

doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(3)
section.right_margin = Cm(2.5)


def h(text, level=2):
    return doc.add_heading(text, level=level)


def para(text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(10)
    return p


def code_block(text):
    p = doc.add_paragraph()
    p.style = doc.styles["No Spacing"]
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p


def add_table(rows_data, headers):
    n = len(headers)
    tbl = doc.add_table(rows=1 + len(rows_data), cols=n)
    tbl.style = "Table Grid"
    # Header row
    for ci, hdr in enumerate(headers):
        cell = tbl.rows[0].cells[ci]
        cell.text = hdr
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(255, 255, 255)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F4E79")
        tcPr.append(shd)
    # Data rows
    for ri, row_data in enumerate(rows_data, 1):
        for ci, val in enumerate(row_data):
            cell = tbl.rows[ri].cells[ci]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
    return tbl


# ── Title ──
title = doc.add_heading("Bug Fix: Report Type Dialog Timeout", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Date: 2026-05-24  |  Commit: 92b226f  |  File: src/chainlit_app/app.py")
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(100, 100, 100)
doc.add_paragraph()

# ── 1. Problem ──
h("1. Problem Description")
para(
    "When the 120-second report-type selection dialog timed out without user interaction, "
    "the system displayed a confirmation message claiming that phases P4, P5, and P6 were "
    "skipped. In reality, the full P0-P6 pipeline was executing."
)
doc.add_paragraph()

h("1.1  Bug A — Wrong Confirmation Message", level=3)
para("Both timeout code paths sent the report_type.selected_normal i18n key, displaying:")
code_block("📄 已選擇：一般報告（跳過 P4 改善建議、P5 交叉詰問、P6 來源驗證）")
para("The message claimed P4, P5, and P6 were skipped. They were not.")
doc.add_paragraph()

h("1.2  Bug B — Incorrect skipped_phases Value", level=3)
para("The timeout path called _fallback_skip_phases(), which contained a Python type-check bug:")
code_block(
    "saved = get_custom_skip_phases()  # returns [] (empty list)\n"
    "if saved is not None:             # True for empty list!\n"
    "    return saved                  # returns [] → skip nothing"
)
para(
    "Because '[] is not None' evaluates to True, the function returned [] (no phases skipped) "
    "instead of reaching the intended default [\"phase_6\"]."
)
doc.add_paragraph()

h("1.3  Combined Effect", level=3)
para("The empty list [] was further collapsed to None by the conditional expression:")
code_block("custom_skip_phases = [] if [] else None  →  None")
para(
    "In pipeline_runner.py, 'if None:' is False, so pipeline.state.skipped_phases was never set. "
    "All phases ran silently."
)
doc.add_paragraph()

# ── 2. Root Cause ──
h("2. Root Cause Trace")
code_block(
    "User: no click within 120s\n"
    "  ↓\n"
    "AskActionMessage timeout → res = None\n"
    "  ↓\n"
    "else branch: _fallback_skip_phases()\n"
    "  ↓\n"
    "get_custom_skip_phases() → []  (no saved config)\n"
    "  ↓\n"
    "if [] is not None: → True  ← BUG\n"
    "  ↓\n"
    "return []  (skip nothing)\n"
    "  ↓\n"
    'Message sent: "一般報告"  ← WRONG label\n'
    "  ↓\n"
    "P4 (改善建議) runs  ✓  (should have been skipped per message)\n"
    "P5 (獨立驗證) runs  ✓  (should have been skipped per message)\n"
    "P6 (來源驗證) runs  ✓  (should have been skipped per message)"
)
doc.add_paragraph()

# ── 3. Report Types ──
h("3. Three Report Types — Definition")
add_table(
    [
        ["📄 一般報告",      "report_type_normal", "phase_4, phase_5, phase_6", "P0 → P0.5 → P1 → P2 → P3"],
        ["⚖️ 風險分析報告",  "report_type_risk",   "phase_5, phase_6",          "P0 → ... → P4"],
        ["📋 深度報告",      "report_type_deep",   "(none)",                    "P0 → ... → P6"],
        ["(timeout/error)", "—",                   "(none, now fixed)",         "P0-P6, correctly shows 深度報告"],
    ],
    headers=["Label", "Action Name", "Skipped Phases", "Pipeline Executed"],
)
doc.add_paragraph()

# ── 4. Fix ──
h("4. Fix Applied")
h("Before (buggy)", level=3)
code_block(
    "except Exception:\n"
    "    _fb = _fallback_skip_phases()\n"
    '    await cl.Message(content=t("report_type.selected_normal")).send()\n'
    "    return _fb\n"
    "\n"
    "else:\n"
    "    # Timeout (res is None) — use persisted config or standard default\n"
    "    _fb = _fallback_skip_phases()\n"
    '    await cl.Message(content=t("report_type.selected_normal")).send()\n'
    "    return _fb"
)

h("After (fixed)", level=3)
code_block(
    "except Exception:\n"
    '    await cl.Message(content=t("report_type.selected_deep")).send()\n'
    "    return []\n"
    "\n"
    "else:\n"
    "    # Timeout (res is None) — default to deep report (run all phases)\n"
    '    await cl.Message(content=t("report_type.selected_deep")).send()\n'
    "    return []"
)

h("Changes Summary", level=3)
add_table(
    [
        ["except Exception branch",
         'Called _fallback_skip_phases(), sent "一般報告" message',
         'Sends "深度報告" message, returns []'],
        ["else (timeout) branch",
         'Called _fallback_skip_phases(), sent "一般報告" message',
         'Sends "深度報告" message, returns []'],
        ["_fallback_skip_phases()",
         "Called from both timeout paths",
         "No longer called from _ask_report_type()"],
    ],
    headers=["Code Path", "Before", "After"],
)
doc.add_paragraph()

# ── 5. After Fix ──
h("5. Behavior After Fix")
para("When the dialog times out, the user now sees:")
code_block("📋 已選擇：深度報告（執行所有分析階段）")
para(
    "The pipeline runs P0-P6, matching the actual behavior. "
    "Message and execution are now fully consistent."
)
doc.add_paragraph()

# ── 6. Report Comparison ──
h("6. Report Content Comparison")
para(
    "Sample reports generated from pipeline run run_8c3ee3a447c1 "
    "(311 rows, regulatory_list, ISO 13485) saved to doc/:"
)
add_table(
    [
        ["Phase progress table",       "✅ P4/P5/P6 marked skip", "✅ P5/P6 marked skip",   "✅ all phases"],
        ["Data quality summary",       "✅",                       "✅",                      "✅"],
        ["Gap analysis (P1)",          "✅ 311 clauses",           "✅ 311 clauses",          "✅ 311 clauses"],
        ["Checklist verification (P2)","✅",                       "✅",                      "✅"],
        ["Risk matrix (P3)",           "✅ 109 immediate / 90 deadline", "✅ same",           "✅ same"],
        ["Remediation (P4)",           "❌ skipped",               "✅ included",             "✅ included"],
        ["Cross-examination (P5)",     "❌ skipped",               "❌ skipped",              "✅ included"],
        ["Source verification (P6)",   "❌ skipped",               "❌ skipped",              "✅ included"],
        ["Word file size",             "~102 KB",                  "~154 KB",                 "larger"],
        ["Excel file size",            "~141 KB",                  "~193 KB",                 "larger"],
    ],
    headers=["Content Section", "一般報告 (P0-P3)", "風險分析報告 (P0-P4)", "深度報告 (P0-P6)"],
)
doc.add_paragraph()

# ── 7. Files Changed ──
h("7. Files Changed")
add_table(
    [["src/chainlit_app/app.py", "Fixed two timeout paths in _ask_report_type()"]],
    headers=["File", "Change"],
)
para(
    "No other files were modified. The _fallback_skip_phases() helper function "
    "remains in place but is no longer called from _ask_report_type()."
)

doc.save(str(OUTPUT))
print("Word saved:", OUTPUT)
