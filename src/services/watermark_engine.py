"""
AI-QMS - Watermark Application Engine
======================================

Applies watermark images to document files (PDF and Word).

Supported formats:
- PDF (.pdf): Uses reportlab to create watermark layer + PyPDF2 to merge/overlay
- Word (.docx): Uses python-docx to insert watermark image in document header

Also provides preview generation (creates a sample PDF with watermark applied).

Unsupported formats (.xlsx, .pptx, .png, etc.) are skipped with a warning.
"""

import io
import logging
import os
import shutil
import tempfile
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# ------------------------------------------------------------------
# Optional dependency checks
# ------------------------------------------------------------------

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas as rl_canvas

    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

try:
    import PyPDF2

    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    from PIL import Image

    HAS_PILLOW = True
except ImportError:
    HAS_PILLOW = False

try:
    import docx as python_docx
    from docx.shared import Inches

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class WatermarkEngine:
    """浮水印施加引擎，支援 PDF 與 Word 格式。"""

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def apply_to_pdf(
        self,
        pdf_path: str,
        watermark_image_path: str,
        settings: dict,
        output_path: Optional[str] = None,
    ) -> dict:
        """
        Apply watermark to a PDF file.

        Creates a transparent watermark PDF layer using reportlab, then
        merges it onto each page of the original PDF using PyPDF2.

        Args:
            pdf_path: Path to the original PDF file.
            watermark_image_path: Path to the watermark image.
            settings: Watermark settings dict (angle, opacity, scale, position, repeat).
            output_path: Output path. None = overwrite original.

        Returns:
            {"success": bool, "pages_processed": int, "error": str|None}
        """
        if not HAS_REPORTLAB:
            return {
                "success": False,
                "pages_processed": 0,
                "error": "Required package 'reportlab' is not installed.",
            }
        if not HAS_PYPDF2:
            return {
                "success": False,
                "pages_processed": 0,
                "error": "Required package 'PyPDF2' is not installed.",
            }
        if not HAS_PILLOW:
            return {
                "success": False,
                "pages_processed": 0,
                "error": "Required package 'Pillow' is not installed.",
            }

        tmp_file = None
        try:
            # Read original PDF
            reader = PyPDF2.PdfReader(pdf_path)
            if reader.is_encrypted:
                try:
                    reader.decrypt("")
                except Exception:
                    return {
                        "success": False,
                        "pages_processed": 0,
                        "error": "PDF is encrypted and cannot be decrypted.",
                    }

            writer = PyPDF2.PdfWriter()
            pages_processed = 0

            for page in reader.pages:
                page_width = float(page.mediabox.width)
                page_height = float(page.mediabox.height)

                # Create watermark layer for this page
                wm_pdf_bytes = self._create_watermark_pdf(
                    watermark_image_path, settings, page_width, page_height
                )

                wm_reader = PyPDF2.PdfReader(io.BytesIO(wm_pdf_bytes))
                wm_page = wm_reader.pages[0]

                # Merge watermark onto original page
                page.merge_page(wm_page)
                writer.add_page(page)
                pages_processed += 1

            # Write result
            final_path = output_path or pdf_path
            if final_path == pdf_path:
                # Write to temp file first, then replace
                tmp_fd, tmp_file = tempfile.mkstemp(suffix=".pdf")
                os.close(tmp_fd)
                with open(tmp_file, "wb") as f:
                    writer.write(f)
                shutil.move(tmp_file, pdf_path)
                tmp_file = None  # prevent cleanup
            else:
                os.makedirs(os.path.dirname(final_path), exist_ok=True)
                with open(final_path, "wb") as f:
                    writer.write(f)

            logger.info(
                "Watermark applied to PDF: %s (%d pages)", pdf_path, pages_processed
            )
            return {"success": True, "pages_processed": pages_processed, "error": None}

        except Exception as e:
            logger.error("Failed to apply watermark to PDF %s: %s", pdf_path, e)
            return {"success": False, "pages_processed": 0, "error": str(e)}
        finally:
            if tmp_file and os.path.isfile(tmp_file):
                try:
                    os.remove(tmp_file)
                except OSError:
                    pass

    def apply_to_docx(
        self,
        docx_path: str,
        watermark_image_path: str,
        settings: dict,
        output_path: Optional[str] = None,
    ) -> dict:
        """
        Apply watermark to a Word (.docx) document.

        Inserts the watermark image into each section's header with
        behind-text positioning.

        Args:
            docx_path: Path to the original .docx file.
            watermark_image_path: Path to the watermark image.
            settings: Watermark settings dict.
            output_path: Output path. None = overwrite original.

        Returns:
            {"success": bool, "error": str|None}
        """
        if not HAS_DOCX:
            return {
                "success": False,
                "error": "Required package 'python-docx' is not installed.",
            }
        if not HAS_PILLOW:
            return {
                "success": False,
                "error": "Required package 'Pillow' is not installed.",
            }

        try:
            doc = python_docx.Document(docx_path)
            opacity = settings.get("opacity", 0.15)
            scale = settings.get("scale", 1.0)

            # Prepare watermark image with opacity
            wm_image_path = self._prepare_image_with_opacity(
                watermark_image_path, opacity, scale
            )

            for section in doc.sections:
                header = section.header
                header.is_linked_to_previous = False

                # Add watermark image paragraph
                paragraph = header.add_paragraph()
                paragraph.alignment = 1  # Center

                run = paragraph.add_run()
                run.add_picture(wm_image_path, width=Inches(5 * scale))

            final_path = output_path or docx_path
            doc.save(final_path)

            # Clean up temp image
            if wm_image_path != watermark_image_path:
                try:
                    os.remove(wm_image_path)
                except OSError:
                    pass

            logger.info("Watermark applied to DOCX: %s", docx_path)
            return {"success": True, "error": None}

        except Exception as e:
            logger.error("Failed to apply watermark to DOCX %s: %s", docx_path, e)
            return {"success": False, "error": str(e)}

    def generate_preview(
        self,
        watermark_image_path: str,
        settings: dict,
        output_path: str,
    ) -> dict:
        """
        Generate a preview PDF showing the watermark effect on a sample page.

        Creates a single A4 page with sample text and the watermark applied.

        Args:
            watermark_image_path: Path to the watermark image.
            settings: Watermark settings dict.
            output_path: Where to save the preview PDF.

        Returns:
            {"success": bool, "preview_path": str, "error": str|None}
        """
        if not HAS_REPORTLAB:
            return {
                "success": False,
                "preview_path": "",
                "error": "Required package 'reportlab' is not installed.",
            }
        if not HAS_PILLOW:
            return {
                "success": False,
                "preview_path": "",
                "error": "Required package 'Pillow' is not installed.",
            }

        try:
            page_w, page_h = A4  # 595.27 x 841.89 points

            # Step 1: Create base page with sample text
            base_buf = io.BytesIO()
            c = rl_canvas.Canvas(base_buf, pagesize=A4)

            # Title
            c.setFont("Helvetica-Bold", 18)
            c.drawCentredString(page_w / 2, page_h - 60, "Quality Manual / 品質手冊")

            # Subtitle
            c.setFont("Helvetica", 12)
            c.drawCentredString(
                page_w / 2, page_h - 85, "Document Preview - Watermark Test"
            )

            # Sample body text
            c.setFont("Helvetica", 10)
            sample_lines = [
                "1. Purpose",
                "   This document establishes the quality management system",
                "   in accordance with ISO 13485:2016 requirements.",
                "",
                "2. Scope",
                "   This quality manual applies to the design, development,",
                "   production, and servicing of medical devices.",
                "",
                "3. Quality Policy",
                "   Our organization is committed to providing safe and",
                "   effective medical devices that meet regulatory requirements",
                "   and customer expectations.",
                "",
                "4. Document Control",
                "   All quality documents shall be reviewed, approved,",
                "   and controlled to ensure current versions are available.",
                "",
                "5. Management Responsibility",
                "   Top management shall demonstrate commitment to the",
                "   quality management system through leadership and planning.",
            ]

            y = page_h - 120
            for line in sample_lines:
                c.drawString(60, y, line)
                y -= 16

            c.save()
            base_buf.seek(0)

            # Step 2: Create watermark layer
            wm_pdf_bytes = self._create_watermark_pdf(
                watermark_image_path, settings, page_w, page_h
            )

            # Step 3: Merge
            if HAS_PYPDF2:
                base_reader = PyPDF2.PdfReader(base_buf)
                wm_reader = PyPDF2.PdfReader(io.BytesIO(wm_pdf_bytes))

                writer = PyPDF2.PdfWriter()
                page = base_reader.pages[0]
                page.merge_page(wm_reader.pages[0])
                writer.add_page(page)

                os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
                with open(output_path, "wb") as f:
                    writer.write(f)
            else:
                # Fallback: just write the base PDF without watermark overlay
                os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
                with open(output_path, "wb") as f:
                    f.write(base_buf.getvalue())

            logger.info("Watermark preview generated: %s", output_path)
            return {"success": True, "preview_path": output_path, "error": None}

        except Exception as e:
            logger.error("Failed to generate watermark preview: %s", e)
            return {"success": False, "preview_path": "", "error": str(e)}

    def apply_watermark(
        self,
        file_path: str,
        watermark_image_path: str,
        settings: dict,
    ) -> dict:
        """
        Auto-detect file format and apply watermark.

        Supported: .pdf, .docx
        Unsupported: returns skip result.

        Args:
            file_path: Path to the file.
            watermark_image_path: Path to the watermark image.
            settings: Watermark settings dict.

        Returns:
            Result dict from format-specific method, or skip result.
        """
        ext = Path(file_path).suffix.lower()

        if ext == ".pdf":
            return self.apply_to_pdf(file_path, watermark_image_path, settings)
        elif ext == ".docx":
            return self.apply_to_docx(file_path, watermark_image_path, settings)
        else:
            return {
                "success": False,
                "skipped": True,
                "reason": "format_not_supported",
                "format": ext,
            }

    # ------------------------------------------------------------------
    # Private helpers
    # ------------------------------------------------------------------

    def _create_watermark_pdf(
        self,
        image_path: str,
        settings: dict,
        page_width: float,
        page_height: float,
    ) -> bytes:
        """
        Create a single-page transparent PDF containing the watermark image.

        Uses reportlab to draw the image with rotation and positioning.
        Opacity is applied by pre-processing the image with Pillow.

        Returns:
            PDF content as bytes.
        """
        angle = settings.get("angle", -45)
        opacity = settings.get("opacity", 0.15)
        scale = settings.get("scale", 1.0)
        position = settings.get("position", "center")
        repeat = settings.get("repeat", False)

        # Pre-process image: apply opacity via Pillow
        processed_path = self._prepare_image_with_opacity(image_path, opacity, scale)

        try:
            img = Image.open(processed_path)
            img_w, img_h = img.size
            img.close()

            # Points per pixel (rough conversion at 72 DPI)
            # Scale image to reasonable size relative to page
            base_size = min(page_width, page_height) * 0.3 * scale
            aspect = img_w / img_h if img_h > 0 else 1.0
            if aspect >= 1:
                draw_w = base_size
                draw_h = base_size / aspect
            else:
                draw_h = base_size
                draw_w = base_size * aspect

            buf = io.BytesIO()
            c = rl_canvas.Canvas(buf, pagesize=(page_width, page_height))

            if repeat:
                # Tile mode: place watermarks in a grid
                step_x = draw_w * 1.5
                step_y = draw_h * 1.5
                x = step_x * 0.25
                while x < page_width:
                    y = step_y * 0.25
                    while y < page_height:
                        self._draw_watermark_at(
                            c, processed_path, x, y, draw_w, draw_h, angle
                        )
                        y += step_y
                    x += step_x
            else:
                # Single placement
                cx, cy = self._get_position_coords(
                    position, page_width, page_height, draw_w, draw_h
                )
                self._draw_watermark_at(
                    c, processed_path, cx, cy, draw_w, draw_h, angle
                )

            c.save()
            return buf.getvalue()

        finally:
            # Clean up temp image
            if processed_path != image_path:
                try:
                    os.remove(processed_path)
                except OSError:
                    pass

    def _draw_watermark_at(
        self,
        canvas_obj,
        image_path: str,
        cx: float,
        cy: float,
        width: float,
        height: float,
        angle: float,
    ) -> None:
        """Draw a single watermark image at the specified position with rotation."""
        canvas_obj.saveState()
        canvas_obj.translate(cx, cy)
        canvas_obj.rotate(angle)
        canvas_obj.drawImage(
            image_path,
            -width / 2,
            -height / 2,
            width=width,
            height=height,
            mask="auto",
        )
        canvas_obj.restoreState()

    @staticmethod
    def _get_position_coords(
        position: str,
        page_w: float,
        page_h: float,
        img_w: float,
        img_h: float,
    ) -> tuple:
        """Calculate center coordinates for the given position name."""
        margin = 50  # points from edge
        positions = {
            "center": (page_w / 2, page_h / 2),
            "top-left": (margin + img_w / 2, page_h - margin - img_h / 2),
            "top-right": (page_w - margin - img_w / 2, page_h - margin - img_h / 2),
            "bottom-left": (margin + img_w / 2, margin + img_h / 2),
            "bottom-right": (page_w - margin - img_w / 2, margin + img_h / 2),
        }
        return positions.get(position, positions["center"])

    @staticmethod
    def _prepare_image_with_opacity(
        image_path: str,
        opacity: float,
        scale: float = 1.0,
    ) -> str:
        """
        Pre-process watermark image: apply opacity and scale via Pillow.

        Creates a temporary file with the processed image.

        Returns:
            Path to the processed image (temp file, or original if no processing needed).
        """
        if not HAS_PILLOW:
            return image_path

        try:
            img = Image.open(image_path)

            # Convert to RGBA if not already
            if img.mode != "RGBA":
                img = img.convert("RGBA")

            # Apply scale
            if scale != 1.0:
                new_w = max(1, int(img.width * scale))
                new_h = max(1, int(img.height * scale))
                img = img.resize((new_w, new_h), Image.LANCZOS)

            # Apply opacity by modifying alpha channel
            if opacity < 1.0:
                alpha = img.split()[3]
                # Multiply existing alpha by opacity
                alpha = alpha.point(lambda p: int(p * opacity))
                img.putalpha(alpha)

            # Save to temp file
            tmp_fd, tmp_path = tempfile.mkstemp(suffix=".png")
            os.close(tmp_fd)
            img.save(tmp_path, format="PNG")
            img.close()

            return tmp_path

        except Exception as e:
            logger.warning("Failed to pre-process watermark image: %s", e)
            return image_path
