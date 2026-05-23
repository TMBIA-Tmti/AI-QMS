"""
AI-QMS — Daily Audit & 10-Day Meta Review Module
==================================================

Third-party LLM agent that evaluates cross-examination quality through two
dimensions:

  Dimension A — Regulation Accuracy:
    MDSAP on  → Compares against all 7-country MDSAP regulation profiles.
    MDSAP off → Compares against TFDA + EU_MDR (2-country) only.

  Dimension B — Cross-Examination Quality Verification:
    MDSAP on  → MDSAP 5-country verification quality (7-country coverage).
    MDSAP off → 2-country quality evaluation with advisory note.

Daily audit runs produce a DailyAuditResult with 0-100 scores for each
dimension plus an overall score.

Every 10 days, a meta review aggregates daily results and produces a
MetaReviewResult with trend analysis and deviation summary.

Uses the bilingual prompt pattern: {"zh": ..., "en": ...}
Follows the pattern from crossexam_qa_agent.py.
"""

from __future__ import annotations

import json
import logging
import time
from datetime import datetime
from pathlib import Path
from typing import Optional

from src.database.daily_crossexam_store import DailySamplingRecord
from src.analysis.compliance_rules import get_audit_question as _get_audit_question, get_checklist as _get_checklist
from src.utils.safe_io import atomic_write_json

logger = logging.getLogger(__name__)

__all__ = [
    "DailyAuditResult",
    "MetaReviewResult",
    "AuditFeedback",
    "run_daily_sampling_crossexam",
    "run_daily_audit",
    "run_10day_meta_review",
    "get_daily_audit_history",
    "get_latest_meta_review",
    "export_daily_audit_word",
    "export_daily_audit_excel",
    "export_meta_review_word",
    "export_meta_review_excel",
    "save_feedback",
    "get_all_feedback",
    "update_feedback",
    "delete_feedback",
    "get_feedback_by_id",
    "get_feedback_for_audit",
    "get_active_feedback_context",
]

AUDIT_DIR = Path("./data/daily_audit")
EXPORT_DIR = Path("data/exports")
FEEDBACK_DIR = AUDIT_DIR / "feedback"

# Deviation thresholds
DEVIATION_OVERALL_THRESHOLD = 70
DEVIATION_DIM_GAP_THRESHOLD = 20


# ============================================================
# i18n helper
# ============================================================

from src.chainlit_app.lang_config import lang_key as _get_prompt_lang  # noqa: E402


# ============================================================
# System prompts — Dimension A (MDSAP regulation accuracy)
# ============================================================


_DIM_A_SYSTEM_PROMPTS_MDSAP_ON = {
    "zh": """你是品質管理系統的「MDSAP 法規準確性稽核專家」。你的任務是比對交叉詰問中 LLM 的回答與 MDSAP 法規原文，評估回答的準確性。

你需要檢查：
1. **法規引用準確性**: LLM 引用的法規條文是否正確？條號、內容是否與原文一致？
2. **要求完整性**: LLM 是否遺漏了法規中的關鍵要求？
3. **解釋正確性**: LLM 對法規的解釋是否正確？有無曲解或過度簡化？
4. **跨國比較準確性**: 當 LLM 進行跨國比較時，差異描述是否正確？
5. **原文一致性**: LLM 的回答是否與法規原文的語義一致？

評分標準（必須嚴格依照此表給分，不得自行裁量）：

| 分數區間 | 條件說明 |
|---------|---------|
| 90–100 | 法規引用完整精確，所有條號與原文一致，無遺漏，跨國比較正確，解釋無誤 |
| 70–89  | 法規引用大致正確，有輕微遺漏或引用不夠精確，但核心要求已涵蓋 |
| 50–69  | 部分符合，遺漏 1–2 項關鍵要求，或有輕微曲解 |
| 30–49  | 僅表面符合，缺乏實質法規依據，或重要條文錯誤 |
| 0–29   | 完全不符合、幻覺性引用、或無法提供有效法規依據 |

回答必須使用以下 JSON 格式：
{
  "dim_a_score": 0-100,
  "score_rationale": "說明依照上表選擇此分數區間的理由",
  "checks": [
    {
      "check_type": "reference_accuracy | completeness | interpretation | cross_comparison | text_consistency",
      "regulation": "被檢查的法規名稱",
      "issue": "發現的問題描述（若無問題則為 null）",
      "severity": "none | low | medium | high | critical",
      "evidence": "支持證據"
    }
  ],
  "summary": "Dim A 評估摘要（2-3 句話）"
}""",
    "en": """You are a "MDSAP Regulation Accuracy Audit Expert" for a quality management system. Your task is to compare LLM cross-examination answers against MDSAP regulation source texts and evaluate answer accuracy.

You need to check:
1. **Reference Accuracy**: Are the regulatory references cited by the LLM correct? Do article numbers and content match the source?
2. **Completeness**: Has the LLM missed critical regulatory requirements?
3. **Interpretation Correctness**: Is the LLM's interpretation of regulations correct? Any distortions or oversimplifications?
4. **Cross-Country Comparison Accuracy**: When the LLM compares across countries, are the differences described correctly?
5. **Source Text Consistency**: Is the LLM's answer semantically consistent with the regulation's original text?

Scoring Rubric (strictly follow this table — no discretionary scoring):

| Score Range | Criteria |
|-------------|----------|
| 90–100 | Fully accurate citations, all article numbers match source, no omissions, correct cross-country comparisons, accurate interpretation |
| 70–89  | Mostly correct, minor omissions or imprecision, but core requirements covered |
| 50–69  | Partial compliance, 1–2 key requirements missing, or minor misinterpretation |
| 30–49  | Surface compliance only, lacks substantive basis, or significant article errors |
| 0–29   | Non-compliant, hallucinated citations, or unable to provide valid regulatory basis |

Respond in the following JSON format:
{
  "dim_a_score": 0-100,
  "score_rationale": "Explain which band applies and why, citing specific evidence",
  "checks": [
    {
      "check_type": "reference_accuracy | completeness | interpretation | cross_comparison | text_consistency",
      "regulation": "Regulation being checked",
      "issue": "Description of issue found (null if no issue)",
      "severity": "none | low | medium | high | critical",
      "evidence": "Supporting evidence"
    }
  ],
  "summary": "Dim A assessment summary (2-3 sentences)"
}""",
    "ja": """あなたは品質管理システムの「MDSAP 法規正確性監査専門家」です。あなたの任務は、交差詰問における LLM の回答を MDSAP 法規原文と比較し、回答の正確性を評価することです。

確認すべき項目：
1. **法規引用の正確性**: LLM が引用した法規条文は正しいか？条番号と内容は原文と一致しているか？
2. **完全性**: LLM は法規の重要な要件を見落としていないか？
3. **解釈の正確性**: LLM の法規解釈は正しいか？歪曲や過度の簡略化はないか？
4. **国間比較の正確性**: LLM が国間比較を行う際、差異の記述は正しいか？
5. **原文との一貫性**: LLM の回答は法規原文の意味と一貫しているか？

採点基準（この表に厳密に従うこと — 裁量的採点は不可）：

| スコア範囲 | 基準 |
|------------|------|
| 90-100 | 引用が完全に正確、全条番号が原文と一致、欠落なし、国間比較が正確、解釈が正確 |
| 70-89  | 概ね正確、軽微な欠落または不正確さがあるが、コア要件はカバー |
| 50-69  | 部分的準拠、1-2の重要要件が欠落、または軽微な誤解釈 |
| 30-49  | 表面的準拠のみ、実質的根拠なし、または重要な条文エラー |
| 0-29   | 不準拠、幻覚引用、または有効な法規根拠を提供できない |

以下の JSON 形式で回答してください：
{
  "dim_a_score": 0-100,
  "score_rationale": "上表のどの区間に該当するかとその理由を説明",
  "checks": [
    {
      "check_type": "reference_accuracy | completeness | interpretation | cross_comparison | text_consistency",
      "regulation": "確認対象の法規名",
      "issue": "発見された問題の説明（問題がない場合は null）",
      "severity": "none | low | medium | high | critical",
      "evidence": "裏付け証拠"
    }
  ],
  "summary": "Dim A 評価の要約（2-3文）"
}""",
}

_DIM_A_SYSTEM_PROMPTS_MDSAP_OFF = {
    "zh": """你是品質管理系統的「法規準確性稽核專家」。你的任務是比對交叉詰問中 LLM 的回答與 TFDA（台灣）及 EU MDR（歐盟）法規原文，評估回答的準確性。

目前系統僅使用 2 國法規（TFDA + EU_MDR）進行交叉詰問，尚未啟用 MDSAP 5 國驗證。請僅針對這 2 國法規進行評估。

你需要檢查：
1. **法規引用準確性**: LLM 引用的 TFDA / EU MDR 條文是否正確？條號、內容是否與原文一致？
2. **要求完整性**: LLM 是否遺漏了這 2 國法規中的關鍵要求？
3. **解釋正確性**: LLM 對法規的解釋是否正確？有無曲解或過度簡化？
4. **跨國比較準確性**: 當 LLM 比較 TFDA 與 EU MDR 差異時，描述是否正確？
5. **原文一致性**: LLM 的回答是否與法規原文的語義一致？

評分標準（必須嚴格依照此表給分，不得自行裁量）：

| 分數區間 | 條件說明 |
|---------|---------|
| 90–100 | 法規引用完整精確，所有條號與原文一致，無遺漏，跨國比較正確，解釋無誤 |
| 70–89  | 法規引用大致正確，有輕微遺漏或引用不夠精確，但核心要求已涵蓋 |
| 50–69  | 部分符合，遺漏 1–2 項關鍵要求，或有輕微曲解 |
| 30–49  | 僅表面符合，缺乏實質法規依據，或重要條文錯誤 |
| 0–29   | 完全不符合、幻覺性引用、或無法提供有效法規依據 |

回答必須使用以下 JSON 格式：
{
  "dim_a_score": 0-100,
  "score_rationale": "說明依照上表選擇此分數區間的理由",
  "checks": [
    {
      "check_type": "reference_accuracy | completeness | interpretation | cross_comparison | text_consistency",
      "regulation": "被檢查的法規名稱（TFDA 或 EU_MDR）",
      "issue": "發現的問題描述（若無問題則為 null）",
      "severity": "none | low | medium | high | critical",
      "evidence": "支持證據"
    }
  ],
  "summary": "Dim A 評估摘要（2-3 句話）"
}""",
    "en": """You are a "Regulation Accuracy Audit Expert" for a quality management system. Your task is to compare LLM cross-examination answers against TFDA (Taiwan) and EU MDR (European Union) regulation source texts and evaluate answer accuracy.

The system currently uses only 2 countries' regulations (TFDA + EU_MDR) for cross-examination. MDSAP 5-country verification is not enabled. Please evaluate only against these 2 regulations.

You need to check:
1. **Reference Accuracy**: Are the TFDA / EU MDR references cited by the LLM correct? Do article numbers and content match the source?
2. **Completeness**: Has the LLM missed critical requirements from these 2 regulations?
3. **Interpretation Correctness**: Is the LLM's interpretation of regulations correct? Any distortions or oversimplifications?
4. **Cross-Country Comparison Accuracy**: When the LLM compares TFDA vs EU MDR differences, are the descriptions correct?
5. **Source Text Consistency**: Is the LLM's answer semantically consistent with the regulation's original text?

Scoring Rubric (strictly follow this table — no discretionary scoring):

| Score Range | Criteria |
|-------------|----------|
| 90–100 | Fully accurate citations, all article numbers match source, no omissions, correct cross-country comparisons, accurate interpretation |
| 70–89  | Mostly correct, minor omissions or imprecision, but core requirements covered |
| 50–69  | Partial compliance, 1–2 key requirements missing, or minor misinterpretation |
| 30–49  | Surface compliance only, lacks substantive basis, or significant article errors |
| 0–29   | Non-compliant, hallucinated citations, or unable to provide valid regulatory basis |

Respond in the following JSON format:
{
  "dim_a_score": 0-100,
  "score_rationale": "Explain which band applies and why, citing specific evidence",
  "checks": [
    {
      "check_type": "reference_accuracy | completeness | interpretation | cross_comparison | text_consistency",
      "regulation": "Regulation being checked (TFDA or EU_MDR)",
      "issue": "Description of issue found (null if no issue)",
      "severity": "none | low | medium | high | critical",
      "evidence": "Supporting evidence"
    }
  ],
  "summary": "Dim A assessment summary (2-3 sentences)"
}""",
    "ja": """あなたは品質管理システムの「法規正確性監査専門家」です。あなたの任務は、交差詰問における LLM の回答を TFDA（台湾）および EU MDR（欧州連合）の法規原文と比較し、回答の正確性を評価することです。

現在システムは2カ国の法規（TFDA + EU_MDR）のみを使用しています。MDSAP 5カ国検証は有効化されていません。これら2つの法規に対してのみ評価してください。

確認すべき項目：
1. **法規引用の正確性**: LLM が引用した TFDA / EU MDR の条文は正しいか？
2. **完全性**: LLM はこれら2カ国の法規の重要な要件を見落としていないか？
3. **解釈の正確性**: LLM の法規解釈は正しいか？
4. **国間比較の正確性**: LLM が TFDA と EU MDR の差異を比較する際、記述は正しいか？
5. **原文との一貫性**: LLM の回答は法規原文の意味と一貫しているか？

採点基準（この表に厳密に従うこと）：

| スコア範囲 | 基準 |
|------------|------|
| 90-100 | 引用が完全に正確、全条番号が原文と一致、欠落なし、国間比較が正確 |
| 70-89  | 概ね正確、軽微な欠落があるがコア要件はカバー |
| 50-69  | 部分的準拠、1-2の重要要件が欠落 |
| 30-49  | 表面的準拠のみ、実質的根拠なし |
| 0-29   | 不準拠、幻覚引用 |

以下の JSON 形式で回答してください：
{
  "dim_a_score": 0-100,
  "score_rationale": "該当するスコア区間とその理由の説明",
  "checks": [
    {
      "check_type": "reference_accuracy | completeness | interpretation | cross_comparison | text_consistency",
      "regulation": "確認対象の法規（TFDA または EU_MDR）",
      "issue": "発見された問題の説明（問題がない場合は null）",
      "severity": "none | low | medium | high | critical",
      "evidence": "裏付け証拠"
    }
  ],
  "summary": "Dim A 評価の要約（2-3文）"
}""",
}


# ============================================================
# System prompts — Dimension B (cross-exam quality)
# ============================================================


_DIM_B_SYSTEM_PROMPTS_MDSAP_ON = {
    "zh": """你是品質管理系統的「7國交叉詰問品質評估專家」。你的任務是評估交叉詰問的整體品質。

你需要評估：
1. **國家覆蓋均衡性**: 7國（US, EU, TW, CA, JP, BR, AU）的覆蓋是否均衡？
2. **問題深度**: 問題是否有足夠深度，不是表面性的？
3. **回答品質**: 分析者和驗證者的回答是否有實質內容？
4. **差異識別**: 跨國法規差異是否被正確識別？
5. **同意率合理性**: 同意率是否在合理範圍內（過高過低都有問題）？
6. **可操作性**: 建議和發現是否具有可操作性？

評分標準（必須嚴格依照此表給分，不得自行裁量）：

| 分數區間 | 條件說明 |
|---------|---------|
| 90–100 | 問題深度充分，各國覆蓋均衡，回答有實質內容，差異識別正確，同意率合理（55–85%），建議可操作 |
| 70–89  | 大致符合品質要求，1–2 項評估面向有輕微缺失，整體詰問有效 |
| 50–69  | 問題深度不足或覆蓋不均衡，有明顯缺失但仍有部分有效詰問 |
| 30–49  | 問題流於表面，回答缺乏實質，同意率異常（<30% 或 >95%） |
| 0–29   | 詰問品質嚴重不足，無法識別差異，建議無法操作 |

回答必須使用以下 JSON 格式：
{
  "dim_b_score": 0-100,
  "score_rationale": "說明依照上表選擇此分數區間的理由",
  "country_scores": {
    "US": 0-100, "EU": 0-100, "TW": 0-100,
    "CA": 0-100, "JP": 0-100, "BR": 0-100, "AU": 0-100
  },
  "findings": [
    {
      "category": "coverage_balance | question_depth | answer_quality | gap_identification | agreement_rate | actionability",
      "severity": "low | medium | high | critical",
      "description": "具體描述",
      "recommendation": "建議改善措施"
    }
  ],
  "summary": "Dim B 評估摘要（2-3 句話）"
}""",
    "en": """You are a "7-Country Cross-Examination Quality Assessment Expert" for a quality management system. Your task is to evaluate overall cross-examination quality.

You need to assess:
1. **Country Coverage Balance**: Is coverage balanced across all 7 countries (US, EU, TW, CA, JP, BR, AU)?
2. **Question Depth**: Are questions sufficiently deep, not superficial?
3. **Answer Quality**: Do analyzer and verifier answers have substantive content?
4. **Gap Identification**: Are cross-country regulatory differences correctly identified?
5. **Agreement Rate Reasonableness**: Is the agreement rate within a reasonable range (both too high and too low are problematic)?
6. **Actionability**: Are recommendations and findings actionable?

Scoring Rubric (strictly follow this table — no discretionary scoring):

| Score Range | Criteria |
|-------------|----------|
| 90–100 | Sufficient question depth, balanced country coverage, substantive answers, correct gap identification, reasonable agreement rate (55–85%), actionable recommendations |
| 70–89  | Generally meets quality requirements, 1–2 minor deficiencies, overall effective cross-examination |
| 50–69  | Insufficient question depth or coverage imbalance, notable deficiencies but some effective examination |
| 30–49  | Superficial questions, lacking substantive answers, abnormal agreement rate (<30% or >95%) |
| 0–29   | Severely inadequate cross-examination quality, unable to identify gaps, unactionable recommendations |

Respond in the following JSON format:
{
  "dim_b_score": 0-100,
  "score_rationale": "Explain which band applies and why, citing specific evidence",
  "country_scores": {
    "US": 0-100, "EU": 0-100, "TW": 0-100,
    "CA": 0-100, "JP": 0-100, "BR": 0-100, "AU": 0-100
  },
  "findings": [
    {
      "category": "coverage_balance | question_depth | answer_quality | gap_identification | agreement_rate | actionability",
      "severity": "low | medium | high | critical",
      "description": "Specific description",
      "recommendation": "Recommended improvement"
    }
  ],
  "summary": "Dim B assessment summary (2-3 sentences)"
}""",
    "ja": """あなたは品質管理システムの「7カ国交差詰問品質評価専門家」です。あなたの任務は交差詰問の全体的な品質を評価することです。

評価すべき項目：
1. **国別カバレッジバランス**: 7カ国（US, EU, TW, CA, JP, BR, AU）のカバレッジは均衡しているか？
2. **質問の深さ**: 質問は十分に深いか、表面的でないか？
3. **回答品質**: 分析者と検証者の回答に実質的な内容があるか？
4. **差異識別**: 国間の法規差異が正しく識別されているか？
5. **同意率の妥当性**: 同意率は妥当な範囲内か（高すぎても低すぎても問題）？
6. **実行可能性**: 推奨事項と発見は実行可能か？

採点基準（この表に厳密に従うこと）：

| スコア範囲 | 基準 |
|------------|------|
| 90-100 | 十分な質問深度、均衡な国別カバレッジ、実質的回答、正確な差異識別、妥当な同意率（55-85%）、実行可能な推奨 |
| 70-89  | 概ね品質要件を満たす、1-2の軽微な不足、全体的に効果的な交差詰問 |
| 50-69  | 質問深度不足またはカバレッジ不均衡、顕著な不足あるが部分的に効果的 |
| 30-49  | 表面的な質問、実質のない回答、異常な同意率（<30%または>95%） |
| 0-29   | 交差詰問品質が深刻に不足、差異識別不能、実行不可能な推奨 |

以下の JSON 形式で回答してください：
{
  "dim_b_score": 0-100,
  "score_rationale": "該当するスコア区間とその理由の説明",
  "country_scores": {
    "US": 0-100, "EU": 0-100, "TW": 0-100,
    "CA": 0-100, "JP": 0-100, "BR": 0-100, "AU": 0-100
  },
  "findings": [
    {
      "category": "coverage_balance | question_depth | answer_quality | gap_identification | agreement_rate | actionability",
      "severity": "low | medium | high | critical",
      "description": "具体的な記述",
      "recommendation": "推奨される改善措置"
    }
  ],
  "summary": "Dim B 評価の要約（2-3文）"
}""",
}

_DIM_B_SYSTEM_PROMPTS_MDSAP_OFF = {
    "zh": """你是品質管理系統的「2國交叉詰問品質評估專家」。你的任務是評估 TFDA（台灣）與 EU MDR（歐盟）2 國法規交叉詰問的整體品質。

目前系統僅使用 2 國法規（TFDA + EU_MDR）進行交叉詰問，尚未啟用 MDSAP 5 國驗證。請僅針對這 2 國法規進行品質評估。

你需要評估：
1. **國家覆蓋均衡性**: TW 與 EU 兩國的覆蓋是否均衡？
2. **問題深度**: 問題是否有足夠深度，不是表面性的？
3. **回答品質**: 分析者和驗證者的回答是否有實質內容？
4. **差異識別**: TFDA 與 EU MDR 的法規差異是否被正確識別？
5. **同意率合理性**: 同意率是否在合理範圍內（過高過低都有問題）？
6. **可操作性**: 建議和發現是否具有可操作性？

評分標準（必須嚴格依照此表給分，不得自行裁量）：

| 分數區間 | 條件說明 |
|---------|---------|
| 90–100 | 問題深度充分，TW/EU 覆蓋均衡，回答有實質內容，差異識別正確，同意率合理（55–85%），建議可操作 |
| 70–89  | 大致符合品質要求，1–2 項評估面向有輕微缺失，整體詰問有效 |
| 50–69  | 問題深度不足或覆蓋不均衡，有明顯缺失但仍有部分有效詰問 |
| 30–49  | 問題流於表面，回答缺乏實質，同意率異常（<30% 或 >95%） |
| 0–29   | 詰問品質嚴重不足，無法識別差異，建議無法操作 |

回答必須使用以下 JSON 格式：
{
  "dim_b_score": 0-100,
  "score_rationale": "說明依照上表選擇此分數區間的理由",
  "country_scores": {
    "EU": 0-100, "TW": 0-100
  },
  "findings": [
    {
      "category": "coverage_balance | question_depth | answer_quality | gap_identification | agreement_rate | actionability",
      "severity": "low | medium | high | critical",
      "description": "具體描述",
      "recommendation": "建議改善措施"
    }
  ],
  "summary": "Dim B 評估摘要（2-3 句話）"
}""",
    "en": """You are a "2-Country Cross-Examination Quality Assessment Expert" for a quality management system. Your task is to evaluate the quality of cross-examination using TFDA (Taiwan) and EU MDR (European Union) regulations.

The system currently uses only 2 countries' regulations (TFDA + EU_MDR) for cross-examination. MDSAP 5-country verification is not enabled. Please evaluate quality only for these 2 regulations.

You need to assess:
1. **Country Coverage Balance**: Is coverage balanced between TW and EU?
2. **Question Depth**: Are questions sufficiently deep, not superficial?
3. **Answer Quality**: Do analyzer and verifier answers have substantive content?
4. **Gap Identification**: Are TFDA vs EU MDR regulatory differences correctly identified?
5. **Agreement Rate Reasonableness**: Is the agreement rate within a reasonable range (both too high and too low are problematic)?
6. **Actionability**: Are recommendations and findings actionable?

Scoring Rubric (strictly follow this table — no discretionary scoring):

| Score Range | Criteria |
|-------------|----------|
| 90–100 | Sufficient question depth, balanced TW/EU coverage, substantive answers, correct gap identification, reasonable agreement rate (55–85%), actionable recommendations |
| 70–89  | Generally meets quality requirements, 1–2 minor deficiencies, overall effective cross-examination |
| 50–69  | Insufficient question depth or coverage imbalance, notable deficiencies but some effective examination |
| 30–49  | Superficial questions, lacking substantive answers, abnormal agreement rate (<30% or >95%) |
| 0–29   | Severely inadequate cross-examination quality, unable to identify gaps, unactionable recommendations |

Respond in the following JSON format:
{
  "dim_b_score": 0-100,
  "score_rationale": "Explain which band applies and why, citing specific evidence",
  "country_scores": {
    "EU": 0-100, "TW": 0-100
  },
  "findings": [
    {
      "category": "coverage_balance | question_depth | answer_quality | gap_identification | agreement_rate | actionability",
      "severity": "low | medium | high | critical",
      "description": "Specific description",
      "recommendation": "Recommended improvement"
    }
  ],
  "summary": "Dim B assessment summary (2-3 sentences)"
}""",
    "ja": """あなたは品質管理システムの「2カ国交差詰問品質評価専門家」です。あなたの任務は TFDA（台湾）と EU MDR（欧州連合）の2カ国法規交差詰問の品質を評価することです。

現在システムは2カ国の法規（TFDA + EU_MDR）のみを使用しています。これら2つの法規に対してのみ品質を評価してください。

評価すべき項目：
1. **国別カバレッジバランス**: TW と EU のカバレッジは均衡しているか？
2. **質問の深さ**: 質問は十分に深いか？
3. **回答品質**: 分析者と検証者の回答に実質的な内容があるか？
4. **差異識別**: TFDA と EU MDR の法規差異が正しく識別されているか？
5. **同意率の妥当性**: 同意率は妥当な範囲内か？
6. **実行可能性**: 推奨事項と発見は実行可能か？

採点基準（この表に厳密に従うこと）：

| スコア範囲 | 基準 |
|------------|------|
| 90-100 | 十分な質問深度、均衡な TW/EU カバレッジ、実質的回答、正確な差異識別、妥当な同意率（55-85%） |
| 70-89  | 概ね品質要件を満たす、1-2の軽微な不足 |
| 50-69  | 質問深度不足またはカバレッジ不均衡 |
| 30-49  | 表面的な質問、実質のない回答、異常な同意率 |
| 0-29   | 交差詰問品質が深刻に不足 |

以下の JSON 形式で回答してください：
{
  "dim_b_score": 0-100,
  "score_rationale": "該当するスコア区間とその理由",
  "country_scores": {
    "EU": 0-100, "TW": 0-100
  },
  "findings": [
    {
      "category": "coverage_balance | question_depth | answer_quality | gap_identification | agreement_rate | actionability",
      "severity": "low | medium | high | critical",
      "description": "具体的な記述",
      "recommendation": "推奨される改善措置"
    }
  ],
  "summary": "Dim B 評価の要約（2-3文）"
}""",
}


# ============================================================
# User prompt templates
# ============================================================


_DIM_A_USER_TEMPLATES_MDSAP_ON = {
    "zh": """## MDSAP 法規準確性稽核任務

以下是最近 {record_count} 份交叉詰問記錄中涉及 MDSAP 法規的回答摘要：

### MDSAP 法規參考資料
{mdsap_references}

### 交叉詰問回答樣本
{exam_samples}

請比對交叉詰問回答與 MDSAP 法規原文，給出準確性評分。""",
    "en": """## MDSAP Regulation Accuracy Audit Task

Below are answer summaries from the most recent {record_count} cross-examination records involving MDSAP regulations:

### MDSAP Regulation References
{mdsap_references}

### Cross-Examination Answer Samples
{exam_samples}

Please compare cross-examination answers against MDSAP regulation source texts and provide an accuracy score.""",
    "ja": """## MDSAP 法規正確性監査タスク

以下は最新の {record_count} 件の交差詰問記録における MDSAP 法規に関する回答の要約です：

### MDSAP 法規参考資料
{mdsap_references}

### 交差詰問回答サンプル
{exam_samples}

交差詰問の回答を MDSAP 法規原文と比較し、正確性スコアを付与してください。""",
}

_DIM_A_USER_TEMPLATES_MDSAP_OFF = {
    "zh": """## 法規準確性稽核任務（2國模式：TFDA + EU_MDR）

以下是最近 {record_count} 份交叉詰問記錄中涉及 TFDA（台灣）及 EU MDR（歐盟）法規的回答摘要：

⚠️ 目前僅使用 2 國法規（TFDA + EU_MDR），尚未啟用 MDSAP 5 國驗證。

### 法規參考資料
{mdsap_references}

### 交叉詰問回答樣本
{exam_samples}

請比對交叉詰問回答與 TFDA / EU MDR 法規原文，給出準確性評分。""",
    "en": """## Regulation Accuracy Audit Task (2-Country Mode: TFDA + EU_MDR)

Below are answer summaries from the most recent {record_count} cross-examination records involving TFDA (Taiwan) and EU MDR (European Union) regulations:

⚠️ Currently using only 2 countries' regulations (TFDA + EU_MDR). MDSAP 5-country verification is not enabled.

### Regulation References
{mdsap_references}

### Cross-Examination Answer Samples
{exam_samples}

Please compare cross-examination answers against TFDA / EU MDR regulation source texts and provide an accuracy score.""",
    "ja": """## 法規正確性監査タスク（2カ国モード：TFDA + EU_MDR）

以下は最新の {record_count} 件の交差詰問記録における TFDA（台湾）および EU MDR（欧州連合）法規に関する回答の要約です：

⚠️ 現在2カ国の法規（TFDA + EU_MDR）のみ使用中。MDSAP 5カ国検証は未有効化。

### 法規参考資料
{mdsap_references}

### 交差詰問回答サンプル
{exam_samples}

交差詰問の回答を TFDA / EU MDR 法規原文と比較し、正確性スコアを付与してください。""",
}


_DIM_B_USER_TEMPLATES_MDSAP_ON = {
    "zh": """## 7國交叉詰問品質評估任務

以下是最近 {record_count} 份交叉詰問記錄的統計數據：

### 統計摘要
- 記錄總數: {record_count}
- 時間範圍: {time_range}
- 平均同意率: {avg_agreement_rate:.1%}
- 涵蓋國家: {countries}

### 國家分布
{country_distribution}

### 問題類型分布
{question_type_distribution}

### 樣本記錄
{sample_records}

請評估整體品質並為每個國家評分。""",
    "en": """## 7-Country Cross-Examination Quality Assessment Task

Below are statistics from the most recent {record_count} cross-examination records:

### Statistical Summary
- Total Records: {record_count}
- Time Range: {time_range}
- Average Agreement Rate: {avg_agreement_rate:.1%}
- Countries Covered: {countries}

### Country Distribution
{country_distribution}

### Question Type Distribution
{question_type_distribution}

### Sample Records
{sample_records}

Please evaluate overall quality and score each country.""",
    "ja": """## 7カ国交差詰問品質評価タスク

以下は最新の {record_count} 件の交差詰問記録の統計データです：

### 統計サマリー
- 記録総数: {record_count}
- 時間範囲: {time_range}
- 平均同意率: {avg_agreement_rate:.1%}
- カバー国: {countries}

### 国別分布
{country_distribution}

### 質問タイプ分布
{question_type_distribution}

### サンプル記録
{sample_records}

全体的な品質を評価し、各国にスコアを付けてください。""",
}

_DIM_B_USER_TEMPLATES_MDSAP_OFF = {
    "zh": """## 2國交叉詰問品質評估任務（TFDA + EU_MDR）

以下是最近 {record_count} 份交叉詰問記錄的統計數據：

⚠️ 目前僅使用 2 國法規（TFDA + EU_MDR），尚未啟用 MDSAP 5 國驗證。請僅針對這 2 國法規進行品質評估。

### 統計摘要
- 記錄總數: {record_count}
- 時間範圍: {time_range}
- 平均同意率: {avg_agreement_rate:.1%}
- 涵蓋國家: {countries}

### 國家分布
{country_distribution}

### 問題類型分布
{question_type_distribution}

### 樣本記錄
{sample_records}

請評估整體品質並為每個國家評分。""",
    "en": """## 2-Country Cross-Examination Quality Assessment Task (TFDA + EU_MDR)

Below are statistics from the most recent {record_count} cross-examination records:

⚠️ Currently using only 2 countries' regulations (TFDA + EU_MDR). MDSAP 5-country verification is not enabled. Please evaluate quality only for these 2 regulations.

### Statistical Summary
- Total Records: {record_count}
- Time Range: {time_range}
- Average Agreement Rate: {avg_agreement_rate:.1%}
- Countries Covered: {countries}

### Country Distribution
{country_distribution}

### Question Type Distribution
{question_type_distribution}

### Sample Records
{sample_records}

Please evaluate overall quality and score each country.""",
    "ja": """## 2カ国交差詰問品質評価タスク（TFDA + EU_MDR）

以下は最新の {record_count} 件の交差詰問記録の統計データです：

⚠️ 現在2カ国の法規（TFDA + EU_MDR）のみ使用中。これら2つの法規に対してのみ品質を評価してください。

### 統計サマリー
- 記録総数: {record_count}
- 時間範囲: {time_range}
- 平均同意率: {avg_agreement_rate:.1%}
- カバー国: {countries}

### 国別分布
{country_distribution}

### 質問タイプ分布
{question_type_distribution}

### サンプル記録
{sample_records}

全体的な品質を評価し、各国にスコアを付けてください。""",
}


# ============================================================
# Meta review prompts
# ============================================================


_META_REVIEW_SYSTEM_PROMPTS = {
    "zh": """你是品質管理系統的「10日總檢專家」。你的任務是分析過去10天的每日稽核結果，提供趨勢分析和綜合建議。

你需要分析：
1. **分數趨勢**: Dim A 和 Dim B 分數是上升、持平還是下降？
2. **偏差模式**: 是否有持續性偏差？偏差是否在加劇或改善？
3. **國家表現趨勢**: 各國評分是否有一致趨勢？
4. **改善建議**: 基於10天數據提出具體改善建議。

回答必須使用以下 JSON 格式：
{
  "avg_dim_a": 0-100,
  "avg_dim_b": 0-100,
  "trend_analysis": "趨勢分析文字（3-5 句話）",
  "recommendations": ["建議1", "建議2", ...],
  "deviation_summary": "偏差摘要（若無偏差為 null）",
  "country_trends": {
    "US": "trend description", ...
  }
}""",
    "en": """You are a "10-Day Meta Review Expert" for a quality management system. Your task is to analyze the past 10 days of daily audit results and provide trend analysis and recommendations.

You need to analyze:
1. **Score Trends**: Are Dim A and Dim B scores rising, stable, or declining?
2. **Deviation Patterns**: Are there persistent deviations? Are deviations worsening or improving?
3. **Country Performance Trends**: Are country scores showing consistent trends?
4. **Improvement Recommendations**: Provide specific recommendations based on 10-day data.

Respond in the following JSON format:
{
  "avg_dim_a": 0-100,
  "avg_dim_b": 0-100,
  "trend_analysis": "Trend analysis text (3-5 sentences)",
  "recommendations": ["Recommendation 1", "Recommendation 2", ...],
  "deviation_summary": "Deviation summary (null if no deviations)",
  "country_trends": {
    "US": "trend description", ...
  }
}""",
    "ja": """あなたは品質管理システムの「10日間メタレビュー専門家」です。あなたの任務は過去10日間の毎日監査結果を分析し、傾向分析と総合的な推奨事項を提供することです。

分析すべき項目：
1. **スコア傾向**: Dim A と Dim B のスコアは上昇、安定、低下のいずれか？
2. **偏差パターン**: 持続的な偏差があるか？偏差は悪化しているか改善しているか？
3. **国別パフォーマンス傾向**: 各国のスコアに一貫した傾向があるか？
4. **改善推奨**: 10日間のデータに基づく具体的な推奨事項を提供。

以下の JSON 形式で回答してください：
{
  "avg_dim_a": 0-100,
  "avg_dim_b": 0-100,
  "trend_analysis": "傾向分析テキスト（3-5文）",
  "recommendations": ["推奨1", "推奨2", ...],
  "deviation_summary": "偏差の要約（偏差がない場合は null）",
  "country_trends": {
    "US": "傾向の説明", ...
  }
}""",
}


# ============================================================
# Data classes
# ============================================================


class DailyAuditResult:
    """Result of a daily audit run."""

    def __init__(
        self,
        *,
        audit_id: str = "",
        audit_date: str = "",
        dim_a_score: float = 0.0,
        dim_a_checks: list[dict] | None = None,
        dim_a_summary: str = "",
        dim_a_score_rationale: str = "",
        dim_b_score: float = 0.0,
        dim_b_country_scores: dict[str, float] | None = None,
        dim_b_findings: list[dict] | None = None,
        dim_b_summary: str = "",
        dim_b_score_rationale: str = "",
        overall_score: float = 0.0,
        deviation_detected: bool = False,
        deviation_details: str = "",
        model: str = "",
        usage: dict | None = None,
        timestamp: str = "",
        summary: str = "",
        cross_validation: dict | None = None,
        incomplete_data_warning: bool = False,
        incomplete_countries: list[str] | None = None,
        sampling_details: dict | None = None,
    ):
        self.audit_id = audit_id or f"audit_{int(time.time())}"
        self.audit_date = audit_date or datetime.now().strftime("%Y-%m-%d")
        self.dim_a_score = dim_a_score
        self.dim_a_checks = dim_a_checks or []
        self.dim_a_summary = dim_a_summary
        self.dim_a_score_rationale = dim_a_score_rationale
        self.dim_b_score = dim_b_score
        self.dim_b_country_scores = dim_b_country_scores or {}
        self.dim_b_findings = dim_b_findings or []
        self.dim_b_summary = dim_b_summary
        self.dim_b_score_rationale = dim_b_score_rationale
        self.overall_score = overall_score
        self.deviation_detected = deviation_detected
        self.deviation_details = deviation_details
        self.model = model
        self.usage = usage or {}
        self.timestamp = timestamp or datetime.now().isoformat()
        self.summary = summary
        self.cross_validation = cross_validation or {}
        self.incomplete_data_warning = incomplete_data_warning
        self.incomplete_countries = incomplete_countries or []
        self.sampling_details = sampling_details or {}

    def to_dict(self) -> dict:
        return {
            "audit_id": self.audit_id,
            "audit_date": self.audit_date,
            "dim_a_score": self.dim_a_score,
            "dim_a_checks": self.dim_a_checks,
            "dim_a_summary": self.dim_a_summary,
            "dim_a_score_rationale": self.dim_a_score_rationale,
            "dim_b_score": self.dim_b_score,
            "dim_b_country_scores": self.dim_b_country_scores,
            "dim_b_findings": self.dim_b_findings,
            "dim_b_summary": self.dim_b_summary,
            "dim_b_score_rationale": self.dim_b_score_rationale,
            "overall_score": self.overall_score,
            "deviation_detected": self.deviation_detected,
            "deviation_details": self.deviation_details,
            "model": self.model,
            "usage": self.usage,
            "timestamp": self.timestamp,
            "summary": self.summary,
            "cross_validation": self.cross_validation,
            "incomplete_data_warning": self.incomplete_data_warning,
            "incomplete_countries": self.incomplete_countries,
            "sampling_details": self.sampling_details,
        }

    @classmethod
    def from_dict(cls, data: dict) -> "DailyAuditResult":
        import inspect

        valid_params = set(inspect.signature(cls.__init__).parameters.keys()) - {"self"}
        return cls(**{k: v for k, v in data.items() if k in valid_params})

    def detect_deviation(self) -> bool:
        """Check if this result shows deviation beyond thresholds."""
        if self.overall_score < DEVIATION_OVERALL_THRESHOLD:
            self.deviation_detected = True
            self.deviation_details = (
                f"Overall score {self.overall_score:.0f} is below threshold "
                f"{DEVIATION_OVERALL_THRESHOLD}"
            )
            return True
        dim_gap = abs(self.dim_a_score - self.dim_b_score)
        if dim_gap > DEVIATION_DIM_GAP_THRESHOLD:
            self.deviation_detected = True
            self.deviation_details = (
                f"Dim A ({self.dim_a_score:.0f}) and Dim B ({self.dim_b_score:.0f}) "
                f"gap ({dim_gap:.0f}) exceeds threshold {DEVIATION_DIM_GAP_THRESHOLD}"
            )
            return True
        return False


class MetaReviewResult:
    """Result of a 10-day meta review."""

    def __init__(
        self,
        *,
        review_id: str = "",
        period_start: str = "",
        period_end: str = "",
        daily_results: list[dict] | None = None,
        avg_dim_a: float = 0.0,
        avg_dim_b: float = 0.0,
        trend_analysis: str = "",
        recommendations: list[str] | None = None,
        deviation_summary: str = "",
        country_trends: dict[str, str] | None = None,
        model: str = "",
        usage: dict | None = None,
        timestamp: str = "",
    ):
        self.review_id = review_id or f"meta_{int(time.time())}"
        self.period_start = period_start
        self.period_end = period_end
        self.daily_results = daily_results or []
        self.avg_dim_a = avg_dim_a
        self.avg_dim_b = avg_dim_b
        self.trend_analysis = trend_analysis
        self.recommendations = recommendations or []
        self.deviation_summary = deviation_summary
        self.country_trends = country_trends or {}
        self.model = model
        self.usage = usage or {}
        self.timestamp = timestamp or datetime.now().isoformat()

    def to_dict(self) -> dict:
        return {
            "review_id": self.review_id,
            "period_start": self.period_start,
            "period_end": self.period_end,
            "daily_results": self.daily_results,
            "avg_dim_a": self.avg_dim_a,
            "avg_dim_b": self.avg_dim_b,
            "trend_analysis": self.trend_analysis,
            "recommendations": self.recommendations,
            "deviation_summary": self.deviation_summary,
            "country_trends": self.country_trends,
            "model": self.model,
            "usage": self.usage,
            "timestamp": self.timestamp,
        }

    @classmethod
    def from_dict(cls, data: dict) -> "MetaReviewResult":
        import inspect

        valid_params = set(inspect.signature(cls.__init__).parameters.keys()) - {"self"}
        return cls(**{k: v for k, v in data.items() if k in valid_params})


# ============================================================
# Core: run_daily_sampling_crossexam  (Phase 5 re-execution)
# ============================================================

DAILY_SAMPLE_RATE = 0.2  # 20% of rows sampled per day
DAILY_DEFAULT_REGULATIONS = ["TFDA", "EU_MDR"]  # MDSAP off → 2-country
DAILY_ALL_REGULATIONS = ["TFDA", "QMSR", "EU_MDR", "HC", "PMDA", "ANVISA", "TGA"]
MDSAP_MEMBER_REGULATIONS = ["QMSR", "HC", "PMDA", "ANVISA", "TGA"]


def _find_latest_pipeline_state() -> Optional[Path]:
    """Find the most recent completed PipelineState JSON file."""
    state_dir = Path("data/analysis_pipeline")
    if not state_dir.exists():
        return None
    state_files = sorted(
        state_dir.glob("*.json"), key=lambda p: p.stat().st_mtime, reverse=True
    )
    for sf in state_files:
        try:
            import json as _json

            data = _json.loads(sf.read_text(encoding="utf-8"))
            if data.get("status") == "completed":
                return sf
        except Exception:
            continue
    return None


def run_daily_sampling_crossexam(
    llm_completion_fn: callable,
    model: str = "default",
    mdsap_enabled: bool = False,
    lang: str = "zh-TW",
) -> Optional[DailySamplingRecord]:
    """Third-party QA audit on a 20% random sample of existing Phase 5
    Analyzer/Verifier debate transcripts from the latest completed pipeline.

    Does NOT re-run the debates — reads the existing debate transcripts and
    sends them to an independent third-party LLM auditor to check for:
    - Question reasonableness
    - Answer factual accuracy
    - Hallucination / fabrication
    - Logic consistency

    When mdsap_enabled=True, additionally samples 20% of MDSAP-specific
    clauses (clauses with delta/exceeds items from MDSAP countries).

    Results go into DailyCrossExamStore (separate from pipeline CrossExamStore).
    """
    import math
    import random
    import time as _time

    from src.analysis.state import PipelineState, RowState
    from src.analysis.verifier import run_qa_audit_document
    from src.database.daily_crossexam_store import get_daily_crossexam_store

    state_path = _find_latest_pipeline_state()
    if state_path is None:
        logger.info("No completed pipeline state found for daily sampling cross-exam")
        return None

    state = PipelineState.load(state_path)
    logger.info(
        "Daily sampling cross-exam: loaded pipeline state %s (%d rows)",
        state.run_id,
        len(state.rows),
    )

    all_rows = [RowState.from_dict(r) for r in state.rows.values()]
    rows_with_debates = [
        r for r in all_rows if r.verification_rounds and len(r.verification_rounds) > 0
    ]
    if not rows_with_debates:
        logger.info("No rows with debate transcripts found in pipeline state")
        return None

    if mdsap_enabled:
        selected_regulations = DAILY_ALL_REGULATIONS
        countries = DAILY_ALL_REGULATIONS
    else:
        selected_regulations = DAILY_DEFAULT_REGULATIONS
        countries = DAILY_DEFAULT_REGULATIONS

    sample_count = max(1, math.ceil(len(rows_with_debates) * DAILY_SAMPLE_RATE))
    sampled_rows = random.sample(
        rows_with_debates, min(sample_count, len(rows_with_debates))
    )

    mdsap_sampled_rows: list[RowState] = []
    if mdsap_enabled:
        mdsap_5_ids = {"QMSR", "HC", "PMDA", "ANVISA", "TGA"}
        rows_with_mdsap_context = []
        for r in rows_with_debates:
            qa_data = r.qa_audit or {}
            phase_5_result = r.phase_results.get("phase_5", {})
            regs_in_result = phase_5_result.get("output", {}).get(
                "selected_regulations", []
            )
            if any(reg in mdsap_5_ids for reg in regs_in_result):
                rows_with_mdsap_context.append(r)

        if not rows_with_mdsap_context:
            rows_with_mdsap_context = rows_with_debates

        mdsap_sample_count = max(
            1, math.ceil(len(rows_with_mdsap_context) * DAILY_SAMPLE_RATE)
        )
        mdsap_sampled_rows = random.sample(
            rows_with_mdsap_context,
            min(mdsap_sample_count, len(rows_with_mdsap_context)),
        )
        existing_ids = {r.row_id for r in sampled_rows}
        for mr in mdsap_sampled_rows:
            if mr.row_id not in existing_ids:
                sampled_rows.append(mr)
                existing_ids.add(mr.row_id)

    logger.info(
        "Daily sampling: %d/%d rows (%.0f%%) + %d MDSAP rows, regulations=%s, mdsap=%s",
        sample_count,
        len(rows_with_debates),
        DAILY_SAMPLE_RATE * 100,
        len(mdsap_sampled_rows),
        selected_regulations,
        mdsap_enabled,
    )

    start_time = _time.time()

    doc_groups: dict[str, list[RowState]] = {}
    for row in sampled_rows:
        doc_groups.setdefault(row.doc_id or "unknown", []).append(row)

    total_usage = {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}
    clause_results: list[dict] = []

    for doc_id, doc_rows in doc_groups.items():
        try:
            qa_result = run_qa_audit_document(
                doc_id=doc_id,
                rows=doc_rows,
                state=state,
                llm_completion_fn=llm_completion_fn,
                model=model,
                selected_regulations=selected_regulations,
                run_id=f"daily_{datetime.now().strftime('%Y%m%d')}",
            )

            for audit_item in qa_result.get("clause_audits", []):
                cid = audit_item.get("clause_id", "")
                matched_row = next((r for r in doc_rows if r.clause_id == cid), None)
                original_rounds = []
                clause_title = ""
                audit_question = ""
                verdict = ""
                if matched_row:
                    original_rounds = matched_row.verification_rounds or []
                    clause_title = getattr(matched_row, "clause_title", "")
                    _checklist = _get_checklist("ISO_13485")
                    _clause_info = _checklist.get(cid, {})
                    audit_question = _get_audit_question(_clause_info, lang=lang) if _clause_info else getattr(matched_row, "audit_question", "")
                    verdict = getattr(matched_row, "verdict", "")

                clause_results.append(
                    {
                        "clause_id": cid,
                        "doc_id": doc_id,
                        "clause_title": clause_title,
                        "audit_question": audit_question,
                        "verdict": verdict,
                        "score": audit_item.get("score", 0),
                        "question_quality": audit_item.get(
                            "question_quality", "unknown"
                        ),
                        "answer_accuracy": audit_item.get("answer_accuracy", "unknown"),
                        "hallucination_detected": audit_item.get(
                            "hallucination_detected", False
                        ),
                        "hallucination_details": audit_item.get(
                            "hallucination_details", ""
                        ),
                        "logic_consistency": audit_item.get(
                            "logic_consistency", "unknown"
                        ),
                        "depth_sufficient": audit_item.get("depth_sufficient", True),
                        "conclusion_reasonable": audit_item.get(
                            "conclusion_reasonable", True
                        ),
                        "issues": audit_item.get("issues", []),
                        "agreed": bool(matched_row and matched_row.verification_agreed),
                        "flagged": bool(matched_row and matched_row.flagged_for_ra),
                        "rounds": original_rounds,
                    }
                )

            qa_summary = qa_result.get("summary", "")
            qa_recommendations = qa_result.get("recommendations", [])
            if qa_summary or qa_recommendations:
                clause_results.append(
                    {
                        "_qa_doc_summary": True,
                        "doc_id": doc_id,
                        "summary": qa_summary,
                        "recommendations": qa_recommendations,
                        "overall_score": qa_result.get("overall_score", 0),
                    }
                )

            qa_usage = qa_result.get("llm_usage", {})
            for k in total_usage:
                total_usage[k] += qa_usage.get(k, 0)

        except Exception as e:
            logger.error("Daily sampling QA audit failed for doc %s: %s", doc_id, e)

    duration = _time.time() - start_time

    total_agreed = sum(1 for c in clause_results if c.get("agreed"))
    total_flagged = sum(1 for c in clause_results if c.get("flagged"))

    record = DailySamplingRecord(
        source_run_id=state.run_id,
        selected_regulations=selected_regulations,
        countries=countries,
        mdsap_enabled=mdsap_enabled,
        sample_rate=DAILY_SAMPLE_RATE,
        total_rows_available=len(rows_with_debates),
        sampled_row_ids=[r.row_id for r in sampled_rows],
        clauses=clause_results,
        total_clauses=len(clause_results),
        total_agreed=total_agreed,
        total_flagged=total_flagged,
        total_rounds=0,
        questions_used=[],
        llm_usage=total_usage,
        llm_model=model,
        duration_seconds=duration,
        lang=lang,
    )

    daily_store = get_daily_crossexam_store()
    daily_store.save_record(record)

    logger.info(
        "Daily sampling QA audit complete: %d clauses audited, %.1fs",
        record.total_clauses,
        duration,
    )
    return record


# ============================================================
# Core: run_daily_audit
# ============================================================


def run_daily_audit(
    llm_completion_fn: callable,
    model: str = "default",
    temperature: float = 0.3,
    max_tokens: int = 8192,
    lang: str = "zh-TW",
    store=None,
    feedback_context: str = "",
    incomplete_countries: list[str] | None = None,
    mdsap_enabled: bool = False,
) -> DailyAuditResult:
    """Run daily audit with Dim A (MDSAP accuracy) + Dim B (cross-exam quality).

    Now reads from DailyCrossExamStore (daily sampling cross-exam records)
    instead of CrossExamStore (pipeline full cross-exam records).

    Args:
        llm_completion_fn: LLM completion function
        model: LLM model name
        temperature: LLM temperature
        max_tokens: Max response tokens
        lang: Language code
        store: DailyCrossExamStore instance (uses singleton if not provided)
        feedback_context: Optional user feedback context to append to prompts
        incomplete_countries: List of country/regulation IDs with incomplete data.
            When non-empty, the result is annotated with a warning.

    Returns:
        DailyAuditResult with scores and findings
    """
    from src.database.daily_crossexam_store import get_daily_crossexam_store

    if store is None:
        store = get_daily_crossexam_store()

    records = store.get_all_records()
    if not records:
        return DailyAuditResult(
            summary="No cross-examination records available for audit.",
            dim_a_summary="No records.",
            dim_b_summary="No records.",
        )

    _lang_key = _get_prompt_lang(lang)
    _incomplete = incomplete_countries or []

    latest_record = records[0]
    clauses = latest_record.clauses or []
    clause_scores = [c.get("score", 0) for c in clauses if c.get("score") is not None]
    hallucination_count = sum(1 for c in clauses if c.get("hallucination_detected"))
    sampling_details = {
        "source_run_id": latest_record.source_run_id,
        "mdsap_enabled": latest_record.mdsap_enabled,
        "selected_regulations": latest_record.selected_regulations,
        "sample_rate": latest_record.sample_rate,
        "total_rows_available": latest_record.total_rows_available,
        "sampled_count": len(latest_record.sampled_row_ids),
        "clauses": clauses,
        "clauses_audited": len(clause_scores),
        "avg_qa_score": sum(clause_scores) / len(clause_scores) if clause_scores else 0,
        "hallucinations_found": hallucination_count,
    }

    result = DailyAuditResult(
        incomplete_data_warning=bool(_incomplete),
        incomplete_countries=_incomplete,
        sampling_details=sampling_details,
    )

    # ---- Dimension A + B: run in parallel ----
    from concurrent.futures import ThreadPoolExecutor, as_completed

    def _exec_dim_a() -> dict:
        return _run_dim_a(
            llm_completion_fn,
            records,
            model,
            temperature,
            max_tokens,
            _lang_key,
            feedback_context=feedback_context,
            mdsap_enabled=mdsap_enabled,
        )

    def _exec_dim_b() -> dict:
        return _run_dim_b(
            llm_completion_fn,
            records,
            store,
            model,
            temperature,
            max_tokens,
            _lang_key,
            feedback_context=feedback_context,
            mdsap_enabled=mdsap_enabled,
        )

    dim_a_result: dict | None = None
    dim_b_result: dict | None = None
    dim_a_err: Exception | None = None
    dim_b_err: Exception | None = None

    with ThreadPoolExecutor(max_workers=2) as executor:
        future_a = executor.submit(_exec_dim_a)
        future_b = executor.submit(_exec_dim_b)
        try:
            dim_a_result = future_a.result()
        except Exception as e:
            dim_a_err = e
        try:
            dim_b_result = future_b.result()
        except Exception as e:
            dim_b_err = e

    if dim_a_result is not None:
        result.dim_a_score = dim_a_result.get("dim_a_score", 0.0)
        result.dim_a_checks = dim_a_result.get("checks", [])
        result.dim_a_summary = dim_a_result.get("summary", "")
        result.dim_a_score_rationale = dim_a_result.get("score_rationale", "")
    if dim_a_err is not None:
        logger.error("Dim A audit failed: %s", dim_a_err)
        result.dim_a_summary = f"Dim A audit failed: {str(dim_a_err)[:200]}"

    if dim_b_result is not None:
        result.dim_b_score = dim_b_result.get("dim_b_score", 0.0)
        result.dim_b_country_scores = dim_b_result.get("country_scores", {})
        result.dim_b_findings = dim_b_result.get("findings", [])
        result.dim_b_summary = dim_b_result.get("summary", "")
        result.dim_b_score_rationale = dim_b_result.get("score_rationale", "")
    if dim_b_err is not None:
        logger.error("Dim B audit failed: %s", dim_b_err)
        result.dim_b_summary = f"Dim B audit failed: {str(dim_b_err)[:200]}"

    # ---- Overall score ----
    result.overall_score = (result.dim_a_score + result.dim_b_score) / 2

    # ---- Cross-validation: 7-country vs 5-country (MDSAP) ----
    try:
        result.cross_validation = _run_cross_validation(records, result)
    except Exception as e:
        logger.error(f"Cross-validation failed: {e}")
        result.cross_validation = {"error": str(e)[:200]}

    _warning_suffix = ""
    if result.incomplete_data_warning:
        _warning_suffix = f" ⚠️ 基於不完整資料 (Based on incomplete data: {', '.join(result.incomplete_countries)})"
    result.summary = (
        f"Dim A: {result.dim_a_score:.0f}, Dim B: {result.dim_b_score:.0f}, "
        f"Overall: {result.overall_score:.0f}{_warning_suffix}"
    )

    # ---- Deviation detection ----
    result.detect_deviation()

    # ---- Save to disk ----
    _save_daily_audit(result)

    logger.info(
        "Daily audit complete: overall=%.0f, dimA=%.0f, dimB=%.0f, deviation=%s",
        result.overall_score,
        result.dim_a_score,
        result.dim_b_score,
        result.deviation_detected,
    )
    return result


def _run_dim_a(
    llm_completion_fn: callable,
    records: list,
    model: str,
    temperature: float,
    max_tokens: int,
    lang_key: str,
    feedback_context: str = "",
    mdsap_enabled: bool = False,
) -> dict:
    """Run Dimension A audit: regulation accuracy check.

    When mdsap_enabled=True  → 7-country MDSAP prompts
    When mdsap_enabled=False → 2-country (TFDA + EU_MDR) prompts
    """
    mdsap_refs = _build_mdsap_reference_context()
    exam_samples = _build_exam_samples(records[:5])

    feedback_section = ""
    if feedback_context:
        feedback_section = (
            f"\n\n### User Feedback Context\n"
            f"The following feedback was provided by the user. Please incorporate it into your evaluation:\n"
            f"{feedback_context}"
        )

    sys_prompts = (
        _DIM_A_SYSTEM_PROMPTS_MDSAP_ON
        if mdsap_enabled
        else _DIM_A_SYSTEM_PROMPTS_MDSAP_OFF
    )
    usr_templates = (
        _DIM_A_USER_TEMPLATES_MDSAP_ON
        if mdsap_enabled
        else _DIM_A_USER_TEMPLATES_MDSAP_OFF
    )

    user_prompt = (
        usr_templates[lang_key].format(
            record_count=len(records),
            mdsap_references=mdsap_refs or "  (No references available)",
            exam_samples=exam_samples or "  (No exam samples available)",
        )
        + feedback_section
    )

    messages = [
        {"role": "system", "content": sys_prompts[lang_key]},
        {"role": "user", "content": user_prompt},
    ]

    response = llm_completion_fn(
        messages=messages,
        model=model,
        temperature=temperature,
        max_tokens=max_tokens,
        stream=False,
    )

    response_text = response.get("content", "")
    parsed = _parse_json_response(response_text)
    return parsed


def _run_dim_b(
    llm_completion_fn: callable,
    records: list,
    store,
    model: str,
    temperature: float,
    max_tokens: int,
    lang_key: str,
    feedback_context: str = "",
    mdsap_enabled: bool = False,
) -> dict:
    """Run Dimension B audit: cross-exam quality verification.

    When mdsap_enabled=True  → MDSAP 5-country verification prompts (7-country coverage)
    When mdsap_enabled=False → 2-country (TFDA + EU_MDR) quality prompts with advisory note
    """
    country_dist = store.get_country_distribution()
    qtype_dist = store.get_question_type_distribution()

    total_clauses = sum(r.total_clauses for r in records)
    total_agreed = sum(r.total_agreed for r in records)
    avg_agreement = total_agreed / max(total_clauses, 1)

    time_range = (
        f"{records[-1].timestamp} ~ {records[0].timestamp}" if records else "N/A"
    )

    countries = (
        ", ".join(
            sorted(
                set(c for r in records for c in getattr(r, "selected_regulations", []))
            )
        )
        or "N/A"
    )

    country_text = "\n".join(
        f"  - {c}: {n} records"
        for c, n in sorted(country_dist.items(), key=lambda x: -x[1])
    )
    qtype_text = "\n".join(
        f"  - {t}: {n} records"
        for t, n in sorted(qtype_dist.items(), key=lambda x: -x[1])
    )

    sample_text = _build_exam_samples(records[:3])

    feedback_section = ""
    if feedback_context:
        feedback_section = (
            f"\n\n### User Feedback Context\n"
            f"The following feedback was provided by the user. Please incorporate it into your evaluation:\n"
            f"{feedback_context}"
        )

    sys_prompts = (
        _DIM_B_SYSTEM_PROMPTS_MDSAP_ON
        if mdsap_enabled
        else _DIM_B_SYSTEM_PROMPTS_MDSAP_OFF
    )
    usr_templates = (
        _DIM_B_USER_TEMPLATES_MDSAP_ON
        if mdsap_enabled
        else _DIM_B_USER_TEMPLATES_MDSAP_OFF
    )

    user_prompt = (
        usr_templates[lang_key].format(
            record_count=len(records),
            time_range=time_range,
            avg_agreement_rate=avg_agreement,
            countries=countries,
            country_distribution=country_text or "  (No country data)",
            question_type_distribution=qtype_text or "  (No question type data)",
            sample_records=sample_text or "  (No sample records)",
        )
        + feedback_section
    )

    messages = [
        {"role": "system", "content": sys_prompts[lang_key]},
        {"role": "user", "content": user_prompt},
    ]

    response = llm_completion_fn(
        messages=messages,
        model=model,
        temperature=temperature,
        max_tokens=max_tokens,
        stream=False,
    )

    response_text = response.get("content", "")
    parsed = _parse_json_response(response_text)
    return parsed


def _run_cross_validation(
    records: list,
    result: "DailyAuditResult",
) -> dict:
    """Cross-validate audit quality via two independent methods.

    Method 1 – Temporal Drift:
        Compare today's scores against the rolling 30-day historical average
        stored in DailyCrossExamStore.  Detects whether the LLM is becoming
        progressively more lenient or strict over time.

    Method 2 – Country Imbalance:
        Inspect dim_b_country_scores for the current run.  If one country's
        score is far below the group average it may indicate systematic
        weakness in that country's regulatory question quality.

    Args:
        records: Unused (kept for API compatibility).
        result:  The DailyAuditResult already populated with dim scores and
                 dim_b_country_scores.

    Returns:
        dict with keys:
          - temporal_drift  : {status, history_count, rolling_avg_*, today_*, delta_*}
          - country_imbalance: {status, country_scores, avg_score, spread, outlier_countries}
          - overall_assessment: 'normal' | 'monitor' | 'action_required' | 'insufficient_data'
    """
    from datetime import datetime, timedelta

    from src.database.daily_crossexam_store import get_daily_crossexam_store

    cross_val: dict = {}

    # ------------------------------------------------------------------
    # Method 1: Temporal drift – rolling 30-day average
    # ------------------------------------------------------------------
    today_str = datetime.now().strftime("%Y-%m-%d")
    cutoff_str = (datetime.now() - timedelta(days=30)).strftime("%Y-%m-%d")

    store = get_daily_crossexam_store()
    historical = [
        r
        for r in store.get_all_records()
        if r.date < today_str and r.date >= cutoff_str and r.overall_score > 0
    ]

    if len(historical) < 3:
        cross_val["temporal_drift"] = {
            "status": "insufficient_data",
            "reason": f"歷史記錄僅 {len(historical)} 筆（需至少 3 筆）",
            "history_count": len(historical),
        }
    else:
        n = len(historical)
        rolling_overall = sum(r.overall_score for r in historical) / n
        rolling_dim_a = sum(r.dim_a_score for r in historical) / n
        rolling_dim_b = sum(r.dim_b_score for r in historical) / n

        delta_overall = result.overall_score - rolling_overall
        delta_dim_a = result.dim_a_score - rolling_dim_a
        delta_dim_b = result.dim_b_score - rolling_dim_b

        abs_delta = abs(delta_overall)
        if abs_delta <= 5.0:
            drift_status = "stable"
        elif abs_delta <= 15.0:
            drift_status = "minor_drift"
        else:
            drift_status = "significant_drift"

        cross_val["temporal_drift"] = {
            "status": drift_status,
            "history_count": n,
            "rolling_window_days": 30,
            "rolling_avg_overall": round(rolling_overall, 2),
            "rolling_avg_dim_a": round(rolling_dim_a, 2),
            "rolling_avg_dim_b": round(rolling_dim_b, 2),
            "today_overall": round(result.overall_score, 2),
            "today_dim_a": round(result.dim_a_score, 2),
            "today_dim_b": round(result.dim_b_score, 2),
            "delta_overall": round(delta_overall, 2),
            "delta_dim_a": round(delta_dim_a, 2),
            "delta_dim_b": round(delta_dim_b, 2),
        }

    # ------------------------------------------------------------------
    # Method 2: Country imbalance – per-country Dim B scores
    # ------------------------------------------------------------------
    country_scores = result.dim_b_country_scores or {}

    if len(country_scores) < 2:
        cross_val["country_imbalance"] = {
            "status": "insufficient_data",
            "reason": f"只有 {len(country_scores)} 個國家分數（需至少 2 個）",
            "country_count": len(country_scores),
        }
    else:
        scores_list = list(country_scores.values())
        avg_score = sum(scores_list) / len(scores_list)
        min_score = min(scores_list)
        max_score = max(scores_list)
        spread = max_score - min_score

        # Countries more than 15 points below the group average are outliers
        outlier_threshold = avg_score - 15
        outlier_countries = {
            c: round(s, 2)
            for c, s in country_scores.items()
            if s < outlier_threshold
        }

        if spread <= 10:
            imbalance_status = "balanced"
        elif spread <= 25:
            imbalance_status = "minor_imbalance"
        else:
            imbalance_status = "significant_imbalance"

        cross_val["country_imbalance"] = {
            "status": imbalance_status,
            "country_count": len(country_scores),
            "country_scores": {c: round(s, 2) for c, s in country_scores.items()},
            "avg_score": round(avg_score, 2),
            "min_score": round(min_score, 2),
            "max_score": round(max_score, 2),
            "spread": round(spread, 2),
            "outlier_countries": outlier_countries,
        }

    # ------------------------------------------------------------------
    # Overall assessment
    # ------------------------------------------------------------------
    drift_status = cross_val.get("temporal_drift", {}).get("status", "insufficient_data")
    imbalance_status = cross_val.get("country_imbalance", {}).get("status", "insufficient_data")

    if drift_status == "significant_drift" or imbalance_status == "significant_imbalance":
        overall_assessment = "action_required"
    elif drift_status in ("minor_drift", "stable") or imbalance_status in (
        "minor_imbalance",
        "balanced",
    ):
        if drift_status == "minor_drift" or imbalance_status == "minor_imbalance":
            overall_assessment = "monitor"
        else:
            overall_assessment = "normal"
    elif drift_status == "insufficient_data" and imbalance_status == "insufficient_data":
        overall_assessment = "insufficient_data"
    else:
        overall_assessment = "normal"

    cross_val["overall_assessment"] = overall_assessment

    logger.info(
        "Cross-validation: temporal=%s (delta=%.1f), country=%s (spread=%.1f) → %s",
        drift_status,
        cross_val.get("temporal_drift", {}).get("delta_overall", 0.0),
        imbalance_status,
        cross_val.get("country_imbalance", {}).get("spread", 0.0),
        overall_assessment,
    )
    return cross_val


# ============================================================
# Core: run_10day_meta_review
# ============================================================


def run_10day_meta_review(
    llm_completion_fn: callable,
    model: str = "default",
    temperature: float = 0.3,
    max_tokens: int = 8192,
    lang: str = "zh-TW",
    feedback_context: str = "",
) -> MetaReviewResult:
    """Run 10-day meta review of daily audit results.

    Aggregates up to 10 most recent daily audit results and analyzes trends.

    Args:
        llm_completion_fn: LLM completion function
        model: LLM model name
        temperature: LLM temperature
        max_tokens: Max response tokens
        lang: Language code
        feedback_context: Optional user feedback context to append to prompts

    Returns:
        MetaReviewResult with trend analysis and recommendations
    """
    _lang_key = _get_prompt_lang(lang)

    # Load recent daily audit results
    daily_results = get_daily_audit_history(limit=10)
    if not daily_results:
        return MetaReviewResult(
            trend_analysis="No daily audit results available for meta review.",
        )

    # Build context for LLM
    results_summary = []
    for r in daily_results:
        results_summary.append(
            f"- {r.audit_date}: DimA={r.dim_a_score:.0f}, DimB={r.dim_b_score:.0f}, "
            f"Overall={r.overall_score:.0f}, Deviation={'Yes' if r.deviation_detected else 'No'}"
        )

    period_start = daily_results[-1].audit_date if daily_results else ""
    period_end = daily_results[0].audit_date if daily_results else ""

    feedback_section = ""
    if feedback_context:
        feedback_section = (
            f"\n\n### User Feedback Context\n"
            f"The following feedback was provided by the user. Please incorporate it into your analysis:\n"
            f"{feedback_context}\n"
        )

    user_prompt = (
        f"## 10-Day Meta Review\n\n"
        f"Period: {period_start} to {period_end}\n"
        f"Number of daily audits: {len(daily_results)}\n\n"
        f"### Daily Results (newest first):\n"
        + "\n".join(results_summary)
        + feedback_section
        + "\n\nPlease analyze trends and provide recommendations."
    )

    messages = [
        {"role": "system", "content": _META_REVIEW_SYSTEM_PROMPTS[_lang_key]},
        {"role": "user", "content": user_prompt},
    ]

    result = MetaReviewResult(
        period_start=period_start,
        period_end=period_end,
        daily_results=[r.to_dict() for r in daily_results],
    )

    try:
        response = llm_completion_fn(
            messages=messages,
            model=model,
            temperature=temperature,
            max_tokens=max_tokens,
            stream=False,
        )

        response_text = response.get("content", "")
        usage = response.get("usage", {})
        result.model = response.get("model", model)
        result.usage = usage

        parsed = _parse_json_response(response_text)
        result.avg_dim_a = parsed.get("avg_dim_a", 0.0)
        result.avg_dim_b = parsed.get("avg_dim_b", 0.0)
        result.trend_analysis = parsed.get("trend_analysis", "")
        result.recommendations = parsed.get("recommendations", [])
        result.deviation_summary = parsed.get("deviation_summary", "")
        result.country_trends = parsed.get("country_trends", {})

    except Exception as e:
        logger.error(f"Meta review failed: {e}")
        result.trend_analysis = f"Meta review failed: {str(e)[:200]}"

    # Save to disk
    _save_meta_review(result)

    logger.info(
        "Meta review complete: avgA=%.0f, avgB=%.0f, deviations=%s",
        result.avg_dim_a,
        result.avg_dim_b,
        bool(result.deviation_summary),
    )
    return result


# ============================================================
# Helper: build context for LLM
# ============================================================


def _build_mdsap_reference_context() -> str:
    """Build MDSAP regulation reference context from predefined profiles.

    Enhanced version: includes unique requirements with original regulation
    text for third-party LLM validation (Dim A). The richer context enables
    the auditing LLM to compare cross-examination answers against actual
    regulation source text.
    """
    try:
        from src.analysis.compliance_rules import PREDEFINED_REGULATIONS

        mdsap_keys = ["HC", "PMDA", "ANVISA", "TGA", "QMSR"]
        parts = []
        for key in mdsap_keys:
            profile = PREDEFINED_REGULATIONS.get(key)
            if not profile:
                continue
            parts.append(
                f"### {profile.name_en} ({profile.country})\n"
                f"- Regulation ID: {profile.regulation_id}\n"
                f"- Source: {profile.source_url}\n"
                f"- Last Updated: {profile.last_updated}\n"
                f"- Mapped clauses: {len(profile.iso_mapped)}\n"
                f"- Unique requirements: {len(profile.unique_requirements)}\n"
            )

            # Include sample clause mappings (up to 5)
            sample_clauses = list(profile.iso_mapped.items())[:5]
            for clause_id, mapping in sample_clauses:
                parts.append(
                    f"  - {clause_id}: {mapping.regulation_ref} "
                    f"[{mapping.status.value}] — {mapping.rationale_en[:150]}...\n"
                )

            # Include ALL unique requirements with original text
            # (critical for third-party validation against actual regulation)
            if profile.unique_requirements:
                parts.append("\n  **Unique Requirements (delta from ISO 13485):**\n")
                for ureq in profile.unique_requirements:
                    parts.append(
                        f"  [{ureq.req_id}] {ureq.title_en} ({ureq.regulation_ref})\n"
                        f"    Impact: {ureq.audit_impact}\n"
                        f"    Requirement: {ureq.requirement_en}\n"
                    )
                    if ureq.original_text:
                        parts.append(
                            f"    Original text ({ureq.original_lang}): "
                            f"{ureq.original_text[:300]}\n"
                        )
                    if ureq.semantic_note:
                        parts.append(
                            f"    Cross-country note: {ureq.semantic_note[:200]}...\n"
                        )
                    parts.append("\n")

        return "\n".join(parts)
    except Exception as e:
        logger.warning(f"Failed to build MDSAP reference context: {e}")
        return ""


def _build_exam_samples(records: list) -> str:
    """Build exam sample text from records."""
    parts = []
    for r in records:
        parts.append(
            f"--- Record {r.record_id} ({r.timestamp}) ---\n"
            f"Regulations: {', '.join(getattr(r, 'selected_regulations', []))}\n"
            f"Clauses: {r.total_clauses}, Agreed: {r.total_agreed}, "
            f"Flagged: {r.total_flagged}\n"
        )
        for clause in getattr(r, "clauses", [])[:2]:
            parts.append(
                f"  Clause {clause.get('clause_id', '')}: "
                f"agreed={clause.get('agreed')}, "
                f"rounds={len(clause.get('rounds', []))}\n"
            )
            for rd in clause.get("rounds", [])[:1]:
                analyzer = rd.get("analyzer", {})
                verifier = rd.get("verifier", {})
                parts.append(
                    f"    Analyzer: confidence={analyzer.get('confidence', 'N/A')}, "
                    f"position={str(analyzer.get('position', ''))[:200]}\n"
                    f"    Verifier: agreement={verifier.get('agreement_level', 'N/A')}\n"
                )
    return "".join(parts)


# ============================================================
# Persistence helpers
# ============================================================


def _save_daily_audit(result: DailyAuditResult) -> None:
    """Save daily audit result to disk."""
    AUDIT_DIR.mkdir(parents=True, exist_ok=True)
    ts_suffix = result.timestamp[11:19].replace(":", "-") if len(result.timestamp) >= 19 else datetime.now().strftime("%H-%M-%S")
    filepath = AUDIT_DIR / f"daily_{result.audit_date}_{ts_suffix}.json"
    atomic_write_json(filepath, result.to_dict())


def _save_meta_review(result: MetaReviewResult) -> None:
    """Save meta review result to disk."""
    AUDIT_DIR.mkdir(parents=True, exist_ok=True)
    period_key = result.period_end or result.timestamp[:10]
    ts_suffix = result.timestamp[11:19].replace(":", "-") if len(result.timestamp) >= 19 else datetime.now().strftime("%H-%M-%S")
    filepath = AUDIT_DIR / f"meta_review_{period_key}_{ts_suffix}.json"
    atomic_write_json(filepath, result.to_dict())


def get_daily_audit_history(limit: int = 30) -> list[DailyAuditResult]:
    """Load daily audit history from disk, most recent first.

    Args:
        limit: Maximum number of records to return

    Returns:
        List of DailyAuditResult sorted by date descending
    """
    AUDIT_DIR.mkdir(parents=True, exist_ok=True)
    files = sorted(AUDIT_DIR.glob("daily_*.json"), reverse=True)

    results = []
    for f in files[:limit]:
        try:
            with open(f, "r", encoding="utf-8") as fh:
                data = json.load(fh)
            results.append(DailyAuditResult.from_dict(data))
        except (json.JSONDecodeError, OSError) as e:
            logger.warning(f"Failed to load audit file {f}: {e}")
    return results


def get_latest_meta_review() -> Optional[MetaReviewResult]:
    """Load the most recent meta review result from disk."""
    AUDIT_DIR.mkdir(parents=True, exist_ok=True)
    files = sorted(AUDIT_DIR.glob("meta_review_*.json"), reverse=True)
    if not files:
        return None

    try:
        with open(files[0], "r", encoding="utf-8") as f:
            data = json.load(f)
        return MetaReviewResult.from_dict(data)
    except (json.JSONDecodeError, OSError):
        return None


# ============================================================
# Response parsing
# ============================================================


def _parse_json_response(response_text: str) -> dict:
    """Parse LLM JSON response, handling code fences, raw JSON, and truncated output.

    Fallback strategy:
    1. Extract from ```json ... ``` code fence
    2. Direct json.loads on stripped text
    3. Find outermost { ... } and parse
    4. Try to extract score values from raw text with regex
    5. Return summary with score 0 only as last resort
    """
    import re

    text = response_text.strip()

    # Try to extract JSON from code fence (greedy to handle large blocks)
    json_match = re.search(r"```(?:json)?\s*\n?(.*?)\n?```", text, re.DOTALL)
    if json_match:
        extracted = json_match.group(1).strip()
        try:
            parsed = json.loads(extracted)
            if isinstance(parsed, dict):
                return parsed
        except json.JSONDecodeError:
            # Code fence found but JSON invalid (possibly truncated)
            text = extracted  # Still use extracted content for further parsing

    # Try direct parse
    try:
        parsed = json.loads(text)
        if isinstance(parsed, dict):
            return parsed
    except json.JSONDecodeError:
        pass

    # Try to find JSON object in the text
    brace_start = text.find("{")
    brace_end = text.rfind("}")
    if brace_start >= 0 and brace_end > brace_start:
        try:
            parsed = json.loads(text[brace_start : brace_end + 1])
            if isinstance(parsed, dict):
                return parsed
        except json.JSONDecodeError:
            pass

    # Fallback: try to extract scores from raw text with regex
    # This catches cases where JSON is malformed but scores are visible
    fallback = {"summary": text[:500]}
    dim_a_match = re.search(r'"dim_a_score"\s*:\s*(\d+(?:\.\d+)?)', text)
    dim_b_match = re.search(r'"dim_b_score"\s*:\s*(\d+(?:\.\d+)?)', text)
    if dim_a_match:
        fallback["dim_a_score"] = float(dim_a_match.group(1))
    else:
        fallback["dim_a_score"] = 0.0
    if dim_b_match:
        fallback["dim_b_score"] = float(dim_b_match.group(1))
    else:
        fallback["dim_b_score"] = 0.0

    logger.warning(
        "JSON parse fallback: extracted dim_a=%.0f, dim_b=%.0f from raw text",
        fallback["dim_a_score"],
        fallback["dim_b_score"],
    )
    return fallback


# ============================================================
# Adjustment Guidance Generator
# ============================================================


def _generate_adjustment_guidance(
    result: DailyAuditResult,
    lang: str = "zh-TW",
) -> tuple[list[str], list[str]]:
    """Generate fixed steps and dynamic recommendations for human adjustment.

    Returns:
        (fixed_steps, dynamic_recommendations)
        - fixed_steps: Generic steps on how to use Eira's feedback mechanism
        - dynamic_recommendations: Specific advice derived from current findings
    """
    lk = _get_prompt_lang(lang)

    # ---- Fixed steps (always the same) ----
    _FIXED_STEPS = {
        "zh": [
            (
                "檢視 Regulation Profile 設定 / Review Regulation Profile settings — "
                "確認各國法規 Profile（TFDA, EU_MDR, QMSR 等）的內容是否完整且為最新版本。"
                "如有法規更新，請至 Eira「法規清單更新」功能更新法規資料。"
            ),
            (
                "確認 MDSAP 模式是否正確 / Verify MDSAP mode — "
                "若貴組織取得或規劃 MDSAP 認證，請確保系統已啟用 MDSAP 模式"
                "（7 國法規交叉詰問）。若僅針對台灣與歐盟市場，使用 2 國模式即可。"
            ),
            (
                "利用 Eira 回饋功能 / Use Eira's Feedback feature — "
                "在每日稽核報告中，對有疑慮的項目提供回饋（同意/不同意/部分同意），"
                "系統將記錄您的意見並納入後續分析改善。"
            ),
            (
                "檢查交叉詰問深度 / Review cross-examination depth — "
                "若交叉詰問輪次不足或過早結束，可能導致品質評分偏低。"
                "建議檢查 Verifier prompt 設定，確保詰問深度足夠。"
            ),
            (
                "定期檢視 10 日總檢報告 / Review 10-Day Meta Review — "
                "透過 10 日總檢報告追蹤趨勢變化，觀察各維度分數是否穩定改善。"
                "若持續偏低，考慮調整 LLM 模型或法規資料。"
            ),
        ],
        "en": [
            (
                "Review Regulation Profile settings — "
                "Ensure all country Regulation Profiles (TFDA, EU_MDR, QMSR, etc.) are complete and up-to-date. "
                "If regulations have been updated, use Eira's 'Regulatory List Update' feature to refresh regulation data."
            ),
            (
                "Verify MDSAP mode — "
                "If your organization holds or plans MDSAP certification, ensure MDSAP mode is enabled "
                "(7-country cross-examination). For Taiwan and EU markets only, 2-country mode is sufficient."
            ),
            (
                "Use Eira's Feedback feature — "
                "In daily audit reports, provide feedback (agree/disagree/partial) on items of concern. "
                "The system will record your input and incorporate it into future analysis improvements."
            ),
            (
                "Review cross-examination depth — "
                "If cross-examination rounds are insufficient or end prematurely, quality scores may be low. "
                "Review Verifier prompt settings to ensure adequate examination depth."
            ),
            (
                "Review 10-Day Meta Review regularly — "
                "Track trend changes through 10-day meta review reports to see if dimension scores are steadily improving. "
                "If persistently low, consider adjusting the LLM model or regulation data."
            ),
        ],
        "ja": [
            (
                "Regulation Profile 設定を確認 — "
                "各国の法規 Profile（TFDA, EU_MDR, QMSR 等）の内容が完全で最新版であることを確認してください。"
                "法規が更新された場合は、Eira の「法規リスト更新」機能で法規データを更新してください。"
            ),
            (
                "MDSAP モードを確認 — "
                "組織が MDSAP 認証を取得または計画している場合、MDSAP モードが有効になっていることを確認してください"
                "（7カ国法規交差詰問）。台湾と EU 市場のみの場合、2カ国モードで十分です。"
            ),
            (
                "Eira のフィードバック機能を活用 — "
                "毎日の監査レポートで、懸念のある項目にフィードバック（同意/不同意/部分同意）を提供してください。"
                "システムがご意見を記録し、今後の分析改善に反映します。"
            ),
            (
                "交差詰問の深さを確認 — "
                "交差詰問のラウンドが不足している、または早期に終了する場合、品質スコアが低くなる可能性があります。"
                "Verifier プロンプト設定を確認し、十分な詰問深度を確保してください。"
            ),
            (
                "10日間メタレビューを定期的に確認 — "
                "10日間メタレビューレポートで傾向変化を追跡し、各ディメンションのスコアが着実に改善しているか確認してください。"
                "持続的に低い場合は、LLM モデルまたは法規データの調整を検討してください。"
            ),
        ],
    }
    fixed_steps = _FIXED_STEPS.get(lk, _FIXED_STEPS["en"])

    # ---- Dynamic recommendations (based on current findings) ----
    dynamic_recs: list[str] = []
    sd = result.sampling_details or {}
    is_mdsap = sd.get("mdsap_enabled", False)

    # Localized message templates
    _msg = {
        "zh": {
            "overall_low": "⚠️ 整體評分 {score:.0f} 低於閾值 {threshold} — 建議立即檢視稽核發現事項，確認是否有系統性問題需處理。",
            "dim_a_critical": "🔴 法規準確性 (Dim A) 有 {count} 項 critical/high 嚴重度問題 — 建議優先檢查法規引用是否正確，更新 Regulation Profile 中的條文內容。",
            "completeness": "📋 法規完整性不足 ({count} 項) — LLM 可能遺漏關鍵法規要求。建議豐富 Regulation Profile 的內容，補充容易被遺漏的條文。",
            "interpretation": "📖 法規解釋偏差 ({count} 項) — LLM 對法規的解釋可能過度簡化或曲解。建議在 Regulation Profile 中加入關鍵條文的正確解釋說明。",
            "dim_b_critical": "🔴 交叉詰問品質 (Dim B) 有 {count} 項 critical/high 問題 — 建議檢查 Analyzer/Verifier 的 prompt 設定，確保詰問品質。",
            "dim_gap": "📊 Dim A 與 Dim B 分數差距達 {gap:.0f} 分 — Dimension {lower_dim} 明顯較弱，建議針對該維度重點改善。",
            "dim_a_label": "A (法規準確性)",
            "dim_b_label": "B (交叉詰問品質)",
            "deviation": "⚠️ 偏差已偵測 — {details}",
            "deviation_default": "詳見偏差詳情",
            "country_imbalance": "🌍 各國評分差距大 (最高 {max:.0f}, 最低 {min:.0f}) — 較弱國家: {countries}。建議補強這些國家的法規 Profile 資料。",
            "flagged": "🚩 {count} 條 ISO 13485 條款被標記 — 包含: {ids}{ellipsis}。建議檢視這些條款的文件是否完整且符合最新法規要求。",
            "incomplete": "📂 以下國家/法規資料不完整: {countries} — 部分評估結果可能不準確。建議補充缺少的法規資料。",
            "all_good": "✅ 本次稽核未發現需要人為介入的明顯問題。建議持續定期檢視以保持品質。",
        },
        "en": {
            "overall_low": "⚠️ Overall score {score:.0f} is below threshold {threshold} — Review audit findings immediately to check for systemic issues.",
            "dim_a_critical": "🔴 Regulation Accuracy (Dim A) has {count} critical/high severity issues — Prioritize checking regulation references and update Regulation Profile content.",
            "completeness": "📋 Regulation completeness gaps ({count} items) — LLM may have missed critical regulatory requirements. Enrich Regulation Profile content.",
            "interpretation": "📖 Regulation interpretation deviations ({count} items) — LLM interpretation may be oversimplified or distorted. Add correct interpretation notes to Regulation Profile.",
            "dim_b_critical": "🔴 Cross-Examination Quality (Dim B) has {count} critical/high issues — Review Analyzer/Verifier prompt settings to ensure examination quality.",
            "dim_gap": "📊 Dim A and Dim B score gap is {gap:.0f} points — Dimension {lower_dim} is notably weaker; focus improvement on that dimension.",
            "dim_a_label": "A (Regulation Accuracy)",
            "dim_b_label": "B (Cross-Exam Quality)",
            "deviation": "⚠️ Deviation detected — {details}",
            "deviation_default": "See deviation details",
            "country_imbalance": "🌍 Large country score gap (max {max:.0f}, min {min:.0f}) — Weaker countries: {countries}. Strengthen Regulation Profile data for these countries.",
            "flagged": "🚩 {count} ISO 13485 clauses flagged — Including: {ids}{ellipsis}. Review whether documents for these clauses are complete and current.",
            "incomplete": "📂 Incomplete data for: {countries} — Some assessment results may be inaccurate. Supplement missing regulation data.",
            "all_good": "✅ No significant issues requiring human intervention found in this audit. Continue regular reviews to maintain quality.",
        },
        "ja": {
            "overall_low": "⚠️ 全体スコア {score:.0f} が閾値 {threshold} を下回っています — 監査発見事項を直ちに確認し、体系的な問題がないか確認してください。",
            "dim_a_critical": "🔴 法規正確性 (Dim A) に {count} 件の critical/high 重大度の問題があります — 法規引用が正しいか優先的に確認し、Regulation Profile の内容を更新してください。",
            "completeness": "📋 法規完全性の不足 ({count} 件) — LLM が重要な法規要件を見落としている可能性があります。Regulation Profile の内容を充実させてください。",
            "interpretation": "📖 法規解釈の偏差 ({count} 件) — LLM の法規解釈が過度に簡略化または歪曲されている可能性があります。Regulation Profile に正確な解釈説明を追加してください。",
            "dim_b_critical": "🔴 交差詰問品質 (Dim B) に {count} 件の critical/high 問題があります — Analyzer/Verifier のプロンプト設定を確認してください。",
            "dim_gap": "📊 Dim A と Dim B のスコア差が {gap:.0f} 点です — Dimension {lower_dim} が著しく弱いため、そのディメンションの重点改善を推奨します。",
            "dim_a_label": "A (法規正確性)",
            "dim_b_label": "B (交差詰問品質)",
            "deviation": "⚠️ 偏差検出 — {details}",
            "deviation_default": "偏差詳細を参照",
            "country_imbalance": "🌍 国別スコアの差が大きい (最高 {max:.0f}, 最低 {min:.0f}) — 弱い国: {countries}。これらの国の Regulation Profile データを強化してください。",
            "flagged": "🚩 {count} 件の ISO 13485 条項がフラグ付き — 含む: {ids}{ellipsis}。これらの条項の文書が完全で最新の法規要件に適合しているか確認してください。",
            "incomplete": "📂 以下の国/法規のデータが不完全: {countries} — 一部の評価結果が不正確な可能性があります。不足している法規データを補充してください。",
            "all_good": "✅ 今回の監査では人的介入を要する重大な問題は発見されませんでした。品質維持のため定期的な確認を続けてください。",
        },
    }
    m = _msg.get(lk, _msg["en"])

    # 1. Overall score too low
    if result.overall_score < DEVIATION_OVERALL_THRESHOLD:
        dynamic_recs.append(
            m["overall_low"].format(score=result.overall_score, threshold=DEVIATION_OVERALL_THRESHOLD)
        )

    # 2. Dim A specific issues
    if result.dim_a_checks:
        severity_counts: dict[str, int] = {}
        for check in result.dim_a_checks:
            sev = check.get("severity", "none")
            if sev not in ("none",):
                severity_counts[sev] = severity_counts.get(sev, 0) + 1

        critical_high = severity_counts.get("critical", 0) + severity_counts.get(
            "high", 0
        )
        if critical_high > 0:
            dynamic_recs.append(m["dim_a_critical"].format(count=critical_high))

        completeness_issues = sum(
            1
            for c in result.dim_a_checks
            if c.get("check_type") == "completeness"
            and c.get("severity") not in ("none",)
        )
        if completeness_issues > 0:
            dynamic_recs.append(m["completeness"].format(count=completeness_issues))

        interpretation_issues = sum(
            1
            for c in result.dim_a_checks
            if c.get("check_type") == "interpretation"
            and c.get("severity") not in ("none",)
        )
        if interpretation_issues > 0:
            dynamic_recs.append(m["interpretation"].format(count=interpretation_issues))

    # 3. Dim B specific issues
    if result.dim_b_findings:
        b_critical_high = sum(
            1
            for f in result.dim_b_findings
            if f.get("severity") in ("critical", "high")
        )
        if b_critical_high > 0:
            dynamic_recs.append(m["dim_b_critical"].format(count=b_critical_high))

    # 4. Large gap between Dim A and Dim B
    dim_gap = abs(result.dim_a_score - result.dim_b_score)
    if dim_gap > DEVIATION_DIM_GAP_THRESHOLD:
        lower_dim = (
            m["dim_a_label"]
            if result.dim_a_score < result.dim_b_score
            else m["dim_b_label"]
        )
        dynamic_recs.append(m["dim_gap"].format(gap=dim_gap, lower_dim=lower_dim))

    # 5. Deviation detected
    if result.deviation_detected:
        dynamic_recs.append(
            m["deviation"].format(
                details=result.deviation_details[:200] if result.deviation_details else m["deviation_default"]
            )
        )

    # 6. Country score imbalance (MDSAP mode)
    if is_mdsap and result.dim_b_country_scores:
        scores = list(result.dim_b_country_scores.values())
        if scores:
            min_score = min(scores)
            max_score = max(scores)
            if max_score - min_score > 25:
                weak_countries = [
                    c
                    for c, s in result.dim_b_country_scores.items()
                    if s < min_score + 10
                ]
                dynamic_recs.append(
                    m["country_imbalance"].format(
                        max=max_score, min=min_score,
                        countries=", ".join(weak_countries),
                    )
                )

    # 7. Clauses with flagged items
    clauses = sd.get("clauses", [])
    flagged_clauses = [c for c in clauses if c.get("flagged")]
    if flagged_clauses:
        flagged_ids = [c.get("clause_id", "?") for c in flagged_clauses[:5]]
        dynamic_recs.append(
            m["flagged"].format(
                count=len(flagged_clauses),
                ids=", ".join(flagged_ids),
                ellipsis="..." if len(flagged_clauses) > 5 else "",
            )
        )

    # 8. Incomplete data
    if result.incomplete_data_warning:
        dynamic_recs.append(
            m["incomplete"].format(countries=", ".join(result.incomplete_countries))
        )

    # If no dynamic issues found, add a positive note
    if not dynamic_recs:
        dynamic_recs.append(m["all_good"])

    return fixed_steps, dynamic_recs


# ============================================================
# Export: Word / Excel
# ============================================================


def export_daily_audit_word(result: DailyAuditResult, lang: str = "zh-TW") -> Path:
    """Export a daily audit result as a Word document.

    Args:
        result: DailyAuditResult to export
        lang: Language code for report content

    Returns:
        Path to the generated .docx file
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    filepath = EXPORT_DIR / f"daily_audit_{result.audit_date}.docx"

    lk = _get_prompt_lang(lang)

    # ---- Localized labels ----
    _L = {
        "zh": {
            "title": "AI-QMS 每日稽核報告",
            "meta": "稽核 ID: {aid}  |  日期: {date}  |  時間: {ts}",
            "incomplete_banner": "⚠️ 基於不完整資料\n以下國家/法規資料不完整: {countries}",
            "mode_mdsap": "MDSAP 7國 (全球)",
            "mode_2country": "2國 (TFDA + EU_MDR)",
            "audit_mode": "稽核模式: {mode}",
            "abbrev_title": "縮寫說明",
            "abbrev_body": (
                "Dim A      — Dimension A：法規準確性評分（0-100），衡量 AI 引用法規條文的精確程度\n"
                "Dim B      — Dimension B：交叉詰問品質評分（0-100），衡量 Analyzer/Verifier 辯論流程品質\n"
                "RA         — 法規事務（Regulatory Affairs）；標記為需 RA 人員審查的高風險條款\n"
                "QA Auditor — 第三方品質稽核員角色，由 AI 模擬獨立第三方對辯論記錄進行評分\n"
                "Analyzer   — 分析者角色：針對法規條款分析文件是否符合，提出立場與證據\n"
                "Verifier   — 驗證者角色：質疑 Analyzer 的論點，提出反證或挑戰\n"
                "Agreement Level — 辯論結論：agree（同意）/ partial（部分同意）/ disagree（不同意）\n"
                "MDSAP      — 醫療器材單一稽核計畫（美國/加拿大/巴西/澳洲/日本 5 國）\n"
                "TFDA       — 臺灣食品藥物管理署；EU MDR — 歐盟醫療器材法規 2017/745\n"
                "ISO 13485  — 醫療器材品質管理系統國際標準（71 條稽核問題的依據）\n"
                "score_rationale — AI 對評分依據的文字說明，解釋落在哪個分數區間及原因"
            ),
            "how_title": "作用原理",
            "how_body": (
                "每日交叉詰問稽核採用三層架構：\n\n"
                "【第一層：Phase 5 交叉詰問】\n"
                "  Analyzer 分析每份 QMS 文件是否符合 ISO 13485 / TFDA / EU MDR 條款，\n"
                "  Verifier 逐條質疑 Analyzer 的論點，雙方進行多輪辯論直到達成共識或標記為爭議。\n\n"
                "【第二層：每日抽樣稽核（本報告）】\n"
                "  從當天所有辯論記錄中隨機抽取 20% 條款，\n"
                "  QA Auditor 以第三方視角評分辯論品質（Dim A：法規準確性，Dim B：詰問品質），\n"
                "  並偵測時間軸分數偏差與國家評分失衡。\n\n"
                "【第三層：10 日 Meta Review】\n"
                "  累積 10 次以上每日稽核後，進行趨勢分析，偵測 overfitting / 系統性偏差。\n\n"
                "【問題輪替機制】\n"
                "  稽核問題以當天日期為 seed 確定性輪替。\n"
                "  預期答案（expected_evidence）為靜態清單，定義各條款應提供的書面證據種類。"
            ),
            "scoring_title": "評分說明",
            "scoring_body": (
                "風險等級說明：\n"
                "🔴 immediate_correction（立即矯正）= 重大不符合，須立即採取行動\n"
                "🟠 deadline_correction（限期矯正）= 不符合，須在指定期限內完成矯正\n"
                "🟡 improvement_plan（改善計畫）= 有缺口，須制定改善計畫\n"
                "🟢 suggested_improvement（建議改善）= 輕微缺失，建議改善\n"
                "✅ compliant（符合）= 符合法規要求\n\n"
                "Dim A 評分（法規準確性）：\n"
                "90-100 引用精確無誤 | 70-89 輕微遺漏 | 50-69 部分符合 | 30-49 表面符合 | 0-29 完全不符\n\n"
                "Dim B 評分（詰問品質）：\n"
                "90-100 深度均衡可操作 | 70-89 輕微缺失 | 50-69 明顯缺失 | 30-49 表面流於形式 | 0-29 嚴重不足"
            ),
            "score_summary": "評分摘要",
            "dim_a_label": "法規準確性",
            "dim_b_label": "交叉詰問品質驗證",
            "regs_label": "比對法規",
            "deviation_title": "偏差詳情",
            "sampling_title": "20% 抽樣明細",
            "source_pipeline": "來源分析",
            "mode_label": "模式",
            "sample_rate": "抽樣率",
            "available_rows": "可用列數",
            "sampled_rows": "抽樣列數",
            "per_clause_title": "逐條分析結果",
            "col_clause": "條號", "col_doc": "文件", "col_question": "稽核問題",
            "col_verdict": "判定", "col_agreed": "同意", "col_qa_score": "QA 分數",
            "col_q_quality": "問題品質", "col_hallucination": "幻覺",
            "debate_title": "逐條辯論紀錄與第三方稽核",
            "audit_question": "稽核問題",
            "expected_evidence": "預期書面證據",
            "analyzer": "分析者 (Analyzer)",
            "position": "立場",
            "evidence": "證據",
            "verifier": "驗證者 (Verifier)",
            "challenge": "質疑",
            "comment": "評語",
            "no_rounds": "（無辯論輪次紀錄）",
            "qa_result": "第三方稽核結果: {score}/100 | 問題品質: {qq} | 回答正確性: {aa} | 邏輯一致性: {lc}",
            "hallucination_detect": "🚨 幻覺偵測",
            "qa_summary_title": "第三方稽核摘要",
            "doc_score": "文件 {doc_id} — 總分: {score}/100",
            "dim_a_mdsap": "Dimension A — MDSAP 法規準確性",
            "dim_a_2c": "Dimension A — 法規準確性 (TFDA + EU_MDR)",
            "score_rationale": "評分依據",
            "check_items": "檢查項目",
            "dim_b_mdsap": "Dimension B — MDSAP 5國交叉詰問品質驗證",
            "dim_b_2c": "Dimension B — 2國交叉詰問品質驗證",
            "country_scores": "各國評分",
            "findings": "發現事項",
            "crossval_title": "交叉驗證",
            "td_stable": "穩定 ✅", "td_minor": "輕微漂移 ⚠️",
            "td_significant": "顯著漂移 🔴", "td_insufficient": "資料不足",
            "method1_title": "方法一：時間軸偏差（30天滾動平均）",
            "insufficient_reason": "資料不足，無法比較",
            "status": "狀態", "history_count": "歷史筆數",
            "rolling_avg": "30天滾動均分", "today_score": "今日分數", "delta": "差距",
            "ci_balanced": "均衡 ✅", "ci_minor": "輕微失衡 ⚠️",
            "ci_significant": "顯著失衡 🔴", "ci_insufficient": "資料不足",
            "method2_title": "方法二：國家間分數失衡",
            "country_count": "國家數", "avg_score": "平均分",
            "max_score": "最高", "min_score": "最低", "spread": "落差",
            "outlier_countries": "異常國家（低於均分15分以上）",
            "none": "無",
            "overall_normal": "正常 ✅", "overall_monitor": "需持續觀察 ⚠️",
            "overall_action": "需立即處理 🔴", "overall_insufficient": "資料不足",
            "overall_assessment": "整體評估",
            "adjust_title": "人為調整指引",
            "general_steps_title": "一般改善步驟",
            "dynamic_recs_title": "本次動態建議",
            "records_suffix": "筆（過去30天）",
        },
        "en": {
            "title": "AI-QMS Daily Audit Report",
            "meta": "Audit ID: {aid}  |  Date: {date}  |  Time: {ts}",
            "incomplete_banner": "⚠️ Based on Incomplete Data\nIncomplete data for: {countries}",
            "mode_mdsap": "MDSAP 7-Country (Global)",
            "mode_2country": "2-Country (TFDA + EU_MDR)",
            "audit_mode": "Audit Mode: {mode}",
            "abbrev_title": "Abbreviation Legend",
            "abbrev_body": (
                "Dim A      — Dimension A: Regulation Accuracy Score (0-100), measures precision of AI regulatory citations\n"
                "Dim B      — Dimension B: Cross-Examination Quality Score (0-100), measures Analyzer/Verifier debate quality\n"
                "RA         — Regulatory Affairs; high-risk clauses flagged for RA personnel review\n"
                "QA Auditor — Third-party quality auditor role, AI-simulated independent third-party scoring of debate records\n"
                "Analyzer   — Analyzer role: analyzes document compliance against regulatory clauses, presents position and evidence\n"
                "Verifier   — Verifier role: challenges Analyzer's arguments, presents counter-evidence\n"
                "Agreement Level — Debate conclusion: agree / partial / disagree\n"
                "MDSAP      — Medical Device Single Audit Program (US/CA/BR/AU/JP)\n"
                "TFDA       — Taiwan FDA; EU MDR — EU Medical Device Regulation 2017/745\n"
                "ISO 13485  — International standard for medical device QMS (basis of 71 audit questions)\n"
                "score_rationale — AI's textual explanation of scoring basis, explaining which score band applies and why"
            ),
            "how_title": "How This Report Works",
            "how_body": (
                "Daily cross-examination audit uses a three-layer architecture:\n\n"
                "[Layer 1: Phase 5 Cross-Examination]\n"
                "  Analyzer reviews each QMS document against ISO 13485 / TFDA / EU MDR clauses.\n"
                "  Verifier challenges Analyzer's arguments clause-by-clause through multiple rounds until consensus or flagged.\n\n"
                "[Layer 2: Daily Sampling Audit (this report)]\n"
                "  Randomly samples 20% of clauses from the day's debate records.\n"
                "  QA Auditor scores debate quality from a third-party perspective (Dim A: Regulation Accuracy, Dim B: Exam Quality).\n\n"
                "[Layer 3: 10-Day Meta Review]\n"
                "  After 10+ daily audits, performs trend analysis to detect overfitting / systematic bias.\n\n"
                "[Question Rotation]\n"
                "  Audit questions rotate deterministically using the date as seed.\n"
                "  Expected evidence is a static checklist defining required documentary evidence per clause."
            ),
            "scoring_title": "Scoring Legend",
            "scoring_body": (
                "Risk Level Legend:\n"
                "🔴 immediate_correction = Major nonconformity, immediate action required\n"
                "🟠 deadline_correction = Nonconformity, must be corrected within deadline\n"
                "🟡 improvement_plan = Gap found, improvement plan required\n"
                "🟢 suggested_improvement = Minor deficiency, improvement suggested\n"
                "✅ compliant = Meets regulatory requirements\n\n"
                "Dim A Scoring (Regulation Accuracy):\n"
                "90-100 Precise citations | 70-89 Minor omissions | 50-69 Partial compliance | 30-49 Surface compliance | 0-29 Non-compliant\n\n"
                "Dim B Scoring (Examination Quality):\n"
                "90-100 Deep, balanced, actionable | 70-89 Minor gaps | 50-69 Notable gaps | 30-49 Superficial | 0-29 Severely inadequate"
            ),
            "score_summary": "Score Summary",
            "dim_a_label": "Regulation Accuracy",
            "dim_b_label": "Cross-Exam Quality",
            "regs_label": "Regulations",
            "deviation_title": "Deviation Details",
            "sampling_title": "20% Sampling Details",
            "source_pipeline": "Source Pipeline",
            "mode_label": "Mode",
            "sample_rate": "Sample Rate",
            "available_rows": "Available Rows",
            "sampled_rows": "Sampled Rows",
            "per_clause_title": "Per-Clause Results",
            "col_clause": "Clause", "col_doc": "Doc", "col_question": "Audit Question",
            "col_verdict": "Verdict", "col_agreed": "Agreed", "col_qa_score": "QA Score",
            "col_q_quality": "Q. Quality", "col_hallucination": "Hallucination",
            "debate_title": "Debate & QA Audit Details",
            "audit_question": "Audit Question",
            "expected_evidence": "Expected Evidence",
            "analyzer": "Analyzer",
            "position": "Position",
            "evidence": "Evidence",
            "verifier": "Verifier",
            "challenge": "Challenge",
            "comment": "Assessment",
            "no_rounds": "(No debate rounds recorded)",
            "qa_result": "QA Audit Result: {score}/100 | Q. Quality: {qq} | Answer Accuracy: {aa} | Logic Consistency: {lc}",
            "hallucination_detect": "🚨 Hallucination Detected",
            "qa_summary_title": "QA Audit Summary",
            "doc_score": "Document {doc_id} — Score: {score}/100",
            "dim_a_mdsap": "Dimension A — MDSAP Regulation Accuracy",
            "dim_a_2c": "Dimension A — Regulation Accuracy (TFDA + EU_MDR)",
            "score_rationale": "Score Rationale",
            "check_items": "Check Items",
            "dim_b_mdsap": "Dimension B — MDSAP 5-Country Cross-Exam Quality",
            "dim_b_2c": "Dimension B — 2-Country Cross-Exam Quality",
            "country_scores": "Country Scores",
            "findings": "Findings",
            "crossval_title": "Cross-Validation",
            "td_stable": "Stable ✅", "td_minor": "Minor Drift ⚠️",
            "td_significant": "Significant Drift 🔴", "td_insufficient": "Insufficient Data",
            "method1_title": "Method 1: Temporal Drift (30-Day Rolling Average)",
            "insufficient_reason": "Insufficient data for comparison",
            "status": "Status", "history_count": "History Count",
            "rolling_avg": "30-Day Rolling Avg", "today_score": "Today's Score", "delta": "Delta",
            "ci_balanced": "Balanced ✅", "ci_minor": "Minor Imbalance ⚠️",
            "ci_significant": "Significant Imbalance 🔴", "ci_insufficient": "Insufficient Data",
            "method2_title": "Method 2: Country Score Imbalance",
            "country_count": "Countries", "avg_score": "Avg Score",
            "max_score": "Max", "min_score": "Min", "spread": "Spread",
            "outlier_countries": "Outlier Countries (>15 pts below avg)",
            "none": "None",
            "overall_normal": "Normal ✅", "overall_monitor": "Monitor ⚠️",
            "overall_action": "Action Required 🔴", "overall_insufficient": "Insufficient Data",
            "overall_assessment": "Overall Assessment",
            "adjust_title": "How to Adjust",
            "general_steps_title": "General Improvement Steps",
            "dynamic_recs_title": "Dynamic Recommendations",
            "records_suffix": " records (past 30 days)",
        },
        "ja": {
            "title": "AI-QMS 毎日監査レポート",
            "meta": "監査 ID: {aid}  |  日付: {date}  |  時刻: {ts}",
            "incomplete_banner": "⚠️ 不完全なデータに基づく\n不完全な国/法規データ: {countries}",
            "mode_mdsap": "MDSAP 7カ国（グローバル）",
            "mode_2country": "2カ国（TFDA + EU_MDR）",
            "audit_mode": "監査モード: {mode}",
            "abbrev_title": "略語説明",
            "abbrev_body": (
                "Dim A      — Dimension A：法規正確性スコア（0-100）、AI の法規引用の精度を測定\n"
                "Dim B      — Dimension B：交差詰問品質スコア（0-100）、Analyzer/Verifier 議論品質を測定\n"
                "RA         — 法規事務（Regulatory Affairs）；RA 担当者レビューが必要な高リスク条項\n"
                "QA Auditor — 第三者品質監査役、AI が独立第三者として議論記録を採点\n"
                "Analyzer   — 分析者：法規条項に対する文書適合性を分析、立場と証拠を提示\n"
                "Verifier   — 検証者：Analyzer の議論に異議を唱え、反証を提示\n"
                "Agreement Level — 議論結論：agree / partial / disagree\n"
                "MDSAP      — 医療機器単一監査プログラム（米/加/伯/豪/日 5カ国）\n"
                "TFDA       — 台湾 FDA；EU MDR — EU 医療機器規則 2017/745\n"
                "ISO 13485  — 医療機器 QMS 国際規格（71監査質問の基準）\n"
                "score_rationale — AI のスコア根拠の説明"
            ),
            "how_title": "このレポートの仕組み",
            "how_body": (
                "毎日交差詰問監査は三層アーキテクチャを採用：\n\n"
                "【第1層：Phase 5 交差詰問】\n"
                "  Analyzer が各 QMS 文書を ISO 13485 / TFDA / EU MDR 条項と照合分析。\n"
                "  Verifier が Analyzer の議論を条項ごとに質問し、合意またはフラグまで複数ラウンド議論。\n\n"
                "【第2層：毎日サンプリング監査（本レポート）】\n"
                "  当日の議論記録から20%の条項をランダムサンプリング。\n"
                "  QA Auditor が第三者視点で議論品質を採点。\n\n"
                "【第3層：10日間メタレビュー】\n"
                "  10回以上の毎日監査後、傾向分析を実施。\n\n"
                "【質問ローテーション】\n"
                "  監査質問は日付をシードとして決定的にローテーション。"
            ),
            "scoring_title": "採点説明",
            "scoring_body": (
                "リスクレベル：\n"
                "🔴 immediate_correction = 重大不適合、即時対応必要\n"
                "🟠 deadline_correction = 不適合、期限内是正必要\n"
                "🟡 improvement_plan = ギャップあり、改善計画必要\n"
                "🟢 suggested_improvement = 軽微な不足、改善推奨\n"
                "✅ compliant = 法規要件適合\n\n"
                "Dim A 採点（法規正確性）：\n"
                "90-100 正確な引用 | 70-89 軽微な欠落 | 50-69 部分適合 | 30-49 表面適合 | 0-29 不適合\n\n"
                "Dim B 採点（詰問品質）：\n"
                "90-100 深く均衡で実行可能 | 70-89 軽微な不足 | 50-69 顕著な不足 | 30-49 表面的 | 0-29 深刻に不足"
            ),
            "score_summary": "スコアサマリー",
            "dim_a_label": "法規正確性",
            "dim_b_label": "交差詰問品質",
            "regs_label": "比較法規",
            "deviation_title": "偏差詳細",
            "sampling_title": "20% サンプリング詳細",
            "source_pipeline": "ソースパイプライン",
            "mode_label": "モード",
            "sample_rate": "サンプリング率",
            "available_rows": "利用可能行数",
            "sampled_rows": "サンプリング行数",
            "per_clause_title": "条項別分析結果",
            "col_clause": "条項", "col_doc": "文書", "col_question": "監査質問",
            "col_verdict": "判定", "col_agreed": "同意", "col_qa_score": "QA スコア",
            "col_q_quality": "質問品質", "col_hallucination": "幻覚",
            "debate_title": "議論記録と第三者監査詳細",
            "audit_question": "監査質問",
            "expected_evidence": "予想される書面証拠",
            "analyzer": "分析者 (Analyzer)",
            "position": "立場",
            "evidence": "証拠",
            "verifier": "検証者 (Verifier)",
            "challenge": "質問",
            "comment": "評価",
            "no_rounds": "（議論ラウンド記録なし）",
            "qa_result": "第三者監査結果: {score}/100 | 質問品質: {qq} | 回答正確性: {aa} | 論理一貫性: {lc}",
            "hallucination_detect": "🚨 幻覚検出",
            "qa_summary_title": "第三者監査サマリー",
            "doc_score": "文書 {doc_id} — スコア: {score}/100",
            "dim_a_mdsap": "Dimension A — MDSAP 法規正確性",
            "dim_a_2c": "Dimension A — 法規正確性 (TFDA + EU_MDR)",
            "score_rationale": "スコア根拠",
            "check_items": "チェック項目",
            "dim_b_mdsap": "Dimension B — MDSAP 5カ国交差詰問品質",
            "dim_b_2c": "Dimension B — 2カ国交差詰問品質",
            "country_scores": "国別スコア",
            "findings": "発見事項",
            "crossval_title": "交差検証",
            "td_stable": "安定 ✅", "td_minor": "軽微なドリフト ⚠️",
            "td_significant": "顕著なドリフト 🔴", "td_insufficient": "データ不足",
            "method1_title": "方法1：時間軸偏差（30日間ローリング平均）",
            "insufficient_reason": "データ不足、比較不可",
            "status": "ステータス", "history_count": "履歴件数",
            "rolling_avg": "30日間ローリング平均", "today_score": "本日スコア", "delta": "差分",
            "ci_balanced": "均衡 ✅", "ci_minor": "軽微な不均衡 ⚠️",
            "ci_significant": "顕著な不均衡 🔴", "ci_insufficient": "データ不足",
            "method2_title": "方法2：国間スコア不均衡",
            "country_count": "国数", "avg_score": "平均スコア",
            "max_score": "最高", "min_score": "最低", "spread": "差",
            "outlier_countries": "異常国（平均より15点以上低い）",
            "none": "なし",
            "overall_normal": "正常 ✅", "overall_monitor": "要観察 ⚠️",
            "overall_action": "要対応 🔴", "overall_insufficient": "データ不足",
            "overall_assessment": "総合評価",
            "adjust_title": "調整ガイド",
            "general_steps_title": "一般改善ステップ",
            "dynamic_recs_title": "今回の動的推奨",
            "records_suffix": " 件（過去30日間）",
        },
    }
    L = _L.get(lk, _L["en"])

    doc = Document()
    title = doc.add_heading(L["title"], level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = meta.add_run(L["meta"].format(aid=result.audit_id, date=result.audit_date, ts=result.timestamp))
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Incomplete data warning banner
    if result.incomplete_data_warning:
        warning_para = doc.add_paragraph()
        warning_run = warning_para.add_run(
            L["incomplete_banner"].format(countries=", ".join(result.incomplete_countries))
        )
        warning_run.font.bold = True
        warning_run.font.color.rgb = RGBColor(204, 102, 0)  # orange

    sd = result.sampling_details or {}
    is_mdsap = sd.get("mdsap_enabled", False)
    regs = sd.get("selected_regulations", [])
    regs_label = ", ".join(regs) if regs else "N/A"
    mode_label = L["mode_mdsap"] if is_mdsap else L["mode_2country"]

    mode_banner = doc.add_paragraph()
    mode_banner_run = mode_banner.add_run(f"  {L['audit_mode'].format(mode=mode_label)}  ")
    mode_banner_run.font.bold = True
    mode_banner_run.font.size = Pt(13)
    if is_mdsap:
        mode_banner_run.font.color.rgb = RGBColor(0, 112, 60)
    else:
        mode_banner_run.font.color.rgb = RGBColor(0, 70, 140)

    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    shading = OxmlElement("w:shd")
    if is_mdsap:
        shading.set(qn("w:fill"), "E6F4EA")
    else:
        shading.set(qn("w:fill"), "E3F2FD")
    shading.set(qn("w:val"), "clear")
    rpr = mode_banner_run._element.get_or_add_rPr()
    rpr.append(shading)

    doc.add_heading(L["abbrev_title"], level=2)
    doc.add_paragraph(L["abbrev_body"])

    doc.add_heading(L["how_title"], level=2)
    doc.add_paragraph(L["how_body"])

    doc.add_heading(L["scoring_title"], level=2)
    doc.add_paragraph(L["scoring_body"])

    doc.add_heading(L["score_summary"], level=2)
    doc.add_paragraph(
        f"Overall Score: {result.overall_score:.0f}/100\n"
        f"Dimension A ({L['dim_a_label']}): {result.dim_a_score:.0f}/100\n"
        f"Dimension B ({L['dim_b_label']}): {result.dim_b_score:.0f}/100\n"
        f"{L['regs_label']}: {regs_label}\n"
        f"Deviation Detected: {'Yes ⚠️' if result.deviation_detected else 'No ✅'}"
    )

    if result.deviation_detected:
        doc.add_heading(L["deviation_title"], level=2)
        doc.add_paragraph(result.deviation_details)

    clauses = sd.get("clauses", [])
    if clauses or regs:
        doc.add_heading(L["sampling_title"], level=2)
        doc.add_paragraph(
            f"{L['source_pipeline']}: {sd.get('source_run_id', 'N/A')}\n"
            f"{L['mode_label']}: {mode_label}\n"
            f"{L['regs_label']}: {regs_label}\n"
            f"{L['sample_rate']}: {sd.get('sample_rate', 0.2):.0%}\n"
            f"{L['available_rows']}: {sd.get('total_rows_available', 0)}\n"
            f"{L['sampled_rows']}: {sd.get('sampled_count', 0)}"
        )

    audit_clauses = [c for c in clauses if not c.get("_qa_doc_summary")]
    summary_entries = [c for c in clauses if c.get("_qa_doc_summary")]

    if audit_clauses:
        doc.add_heading(L["per_clause_title"], level=3)
        table = doc.add_table(rows=1, cols=8)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = L["col_clause"]
        hdr[1].text = L["col_doc"]
        hdr[2].text = L["col_question"]
        hdr[3].text = L["col_verdict"]
        hdr[4].text = L["col_agreed"]
        hdr[5].text = L["col_qa_score"]
        hdr[6].text = L["col_q_quality"]
        hdr[7].text = L["col_hallucination"]
        for clause in audit_clauses:
            row = table.add_row().cells
            row[0].text = str(clause.get("clause_id", ""))
            row[1].text = str(clause.get("doc_id", ""))
            row[2].text = str(clause.get("audit_question", ""))[:60]
            row[3].text = str(clause.get("verdict", ""))[:12]
            row[4].text = "✅" if clause.get("agreed") else "❌"
            row[5].text = f"{clause.get('score', 0)}/100"
            row[6].text = str(clause.get("question_quality", ""))
            row[7].text = "🚨" if clause.get("hallucination_detected") else "—"

        doc.add_heading(L["debate_title"], level=3)

        # Load ISO checklist for expected_evidence lookup
        try:
            from src.analysis.compliance_rules import ISO_13485_CHECKLIST as _ISO_CL
        except Exception:
            _ISO_CL = {}

        for clause in audit_clauses:
            cid = clause.get("clause_id", "")
            ctitle = clause.get("clause_title", "")
            header_text = f"▸ {cid}: {ctitle}" if ctitle else f"▸ {cid}"
            p_header = doc.add_paragraph(header_text)
            p_header.runs[0].font.bold = True

            aq = clause.get("audit_question", "")
            if aq:
                doc.add_paragraph(f"{L['audit_question']}: {aq}")

            # Expected evidence from compliance rules — language-aware
            _clause_def = _ISO_CL.get(cid, {})
            _exp_ev = (
                _clause_def.get("expected_evidence_ja") if lang.startswith("ja")
                else _clause_def.get("expected_evidence_en") if not lang.startswith("zh")
                else _clause_def.get("expected_evidence")
            ) or _clause_def.get("expected_evidence", [])
            if _exp_ev:
                doc.add_paragraph(
                    f"{L['expected_evidence']}:\n"
                    + "\n".join(f"  • {e}" for e in _exp_ev),
                    style="Quote",
                )

            rounds = clause.get("rounds", [])
            if rounds:
                for rd in rounds:
                    round_num = rd.get("round", "?")
                    analyzer = rd.get("analyzer", {})
                    verifier = rd.get("verifier", {})

                    a_position = str(
                        analyzer.get("position", analyzer.get("response", ""))
                    )[:500]
                    a_confidence = analyzer.get(
                        "confidence", analyzer.get("revised_confidence", "N/A")
                    )
                    a_evidence = analyzer.get(
                        "key_evidence", analyzer.get("additional_evidence", [])
                    )

                    v_agreement = verifier.get("agreement_level", "N/A")
                    v_challenges = verifier.get(
                        "challenges", verifier.get("remaining_concerns", [])
                    )
                    v_assessment = verifier.get("overall_assessment", "")[:400]

                    doc.add_paragraph(f"── Round {round_num} ──").runs[
                        0
                    ].font.bold = True
                    doc.add_paragraph(
                        f"{L['analyzer']} — confidence: {a_confidence}\n"
                        f"{L['position']}: {a_position}"
                    )
                    if a_evidence:
                        evidence_str = ", ".join(str(e)[:100] for e in a_evidence[:5])
                        doc.add_paragraph(f"{L['evidence']}: {evidence_str}")

                    challenge_lines = f"{L['verifier']} — agreement: {v_agreement}"
                    if isinstance(v_challenges, list) and v_challenges:
                        for ch in v_challenges[:3]:
                            if isinstance(ch, dict):
                                challenge_lines += (
                                    f"\n  {L['challenge']}: {ch.get('point', str(ch))[:200]}"
                                )
                            else:
                                challenge_lines += f"\n  {L['challenge']}: {str(ch)[:200]}"
                    doc.add_paragraph(challenge_lines)
                    if v_assessment:
                        doc.add_paragraph(f"{L['comment']}: {v_assessment}")
            else:
                doc.add_paragraph(L["no_rounds"])

            qa_score = clause.get("score", 0)
            qq = clause.get("question_quality", "unknown")
            aa = clause.get("answer_accuracy", "unknown")
            lc = clause.get("logic_consistency", "unknown")
            hal = clause.get("hallucination_detected", False)
            hal_detail = clause.get("hallucination_details", "")
            issues = clause.get("issues", [])

            qa_text = L["qa_result"].format(score=qa_score, qq=qq, aa=aa, lc=lc)
            if hal:
                qa_text += (
                    f"\n{L['hallucination_detect']}: {hal_detail}" if hal_detail else f"\n{L['hallucination_detect']}"
                )
            p_qa = doc.add_paragraph(qa_text)
            p_qa.runs[0].font.italic = True

            if issues:
                for issue in issues:
                    doc.add_paragraph(f"  • {issue}")

            doc.add_paragraph("")

    if summary_entries:
        doc.add_heading(L["qa_summary_title"], level=3)
        for entry in summary_entries:
            doc_id_s = entry.get("doc_id", "")
            score_s = entry.get("overall_score", 0)
            summary_s = entry.get("summary", "")
            recs_s = entry.get("recommendations", [])
            doc.add_paragraph(L["doc_score"].format(doc_id=doc_id_s, score=score_s)).runs[
                0
            ].font.bold = True
            if summary_s:
                doc.add_paragraph(summary_s)
            if recs_s:
                for rec in recs_s:
                    doc.add_paragraph(f"  💡 {rec}")

    dim_a_title = L["dim_a_mdsap"] if is_mdsap else L["dim_a_2c"]
    doc.add_heading(dim_a_title, level=2)
    doc.add_paragraph(result.dim_a_summary or "N/A")
    if result.dim_a_score_rationale:
        doc.add_paragraph(f"{L['score_rationale']}: {result.dim_a_score_rationale}", style="Quote")
    if result.dim_a_checks:
        doc.add_heading(L["check_items"], level=3)
        for check in result.dim_a_checks:
            doc.add_paragraph(
                f"• [{check.get('severity', 'N/A')}] {check.get('check_type', '')}: "
                f"{check.get('issue', 'No issue')} "
                f"(Regulation: {check.get('regulation', '')})"
            )

    dim_b_title = L["dim_b_mdsap"] if is_mdsap else L["dim_b_2c"]
    doc.add_heading(dim_b_title, level=2)
    doc.add_paragraph(result.dim_b_summary or "N/A")
    if result.dim_b_score_rationale:
        doc.add_paragraph(f"{L['score_rationale']}: {result.dim_b_score_rationale}", style="Quote")
    if result.dim_b_country_scores:
        doc.add_heading(L["country_scores"], level=3)
        for country, score in sorted(result.dim_b_country_scores.items()):
            doc.add_paragraph(f"  {country}: {score:.0f}/100")
    if result.dim_b_findings:
        doc.add_heading(L["findings"], level=3)
        for finding in result.dim_b_findings:
            doc.add_paragraph(
                f"• [{finding.get('severity', 'N/A')}] {finding.get('category', '')}: "
                f"{finding.get('description', '')}"
            )

    cv = result.cross_validation or {}
    if cv and not cv.get("error"):
        doc.add_heading(L["crossval_title"], level=2)

        # Method 1: Temporal drift
        td = cv.get("temporal_drift", {})
        td_status = td.get("status", "insufficient_data")
        td_labels = {
            "stable": L["td_stable"], "minor_drift": L["td_minor"],
            "significant_drift": L["td_significant"], "insufficient_data": L["td_insufficient"],
        }
        doc.add_heading(L["method1_title"], level=3)
        if td_status == "insufficient_data":
            doc.add_paragraph(td.get("reason", L["insufficient_reason"]))
        else:
            doc.add_paragraph(
                f"{L['status']}: {td_labels.get(td_status, td_status)}\n"
                f"{L['history_count']}: {td.get('history_count', 0)}{L['records_suffix']}\n"
                f"{L['rolling_avg']}: {td.get('rolling_avg_overall', 0):.1f} "
                f"(DimA {td.get('rolling_avg_dim_a', 0):.1f} / "
                f"DimB {td.get('rolling_avg_dim_b', 0):.1f})\n"
                f"{L['today_score']}: {td.get('today_overall', 0):.1f} "
                f"(DimA {td.get('today_dim_a', 0):.1f} / "
                f"DimB {td.get('today_dim_b', 0):.1f})\n"
                f"{L['delta']}: {td.get('delta_overall', 0):+.1f} "
                f"(DimA {td.get('delta_dim_a', 0):+.1f} / "
                f"DimB {td.get('delta_dim_b', 0):+.1f})"
            )

        # Method 2: Country imbalance
        ci = cv.get("country_imbalance", {})
        ci_status = ci.get("status", "insufficient_data")
        ci_labels = {
            "balanced": L["ci_balanced"], "minor_imbalance": L["ci_minor"],
            "significant_imbalance": L["ci_significant"], "insufficient_data": L["ci_insufficient"],
        }
        doc.add_heading(L["method2_title"], level=3)
        if ci_status == "insufficient_data":
            doc.add_paragraph(ci.get("reason", L["insufficient_reason"]))
        else:
            outliers = ci.get("outlier_countries", {})
            outlier_text = (
                ", ".join(f"{c}({s:.0f})" for c, s in outliers.items())
                if outliers
                else L["none"]
            )
            doc.add_paragraph(
                f"{L['status']}: {ci_labels.get(ci_status, ci_status)}\n"
                f"{L['country_count']}: {ci.get('country_count', 0)}\n"
                f"{L['avg_score']}: {ci.get('avg_score', 0):.1f}  "
                f"{L['max_score']}: {ci.get('max_score', 0):.1f}  "
                f"{L['min_score']}: {ci.get('min_score', 0):.1f}  "
                f"{L['spread']}: {ci.get('spread', 0):.1f}\n"
                f"{L['outlier_countries']}: {outlier_text}"
            )

        overall_labels = {
            "normal": L["overall_normal"], "monitor": L["overall_monitor"],
            "action_required": L["overall_action"], "insufficient_data": L["overall_insufficient"],
        }
        overall = cv.get("overall_assessment", "insufficient_data")
        doc.add_paragraph(f"{L['overall_assessment']}: {overall_labels.get(overall, overall)}")

    fixed_steps, dynamic_recs = _generate_adjustment_guidance(result, lang=lang)

    doc.add_heading(L["adjust_title"], level=2)

    doc.add_heading(L["general_steps_title"], level=3)
    for i, step in enumerate(fixed_steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_heading(L["dynamic_recs_title"], level=3)
    for rec in dynamic_recs:
        doc.add_paragraph(f"• {rec}")

    try:
        from src.utils.crossref_export import append_crossref_table_word
        append_crossref_table_word(doc, lang=lang)
    except Exception:
        pass

    from src.utils.safe_io import safe_save_binary

    safe_save_binary(filepath, doc.save)
    return filepath


def export_daily_audit_excel(result: DailyAuditResult, lang: str = "zh-TW") -> Path:
    """Export a daily audit result as an Excel file.

    Args:
        result: DailyAuditResult to export
        lang: Language code for report content

    Returns:
        Path to the generated .xlsx file
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill

    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    filepath = EXPORT_DIR / f"daily_audit_{result.audit_date}.xlsx"

    wb = Workbook()

    # Sheet 1: Summary
    ws_sum = wb.active
    ws_sum.title = "Summary"
    header_fill = PatternFill(
        start_color="4472C4", end_color="4472C4", fill_type="solid"
    )
    header_text_font = Font(bold=True, color="FFFFFF", size=11)

    sd = result.sampling_details or {}
    is_mdsap = sd.get("mdsap_enabled", False)
    regs = sd.get("selected_regulations", [])
    regs_label = ", ".join(regs) if regs else "N/A"
    mode_label = "MDSAP 7-Country" if is_mdsap else "2-Country (TFDA + EU_MDR)"

    summary_data = [
        ("Audit ID", result.audit_id),
        ("Date", result.audit_date),
        ("Timestamp", result.timestamp),
        ("Overall Score", f"{result.overall_score:.0f}/100"),
        ("Dim A Score (Regulation Accuracy)", f"{result.dim_a_score:.0f}/100"),
        ("Dim B Score (Cross-Exam Quality)", f"{result.dim_b_score:.0f}/100"),
        ("Mode", mode_label),
        ("Regulations", regs_label),
        ("Deviation Detected", "Yes" if result.deviation_detected else "No"),
        ("Deviation Details", result.deviation_details or "N/A"),
        ("Model", result.model),
        ("Summary", result.summary),
        (
            "Incomplete Data Warning",
            f"Yes — {', '.join(result.incomplete_countries)}"
            if result.incomplete_data_warning
            else "No",
        ),
    ]

    for row_idx, (key, val) in enumerate(summary_data, start=1):
        ws_sum.cell(row=row_idx, column=1, value=key).font = Font(bold=True)
        ws_sum.cell(row=row_idx, column=2, value=str(val))

    ws_sum.column_dimensions["A"].width = 30
    ws_sum.column_dimensions["B"].width = 60
    ws_a = wb.create_sheet("Dim A Checks")
    a_headers = ["Check Type", "Regulation", "Issue", "Severity", "Evidence"]
    for col, h in enumerate(a_headers, 1):
        cell = ws_a.cell(row=1, column=col, value=h)
        cell.font = header_text_font
        cell.fill = header_fill

    for row_idx, check in enumerate(result.dim_a_checks, start=2):
        ws_a.cell(row=row_idx, column=1, value=check.get("check_type", ""))
        ws_a.cell(row=row_idx, column=2, value=check.get("regulation", ""))
        ws_a.cell(row=row_idx, column=3, value=check.get("issue", ""))
        ws_a.cell(row=row_idx, column=4, value=check.get("severity", ""))
        ws_a.cell(row=row_idx, column=5, value=check.get("evidence", ""))

    # Sheet 3: Dim B Findings
    ws_b = wb.create_sheet("Dim B Findings")
    b_headers = ["Category", "Severity", "Description", "Recommendation"]
    for col, h in enumerate(b_headers, 1):
        cell = ws_b.cell(row=1, column=col, value=h)
        cell.font = header_text_font
        cell.fill = header_fill

    for row_idx, finding in enumerate(result.dim_b_findings, start=2):
        ws_b.cell(row=row_idx, column=1, value=finding.get("category", ""))
        ws_b.cell(row=row_idx, column=2, value=finding.get("severity", ""))
        ws_b.cell(row=row_idx, column=3, value=finding.get("description", ""))
        ws_b.cell(row=row_idx, column=4, value=finding.get("recommendation", ""))

    # Sheet 4: Country Scores
    if result.dim_b_country_scores:
        ws_c = wb.create_sheet("Country Scores")
        c_headers = ["Country", "Score"]
        for col, h in enumerate(c_headers, 1):
            cell = ws_c.cell(row=1, column=col, value=h)
            cell.font = header_text_font
            cell.fill = header_fill
        for row_idx, (country, score) in enumerate(
            sorted(result.dim_b_country_scores.items()), start=2
        ):
            ws_c.cell(row=row_idx, column=1, value=country)
            ws_c.cell(row=row_idx, column=2, value=score)

    cv = result.cross_validation or {}
    if cv and not cv.get("error"):
        ws_cv = wb.create_sheet("Cross-Validation")
        cv_headers = ["Metric", "Value"]
        for col, h in enumerate(cv_headers, 1):
            cell = ws_cv.cell(row=1, column=col, value=h)
            cell.font = header_text_font
            cell.fill = header_fill

        td = cv.get("temporal_drift", {})
        ci = cv.get("country_imbalance", {})
        overall = cv.get("overall_assessment", "insufficient_data")

        drift_labels = {
            "stable": "Stable",
            "minor_drift": "Minor Drift",
            "significant_drift": "Significant Drift",
            "insufficient_data": "Insufficient Data",
        }
        imbalance_labels = {
            "balanced": "Balanced",
            "minor_imbalance": "Minor Imbalance",
            "significant_imbalance": "Significant Imbalance",
            "insufficient_data": "Insufficient Data",
        }
        overall_labels = {
            "normal": "Normal",
            "monitor": "Monitor",
            "action_required": "Action Required",
            "insufficient_data": "Insufficient Data",
        }

        outliers = ci.get("outlier_countries", {})
        outlier_str = (
            ", ".join(f"{c}({s:.0f})" for c, s in outliers.items()) if outliers else "None"
        )

        cv_data = [
            # Method 1: Temporal drift
            ("[Method 1] Temporal Drift Status", drift_labels.get(td.get("status", ""), td.get("status", "N/A"))),
            ("[Method 1] History Count (30-day)", td.get("history_count", "N/A")),
            ("[Method 1] Rolling Avg Overall", td.get("rolling_avg_overall", "N/A")),
            ("[Method 1] Rolling Avg Dim A", td.get("rolling_avg_dim_a", "N/A")),
            ("[Method 1] Rolling Avg Dim B", td.get("rolling_avg_dim_b", "N/A")),
            ("[Method 1] Today Overall", td.get("today_overall", "N/A")),
            ("[Method 1] Delta Overall", td.get("delta_overall", "N/A")),
            ("[Method 1] Delta Dim A", td.get("delta_dim_a", "N/A")),
            ("[Method 1] Delta Dim B", td.get("delta_dim_b", "N/A")),
            # Method 2: Country imbalance
            ("[Method 2] Country Imbalance Status", imbalance_labels.get(ci.get("status", ""), ci.get("status", "N/A"))),
            ("[Method 2] Country Count", ci.get("country_count", "N/A")),
            ("[Method 2] Avg Score", ci.get("avg_score", "N/A")),
            ("[Method 2] Max Score", ci.get("max_score", "N/A")),
            ("[Method 2] Min Score", ci.get("min_score", "N/A")),
            ("[Method 2] Spread", ci.get("spread", "N/A")),
            ("[Method 2] Outlier Countries", outlier_str),
            # Overall
            ("Overall Assessment", overall_labels.get(overall, overall)),
        ]
        for row_idx, (metric, val) in enumerate(cv_data, start=2):
            ws_cv.cell(row=row_idx, column=1, value=metric)
            ws_cv.cell(
                row=row_idx,
                column=2,
                value=round(val, 2) if isinstance(val, float) else val,
            )
        ws_cv.column_dimensions["A"].width = 40
        ws_cv.column_dimensions["B"].width = 25

    clauses = sd.get("clauses", [])
    if clauses:
        ws_sd = wb.create_sheet("Sampling Details")
        sd_headers = [
            "ISO 13485 Clause",
            "Clause Title",
            "Doc ID",
            "Regulations",
            "Verdict",
            "Agreed",
            "Flagged",
            "Rounds",
            "QA Score",
            "Question Quality",
            "Answer Accuracy",
            "Hallucination",
        ]
        for col, h in enumerate(sd_headers, 1):
            cell = ws_sd.cell(row=1, column=col, value=h)
            cell.font = header_text_font
            cell.fill = header_fill
        for row_idx, clause in enumerate(clauses, start=2):
            ws_sd.cell(row=row_idx, column=1, value=clause.get("clause_id", ""))
            ws_sd.cell(row=row_idx, column=2, value=clause.get("clause_title", ""))
            ws_sd.cell(row=row_idx, column=3, value=clause.get("doc_id", ""))
            ws_sd.cell(row=row_idx, column=4, value=regs_label)
            ws_sd.cell(row=row_idx, column=5, value=clause.get("verdict", ""))
            ws_sd.cell(
                row=row_idx, column=6, value="Yes" if clause.get("agreed") else "No"
            )
            ws_sd.cell(
                row=row_idx, column=7, value="Yes" if clause.get("flagged") else "No"
            )
            ws_sd.cell(row=row_idx, column=8, value=len(clause.get("rounds", [])))
            qa = clause.get("qa_audit", {})
            ws_sd.cell(row=row_idx, column=9, value=qa.get("score", "") if qa else "")
            ws_sd.cell(
                row=row_idx,
                column=10,
                value=qa.get("question_quality", "") if qa else "",
            )
            ws_sd.cell(
                row=row_idx,
                column=11,
                value=qa.get("answer_accuracy", "") if qa else "",
            )
            ws_sd.cell(
                row=row_idx,
                column=12,
                value="Yes"
                if qa and qa.get("hallucination_detected")
                else ("No" if qa else ""),
            )
        ws_sd.column_dimensions["A"].width = 20
        ws_sd.column_dimensions["B"].width = 30
        ws_sd.column_dimensions["C"].width = 20
        ws_sd.column_dimensions["D"].width = 40

        ws_debate = wb.create_sheet("Debate Details")
        db_headers = [
            "Clause ID",
            "Doc ID",
            "Round",
            "Analyzer Position",
            "Analyzer Confidence",
            "Verifier Assessment",
            "Agreement",
        ]
        for col, h in enumerate(db_headers, 1):
            cell = ws_debate.cell(row=1, column=col, value=h)
            cell.font = header_text_font
            cell.fill = header_fill
        db_row = 2
        for clause in clauses:
            cid = clause.get("clause_id", "")
            did = clause.get("doc_id", "")
            for rd in clause.get("rounds", []):
                a = rd.get("analyzer", {})
                v = rd.get("verifier", {})
                ws_debate.cell(row=db_row, column=1, value=cid)
                ws_debate.cell(row=db_row, column=2, value=did)
                ws_debate.cell(row=db_row, column=3, value=rd.get("round", 0))
                a_pos = a.get("position", "")
                if isinstance(a_pos, list):
                    a_pos = "; ".join(str(x) for x in a_pos)
                ws_debate.cell(row=db_row, column=4, value=str(a_pos)[:1000])
                ws_debate.cell(
                    row=db_row,
                    column=5,
                    value=str(a.get("confidence", a.get("confidence_score", ""))),
                )
                v_assess = v.get("assessment", "")
                if isinstance(v_assess, list):
                    v_assess = "; ".join(str(x) for x in v_assess)
                ws_debate.cell(row=db_row, column=6, value=str(v_assess)[:1000])
                ws_debate.cell(row=db_row, column=7, value=v.get("agreement_level", ""))
                db_row += 1
        ws_debate.column_dimensions["A"].width = 20
        ws_debate.column_dimensions["B"].width = 20
        ws_debate.column_dimensions["D"].width = 60
        ws_debate.column_dimensions["F"].width = 60

    fixed_steps, dynamic_recs = _generate_adjustment_guidance(result, lang=lang)

    _guide_labels = {
        "zh": ("人為調整指引", "一般改善步驟", "本次動態建議"),
        "en": ("How to Adjust", "General Improvement Steps", "Dynamic Recommendations"),
        "ja": ("調整ガイド", "一般改善ステップ", "今回の動的推奨"),
    }
    _glk = _get_prompt_lang(lang)
    _gl = _guide_labels.get(_glk, _guide_labels["en"])

    ws_guide = wb.create_sheet("Adjustment Guidance")
    guide_section_font = Font(bold=True, size=12)
    guide_fill = PatternFill(
        start_color="FFF2CC", end_color="FFF2CC", fill_type="solid"
    )

    ws_guide.cell(row=1, column=1, value=_gl[0]).font = Font(
        bold=True, size=14
    )
    ws_guide.merge_cells("A1:B1")

    ws_guide.cell(
        row=3, column=1, value=_gl[1]
    ).font = guide_section_font
    for i, step in enumerate(fixed_steps, 1):
        row_idx = 3 + i
        cell_num = ws_guide.cell(row=row_idx, column=1, value=f"Step {i}")
        cell_num.font = Font(bold=True)
        cell_num.fill = guide_fill
        ws_guide.cell(row=row_idx, column=2, value=step)

    dynamic_start = 3 + len(fixed_steps) + 2
    ws_guide.cell(
        row=dynamic_start, column=1, value=_gl[2]
    ).font = guide_section_font
    for i, rec in enumerate(dynamic_recs, 1):
        row_idx = dynamic_start + i
        cell_num = ws_guide.cell(row=row_idx, column=1, value=f"#{i}")
        cell_num.font = Font(bold=True)
        ws_guide.cell(row=row_idx, column=2, value=rec)

    ws_guide.column_dimensions["A"].width = 15
    ws_guide.column_dimensions["B"].width = 80

    try:
        from src.utils.crossref_export import append_crossref_table_excel
        append_crossref_table_excel(wb, lang=lang)
    except Exception:
        pass

    from src.utils.safe_io import safe_save_binary

    safe_save_binary(filepath, wb.save)
    return filepath


def export_meta_review_word(result: MetaReviewResult, lang: str = "zh-TW") -> Path:
    """Export a meta review result as a Word document.

    Args:
        result: MetaReviewResult to export
        lang: Language code for report content

    Returns:
        Path to the generated .docx file
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    filepath = EXPORT_DIR / f"meta_review_{result.period_end or 'latest'}.docx"

    _mr_labels = {
        "zh": {
            "title": "AI-QMS 10日總檢報告",
            "score_summary": "評分摘要",
            "trend_analysis": "趨勢分析",
            "deviation_summary": "偏差摘要",
            "country_trends": "各國趨勢",
            "recommendations": "建議",
            "daily_results": "每日結果",
        },
        "en": {
            "title": "AI-QMS 10-Day Meta Review Report",
            "score_summary": "Score Summary",
            "trend_analysis": "Trend Analysis",
            "deviation_summary": "Deviation Summary",
            "country_trends": "Country Trends",
            "recommendations": "Recommendations",
            "daily_results": "Daily Results",
        },
        "ja": {
            "title": "AI-QMS 10日間メタレビューレポート",
            "score_summary": "スコアサマリー",
            "trend_analysis": "傾向分析",
            "deviation_summary": "偏差サマリー",
            "country_trends": "国別傾向",
            "recommendations": "推奨事項",
            "daily_results": "毎日の結果",
        },
    }
    _mr_lk = _get_prompt_lang(lang)
    _mr = _mr_labels.get(_mr_lk, _mr_labels["en"])

    doc = Document()
    title = doc.add_heading(_mr["title"], level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = meta.add_run(
        f"Review ID: {result.review_id}  |  "
        f"Period: {result.period_start} ~ {result.period_end}  |  "
        f"Generated: {result.timestamp}"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Summary
    doc.add_heading(_mr["score_summary"], level=2)
    doc.add_paragraph(
        f"Average Dim A (MDSAP Accuracy): {result.avg_dim_a:.0f}/100\n"
        f"Average Dim B (Cross-Exam Quality): {result.avg_dim_b:.0f}/100\n"
        f"Daily Audits Analyzed: {len(result.daily_results)}"
    )

    # Trend Analysis
    doc.add_heading(_mr["trend_analysis"], level=2)
    doc.add_paragraph(result.trend_analysis or "N/A")

    # Deviation Summary
    if result.deviation_summary:
        doc.add_heading(_mr["deviation_summary"], level=2)
        doc.add_paragraph(result.deviation_summary)

    # Country Trends
    if result.country_trends:
        doc.add_heading(_mr["country_trends"], level=2)
        for country, trend in sorted(result.country_trends.items()):
            doc.add_paragraph(f"  {country}: {trend}")

    # Recommendations
    if result.recommendations:
        doc.add_heading(_mr["recommendations"], level=2)
        for i, rec in enumerate(result.recommendations, 1):
            doc.add_paragraph(f"{i}. {rec}")

    # Daily Results Table
    doc.add_heading(_mr["daily_results"], level=2)
    if result.daily_results:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Date"
        hdr[1].text = "Dim A"
        hdr[2].text = "Dim B"
        hdr[3].text = "Overall"
        hdr[4].text = "Deviation"
        for dr in result.daily_results:
            row = table.add_row().cells
            row[0].text = str(dr.get("audit_date", ""))
            row[1].text = f"{dr.get('dim_a_score', 0):.0f}"
            row[2].text = f"{dr.get('dim_b_score', 0):.0f}"
            row[3].text = f"{dr.get('overall_score', 0):.0f}"
            row[4].text = "⚠️" if dr.get("deviation_detected") else "✅"

    try:
        from src.utils.crossref_export import append_crossref_table_word
        append_crossref_table_word(doc, lang=lang)
    except Exception:
        pass

    from src.utils.safe_io import safe_save_binary

    safe_save_binary(filepath, doc.save)
    return filepath


def export_meta_review_excel(result: MetaReviewResult, lang: str = "zh-TW") -> Path:
    """Export a meta review result as an Excel file.

    Args:
        result: MetaReviewResult to export
        lang: Language code for report content

    Returns:
        Path to the generated .xlsx file
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill

    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    filepath = EXPORT_DIR / f"meta_review_{result.period_end or 'latest'}.xlsx"

    wb = Workbook()
    header_fill = PatternFill(
        start_color="4472C4", end_color="4472C4", fill_type="solid"
    )
    header_font = Font(bold=True, color="FFFFFF", size=11)

    # Sheet 1: Summary
    ws_sum = wb.active
    ws_sum.title = "Summary"
    summary_data = [
        ("Review ID", result.review_id),
        ("Period", f"{result.period_start} ~ {result.period_end}"),
        ("Generated", result.timestamp),
        ("Avg Dim A", f"{result.avg_dim_a:.0f}/100"),
        ("Avg Dim B", f"{result.avg_dim_b:.0f}/100"),
        ("Daily Audits", str(len(result.daily_results))),
        ("Trend Analysis", result.trend_analysis),
        ("Deviation Summary", result.deviation_summary or "None"),
    ]
    for row_idx, (key, val) in enumerate(summary_data, start=1):
        ws_sum.cell(row=row_idx, column=1, value=key).font = Font(bold=True)
        ws_sum.cell(row=row_idx, column=2, value=str(val))
    ws_sum.column_dimensions["A"].width = 25
    ws_sum.column_dimensions["B"].width = 80

    # Sheet 2: Daily Results
    ws_daily = wb.create_sheet("Daily Results")
    d_headers = ["Date", "Dim A", "Dim B", "Overall", "Deviation", "Summary"]
    for col, h in enumerate(d_headers, 1):
        cell = ws_daily.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill

    for row_idx, dr in enumerate(result.daily_results, start=2):
        ws_daily.cell(row=row_idx, column=1, value=dr.get("audit_date", ""))
        ws_daily.cell(row=row_idx, column=2, value=dr.get("dim_a_score", 0))
        ws_daily.cell(row=row_idx, column=3, value=dr.get("dim_b_score", 0))
        ws_daily.cell(row=row_idx, column=4, value=dr.get("overall_score", 0))
        ws_daily.cell(
            row=row_idx, column=5, value="Yes" if dr.get("deviation_detected") else "No"
        )
        ws_daily.cell(row=row_idx, column=6, value=dr.get("summary", ""))

    # Sheet 3: Recommendations
    if result.recommendations:
        ws_rec = wb.create_sheet("Recommendations")
        ws_rec.cell(row=1, column=1, value="#").font = header_font
        ws_rec.cell(row=1, column=1).fill = header_fill
        ws_rec.cell(row=1, column=2, value="Recommendation").font = header_font
        ws_rec.cell(row=1, column=2).fill = header_fill
        for row_idx, rec in enumerate(result.recommendations, start=2):
            ws_rec.cell(row=row_idx, column=1, value=row_idx - 1)
            ws_rec.cell(row=row_idx, column=2, value=rec)
        ws_rec.column_dimensions["B"].width = 80

    # Sheet 4: Country Trends
    if result.country_trends:
        ws_ct = wb.create_sheet("Country Trends")
        ws_ct.cell(row=1, column=1, value="Country").font = header_font
        ws_ct.cell(row=1, column=1).fill = header_fill
        ws_ct.cell(row=1, column=2, value="Trend").font = header_font
        ws_ct.cell(row=1, column=2).fill = header_fill
        for row_idx, (country, trend) in enumerate(
            sorted(result.country_trends.items()), start=2
        ):
            ws_ct.cell(row=row_idx, column=1, value=country)
            ws_ct.cell(row=row_idx, column=2, value=trend)
        ws_ct.column_dimensions["B"].width = 60

    try:
        from src.utils.crossref_export import append_crossref_table_excel
        append_crossref_table_excel(wb, lang=lang)
    except Exception:
        pass

    from src.utils.safe_io import safe_save_binary

    safe_save_binary(filepath, wb.save)
    return filepath


# ============================================================
# User Feedback — Data class, Storage, CRUD
# ============================================================


class AuditFeedback:
    """User feedback record for daily audit or meta review."""

    def __init__(
        self,
        feedback_id: str,
        audit_type: str,  # 'daily' | 'meta'
        target_id: str,  # audit_id or meta review date
        feedback_text: str,
        created_at: str,
        updated_at: str,
        status: str = "active",  # 'active' | 'deleted'
        re_evaluation_id: Optional[str] = None,
        re_evaluation_score: Optional[int] = None,
    ):
        self.feedback_id = feedback_id
        self.audit_type = audit_type
        self.target_id = target_id
        self.feedback_text = feedback_text
        self.created_at = created_at
        self.updated_at = updated_at
        self.status = status
        self.re_evaluation_id = re_evaluation_id
        self.re_evaluation_score = re_evaluation_score

    def to_dict(self) -> dict:
        return {
            "feedback_id": self.feedback_id,
            "audit_type": self.audit_type,
            "target_id": self.target_id,
            "feedback_text": self.feedback_text,
            "created_at": self.created_at,
            "updated_at": self.updated_at,
            "status": self.status,
            "re_evaluation_id": self.re_evaluation_id,
            "re_evaluation_score": self.re_evaluation_score,
        }

    @classmethod
    def from_dict(cls, d: dict) -> "AuditFeedback":
        return cls(
            feedback_id=d["feedback_id"],
            audit_type=d.get("audit_type", "daily"),
            target_id=d.get("target_id", ""),
            feedback_text=d.get("feedback_text", ""),
            created_at=d.get("created_at", ""),
            updated_at=d.get("updated_at", ""),
            status=d.get("status", "active"),
            re_evaluation_id=d.get("re_evaluation_id"),
            re_evaluation_score=d.get("re_evaluation_score"),
        )


def _ensure_feedback_dir() -> Path:
    FEEDBACK_DIR.mkdir(parents=True, exist_ok=True)
    return FEEDBACK_DIR


def save_feedback(
    audit_type: str,
    target_id: str,
    feedback_text: str,
) -> AuditFeedback:
    """Create and persist a new feedback record."""
    _ensure_feedback_dir()
    ts = datetime.now().isoformat()
    fb_id = f"fb_{int(time.time())}_{audit_type}"
    fb = AuditFeedback(
        feedback_id=fb_id,
        audit_type=audit_type,
        target_id=target_id,
        feedback_text=feedback_text,
        created_at=ts,
        updated_at=ts,
    )
    filepath = FEEDBACK_DIR / f"{fb_id}.json"
    atomic_write_json(filepath, fb.to_dict())
    logger.info(f"Saved feedback {fb_id} for {audit_type}/{target_id}")
    return fb


def get_all_feedback(include_deleted: bool = False) -> list[AuditFeedback]:
    """Return all feedback records, sorted newest-first."""
    _ensure_feedback_dir()
    results: list[AuditFeedback] = []
    for f in sorted(FEEDBACK_DIR.glob("fb_*.json"), reverse=True):
        try:
            data = json.loads(f.read_text(encoding="utf-8"))
            fb = AuditFeedback.from_dict(data)
            if not include_deleted and fb.status == "deleted":
                continue
            results.append(fb)
        except Exception as exc:
            logger.warning(f"Failed to load feedback {f}: {exc}")
    return results


def get_feedback_by_id(feedback_id: str) -> Optional[AuditFeedback]:
    """Load a single feedback record by ID."""
    filepath = FEEDBACK_DIR / f"{feedback_id}.json"
    if not filepath.exists():
        return None
    try:
        data = json.loads(filepath.read_text(encoding="utf-8"))
        return AuditFeedback.from_dict(data)
    except Exception as exc:
        logger.warning(f"Failed to load feedback {feedback_id}: {exc}")
        return None


def update_feedback(feedback_id: str, new_text: str) -> Optional[AuditFeedback]:
    """Update the text of an existing feedback record."""
    fb = get_feedback_by_id(feedback_id)
    if fb is None or fb.status == "deleted":
        return None
    fb.feedback_text = new_text
    fb.updated_at = datetime.now().isoformat()
    filepath = FEEDBACK_DIR / f"{feedback_id}.json"
    atomic_write_json(filepath, fb.to_dict())
    logger.info(f"Updated feedback {feedback_id}")
    return fb


def delete_feedback(feedback_id: str) -> bool:
    """Soft-delete a feedback record (sets status='deleted')."""
    fb = get_feedback_by_id(feedback_id)
    if fb is None:
        return False
    fb.status = "deleted"
    fb.updated_at = datetime.now().isoformat()
    filepath = FEEDBACK_DIR / f"{feedback_id}.json"
    atomic_write_json(filepath, fb.to_dict())
    logger.info(f"Deleted feedback {feedback_id}")
    return True


def get_feedback_for_audit(target_id: str) -> list[AuditFeedback]:
    """Get all active feedback for a specific audit/meta review."""
    all_fb = get_all_feedback()
    return [fb for fb in all_fb if fb.target_id == target_id]


def get_active_feedback_context() -> str:
    """Build a context string from all active feedback for LLM re-evaluation."""
    all_fb = get_all_feedback()
    if not all_fb:
        return ""
    lines = ["--- User feedback history (for reference during audit) ---"]
    for fb in all_fb[:20]:  # limit to 20 most recent
        lines.append(
            f"[{fb.created_at[:10]}] ({fb.audit_type}/{fb.target_id}): {fb.feedback_text}"
        )
    lines.append("--- End of user feedback ---")
    return "\n".join(lines)
