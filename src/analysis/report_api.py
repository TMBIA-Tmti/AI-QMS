"""
AI-QMS — Report API (Phase D)
==============================

REST API endpoints for the interactive compliance report page.
Mounted on the Chainlit FastAPI app at /api/report/...

Endpoints:
  GET  /api/report/runs                        — List all pipeline runs
  GET  /api/report/{run_id}                    — Get full report data
  GET  /api/report/{run_id}/summary            — Get summary statistics
  GET  /api/report/{run_id}/rows               — Get rows with filters
  GET  /api/report/{run_id}/row/{row_id}       — Get single row detail
  POST /api/report/{run_id}/row/{row_id}/override  — RA override verdict
  POST /api/report/{run_id}/row/{row_id}/note      — Add clause note
  POST /api/report/{run_id}/row/{row_id}/restore   — Restore LLM original
  POST /api/report/{run_id}/row/{row_id}/rerun     — Reset row for re-run
  GET  /api/report/{run_id}/row/{row_id}/history   — Get version history
  GET  /api/report/{run_id}/export/{format}        — Export Word/Excel
  GET  /api/report/crossref/regulations            — List available regulations
  GET  /api/report/crossref/table                  — Cross-reference table
  GET  /api/report/crossref/questions              — Cross-exam questions
  GET  /api/report/standards/list                  — List supplemental standards
  POST /api/report/standards/applicable            — Determine applicable standards
  POST /api/report/standards/adjust                — Adjust standard-to-clause mapping
  GET  /api/report/{run_id}/stream                 — SSE cross-examination events
  POST /api/report/{run_id}/inject                 — Human message injection
  POST /api/report/{run_id}/pause                  — Pause cross-examination
  POST /api/report/{run_id}/resume                 — Resume cross-examination
  POST /api/report/crossref/mdsap-verify            — Toggle MDSAP verification
  GET  /api/report/crossref/mdsap-verify            — Get MDSAP verification state
  GET  /api/report/daily-audit/history               — List daily audit records
  POST /api/report/daily-audit/run                   — Trigger daily audit
  GET  /api/report/daily-audit/meta-review           — Get latest 10-day meta review
  POST /api/report/daily-audit/meta-review           — Trigger 10-day meta review
  GET  /api/report/daily-audit/history/{id}/export   — Export audit record Word/Excel
  GET  /api/report/daily-audit/export/{fmt}          — Export latest audit Word/Excel
  GET  /api/report/daily-audit/meta-review/export    — Export meta review Word/Excel
"""

from __future__ import annotations

import asyncio
import json
import logging
import re
import time
import threading
from collections import defaultdict
from html import escape
from pathlib import Path
from typing import Optional

from src.utils.safe_io import safe_save_binary

from fastapi import APIRouter, HTTPException, Query, Request
from fastapi.responses import (
    JSONResponse,
    FileResponse,
    HTMLResponse,
    StreamingResponse,
)

from src.analysis.state import Phase
from src.analysis.comparison_table import ComparisonTable
from src.analysis.risk_matrix import (
    RISK_LEVEL_DISPLAY,
    VERDICT_DISPLAY,
    Verdict,
    RiskLevel,
    assess_risk,
)

logger = logging.getLogger(__name__)

__all__ = ["report_router", "REPORT_STATIC_DIR"]

# Static files directory for report HTML/JS/CSS
REPORT_STATIC_DIR = Path(__file__).parent.parent.parent / "report_ui"

# Pipeline state storage directory
_PIPELINE_DIR = Path("data/analysis_pipeline")

# FastAPI router
report_router = APIRouter(prefix="/api/report", tags=["report"])


# ── Phoenix tracing helper ──


def _phoenix_report_span(name: str, attributes: dict = None):
    """Create a Phoenix span for report API operations.

    Falls back to no-op if Phoenix is not available.
    """
    try:
        from src.chainlit_app.app import phoenix_span

        return phoenix_span(name, profile="主系統 (Main Agent)", attributes=attributes)
    except (ImportError, Exception):
        from contextlib import nullcontext

        return nullcontext()


# ============================================================
# Helpers
# ============================================================


def _load_table(run_id: str) -> ComparisonTable:
    """Load a comparison table by run_id. Raises 404 if not found."""
    filepath = _PIPELINE_DIR / f"{run_id}.json"
    if not filepath.exists():
        raise HTTPException(status_code=404, detail=f"Run '{run_id}' not found")
    try:
        return ComparisonTable.load(run_id, _PIPELINE_DIR)
    except Exception as e:
        logger.error(f"Failed to load pipeline state {run_id}: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")


def _save_table(table: ComparisonTable) -> None:
    """Save the comparison table back to disk."""
    try:
        table.save()
    except Exception as e:
        logger.error(f"Failed to save pipeline state: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")


def _row_to_api(row_dict: dict) -> dict:
    """Enrich a flat row dict with display labels for the frontend."""
    verdict = row_dict.get("verdict")
    risk = row_dict.get("risk_level")

    v_disp = VERDICT_DISPLAY.get(verdict or "", {})
    r_disp = RISK_LEVEL_DISPLAY.get(risk or "", {})

    row_dict["verdict_icon"] = v_disp.get("icon", "")
    row_dict["verdict_label_zh"] = v_disp.get("label_zh", verdict or "")
    row_dict["verdict_label_en"] = v_disp.get("label_en", verdict or "")
    row_dict["risk_icon"] = r_disp.get("icon", "")
    row_dict["risk_label_zh"] = r_disp.get("label_zh", risk or "")
    row_dict["risk_label_en"] = r_disp.get("label_en", risk or "")
    row_dict["risk_action_zh"] = r_disp.get("action_zh", "")

    return row_dict


# ============================================================
# Report page serving
# ============================================================


@report_router.get("/page/{run_id}", response_class=HTMLResponse)
async def serve_report_page(run_id: str):
    """Serve the report HTML page for a specific run.

    Special alias: run_id='latest' resolves to the most recent pipeline run.
    """
    if run_id == "latest":
        runs = ComparisonTable.list_runs(_PIPELINE_DIR)
        if not runs:
            raise HTTPException(status_code=404, detail="No pipeline runs found")
        run_id = runs[0].get("run_id", "")
        if not run_id:
            raise HTTPException(status_code=404, detail="No valid run_id found")

    if not re.match(r'^[a-zA-Z0-9_\-]+$', run_id):
        raise HTTPException(status_code=400, detail="Invalid run_id format")

    html_path = REPORT_STATIC_DIR / "report.html"
    if not html_path.exists():
        raise HTTPException(status_code=404, detail="Report page not found")

    # Verify run exists
    filepath = _PIPELINE_DIR / f"{run_id}.json"
    if not filepath.exists():
        raise HTTPException(status_code=404, detail=f"Run '{run_id}' not found")

    # Read and inject run_id into the HTML template
    html_content = html_path.read_text(encoding="utf-8")
    html_content = html_content.replace("{{RUN_ID}}", escape(run_id))
    return HTMLResponse(content=html_content)


@report_router.get("/static/{filename}")
async def serve_report_static(filename: str):
    """Serve static files (JS, CSS) for the report page."""
    allowed_extensions = {".js", ".css", ".svg", ".png", ".ico"}
    ext = Path(filename).suffix.lower()
    if ext not in allowed_extensions:
        raise HTTPException(status_code=403, detail="File type not allowed")

    filepath = REPORT_STATIC_DIR / filename
    filepath = filepath.resolve()
    if not str(filepath).startswith(str(REPORT_STATIC_DIR.resolve())):
        raise HTTPException(status_code=403, detail="Access denied")
    if not filepath.exists():
        raise HTTPException(status_code=404, detail=f"File not found: {filename}")

    content_types = {
        ".js": "application/javascript",
        ".css": "text/css",
        ".svg": "image/svg+xml",
        ".png": "image/png",
        ".ico": "image/x-icon",
    }
    return FileResponse(
        filepath, media_type=content_types.get(ext, "application/octet-stream")
    )


# ============================================================
# API Endpoints — Read
# ============================================================


@report_router.get("/runs")
async def list_runs():
    """List all saved pipeline runs."""
    runs = ComparisonTable.list_runs(_PIPELINE_DIR)
    return JSONResponse(content={"runs": runs})


@report_router.get("/latest")
async def redirect_to_latest_run():
    """Redirect to the most recent pipeline run's report page.

    Finds the newest run_*.json file and redirects to /api/report/page/{run_id}.
    Returns 404 if no runs exist.
    """
    runs = ComparisonTable.list_runs(_PIPELINE_DIR)
    if not runs:
        raise HTTPException(status_code=404, detail="No pipeline runs found")

    # Runs are sorted by modified time (newest first) by list_runs()
    latest_run_id = runs[0].get("run_id", "")
    if not latest_run_id:
        raise HTTPException(status_code=404, detail="No valid run_id found")

    from starlette.responses import RedirectResponse

    return RedirectResponse(url=f"/api/report/page/{latest_run_id}")


# ============================================================
# Cross-Reference Comparison API
# ============================================================


def _get_cross_ref_modules():
    """Lazy import cross-reference modules to avoid circular imports."""
    from src.analysis.compliance_rules import (
        ISO_13485_CHECKLIST,
        get_all_regulations,
        get_regulation,
        get_overlap_analysis,
        generate_cross_exam_questions,
        MappingStatus,
    )

    return {
        "ISO_13485_CHECKLIST": ISO_13485_CHECKLIST,
        "get_all_regulations": get_all_regulations,
        "get_regulation": get_regulation,
        "get_overlap_analysis": get_overlap_analysis,
        "generate_cross_exam_questions": generate_cross_exam_questions,
        "MappingStatus": MappingStatus,
    }


@report_router.get("/crossref/regulations")
async def list_regulations():
    """List all available regulations (predefined + crawled) for country selector."""
    mods = _get_cross_ref_modules()
    all_regs = mods["get_all_regulations"]()

    regulations = []
    for reg_id, profile in all_regs.items():
        iso_mapped_count = len(profile.iso_mapped)
        unique_count = len(profile.unique_requirements)

        # Count statuses
        status_counts = defaultdict(int)
        for cm in profile.iso_mapped.values():
            status_counts[cm.status.value] += 1

        regulations.append(
            {
                "regulation_id": reg_id,
                "name_en": profile.name_en,
                "name_zh": profile.name_zh,
                "country": profile.country,
                "country_name_en": profile.country_name_en,
                "country_name_zh": profile.country_name_zh,
                "source": profile.source,
                "source_url": profile.source_url,
                "last_updated": profile.last_updated,
                "effective_date": profile.effective_date,
                "iso_mapped_count": iso_mapped_count,
                "unique_requirements_count": unique_count,
                "status_counts": dict(status_counts),
            }
        )

    failed_regions = []
    try:
        crawl_results_path = Path("data") / "regulatory_crawl_results.json"
        if crawl_results_path.exists():
            with open(crawl_results_path, "r", encoding="utf-8") as f:
                crawl_data = json.load(f)
            available_profile_regions = set()
            for profile in all_regs.values():
                if profile.country_name_zh:
                    available_profile_regions.add(profile.country_name_zh)
                if profile.country_name_en:
                    available_profile_regions.add(profile.country_name_en)
            for cr in crawl_data.get("results", []):
                if cr.get("crawl_status") != "success":
                    region = cr.get("region", "")
                    has_profile = any(
                        region in str(available_profile_regions)
                        or cr.get("agency", "") in (p.source or "")
                        for p in all_regs.values()
                    )
                    if not has_profile:
                        failed_regions.append(
                            {
                                "region": region,
                                "agency": cr.get("agency", ""),
                                "reason": cr.get("failure_reason", "未知錯誤"),
                                "url": cr.get("url", ""),
                            }
                        )
    except Exception:
        pass

    return JSONResponse(
        content={
            "regulations": regulations,
            "failed_regions": failed_regions,
        }
    )


@report_router.get("/crossref/table")
async def get_crossref_table(
    regulations: str = Query(
        ..., description="Comma-separated regulation IDs, e.g. QMSR,EU_MDR,TFDA"
    ),
):
    """Get the full cross-reference comparison table.

    Returns ISO 13485 clauses as rows, with each selected regulation's
    mapping status, rationale, method, and confidence.
    Also returns unique requirements (delta items) per regulation.
    """
    mods = _get_cross_ref_modules()
    reg_ids = [r.strip() for r in regulations.split(",") if r.strip()]

    if not reg_ids:
        raise HTTPException(status_code=400, detail="No regulations specified")

    # Validate regulation IDs
    all_regs = mods["get_all_regulations"]()
    for rid in reg_ids:
        if rid not in all_regs:
            raise HTTPException(status_code=404, detail=f"Regulation '{rid}' not found")

    # Build cross-reference rows from ISO 13485 checklist
    checklist = mods["ISO_13485_CHECKLIST"]
    rows = []

    for clause_id, clause_info in checklist.items():
        row = {
            "clause_id": clause_id,
            "clause_title": clause_info.get("title", ""),
            "audit_impact": clause_info.get("audit_impact", ""),
            "regulations": {},
        }

        for rid in reg_ids:
            analysis = mods["get_overlap_analysis"](rid, clause_id)
            profile = all_regs[rid]

            # Flatten: mapping details are nested inside analysis["mapping"]
            mapping = analysis.get("mapping") or {}
            reg_data = {
                "status": analysis.get("status", "na"),
                "is_delta": analysis.get("is_delta", False),
                "regulation_ref": mapping.get("regulation_ref", ""),
                "rationale_en": mapping.get("rationale_en", ""),
                "rationale_zh": mapping.get("rationale_zh", ""),
                "method": mapping.get("method", ""),
                "confidence": mapping.get("confidence", 0.0),
                "notes": mapping.get("notes", ""),
                "original_text": mapping.get("original_text", ""),
                "original_lang": mapping.get("original_lang", ""),
                "english_translation": mapping.get("english_translation", ""),
                "semantic_note": mapping.get("semantic_note", ""),
                "delta_items": analysis.get("delta_items", []),
            }
            row["regulations"][rid] = reg_data

        rows.append(row)

    # Collect unique requirements (delta) per regulation
    unique_reqs = {}
    for rid in reg_ids:
        profile = all_regs[rid]
        reqs = []
        for req in profile.unique_requirements:
            reqs.append(
                {
                    "req_id": req.req_id,
                    "regulation_ref": req.regulation_ref,
                    "title_en": req.title_en,
                    "title_zh": req.title_zh,
                    "requirement_en": req.requirement_en,
                    "requirement_zh": req.requirement_zh,
                    "related_iso_clauses": req.related_iso_clauses,
                    "audit_impact": req.audit_impact,
                    "audit_question_en": req.audit_question_en,
                    "audit_question_zh": req.audit_question_zh,
                    "expected_evidence": req.expected_evidence,
                    "rationale_en": req.rationale_en,
                    "rationale_zh": req.rationale_zh,
                    "method": req.method.value
                    if hasattr(req.method, "value")
                    else str(req.method),
                    "confidence": req.confidence,
                    "original_text": req.original_text,
                    "original_lang": req.original_lang,
                    "english_translation": req.english_translation,
                    "semantic_note": req.semantic_note,
                }
            )
        unique_reqs[rid] = reqs

    # Regulation metadata for the header
    reg_meta = {}
    for rid in reg_ids:
        p = all_regs[rid]
        reg_meta[rid] = {
            "name_en": p.name_en,
            "name_zh": p.name_zh,
            "country": p.country,
            "country_name_zh": p.country_name_zh,
            "source": p.source,
            "effective_date": p.effective_date,
        }

    return JSONResponse(
        content={
            "regulation_ids": reg_ids,
            "regulation_meta": reg_meta,
            "iso_clause_count": len(checklist),
            "rows": rows,
            "unique_requirements": unique_reqs,
        }
    )


@report_router.get("/crossref/questions")
async def get_crossref_questions(
    doc_id: str = Query(..., description="Document ID, e.g. QP-852"),
    doc_title: str = Query("", description="Document title"),
    baseline_clause: str = Query(
        "", description="Primary ISO 13485 clause, e.g. 8.5.2"
    ),
    regulations: str = Query(..., description="Comma-separated regulation IDs"),
    doc_content_summary: str = Query("", description="Brief document content summary"),
):
    """Generate cross-examination questions for a specific document.

    Returns prioritized questions: delta (highest) > exceeds > overlap.
    """
    mods = _get_cross_ref_modules()
    reg_ids = [r.strip() for r in regulations.split(",") if r.strip()]

    if not reg_ids:
        raise HTTPException(status_code=400, detail="No regulations specified")

    questions = mods["generate_cross_exam_questions"](
        doc_id=doc_id,
        doc_title=doc_title,
        baseline_clause=baseline_clause,
        selected_regulations=reg_ids,
        doc_content_summary=doc_content_summary,
    )

    return JSONResponse(
        content={
            "doc_id": doc_id,
            "baseline_clause": baseline_clause,
            "selected_regulations": reg_ids,
            "total_questions": len(questions),
            "questions": questions,
        }
    )


# ============================================================
# Cross-Reference Validation Report Endpoints
# ============================================================


@report_router.get("/crossref/validation")
async def get_crossref_validation(
    regulations: str = "QMSR,EU_MDR,TFDA",
    lang: str = "zh-en",
):
    """Generate a static 3-country × ISO 13485 cross-reference validation report.

    No LLM calls — pure Python data assembly from predefined regulation profiles.

    Args:
        regulations: Comma-separated regulation IDs (default: all 3)
        lang: Language mode: 'zh-en' (default), 'local' (add native text), 'zh-only'
    """
    from src.analysis.crossref_report import (
        generate_crossref_validation_report,
        format_crossref_report_markdown,
    )

    reg_ids = [r.strip() for r in regulations.split(",") if r.strip()]
    report = generate_crossref_validation_report(reg_ids)

    return JSONResponse(
        content={
            "report": report,
            "markdown": format_crossref_report_markdown(report, language_mode=lang),
        }
    )


# ============================================================
# User Language & Regulation Upload Reminders
# ============================================================


@report_router.get("/user/language")
async def get_user_language():
    """Get the current user's language setting from Chainlit user_settings.

    Returns the language code (e.g., 'zh-TW', 'en-US', 'ja-JP').
    The crossref table uses this to set the default display language.
    """
    try:
        from src.utils.user_settings import load_user_settings

        settings = load_user_settings()
        language = settings.get("language", "zh-TW")
    except Exception:
        language = "zh-TW"
    return JSONResponse(content={"language": language})


@report_router.get("/crossref/upload-reminders")
async def get_regulation_upload_reminders():
    """Get list of regulations that need user-uploaded full text.

    Similar to ISO 13485: when official regulation text cannot be
    auto-retrieved, the system reminds the user to upload it.
    Applies to all 7 countries (US, EU, TW, CA, JP, BR, AU).
    """
    try:
        from src.storage.mdsap_markdown_storage import get_mdsap_markdown_store

        store = get_mdsap_markdown_store()
        reminders = store.get_upload_reminders()
        all_status = store.list_all_regulations()
        return JSONResponse(
            content={
                "reminders": reminders,
                "regulations": all_status,
                "total_countries": len(all_status),
                "needs_upload_count": len(reminders),
            }
        )
    except ImportError:
        return JSONResponse(
            content={
                "reminders": [],
                "regulations": [],
                "total_countries": 0,
                "needs_upload_count": 0,
                "error": "mdsap_markdown_storage module not available",
            }
        )
    except Exception as e:
        logger.error(f"Failed to get upload reminders: {e}")
        return JSONResponse(
            content={
                "reminders": [],
                "regulations": [],
                "total_countries": 0,
                "needs_upload_count": 0,
                "error": str(e),
            }
        )


@report_router.post("/crossref/upload-regulation")
async def upload_regulation_text(request: Request):
    """Upload full regulation text for a specific country.

    Body JSON:
        regulation_id: str (e.g., 'PMDA', 'ANVISA')
        filename: str (original filename)
        content: str (full Markdown/text content)
    """
    try:
        body = await request.json()
        regulation_id = body.get("regulation_id", "")
        filename = body.get("filename", "")
        content = body.get("content", "")

        if not regulation_id or not content:
            raise HTTPException(
                status_code=400, detail="regulation_id and content are required"
            )

        from src.storage.mdsap_markdown_storage import get_mdsap_markdown_store

        store = get_mdsap_markdown_store()
        result = store.save_uploaded_regulation(regulation_id, filename, content)

        if result.get("success"):
            return JSONResponse(content=result)
        else:
            raise HTTPException(
                status_code=400, detail=result.get("error", "Upload failed")
            )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to upload regulation: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================
# Supplemental Standards API Endpoints
# ============================================================


def _get_standards_modules():
    """Lazy import supplemental standards to avoid circular imports."""
    from src.analysis.compliance_rules import (
        get_all_standards,
        get_standard,
        get_applicable_standards,
        adjust_standard_clause_mapping,
        ProductProfile,
        StandardCategory,
    )

    return {
        "get_all_standards": get_all_standards,
        "get_standard": get_standard,
        "get_applicable_standards": get_applicable_standards,
        "adjust_standard_clause_mapping": adjust_standard_clause_mapping,
        "ProductProfile": ProductProfile,
        "StandardCategory": StandardCategory,
    }


@report_router.get("/standards/list")
async def list_supplemental_standards():
    """List all available supplemental standards.

    Returns all predefined standards with their categories,
    detection keywords, ISO 13485 clause links, and regulatory references.
    """
    mods = _get_standards_modules()
    all_stds = mods["get_all_standards"]()

    standards = []
    for std_id, std in all_stds.items():
        standards.append(
            {
                "standard_id": std.standard_id,
                "name_en": std.name_en,
                "name_zh": std.name_zh,
                "category": std.category.value,
                "version": std.version,
                "is_universal": std.is_universal,
                "primary_iso_clauses": std.primary_iso_clauses,
                "clause_links": [
                    {
                        "standard_clause": cl.standard_clause,
                        "iso_13485_clause": cl.iso_13485_clause,
                        "relationship": cl.relationship,
                        "description_en": cl.description_en,
                        "description_zh": cl.description_zh,
                    }
                    for cl in std.clause_links
                ],
                "regulatory_references": std.regulatory_references,
                "audit_questions": std.audit_questions,
                "detection_keywords_en": std.detection_keywords_en,
                "detection_keywords_zh": std.detection_keywords_zh,
            }
        )

    return JSONResponse(content={"standards": standards})


@report_router.post("/standards/applicable")
async def get_applicable_standards_endpoint(request: Request):
    """Determine which supplemental standards apply based on product profile.

    Accepts a JSON body with product characteristics:
    {
        "has_software": true/false,
        "has_electrical": true/false,
        "is_implantable": true/false,
        "is_sterile": true/false,
        "sterilization_method": "eo"/"radiation"/"steam"/"",
        "has_biological_contact": true/false,
        "user_confirmed_standards": ["ISO_14971", ...],
        "user_rejected_standards": [...],
        "detected_standard_refs": ["ISO 14971", "IEC 62304", ...],
        "uploaded_standard_files": ["ISO_11135.pdf", ...]
    }

    Returns applicable standards list with ISO 13485 clause mappings.
    """
    mods = _get_standards_modules()
    body = await request.json()

    # Build ProductProfile from request body
    ProfileCls = mods["ProductProfile"]
    profile = ProfileCls(
        has_software=(
            body.get("has_software", False),
            body.get("has_software_confidence", 0.5),
            body.get("has_software_source", "user_manual"),
        ),
        has_electrical=(
            body.get("has_electrical", False),
            body.get("has_electrical_confidence", 0.5),
            body.get("has_electrical_source", "user_manual"),
        ),
        is_implantable=(
            body.get("is_implantable", False),
            body.get("is_implantable_confidence", 0.5),
            body.get("is_implantable_source", "user_manual"),
        ),
        is_sterile=(
            body.get("is_sterile", False),
            body.get("is_sterile_confidence", 0.5),
            body.get("is_sterile_source", "user_manual"),
        ),
        sterilization_method=body.get("sterilization_method", ""),
        has_biological_contact=(
            body.get("has_biological_contact", False),
            body.get("has_biological_contact_confidence", 0.5),
            body.get("has_biological_contact_source", "user_manual"),
        ),
        is_ivd=(
            body.get("is_ivd", False),
            body.get("is_ivd_confidence", 0.5),
            body.get("is_ivd_source", "user_manual"),
        ),
        has_clinical_investigation=(
            body.get("has_clinical_investigation", False),
            body.get("has_clinical_investigation_confidence", 0.5),
            body.get("has_clinical_investigation_source", "user_manual"),
        ),
        user_confirmed_standards=body.get("user_confirmed_standards", []),
        user_rejected_standards=body.get("user_rejected_standards", []),
        detected_standard_refs=body.get("detected_standard_refs", []),
        uploaded_standard_files=body.get("uploaded_standard_files", []),
    )

    applicable = mods["get_applicable_standards"](profile)

    results = []
    for std in applicable:
        results.append(
            {
                "standard_id": std.standard_id,
                "name_en": std.name_en,
                "name_zh": std.name_zh,
                "category": std.category.value,
                "is_universal": std.is_universal,
                "primary_iso_clauses": std.primary_iso_clauses,
                "clause_links": [
                    {
                        "standard_clause": cl.standard_clause,
                        "iso_13485_clause": cl.iso_13485_clause,
                        "relationship": cl.relationship,
                        "description_en": cl.description_en,
                        "description_zh": cl.description_zh,
                    }
                    for cl in std.clause_links
                ],
                "regulatory_references": std.regulatory_references,
                "audit_questions": std.audit_questions,
            }
        )

    return JSONResponse(
        content={
            "product_profile": {
                "has_software": body.get("has_software", False),
                "has_electrical": body.get("has_electrical", False),
                "is_implantable": body.get("is_implantable", False),
                "is_sterile": body.get("is_sterile", False),
                "sterilization_method": body.get("sterilization_method", ""),
                "has_biological_contact": body.get("has_biological_contact", False),
            },
            "applicable_standards_count": len(results),
            "applicable_standards": results,
        }
    )


@report_router.post("/standards/adjust")
async def adjust_standard_mapping(request: Request):
    """Adjust a supplemental standard's clause-to-ISO-13485 mapping.

    Accepts a JSON body:
    {
        "standard_id": "ISO_14971",
        "standard_clause": "ISO 14971 Clause 4",
        "old_iso_clause": "7.1",
        "new_iso_clause": "7.3.3"
    }

    The old_iso_clause is optional but recommended for safety verification.
    Changes persist for the server session duration.
    """
    with _phoenix_report_span("report_adjust_standard"):
        mods = _get_standards_modules()
        body = await request.json()

        standard_id = body.get("standard_id", "")
        standard_clause = body.get("standard_clause", "")
        old_iso_clause = body.get("old_iso_clause", "")
        new_iso_clause = body.get("new_iso_clause", "")

        if not standard_id or not standard_clause or not new_iso_clause:
            return JSONResponse(
                status_code=400,
                content={
                    "success": False,
                    "message": "Required fields: standard_id, standard_clause, new_iso_clause",
                },
            )

        result = mods["adjust_standard_clause_mapping"](
            standard_id=standard_id,
            standard_clause=standard_clause,
            old_iso_clause=old_iso_clause,
            new_iso_clause=new_iso_clause,
        )

        status_code = 200 if result["success"] else 400
        return JSONResponse(status_code=status_code, content=result)


# ============================================================
# API Endpoints — Read (path-parameter routes MUST come AFTER
# all static-path routes to avoid /{run_id} catching /crossref etc.)
# ============================================================


@report_router.get("/{run_id}")
async def get_report(run_id: str):
    """Get full report data for a specific run."""
    table = _load_table(run_id)
    state = table.state

    # Build enriched row list
    flat_rows = table.to_flat_rows()
    enriched = [_row_to_api(r) for r in flat_rows]

    # Summary
    summary = table.summary()
    progress = state.progress_summary()

    # Data quality info
    dq = state.data_quality_summary

    # Source check info
    sc = state.source_check_summary

    # LLM budget
    budget = state.get_budget().to_dict()

    return JSONResponse(
        content={
            "run_id": state.run_id,
            "created_at": state.created_at,
            "updated_at": state.updated_at,
            "completed_at": state.completed_at,
            "status": state.status,
            "mode": state.mode,
            "standard": state.standard,
            "source_command": getattr(state, "source_command", "regulatory_list"),
            "current_phase": state.current_phase,
            "rows": enriched,
            "summary": summary,
            "progress": progress,
            "data_quality": dq,
            "source_check": sc,
            "llm_budget": budget,
            "verdict_options": [
                {"value": v, **VERDICT_DISPLAY.get(v, {})} for v in Verdict.ALL
            ],
            "risk_options": [
                {"value": r, **RISK_LEVEL_DISPLAY.get(r, {})} for r in RiskLevel.ALL
            ],
        }
    )


@report_router.get("/{run_id}/summary")
async def get_summary(run_id: str):
    """Get summary statistics only."""
    table = _load_table(run_id)
    summary = table.summary()
    progress = table.state.progress_summary()
    budget = table.state.get_budget().to_dict()
    qa_audit_summary = table.state.qa_audit_summary

    resp = {
        "summary": summary,
        "progress": progress,
        "llm_budget": budget,
    }
    if qa_audit_summary:
        resp["qa_audit_summary"] = qa_audit_summary

    return JSONResponse(content=resp)


@report_router.get("/{run_id}/rows")
async def get_rows(
    run_id: str,
    doc_id: Optional[str] = Query(None, description="Filter by document ID"),
    clause_id: Optional[str] = Query(None, description="Filter by clause ID"),
    verdict: Optional[str] = Query(None, description="Filter by verdict"),
    risk_level: Optional[str] = Query(None, description="Filter by risk level"),
    flagged: Optional[bool] = Query(None, description="Filter flagged rows only"),
    search: Optional[str] = Query(None, description="Search in clause/doc titles"),
):
    """Get rows with optional filters."""
    table = _load_table(run_id)
    flat_rows = table.to_flat_rows()

    # Apply filters
    if doc_id:
        flat_rows = [r for r in flat_rows if r["doc_id"] == doc_id]
    if clause_id:
        flat_rows = [r for r in flat_rows if r["clause_id"] == clause_id]
    if verdict:
        flat_rows = [r for r in flat_rows if r["verdict"] == verdict]
    if risk_level:
        flat_rows = [r for r in flat_rows if r.get("risk_level") == risk_level]
    if flagged is True:
        flat_rows = [r for r in flat_rows if r.get("flagged_for_ra")]
    if search:
        search_lower = search.lower()
        flat_rows = [
            r
            for r in flat_rows
            if search_lower
            in (r.get("clause_title", "") + r.get("doc_title", "")).lower()
            or search_lower in r.get("clause_id", "").lower()
            or search_lower in r.get("doc_id", "").lower()
        ]

    enriched = [_row_to_api(r) for r in flat_rows]

    return JSONResponse(
        content={
            "total": len(enriched),
            "rows": enriched,
        }
    )


@report_router.get("/{run_id}/row/{row_id}")
async def get_row_detail(run_id: str, row_id: str):
    """Get detailed data for a single row, including all phase results."""
    table = _load_table(run_id)
    row = table.state.get_row(row_id)
    if row is None:
        raise HTTPException(status_code=404, detail=f"Row '{row_id}' not found")

    row_dict = row.to_dict()
    row_dict = _row_to_api(row_dict)

    return JSONResponse(content={"row": row_dict})


@report_router.get("/{run_id}/row/{row_id}/history")
async def get_row_history(run_id: str, row_id: str):
    """Get version history for a specific row."""
    table = _load_table(run_id)
    row = table.state.get_row(row_id)
    if row is None:
        raise HTTPException(status_code=404, detail=f"Row '{row_id}' not found")

    return JSONResponse(
        content={
            "row_id": row_id,
            "version_history": row.version_history,
            "ra_override": row.ra_override,
            "ra_notes": row.ra_notes,
        }
    )


# ============================================================
# Unified LLM-Assist Analysis (Human Intervention + LLM)
# ============================================================


async def _llm_assist_analyze(
    request: Request,
    user_input: str,
    context_type: str,
    context_data: dict = None,
) -> dict:
    """Unified LLM analysis for human intervention inputs.

    Handles:
    - Plain text input → LLM analyzes directly
    - URLs → fetches content, then LLM analyzes
    - Mixed → extracts URLs, fetches content, combines with text
    """
    import re as _re

    with _phoenix_report_span(
        "llm_assist_analyze",
        {
            "context_type": context_type,
            "input_length": len(user_input),
        },
    ):
        # 1. Extract URLs from input
        url_pattern = _re.compile(r'https?://[^\s<>"\')\]]+')
        found_urls = url_pattern.findall(user_input)
        text_without_urls = url_pattern.sub("", user_input).strip()

        # 2. Fetch URL content if any
        fetched_content = {}
        if found_urls:
            try:
                from src.chainlit_app.app import _fetch_web_full_content

                fetched_content = await _fetch_web_full_content(
                    found_urls, max_urls=3, timeout=15.0
                )
            except Exception as e:
                logger.warning(f"URL fetch failed in llm_assist: {e}")
                try:
                    import httpx

                    async with httpx.AsyncClient(
                        timeout=10.0, follow_redirects=True
                    ) as client:
                        for url in found_urls[:3]:
                            try:
                                resp = await client.get(url)
                                if resp.status_code == 200:
                                    fetched_content[url] = resp.text[:5000]
                            except Exception:
                                pass
                except Exception:
                    pass

        # 3. Build LLM prompt based on context_type
        context_prompts = {
            "override": "你是品質管理系統的稽核助手。使用者想要覆寫一個條款的判定結果。請根據使用者提供的資訊（文字、URL內容、文章、算式），分析是否有合理依據支持覆寫，並建議新的判定結果與理由。",
            "note": "你是品質管理系統的稽核助手。使用者想要為一個條款新增備註。請根據使用者提供的資訊，整理成專業的稽核備註內容。",
            "evidence": "你是品質管理系統的稽核助手。使用者想要補充或修改風險證據項目。請分析並建議新的證據項目，包含 evidence_name、found、source_doc_id、relevance_score、is_inadequate、is_outdated、llm_reasoning。返回 JSON 格式。",
            "inject": "你是品質管理系統的交叉詰問助手。使用者想要注入一段內容到交叉詰問過程中。請整理成適合交叉詰問的專業提問或回應。",
            "feedback": "你是品質管理系統的稽核助手。使用者想要對每日稽核結果提出回饋意見。請整理成結構化的稽核改善建議。",
            "standards": "你是品質管理系統的法規標準助手。使用者想要調整標準條款對應關係。請分析並建議合適的 ISO 13485 條款對應。",
        }

        system_prompt = context_prompts.get(context_type, context_prompts["note"])

        if context_data:
            system_prompt += f"\n\n目前的條款/項目資訊：\n```json\n{json.dumps(context_data, ensure_ascii=False, indent=2)}\n```"

        # Build user message
        user_message_parts = []
        if text_without_urls:
            user_message_parts.append(f"使用者輸入：\n{text_without_urls}")
        if fetched_content:
            user_message_parts.append("\n\n--- 以下為使用者提供的 URL 內容 ---")
            for url, content in fetched_content.items():
                user_message_parts.append(f"\n### {url}\n{content[:3000]}")

        user_message = (
            "\n".join(user_message_parts) if user_message_parts else user_input
        )
        user_message += "\n\n請分析以上資訊並提供你的建議。如果使用者提供了算式，請驗證計算結果。回覆請使用繁體中文。"

        # 4. Call LLM
        try:
            llm_fn = _get_llm_completion_fn(request)
            response = llm_fn(
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_message},
                ],
                temperature=0.3,
                max_tokens=2000,
            )
            if hasattr(response, "choices") and response.choices:
                analysis_text = response.choices[0].message.content
            else:
                analysis_text = str(response)
        except Exception as e:
            logger.error(f"LLM assist analysis failed: {e}")
            analysis_text = f"LLM 分析失敗：{str(e)}。您提供的原始內容已保留。"

        return {
            "analysis": analysis_text,
            "fetched_urls": list(fetched_content.keys()),
            "fetched_content_lengths": {
                url: len(c) for url, c in fetched_content.items()
            },
            "raw_input": user_input,
            "context_type": context_type,
            "timestamp": time.time(),
        }


@report_router.post("/llm-assist")
async def llm_assist_endpoint(request: Request):
    """Unified LLM-assist analysis endpoint for all human intervention points.

    Body:
        {
            "user_input": "使用者的自由輸入（文字、URL、文章、算式皆可）",
            "context_type": "override|note|evidence|inject|feedback|standards",
            "context_data": { ... optional ... }
        }
    """
    try:
        body = await request.json()
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid JSON body")

    user_input = body.get("user_input", "").strip()
    context_type = body.get("context_type", "note")
    context_data = body.get("context_data")

    if not user_input:
        raise HTTPException(status_code=400, detail="Missing 'user_input' field")

    valid_types = ("override", "note", "evidence", "inject", "feedback", "standards")
    if context_type not in valid_types:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid context_type: {context_type}. Must be one of {valid_types}",
        )

    with _phoenix_report_span(
        "llm_assist_endpoint",
        {
            "context_type": context_type,
            "input_length": len(user_input),
        },
    ):
        result = await _llm_assist_analyze(
            request, user_input, context_type, context_data
        )

    return JSONResponse(content={"success": True, "result": result})


# ============================================================
# API Endpoints — Write (RA Modifications)
# ============================================================


@report_router.post("/{run_id}/row/{row_id}/override")
async def override_verdict(run_id: str, row_id: str, body: dict):
    """RA overrides a row's verdict.

    Body:
        {"verdict": "full_compliance", "reason": "已確認文件內容符合要求"}
    """
    with _phoenix_report_span(
        "report_override_verdict", {"run_id": run_id, "row_id": row_id}
    ):
        new_verdict = body.get("verdict")
        reason = body.get("reason", "")
        user_id = body.get("user_id", "ra_user")

        if not new_verdict:
            raise HTTPException(status_code=400, detail="Missing 'verdict' field")
        if new_verdict not in Verdict.ALL:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid verdict: {new_verdict}. Must be one of {Verdict.ALL}",
            )
        if not reason:
            raise HTTPException(status_code=400, detail="Missing 'reason' field")

        table = _load_table(run_id)
        updated = table.override_verdict(row_id, new_verdict, reason, user_id)
        if updated is None:
            raise HTTPException(status_code=404, detail=f"Row '{row_id}' not found")

        # Auto-recalculate risk based on new verdict
        # Verdict → gap_severity reverse mapping, then risk matrix lookup
        if updated.audit_impact and new_verdict:
            new_risk = _verdict_to_risk(new_verdict, updated.audit_impact)
            if new_risk:
                updated.risk_level = new_risk
                table.state.update_row(updated)

        _save_table(table)

        row_dict = _row_to_api(updated.to_dict())
        return JSONResponse(content={"success": True, "row": row_dict})


@report_router.post("/{run_id}/row/{row_id}/note")
async def add_note(run_id: str, row_id: str, body: dict):
    """Add a permanent note to a clause.

    Body:
        {"note": "此條款在本公司不適用，因為..."}
    """
    with _phoenix_report_span("report_add_note", {"run_id": run_id, "row_id": row_id}):
        note = body.get("note", "")
        user_id = body.get("user_id", "ra_user")

        if not note:
            raise HTTPException(status_code=400, detail="Missing 'note' field")

        table = _load_table(run_id)
        updated = table.add_clause_note(row_id, note, user_id)
        if updated is None:
            raise HTTPException(status_code=404, detail=f"Row '{row_id}' not found")

        _save_table(table)

        row_dict = _row_to_api(updated.to_dict())
        return JSONResponse(content={"success": True, "row": row_dict})


@report_router.post("/{run_id}/row/{row_id}/restore")
async def restore_original(run_id: str, row_id: str):
    """Restore the LLM's original verdict (undo RA override)."""
    with _phoenix_report_span(
        "report_restore_original", {"run_id": run_id, "row_id": row_id}
    ):
        table = _load_table(run_id)
        updated = table.restore_llm_original(row_id)
        if updated is None:
            raise HTTPException(status_code=404, detail=f"Row '{row_id}' not found")

        _save_table(table)

        row_dict = _row_to_api(updated.to_dict())
        return JSONResponse(content={"success": True, "row": row_dict})


# ============================================================
# API Endpoints — Evidence Editing (Human-in-the-Loop)
# ============================================================


@report_router.post("/{run_id}/row/{row_id}/evidence/preview")
async def preview_evidence_recalc(run_id: str, row_id: str, body: dict):
    """Preview risk recalculation with modified evidence items (no save)."""
    evidence_items = body.get("evidence_items")
    if evidence_items is None:
        raise HTTPException(status_code=400, detail="Missing 'evidence_items' field")

    with _phoenix_report_span(
        "report_evidence_preview", {"run_id": run_id, "row_id": row_id}
    ):
        table = _load_table(run_id)
        result = table.preview_evidence_recalc(row_id, evidence_items)
        if result is None:
            raise HTTPException(status_code=404, detail=f"Row '{row_id}' not found")

    return JSONResponse(content={"success": True, "preview": result})


@report_router.post("/{run_id}/row/{row_id}/evidence/confirm")
async def confirm_evidence_update(run_id: str, row_id: str, body: dict):
    """Confirm and save evidence item changes with rule-engine recalculation."""
    evidence_items = body.get("evidence_items")
    if evidence_items is None:
        raise HTTPException(status_code=400, detail="Missing 'evidence_items' field")

    user_id = body.get("user_id", "ra_user")

    with _phoenix_report_span(
        "report_evidence_confirm", {"run_id": run_id, "row_id": row_id}
    ):
        table = _load_table(run_id)
        updated = table.update_evidence_items(row_id, evidence_items, user_id)
        if updated is None:
            raise HTTPException(status_code=404, detail=f"Row '{row_id}' not found")
        _save_table(table)

    row_dict = _row_to_api(updated.to_dict())
    return JSONResponse(content={"success": True, "row": row_dict})


@report_router.post("/{run_id}/row/{row_id}/evidence/deep-recalc")
async def deep_recalc_evidence(run_id: str, row_id: str, body: dict = None):
    """Deep LLM recalculation — re-runs Phase 1+2+3 for this row."""
    user_id = (body or {}).get("user_id", "ra_user")

    with _phoenix_report_span(
        "report_evidence_deep_recalc", {"run_id": run_id, "row_id": row_id}
    ):
        table = _load_table(run_id)
        row = table._state.get_row(row_id)
        if row is None:
            raise HTTPException(status_code=404, detail=f"Row '{row_id}' not found")

        row.version_history.append(
            {
                "action": "deep_recalc_requested",
                "previous_verdict": row.verdict,
                "previous_risk_level": row.risk_level,
                "by": user_id,
                "at": time.time(),
            }
        )

        updated = table.reset_row_for_rerun(row_id, Phase.GAP_SCAN)
        if updated is None:
            raise HTTPException(status_code=404, detail=f"Row '{row_id}' not found")
        _save_table(table)

    row_dict = _row_to_api(updated.to_dict())
    return JSONResponse(
        content={
            "success": True,
            "row": row_dict,
            "message": "Row reset to Phase 1. Re-run from Chainlit for deep recalculation.",
        }
    )


@report_router.post("/{run_id}/row/{row_id}/rerun")
async def reset_for_rerun(run_id: str, row_id: str, body: dict = None):
    """Reset a row to re-run from Phase 1 (Gap Scan).

    Body (optional):
        {"from_phase": "phase_1"}  (default: phase_1)
    """
    with _phoenix_report_span(
        "report_reset_for_rerun", {"run_id": run_id, "row_id": row_id}
    ):
        from_phase_str = (body or {}).get("from_phase", Phase.GAP_SCAN.value)
        try:
            from_phase = Phase(from_phase_str)
        except ValueError:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid phase: {from_phase_str}",
            )

        table = _load_table(run_id)
        updated = table.reset_row_for_rerun(row_id, from_phase)
        if updated is None:
            raise HTTPException(status_code=404, detail=f"Row '{row_id}' not found")

        _save_table(table)

        row_dict = _row_to_api(updated.to_dict())
        return JSONResponse(
            content={
                "success": True,
                "row": row_dict,
                "message": f"Row reset to {from_phase.display_name}. "
                f"Re-run the pipeline from Chainlit to execute the pending phases.",
            }
        )


# ============================================================
# API Endpoints — Export
# ============================================================


@report_router.get("/{run_id}/export/{fmt}")
async def export_report(run_id: str, fmt: str):
    """Export summary report (摘要匯出) as Word or Excel.

    fmt: "word" or "excel"
    Full report formats ("deep_word", "deep_excel") are forwarded
    to export_deep_report (完整匯出) to avoid FastAPI route-ordering conflicts.
    """
    with _phoenix_report_span("report_export", {"run_id": run_id, "fmt": fmt}):
        pass  # span covers the export operation
    # Forward deep report requests to the dedicated handler
    if fmt in ("deep_word", "deep_excel"):
        deep_fmt = fmt.replace("deep_", "")
        return await export_deep_report(run_id, deep_fmt)

    if fmt not in ("word", "excel"):
        raise HTTPException(status_code=400, detail="Format must be 'word' or 'excel'")

    table = _load_table(run_id)
    flat_rows = table.to_flat_rows()
    summary = table.summary()

    # Use existing export utilities
    export_dir = Path("data/exports")
    export_dir.mkdir(parents=True, exist_ok=True)

    source_cmd = getattr(table.state, "source_command", "regulatory_list")
    source_labels = {
        "regulatory_list": "法規清單",
        "regulatory_update": "法規清單更新",
    }
    src_label = source_labels.get(source_cmd, source_cmd)

    try:
        if fmt == "word":
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            filepath = export_dir / f"compliance_report_{source_cmd}_{run_id}.docx"
            assessment = _build_export_assessment(flat_rows, summary)

            doc = Document()
            title = doc.add_heading(f"AI-QMS 合規性分析報告（{src_label}）", level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            from datetime import datetime

            meta = doc.add_paragraph()
            meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = meta.add_run(
                f"分析 ID: {run_id}  |  來源指令: {src_label}  |  匯出時間: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            )
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)

            # Summary section
            doc.add_heading("摘要", level=2)
            doc.add_paragraph(assessment)

            # Detail table
            doc.add_heading("詳細分析結果", level=2)
            if flat_rows:
                headers = [
                    "條款",
                    "文件",
                    "稽核影響",
                    "判定",
                    "風險",
                    "差距",
                    "RA 標記",
                ]
                tbl = doc.add_table(rows=1 + len(flat_rows), cols=len(headers))
                tbl.style = "Table Grid"
                for i, h in enumerate(headers):
                    tbl.rows[0].cells[i].text = h
                _incomplete_label = "（Pipeline 未完成）"
                for ri, row in enumerate(flat_rows, 1):
                    tbl.rows[ri].cells[
                        0
                    ].text = f"{row.get('clause_id', '')} {row.get('clause_title', '')}"
                    tbl.rows[ri].cells[1].text = f"{row.get('doc_id', '')}"
                    tbl.rows[ri].cells[2].text = row.get("audit_impact", "")
                    _v_icon = row.get("verdict_icon", "")
                    _v_label = row.get("verdict_label", "")
                    tbl.rows[ri].cells[3].text = (
                        f"{_v_icon} {_v_label}"
                        if _v_icon or _v_label
                        else _incomplete_label
                    )
                    _r_icon = row.get("risk_icon", "")
                    _r_label = row.get("risk_label", "")
                    tbl.rows[ri].cells[4].text = (
                        f"{_r_icon} {_r_label}"
                        if _r_icon or _r_label
                        else _incomplete_label
                    )
                    tbl.rows[ri].cells[5].text = row.get("gap_severity", "") or ""
                    tbl.rows[ri].cells[6].text = (
                        "⚠️" if row.get("flagged_for_ra") else ""
                    )

            _qa_sum = getattr(table.state, "qa_audit_summary", None)
            if _qa_sum and not _qa_sum.get("skipped"):
                doc.add_heading("交叉詰問品質稽核摘要", level=2)
                doc.add_paragraph(
                    f"整體品質分數: {_qa_sum.get('overall_score', 0)}/100\n"
                    f"稽核條款數: {_qa_sum.get('clause_count', 0)}"
                )
                qa_sum_text = _qa_sum.get("summary", "")
                if qa_sum_text:
                    doc.add_paragraph(qa_sum_text)
                qa_recs = _qa_sum.get("recommendations", [])
                if qa_recs:
                    for qr in qa_recs:
                        doc.add_paragraph(f"• {qr}")

            safe_save_binary(filepath, doc.save)

        else:  # excel
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, PatternFill

            filepath = export_dir / f"compliance_report_{source_cmd}_{run_id}.xlsx"
            assessment = _build_export_assessment(flat_rows, summary)

            wb = Workbook()
            ws = wb.active
            ws.title = "合規分析"

            from datetime import datetime as _dt_xl

            title_cell = ws.cell(
                row=1,
                column=1,
                value=f"AI-QMS 合規性分析報告（{src_label}）",
            )
            title_cell.font = Font(bold=True, size=14)
            ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=13)

            meta_cell = ws.cell(
                row=2,
                column=1,
                value=f"分析 ID: {run_id}  |  來源指令: {src_label}  |  匯出時間: {_dt_xl.now().strftime('%Y-%m-%d %H:%M:%S')}",
            )
            meta_cell.font = Font(size=9, color="808080")
            ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=13)

            TITLE_ROWS = 3
            data_start_row = TITLE_ROWS + 1

            # Headers
            headers = [
                "條款 ID",
                "條款名稱",
                "文件 ID",
                "文件標題",
                "稽核影響",
                "稽核問題",
                "判定",
                "風險等級",
                "差距嚴重度",
                "證據 (找到/總計)",
                "RA 標記",
                "RA 覆寫",
                "RA 備註",
            ]
            header_fill = PatternFill(
                start_color="4472C4", end_color="4472C4", fill_type="solid"
            )
            header_font = Font(bold=True, color="FFFFFF", size=10)
            for ci, h in enumerate(headers, 1):
                cell = ws.cell(row=data_start_row, column=ci, value=h)
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal="center")

            _xl_incomplete = "（Pipeline 未完成）"
            for ri, row in enumerate(flat_rows, data_start_row + 1):
                ws.cell(row=ri, column=1, value=row.get("clause_id", ""))
                ws.cell(row=ri, column=2, value=row.get("clause_title", ""))
                ws.cell(row=ri, column=3, value=row.get("doc_id", ""))
                ws.cell(row=ri, column=4, value=row.get("doc_title", ""))
                ws.cell(row=ri, column=5, value=row.get("audit_impact", ""))
                ws.cell(row=ri, column=6, value=row.get("audit_question", ""))
                _xv_icon = row.get("verdict_icon", "")
                _xv_label = row.get("verdict_label", "")
                ws.cell(
                    row=ri,
                    column=7,
                    value=f"{_xv_icon} {_xv_label}"
                    if _xv_icon or _xv_label
                    else _xl_incomplete,
                )
                _xr_icon = row.get("risk_icon", "")
                _xr_label = row.get("risk_label", "")
                ws.cell(
                    row=ri,
                    column=8,
                    value=f"{_xr_icon} {_xr_label}"
                    if _xr_icon or _xr_label
                    else _xl_incomplete,
                )
                ws.cell(row=ri, column=9, value=row.get("gap_severity", "") or "")
                ws.cell(
                    row=ri,
                    column=10,
                    value=f"{row.get('evidence_found', 0)}/{row.get('evidence_total', 0)}",
                )
                ws.cell(
                    row=ri, column=11, value="Y" if row.get("flagged_for_ra") else ""
                )
                override = row.get("ra_override")
                ws.cell(
                    row=ri,
                    column=12,
                    value=override.get("reason", "")
                    if isinstance(override, dict)
                    else "",
                )
                ws.cell(row=ri, column=13, value=row.get("ra_notes", "") or "")

            # Auto-width
            for col in ws.columns:
                max_len = max((len(str(c.value or "")) for c in col), default=8)
                ws.column_dimensions[col[0].column_letter].width = min(max_len + 2, 50)

            _qa_sum_xl = getattr(table.state, "qa_audit_summary", None)
            if _qa_sum_xl and not _qa_sum_xl.get("skipped"):
                ws_qa = wb.create_sheet("QA 稽核摘要")
                qa_summary_data = [
                    ("整體品質分數", f"{_qa_sum_xl.get('overall_score', 0)}/100"),
                    ("稽核條款數", str(_qa_sum_xl.get("clause_count", 0))),
                    ("摘要", _qa_sum_xl.get("summary", "")),
                ]
                qa_recs_xl = _qa_sum_xl.get("recommendations", [])
                if qa_recs_xl:
                    qa_summary_data.append(("建議", "; ".join(qa_recs_xl)))
                for ri_qa, (k, v) in enumerate(qa_summary_data, 1):
                    ws_qa.cell(row=ri_qa, column=1, value=k).font = Font(bold=True)
                    ws_qa.cell(row=ri_qa, column=2, value=str(v))
                ws_qa.column_dimensions["A"].width = 20
                ws_qa.column_dimensions["B"].width = 80

            safe_save_binary(filepath, wb.save)
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="Export utilities not available. Check regulatory_export module.",
        )
    except Exception as e:
        logger.error(f"Export failed: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

    if not filepath.exists():
        raise HTTPException(status_code=500, detail="Export file was not created")

    content_type = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if fmt == "word"
        else "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    return FileResponse(
        filepath,
        media_type=content_type,
        filename=filepath.name,
    )


# ============================================================
# Metadata endpoints
# ============================================================


@report_router.get("/{run_id}/filters")
async def get_filter_options(run_id: str):
    """Get available filter options for the report (documents, clauses, verdicts, risks)."""
    table = _load_table(run_id)
    flat_rows = table.to_flat_rows()

    # Unique documents
    docs = {}
    clauses = {}
    verdicts_seen = set()
    risks_seen = set()

    for r in flat_rows:
        did = r.get("doc_id", "")
        if did and did not in docs:
            docs[did] = r.get("doc_title", did)

        cid = r.get("clause_id", "")
        if cid and cid not in clauses:
            clauses[cid] = r.get("clause_title", cid)

        v = r.get("verdict")
        if v:
            verdicts_seen.add(v)

        rl = r.get("risk_level")
        if rl:
            risks_seen.add(rl)

    return JSONResponse(
        content={
            "documents": [{"id": k, "title": v} for k, v in sorted(docs.items())],
            "clauses": [
                {"id": k, "title": v}
                for k, v in sorted(
                    clauses.items(),
                    key=lambda x: [int(n) for n in x[0].split(".") if n.isdigit()],
                )
            ],
            "verdicts": [
                {"value": v, **VERDICT_DISPLAY.get(v, {})}
                for v in Verdict.ALL
                if v in verdicts_seen
            ],
            "risk_levels": [
                {"value": r, **RISK_LEVEL_DISPLAY.get(r, {})}
                for r in RiskLevel.ALL
                if r in risks_seen
            ],
        }
    )


# ============================================================
# Internal helpers
# ============================================================


def _verdict_to_risk(verdict: str, audit_impact: str) -> Optional[str]:
    """Reverse-map verdict to risk level using the risk matrix.

    When RA overrides verdict, we recalculate risk deterministically.
    """
    # Verdict → gap_severity mapping (simplified)
    verdict_gap_map = {
        Verdict.FULL_COMPLIANCE: "none",
        Verdict.PARTIAL_COMPLIANCE: "incomplete",
        Verdict.NON_COMPLIANCE: "missing",
        Verdict.INSUFFICIENT_DATA: "missing",
    }

    gap_severity = verdict_gap_map.get(verdict)
    if gap_severity is None:
        return None

    result = assess_risk(audit_impact, gap_severity)
    return result  # assess_risk returns risk level string directly


def _build_export_assessment(flat_rows: list[dict], summary: dict) -> str:
    """Build a markdown assessment text from flat rows for Word/Excel export."""
    lines = [
        "# 合規性分析報告",
        "",
        f"**分析項目數**: {summary.get('total_rows', 0)}",
        f"**文件數**: {summary.get('documents_analyzed', 0)}",
        f"**需 RA 審查**: {summary.get('flagged_for_ra', 0)} 項",
        "",
        "## 判定結果分布",
        "",
    ]

    vd = summary.get("verdict_distribution", {})
    for v, count in vd.items():
        disp = VERDICT_DISPLAY.get(v, {})
        lines.append(f"- {disp.get('icon', '')} {disp.get('label_zh', v)}: {count} 項")

    lines.append("")
    lines.append("## 風險等級分布")
    lines.append("")

    rd = summary.get("risk_distribution", {})
    for r, count in rd.items():
        disp = RISK_LEVEL_DISPLAY.get(r, {})
        lines.append(f"- {disp.get('icon', '')} {disp.get('label_zh', r)}: {count} 項")

    lines.append("")
    lines.append("## 詳細結果")
    lines.append("")

    for row in flat_rows:
        lines.append(f"### {row.get('clause_id', '')} — {row.get('clause_title', '')}")
        lines.append(
            f"- **文件**: {row.get('doc_title', '')} ({row.get('doc_id', '')})"
        )
        _md_v = f"{row.get('verdict_icon', '')} {row.get('verdict_label', '')}".strip()
        _md_r = f"{row.get('risk_icon', '')} {row.get('risk_label', '')}".strip()
        lines.append(f"- **判定**: {_md_v or '（Pipeline 未完成）'}")
        lines.append(f"- **風險**: {_md_r or '（Pipeline 未完成）'}")
        lines.append(f"- **稽核影響**: {row.get('audit_impact', '')}")
        ev_found = row.get("evidence_found", 0)
        ev_total = row.get("evidence_total", 0)
        lines.append(f"- **證據**: {ev_found}/{ev_total} 項找到")

        if row.get("remediation"):
            lines.append(f"- **改善建議**: {row['remediation']}")

        if row.get("ra_override"):
            override = row["ra_override"]
            lines.append(
                f"- **RA 覆寫**: {override.get('verdict', '')} — {override.get('reason', '')}"
            )

        if row.get("ra_notes"):
            lines.append(f"- **RA 備註**: {row['ra_notes']}")

        lines.append("")

    return "\n".join(lines)


# ============================================================
# SSE Event Streaming for Real-Time Cross-Examination
# ============================================================


# Global event bus: run_id → asyncio.Queue
# Each queue holds dicts like {"type": "analyzer", "content": "...", ...}
_event_queues: dict[str, list[asyncio.Queue]] = defaultdict(list)

# ── Pipeline Control Registry (pause/resume/inject) ──
# run_id → {"pause_event": threading.Event, "injected_messages": list}
# The pause_event is SET when running, CLEARED when paused.
_pipeline_controls: dict[str, dict] = {}
_pipeline_controls_lock = threading.Lock()


def register_pipeline_control(run_id: str) -> threading.Event:
    """Register a new pipeline run for pause/resume/inject control.

    Returns a threading.Event that is initially SET (running).
    The pipeline should call event.wait() to check for pauses.
    """
    event = threading.Event()
    event.set()  # Start in running state
    with _pipeline_controls_lock:
        _pipeline_controls[run_id] = {
            "pause_event": event,
            "injected_messages": [],
        }
    return event


def unregister_pipeline_control(run_id: str) -> None:
    """Remove pipeline control entry when run completes."""
    with _pipeline_controls_lock:
        _pipeline_controls.pop(run_id, None)


def get_pause_event(run_id: str) -> threading.Event | None:
    """Get the pause event for a run_id, or None if not registered."""
    ctrl = _pipeline_controls.get(run_id)
    return ctrl["pause_event"] if ctrl else None


def get_injected_messages(run_id: str) -> list[str]:
    """Drain and return all pending injected messages for a run.

    Thread-safe: uses lock to ensure atomic drain.
    Returns a list of message strings, clearing the buffer.
    """
    with _pipeline_controls_lock:
        ctrl = _pipeline_controls.get(run_id)
        if not ctrl:
            return []
        messages = list(ctrl["injected_messages"])
        ctrl["injected_messages"].clear()
        return messages


def emit_cross_exam_event(run_id: str, event: dict) -> None:
    """Emit an event to all SSE listeners for a given run.

    Call this from the pipeline/verifier when cross-examination
    messages are generated.

    Event types:
        - round_start: {"type": "round_start", "round": 1}
        - analyzer:    {"type": "analyzer", "round": 1, "content": "...", "regulation": "..."}
        - verifier:    {"type": "verifier", "round": 1, "content": "...", "regulation": "..."}
        - round_end:   {"type": "round_end", "round": 1, "agreed": true/false}
        - complete:    {"type": "complete", "verdict": "...", "flagged": true/false}
        - error:       {"type": "error", "message": "..."}
        - human_ack:   {"type": "human_ack", "message": "..."}
    """
    event["timestamp"] = time.time()
    for queue in _event_queues.get(run_id, []):
        try:
            queue.put_nowait(event)
        except asyncio.QueueFull:
            logger.warning(f"SSE queue full for run {run_id}, dropping event")


async def _sse_generator(run_id: str, queue: asyncio.Queue):
    """Async generator that yields SSE events from the queue."""
    try:
        # Send initial connection event
        yield f"data: {json.dumps({'type': 'connected', 'run_id': run_id, 'timestamp': time.time()})}\n\n"

        while True:
            try:
                event = await asyncio.wait_for(queue.get(), timeout=30.0)
                yield f"data: {json.dumps(event, ensure_ascii=False)}\n\n"

                # If complete or error, end the stream
                if event.get("type") in ("complete", "error"):
                    break
            except asyncio.TimeoutError:
                # Send heartbeat to keep connection alive
                yield f"data: {json.dumps({'type': 'heartbeat', 'timestamp': time.time()})}\n\n"
    finally:
        # Cleanup: remove this queue from the listeners
        if run_id in _event_queues:
            try:
                _event_queues[run_id].remove(queue)
            except ValueError:
                pass
            if not _event_queues[run_id]:
                del _event_queues[run_id]


@report_router.get("/{run_id}/stream")
async def stream_cross_examination(run_id: str):
    """SSE endpoint for real-time cross-examination events.

    Client connects with EventSource and receives events as they happen.
    Events include analyzer/verifier messages, round results, and final verdict.
    """
    queue: asyncio.Queue = asyncio.Queue(maxsize=100)
    _event_queues[run_id].append(queue)

    return StreamingResponse(
        _sse_generator(run_id, queue),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@report_router.post("/{run_id}/inject")
async def inject_human_message(run_id: str, body: dict):
    """Inject a human message into the cross-examination.

    Body:
        {"message": "你問錯了，應該問...", "user_id": "ra_user"}
    """
    with _phoenix_report_span("report_inject_human", {"run_id": run_id}):
        message = body.get("message", "").strip()
        user_id = body.get("user_id", "human")

        if not message:
            raise HTTPException(status_code=400, detail="Missing 'message' field")

        # Store message for the pipeline verifier to pick up
        with _pipeline_controls_lock:
            ctrl = _pipeline_controls.get(run_id)
            if ctrl:
                ctrl["injected_messages"].append(message)

        # Broadcast the human injection event to all SSE listeners
        event = {
            "type": "human_injection",
            "message": message,
            "user_id": user_id,
            "run_id": run_id,
        }
        emit_cross_exam_event(run_id, event)

        return JSONResponse(
            content={
                "success": True,
                "message": "Human injection stored and broadcast",
            }
        )


@report_router.post("/{run_id}/pause")
async def pause_cross_examination(run_id: str):
    """Pause the pipeline/cross-examination for this run."""
    event = get_pause_event(run_id)
    if event:
        event.clear()  # Block pipeline threads waiting on this event
    emit_cross_exam_event(run_id, {"type": "pause", "run_id": run_id})
    return JSONResponse(content={"success": True, "message": "Pipeline paused"})


@report_router.post("/{run_id}/resume")
async def resume_cross_examination(run_id: str):
    """Resume the pipeline/cross-examination for this run."""
    event = get_pause_event(run_id)
    if event:
        event.set()  # Unblock pipeline threads
    emit_cross_exam_event(run_id, {"type": "resume", "run_id": run_id})
    return JSONResponse(content={"success": True, "message": "Pipeline resumed"})


# ============================================================
# Cross-Examination History Endpoints
# ============================================================


@report_router.get("/crossexam/history")
async def get_crossexam_history():
    """List all cross-examination history records (pipeline P5 + daily sampling)."""
    try:
        from src.database.crossexam_store import get_crossexam_store

        store = get_crossexam_store()
        pipeline_records = store.get_all_records()
        all_records = []
        for r in pipeline_records:
            d = r.to_dict()
            d["source"] = "pipeline"
            all_records.append(d)

        try:
            from src.database.daily_crossexam_store import get_daily_crossexam_store

            daily_store = get_daily_crossexam_store()
            daily_records = daily_store.get_all_records()
            for r in daily_records:
                d = r.to_dict()
                d["source"] = "daily"
                all_records.append(d)
        except Exception:
            pass

        all_records.sort(key=lambda x: x.get("timestamp", ""), reverse=True)

        return JSONResponse(
            content={
                "records": all_records,
                "total": len(all_records),
                "needs_meta_analysis": store.needs_meta_analysis(),
            }
        )
    except Exception as e:
        logger.error(f"Failed to load crossexam history: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@report_router.get("/crossexam/history/{record_id}")
async def get_crossexam_record(record_id: str):
    """Get a single cross-examination record by ID."""
    try:
        from src.database.crossexam_store import get_crossexam_store

        store = get_crossexam_store()
        record = store.get_record(record_id)
        if not record:
            raise HTTPException(status_code=404, detail=f"Record {record_id} not found")
        return JSONResponse(content=record.to_dict())
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@report_router.get("/crossexam/history/{record_id}/export/{fmt}")
async def export_crossexam_record(record_id: str, fmt: str):
    """Export a single cross-exam record as Word or Excel.

    fmt: 'word' or 'excel'
    """
    if fmt not in ("word", "excel"):
        raise HTTPException(status_code=400, detail="Format must be 'word' or 'excel'")

    try:
        from src.database.crossexam_store import get_crossexam_store

        store = get_crossexam_store()
        record = store.get_record(record_id)
        if not record:
            raise HTTPException(status_code=404, detail=f"Record {record_id} not found")

        from src.utils.crossexam_export import (
            export_crossexam_record_word,
            export_crossexam_record_excel,
        )

        if fmt == "word":
            filepath = export_crossexam_record_word(record.to_dict())
        else:
            filepath = export_crossexam_record_excel(record.to_dict())

        content_type = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            if fmt == "word"
            else "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        return FileResponse(filepath, media_type=content_type, filename=filepath.name)

    except HTTPException:
        raise
    except ImportError:
        raise HTTPException(status_code=500, detail="Export dependencies not available")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================
# Deep Report Export Endpoints
# ============================================================


@report_router.get("/{run_id}/export/deep_{fmt}")
async def export_deep_report(run_id: str, fmt: str):
    """Export a full report including ALL LLM interactions.

    This is the "完整匯出" (Full Export) — includes comparison table,
    LLM interaction logs, cross-examination dialogue, and meta analysis.

    fmt: 'word' or 'excel' (URL: /export/deep_word or /export/deep_excel)
    """
    if fmt not in ("word", "excel"):
        raise HTTPException(status_code=400, detail="Format must be 'word' or 'excel'")

    table = _load_table(run_id)
    flat_rows = table.to_flat_rows()
    summary = table.summary()

    # Load interaction log
    interactions = None
    try:
        from src.database.interaction_log import load_interaction_log

        ilog = load_interaction_log(run_id)
        if ilog:
            interactions = ilog.get_interactions()
    except Exception:
        pass

    # Load crossexam record
    crossexam_record = None
    try:
        from src.database.crossexam_store import get_crossexam_store

        store = get_crossexam_store()
        record = store.get_record_by_run_id(run_id)
        if record:
            crossexam_record = record.to_dict()
    except Exception:
        pass

    # Load meta-analysis (if available)
    meta_analysis = None
    try:
        from src.analysis.crossexam_qa_agent import get_latest_meta_analysis

        ma = get_latest_meta_analysis()
        if ma:
            meta_analysis = ma.llm_response
    except Exception:
        pass

    qa_audit_summary = getattr(table.state, "qa_audit_summary", None)

    try:
        from src.utils.crossexam_export import (
            export_deep_report_word,
            export_deep_report_excel,
        )

        if fmt == "word":
            filepath = export_deep_report_word(
                run_id,
                flat_rows,
                summary,
                interactions,
                crossexam_record,
                meta_analysis,
                qa_audit_summary,
            )
        else:
            filepath = export_deep_report_excel(
                run_id,
                flat_rows,
                summary,
                interactions,
                crossexam_record,
                meta_analysis,
                qa_audit_summary,
            )

        content_type = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            if fmt == "word"
            else "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        return FileResponse(filepath, media_type=content_type, filename=filepath.name)

    except ImportError:
        raise HTTPException(status_code=500, detail="Export dependencies not available")
    except Exception as e:
        logger.error(f"Deep report export failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================
# Meta-Analysis Endpoints
# ============================================================


@report_router.get("/crossexam/meta-analysis")
async def get_meta_analysis():
    """Get the latest meta-analysis results."""
    try:
        from src.analysis.crossexam_qa_agent import get_latest_meta_analysis

        result = get_latest_meta_analysis()
        if not result:
            return JSONResponse(
                content={"available": False, "message": "No meta-analysis available"}
            )
        return JSONResponse(content={"available": True, **result.to_dict()})
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@report_router.get("/crossexam/meta-analysis/export/{fmt}")
async def export_meta_analysis(fmt: str):
    """Export meta-analysis report as Word or Excel.

    fmt: 'word' or 'excel'
    """
    if fmt not in ("word", "excel"):
        raise HTTPException(status_code=400, detail="Format must be 'word' or 'excel'")

    try:
        from src.analysis.crossexam_qa_agent import get_latest_meta_analysis
        from src.utils.crossexam_export import (
            export_deep_report_word,
            export_deep_report_excel,
        )

        result = get_latest_meta_analysis()
        if not result:
            raise HTTPException(status_code=404, detail="No meta-analysis available")

        # Export as a standalone report with just the meta-analysis section
        meta_analysis = result.llm_response
        if fmt == "word":
            filepath = export_deep_report_word(
                f"meta_analysis_{result.analysis_id}", [], {}, None, None, meta_analysis
            )
        else:
            filepath = export_deep_report_excel(
                f"meta_analysis_{result.analysis_id}", [], {}, None, None, meta_analysis
            )

        content_type = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            if fmt == "word"
            else "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        return FileResponse(filepath, media_type=content_type, filename=filepath.name)

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@report_router.get("/{run_id}/interactions")
async def get_interaction_log(run_id: str):
    """Get all LLM interaction logs for a pipeline run."""
    try:
        from src.database.interaction_log import load_interaction_log

        ilog = load_interaction_log(run_id)
        if not ilog:
            return JSONResponse(content={"available": False, "interactions": []})
        return JSONResponse(content={"available": True, **ilog.to_dict()})
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================
# Phase Configuration (custom phase skip)
# ============================================================

_custom_skip_phases: list[str] = []

SKIPPABLE_PHASES = {"phase_2", "phase_4", "phase_5", "phase_6"}


@report_router.post("/phase-config")
async def set_phase_config(request: Request):
    """Set which phases to skip. Only optional phases (P2, P4, P5, P6) allowed."""
    global _custom_skip_phases
    try:
        body = await request.json()
        requested = body.get("skip_phases", [])
        validated = [p for p in requested if p in SKIPPABLE_PHASES]
        _custom_skip_phases = validated
        try:
            from src.utils.app_settings import set_app_setting

            set_app_setting("custom_skip_phases", _custom_skip_phases)
        except Exception:
            pass
        return JSONResponse(content={"skip_phases": _custom_skip_phases})
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@report_router.get("/phase-config")
async def get_phase_config():
    """Get current phase skip configuration."""
    return JSONResponse(content={"skip_phases": _custom_skip_phases})


def get_custom_skip_phases() -> list[str]:
    """Get current custom skip phases (for pipeline runner integration)."""
    return list(_custom_skip_phases)


# ============================================================
# MDSAP Cross-Exam Toggle
# ============================================================


def _load_mdsap_from_settings() -> bool:
    """Load MDSAP verify state from persisted app settings."""
    try:
        from src.utils.app_settings import get_app_setting

        return bool(get_app_setting("mdsap_verify_enabled", True))
    except Exception:
        return True


_mdsap_verify_enabled: bool = _load_mdsap_from_settings()


def _load_skip_phases_from_settings() -> list[str]:
    """Load custom skip phases from persisted app settings."""
    try:
        from src.utils.app_settings import get_app_setting

        saved = get_app_setting("custom_skip_phases", [])
        return [p for p in saved if p in SKIPPABLE_PHASES]
    except Exception:
        return []


_custom_skip_phases = _load_skip_phases_from_settings()


@report_router.post("/crossref/mdsap-verify")
async def set_mdsap_verify(request: Request):
    """Enable or disable MDSAP cross-examination verification."""
    global _mdsap_verify_enabled
    try:
        body = await request.json()
        _mdsap_verify_enabled = bool(body.get("enabled", True))
        try:
            from src.utils.app_settings import set_app_setting

            set_app_setting("mdsap_verify_enabled", _mdsap_verify_enabled)
        except Exception:
            pass
        return JSONResponse(content={"enabled": _mdsap_verify_enabled})
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@report_router.get("/crossref/mdsap-verify")
async def get_mdsap_verify():
    """Get current MDSAP cross-examination verification state."""
    return JSONResponse(content={"enabled": _mdsap_verify_enabled})


# ============================================================
# Daily Audit Endpoints
# ============================================================


@report_router.get("/daily-audit/history")
async def get_daily_audit_history(limit: int = Query(30, ge=1, le=365)):
    """Get daily audit history records."""
    try:
        from src.analysis.daily_audit import get_daily_audit_history

        records = get_daily_audit_history(limit=limit)
        return JSONResponse(
            content={
                "records": [r.to_dict() for r in records],
                "count": len(records),
            }
        )
    except ImportError:
        return JSONResponse(content={"records": [], "count": 0})
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@report_router.post("/daily-audit/run")
async def run_daily_audit_endpoint(request: Request):
    """Trigger a daily audit run with pre-examination regulation freshness check.

    Flow:
      1. Freshness check (regulation versions)
      2. Read MDSAP toggle from app_settings
      3. Run daily sampling cross-exam (Phase 5 on 20% sample)
      4. Run daily audit (Dim A + Dim B evaluation)
    """
    try:
        from src.analysis.daily_audit import (
            run_daily_audit,
            run_daily_sampling_crossexam,
        )
        from src.database.crossexam_store import get_crossexam_store
        from src.services.regulatory_crawler import check_regulation_freshness
        from src.utils.app_settings import get_app_setting

        # Step 0: Pre-cross-examination regulation freshness check
        # Requirement: 交叉詰問前要確認當前當作基準的ISO 13485與MDSAP為最新版
        freshness = await check_regulation_freshness()
        freshness_announcement = None
        if freshness.get("announcement_needed"):
            freshness_announcement = {
                "en": freshness.get("announcement_text", ""),
                "zh": freshness.get("announcement_text_zh", ""),
            }
            logger.warning(
                "Regulation freshness check: announcement needed. %s",
                freshness.get("announcement_text", ""),
            )

        # Get LLM completion function from app state
        llm_fn = _get_llm_completion_fn(request)
        store = get_crossexam_store()

        # Get user language and model name
        lang = "zh-TW"
        model_name = "default"
        try:
            from src.utils.user_settings import load_user_settings

            settings = load_user_settings()
            lang = settings.get("language", "zh-TW")
            model_name = settings.get("model_name", "default")
        except Exception:
            pass

        # Read MDSAP toggle state from persistent app settings
        mdsap_on = get_app_setting("mdsap_verify_enabled", False)

        # Extract incomplete country data from freshness check
        _incomplete_countries: list[str] = []
        country_data = freshness.get("country_completeness", {})
        _incomplete_countries = country_data.get("incomplete_countries", [])

        # Step 1: Run daily sampling cross-exam (Phase 5 on 20% sample)
        # This produces fresh DailyCrossExamStore records with correct MDSAP flag
        sampling_record = run_daily_sampling_crossexam(
            llm_completion_fn=llm_fn,
            model=model_name,
            mdsap_enabled=mdsap_on,
            lang=lang,
        )

        # Step 2: Run daily audit (Dim A + Dim B evaluation)
        result = run_daily_audit(
            llm_completion_fn=llm_fn,
            lang=lang,
            store=store,
            incomplete_countries=_incomplete_countries,
            mdsap_enabled=mdsap_on,
        )

        # Check for deviation and send Chainlit announcement if needed
        if result.deviation_detected:
            _send_deviation_announcement(result)

        # Auto-trigger meta review if 10+ daily records since last meta review
        meta_auto = None
        try:
            meta_auto = _maybe_auto_trigger_meta_review(llm_fn, lang)
        except Exception as e:
            logger.warning(f"Auto meta-review trigger failed (non-fatal): {e}")

        response = result.to_dict()
        # Attach freshness check results and announcement
        response["regulation_freshness"] = freshness
        if freshness_announcement:
            response["freshness_announcement"] = freshness_announcement
        # Attach per-country upload reminders from storage layer
        try:
            from src.storage.mdsap_markdown_storage import get_mdsap_markdown_store

            upload_reminders = get_mdsap_markdown_store().get_upload_reminders()
            if upload_reminders:
                response["upload_reminders"] = upload_reminders
        except Exception:
            pass  # non-fatal
        if meta_auto:
            response["meta_review_auto_triggered"] = True
            response["meta_review_summary"] = meta_auto
        return JSONResponse(content=response)
    except ImportError as e:
        raise HTTPException(
            status_code=501, detail=f"Daily audit module not available: {e}"
        )
    except Exception as e:
        logger.error(f"Daily audit failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@report_router.get("/regulation-freshness")
async def check_freshness_endpoint():
    """Check regulation freshness (ISO 13485 + MDSAP) before cross-examination."""
    try:
        from src.services.regulatory_crawler import check_regulation_freshness

        result = await check_regulation_freshness()
        return JSONResponse(content=result)
    except Exception as e:
        logger.error(f"Regulation freshness check failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@report_router.get("/daily-audit/meta-review")
async def get_meta_review():
    """Get the latest 10-day meta review results."""
    try:
        from src.analysis.daily_audit import get_latest_meta_review

        result = get_latest_meta_review()
        if not result:
            return JSONResponse(
                content={"available": False, "message": "No meta review available"}
            )
        return JSONResponse(content={"available": True, **result.to_dict()})
    except ImportError:
        return JSONResponse(
            content={"available": False, "message": "Module not available"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@report_router.post("/daily-audit/meta-review")
async def run_meta_review_endpoint(request: Request):
    """Trigger a 10-day meta review."""
    try:
        from src.analysis.daily_audit import run_10day_meta_review

        llm_fn = _get_llm_completion_fn(request)

        lang = "zh-TW"
        try:
            from src.utils.user_settings import load_user_settings

            settings = load_user_settings()
            lang = settings.get("language", "zh-TW")
        except Exception:
            pass

        result = run_10day_meta_review(
            llm_completion_fn=llm_fn,
            lang=lang,
        )

        # Send announcement if deviations found
        if result.deviation_summary:
            _send_meta_review_announcement(result)

        return JSONResponse(content=result.to_dict())
    except ImportError as e:
        raise HTTPException(
            status_code=501, detail=f"Daily audit module not available: {e}"
        )
    except Exception as e:
        logger.error(f"Meta review failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@report_router.get("/daily-audit/history/{audit_id}/export/{fmt}")
async def export_audit_record(audit_id: str, fmt: str):
    """Export a single daily audit record as Word or Excel."""
    if fmt not in ("word", "excel"):
        raise HTTPException(status_code=400, detail="Format must be 'word' or 'excel'")

    try:
        from src.analysis.daily_audit import (
            get_daily_audit_history,
            export_daily_audit_word,
            export_daily_audit_excel,
        )

        # Find the record
        records = get_daily_audit_history(limit=365)
        target = None
        for r in records:
            if r.audit_id == audit_id:
                target = r
                break

        if not target:
            raise HTTPException(
                status_code=404, detail=f"Audit record {audit_id} not found"
            )

        if fmt == "word":
            filepath = export_daily_audit_word(target)
        else:
            filepath = export_daily_audit_excel(target)

        content_type = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            if fmt == "word"
            else "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        return FileResponse(filepath, media_type=content_type, filename=filepath.name)

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@report_router.get("/daily-audit/export/{fmt}")
async def export_latest_audit(fmt: str):
    """Export the latest daily audit record as Word or Excel."""
    if fmt not in ("word", "excel"):
        raise HTTPException(status_code=400, detail="Format must be 'word' or 'excel'")

    try:
        from src.analysis.daily_audit import (
            get_daily_audit_history,
            export_daily_audit_word,
            export_daily_audit_excel,
        )

        records = get_daily_audit_history(limit=1)
        if not records:
            raise HTTPException(status_code=404, detail="No audit records available")

        if fmt == "word":
            filepath = export_daily_audit_word(records[0])
        else:
            filepath = export_daily_audit_excel(records[0])

        content_type = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            if fmt == "word"
            else "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        return FileResponse(filepath, media_type=content_type, filename=filepath.name)

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@report_router.get("/daily-audit/meta-review/export/{fmt}")
async def export_meta_review_report(fmt: str):
    """Export the latest 10-day meta review as Word or Excel."""
    if fmt not in ("word", "excel"):
        raise HTTPException(status_code=400, detail="Format must be 'word' or 'excel'")

    try:
        from src.analysis.daily_audit import (
            get_latest_meta_review,
            export_meta_review_word,
            export_meta_review_excel,
        )

        result = get_latest_meta_review()
        if not result:
            raise HTTPException(status_code=404, detail="No meta review available")

        if fmt == "word":
            filepath = export_meta_review_word(result)
        else:
            filepath = export_meta_review_excel(result)

        content_type = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            if fmt == "word"
            else "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        return FileResponse(filepath, media_type=content_type, filename=filepath.name)

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================
# Deviation Announcement Helpers (Chainlit push)
# ============================================================


def _get_llm_completion_fn(request: Request):
    """Extract LLM completion function from app state."""
    # Try to get from app state (set by pipeline runner)
    app = request.app
    llm_fn = getattr(app.state, "llm_completion_fn", None)
    if llm_fn:
        return llm_fn

    # Fallback: create a provider manager and use its completion method
    try:
        from src.llm_providers import create_provider_manager

        manager = create_provider_manager()
        return manager.completion
    except Exception:
        raise HTTPException(
            status_code=503,
            detail="LLM completion function not available. Run an analysis first.",
        )


def _get_llm_completion_fn_standalone():
    """Get LLM completion function without a Request object.

    Used by background schedulers that don't have an active HTTP request.
    Creates an LLMProviderManager from saved user settings.
    Returns None if no LLM function is available.
    """
    try:
        from src.utils.user_settings import load_user_settings
        from src.llm_providers import create_provider_manager

        settings = load_user_settings()
        provider_id = settings.get("provider_id") if settings else None
        if not provider_id:
            logger.debug("No provider_id in saved settings for standalone LLM fn")
            return None

        manager = create_provider_manager(provider_id)
        return manager.completion
    except Exception as e:
        logger.debug(f"Failed to create standalone LLM fn: {e}")
        return None


def _maybe_auto_trigger_meta_review(llm_fn, lang: str = "zh-TW") -> dict | None:
    """Auto-trigger 10-day meta review if enough daily records have accumulated.

    Requirement: 10筆後的當天要顯示當天的交叉詰問與10天的交叉詰問報告
    (After 10 records, display today's cross-examination AND
     10-day cross-examination report with improvement recommendations.)

    Returns meta review summary dict, or None if not triggered.
    """
    from src.analysis.daily_audit import (
        get_daily_audit_history,
        get_latest_meta_review,
        run_10day_meta_review,
    )

    daily_records = get_daily_audit_history(limit=30)
    if len(daily_records) < 10:
        return None

    # Check if a meta review has already been run that covers recent records
    latest_meta = get_latest_meta_review()
    if latest_meta:
        # If the latest meta review's period_end covers the most recent daily record,
        # no need to re-run. Only trigger if 10+ new records exist since last meta.
        meta_end = latest_meta.period_end
        records_since_meta = [r for r in daily_records if r.audit_date > meta_end]
        if len(records_since_meta) < 10:
            return None

    # We have >= 10 records since last meta review (or no prior meta review)
    logger.info(
        "Auto-triggering 10-day meta review: %d daily records available",
        len(daily_records),
    )

    meta_result = run_10day_meta_review(
        llm_completion_fn=llm_fn,
        lang=lang,
    )

    # Send the meta review announcement via Chainlit
    try:
        _send_meta_review_announcement(meta_result)
    except Exception as e:
        logger.warning(f"Meta review announcement failed: {e}")

    return {
        "period_start": meta_result.period_start,
        "period_end": meta_result.period_end,
        "avg_dim_a": meta_result.avg_dim_a,
        "avg_dim_b": meta_result.avg_dim_b,
        "trend_analysis": meta_result.trend_analysis[:300],
        "recommendations": meta_result.recommendations[:5],
        "deviation_summary": meta_result.deviation_summary[:200],
    }


def _send_deviation_announcement(result) -> None:
    """Send deviation alert to Chainlit (same pattern as upload reminder)."""
    try:
        from src.analysis.pipeline_runner import _pipeline_send_message_fn

        send_fn = _pipeline_send_message_fn
        if send_fn is None:
            logger.info("No Chainlit send_fn available for deviation alert")
            return

        import asyncio

        msg = (
            f"\n\n⚠️ **稽核分數差異警告**\n\n"
            f"每日稽核分數出現偏差，請注意以下細節：\n"
            f"- 總分: {result.overall_score}/100\n"
            f"- 法規準確度 (Dim A): {result.dim_a_score}/100\n"
            f"- 交叉詰問品質 (Dim B): {result.dim_b_score}/100\n"
            f"- 偏差說明: {result.deviation_details}\n\n"
            f"系統將根據差異分析結果調整交叉詰問參數。"
            f"如您有任何意見，請在此提出。\n"
        )

        loop = asyncio.get_event_loop()
        if loop.is_running():
            asyncio.ensure_future(send_fn(msg))
        else:
            loop.run_until_complete(send_fn(msg))
    except Exception as e:
        logger.warning(f"Failed to send deviation announcement: {e}")


def _send_meta_review_announcement(result) -> None:
    """Send meta review deviation summary to Chainlit."""
    try:
        from src.analysis.pipeline_runner import _pipeline_send_message_fn

        send_fn = _pipeline_send_message_fn
        if send_fn is None:
            return

        import asyncio

        msg = (
            f"\n\n🧠 **10日總檢報告**\n\n"
            f"過去 10 天的稽核結果已完成總檢：\n"
            f"- 平均 Dim A (法規準確度): {result.avg_dim_a:.0f}/100\n"
            f"- 平均 Dim B (交叉詰問品質): {result.avg_dim_b:.0f}/100\n"
            f"\n{result.deviation_summary}\n\n"
            f"如您有任何意見或建議，請在此提出。\n"
        )

        loop = asyncio.get_event_loop()
        if loop.is_running():
            asyncio.ensure_future(send_fn(msg))
        else:
            loop.run_until_complete(send_fn(msg))
    except Exception as e:
        logger.warning(f"Failed to send meta review announcement: {e}")


# ============================================================
# User Feedback CRUD + Re-evaluation Endpoints
# ============================================================


@report_router.post("/daily-audit/feedback")
async def submit_feedback(request: Request):
    """Submit user feedback for a daily audit or meta review, then re-evaluate."""
    with _phoenix_report_span("report_submit_feedback"):
        try:
            body = await request.json()
            audit_type = body.get("audit_type", "daily")  # 'daily' | 'meta'
            target_id = body.get("target_id", "")
            feedback_text = body.get("feedback_text", "").strip()

            if not feedback_text:
                raise HTTPException(status_code=400, detail="feedback_text is required")
            if audit_type not in ("daily", "meta"):
                raise HTTPException(
                    status_code=400, detail="audit_type must be 'daily' or 'meta'"
                )

            from src.analysis.daily_audit import save_feedback

            fb = save_feedback(
                audit_type=audit_type, target_id=target_id, feedback_text=feedback_text
            )

            # Trigger re-evaluation with feedback context
            re_eval_result = None
            try:
                re_eval_result = await _run_reeval_with_feedback(
                    request, audit_type, fb
                )
            except Exception as e:
                logger.warning(f"Re-evaluation failed after feedback: {e}")

            return JSONResponse(
                content={
                    "success": True,
                    "feedback": fb.to_dict(),
                    "re_evaluation": re_eval_result,
                }
            )
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Submit feedback failed: {e}")
            raise HTTPException(status_code=500, detail=str(e))


@report_router.get("/daily-audit/feedback")
async def list_feedback():
    """List all active feedback records."""
    try:
        from src.analysis.daily_audit import get_all_feedback

        records = get_all_feedback()
        return JSONResponse(
            content={
                "records": [fb.to_dict() for fb in records],
                "total": len(records),
            }
        )
    except ImportError:
        return JSONResponse(content={"records": [], "total": 0})
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@report_router.put("/daily-audit/feedback/{feedback_id}")
async def update_feedback_endpoint(feedback_id: str, request: Request):
    """Update an existing feedback record's text."""
    try:
        body = await request.json()
        new_text = body.get("feedback_text", "").strip()
        if not new_text:
            raise HTTPException(status_code=400, detail="feedback_text is required")

        from src.analysis.daily_audit import update_feedback

        fb = update_feedback(feedback_id, new_text)
        if fb is None:
            raise HTTPException(
                status_code=404, detail=f"Feedback {feedback_id} not found"
            )

        return JSONResponse(content={"success": True, "feedback": fb.to_dict()})
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@report_router.delete("/daily-audit/feedback/{feedback_id}")
async def delete_feedback_endpoint(feedback_id: str):
    """Delete (soft) a feedback record."""
    try:
        from src.analysis.daily_audit import delete_feedback

        ok = delete_feedback(feedback_id)
        if not ok:
            raise HTTPException(
                status_code=404, detail=f"Feedback {feedback_id} not found"
            )
        return JSONResponse(content={"success": True, "deleted": feedback_id})
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@report_router.get("/daily-audit/feedback/export/{fmt}")
async def export_feedback_records(fmt: str):
    """Export all active feedback records as Word or Excel."""
    if fmt not in ("word", "excel"):
        raise HTTPException(status_code=400, detail="Format must be 'word' or 'excel'")

    try:
        from datetime import datetime
        from src.analysis.daily_audit import get_all_feedback

        records = get_all_feedback()
        if not records:
            raise HTTPException(status_code=404, detail="No feedback records available")

        EXPORT_DIR = Path("data/exports")
        EXPORT_DIR.mkdir(parents=True, exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")

        if fmt == "word":
            from docx import Document

            doc = Document()
            doc.add_heading("使用者意見紀錄 (User Feedback Records)", level=1)
            doc.add_paragraph(f"匯出時間: {datetime.now().isoformat()[:19]}")
            doc.add_paragraph(f"總計: {len(records)} 筆紀錄")
            doc.add_paragraph("")
            for fb in records:
                doc.add_heading(
                    f"[{fb.audit_type.upper()}] {fb.feedback_id}",
                    level=2,
                )
                doc.add_paragraph(f"建立時間: {fb.created_at}")
                doc.add_paragraph(f"更新時間: {fb.updated_at}")
                doc.add_paragraph(f"狀態: {fb.status}")
                doc.add_paragraph(f"目標: {fb.target_id or 'N/A'}")
                doc.add_paragraph(f"意見內容: {fb.feedback_text}")
                if fb.re_evaluation_score is not None:
                    doc.add_paragraph(f"重新評估分數: {fb.re_evaluation_score}")
                if fb.re_evaluation_id:
                    doc.add_paragraph(f"重新評估 ID: {fb.re_evaluation_id}")
                doc.add_paragraph("")

            filepath = EXPORT_DIR / f"feedback_records_{ts}.docx"
            doc.save(str(filepath))
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        else:  # excel
            from openpyxl import Workbook

            wb = Workbook()
            ws = wb.active
            ws.title = "Feedback Records"
            headers = [
                "ID",
                "類型",
                "目標",
                "意見內容",
                "建立時間",
                "更新時間",
                "狀態",
                "重新評估分數",
                "重新評估 ID",
            ]
            ws.append(headers)
            for fb in records:
                ws.append(
                    [
                        fb.feedback_id,
                        fb.audit_type,
                        fb.target_id or "",
                        fb.feedback_text,
                        fb.created_at,
                        fb.updated_at,
                        fb.status,
                        fb.re_evaluation_score
                        if fb.re_evaluation_score is not None
                        else "",
                        fb.re_evaluation_id or "",
                    ]
                )
            # Auto-width
            for col in ws.columns:
                max_len = 0
                col_letter = col[0].column_letter
                for cell in col:
                    if cell.value:
                        max_len = max(max_len, len(str(cell.value)))
                ws.column_dimensions[col_letter].width = min(max_len + 2, 60)

            filepath = EXPORT_DIR / f"feedback_records_{ts}.xlsx"
            wb.save(str(filepath))
            content_type = (
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

        return FileResponse(filepath, media_type=content_type, filename=filepath.name)

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Export feedback failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


async def _run_reeval_with_feedback(
    request: Request, audit_type: str, fb
) -> dict | None:
    """Re-run daily audit with user feedback context appended to prompts."""
    try:
        llm_fn = _get_llm_completion_fn(request)
        lang = "zh-TW"
        try:
            from src.utils.user_settings import load_user_settings

            settings = load_user_settings()
            lang = settings.get("language", "zh-TW")
        except Exception:
            pass

        if audit_type == "daily":
            from src.analysis.daily_audit import (
                run_daily_audit,
                get_active_feedback_context,
            )

            feedback_ctx = get_active_feedback_context()
            result = run_daily_audit(
                llm_completion_fn=llm_fn,
                lang=lang,
                feedback_context=feedback_ctx,
            )
            # Update the feedback record with re-evaluation result
            fb.re_evaluation_id = result.audit_id
            fb.re_evaluation_score = result.overall_score
            from src.analysis.daily_audit import FEEDBACK_DIR
            from src.utils.safe_io import atomic_write_json

            atomic_write_json(FEEDBACK_DIR / f"{fb.feedback_id}.json", fb.to_dict())

            return {
                "audit_id": result.audit_id,
                "overall_score": result.overall_score,
                "dim_a_score": result.dim_a_score,
                "dim_b_score": result.dim_b_score,
            }
        elif audit_type == "meta":
            from src.analysis.daily_audit import (
                run_10day_meta_review,
                get_active_feedback_context,
            )

            feedback_ctx = get_active_feedback_context()
            result = run_10day_meta_review(
                llm_completion_fn=llm_fn,
                lang=lang,
                feedback_context=feedback_ctx,
            )
            # Update the feedback record with re-evaluation result
            fb.re_evaluation_id = f"meta_{result.period_end}"
            fb.re_evaluation_score = int((result.avg_dim_a + result.avg_dim_b) / 2)
            from src.analysis.daily_audit import FEEDBACK_DIR
            from src.utils.safe_io import atomic_write_json

            atomic_write_json(FEEDBACK_DIR / f"{fb.feedback_id}.json", fb.to_dict())

            return {
                "avg_dim_a": result.avg_dim_a,
                "avg_dim_b": result.avg_dim_b,
            }
    except Exception as e:
        logger.warning(f"Re-evaluation failed: {e}")
        return None
