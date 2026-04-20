"""
AI-QMS Phase 1 - 文件更動紀錄匯出模組
Export audit records to Word (.docx) and Excel (.xlsx) formats.
"""



from datetime import datetime
from pathlib import Path


from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side


# Action name mapping per language
_ACTION_LABELS_ZH = {
    "document_created": "文件建立",
    "document_version_updated": "文件進版",
    "bulk_delete": "批次刪除",
    "FILE_UPLOADED": "文件上傳",
    "VERSION_CONFIRMED": "版本確認",
}
_ACTION_LABELS_EN = {
    "document_created": "Document Created",
    "document_version_updated": "Version Updated",
    "bulk_delete": "Bulk Delete",
    "FILE_UPLOADED": "File Uploaded",
    "VERSION_CONFIRMED": "Version Confirmed",
}
# Keep legacy name for backward compat
ACTION_LABELS = _ACTION_LABELS_ZH

# UI strings per language
_UI = {
    "zh-TW": {
        "title": "AI-QMS 文件更動紀錄報告",
        "export_time": "匯出時間",
        "total_records": "紀錄總數",
        "no_records": "目前沒有任何文件更動紀錄。",
        "headers": ["#", "時間", "操作", "文件編號", "操作者", "詳情", "下載次數", "下載者（最近5筆）"],
        "sheet_name": "文件更動紀錄",
        "integrity": "紀錄鏈完整性",
        "hash_protected": "SHA-256 雜湊鏈保護",
    },
    "en-US": {
        "title": "AI-QMS Document Change Records Report",
        "export_time": "Export Time",
        "total_records": "Total Records",
        "no_records": "No document change records found.",
        "headers": ["#", "Time", "Action", "Document ID", "Operator", "Details", "Downloads", "Recent Downloaders"],
        "sheet_name": "Change Records",
        "integrity": "Record Chain Integrity",
        "hash_protected": "SHA-256 Hash Chain Protected",
    },
    "ja-JP": {
        "title": "AI-QMS 文書変更記録レポート",
        "export_time": "エクスポート時刻",
        "total_records": "レコード総数",
        "no_records": "文書変更記録がありません。",
        "headers": ["#", "時刻", "操作", "文書番号", "担当者", "詳細", "ダウンロード数", "最近のDL者（最大5件）"],
        "sheet_name": "変更記録",
        "integrity": "記録チェーン整合性",
        "hash_protected": "SHA-256ハッシュチェーン保護",
    },
}


def _ui(lang: str = "zh-TW") -> dict:
    return _UI.get(lang, _UI["zh-TW"])


def _action_label_lang(action: str, lang: str = "zh-TW") -> str:
    labels = _ACTION_LABELS_EN if lang == "en-US" else _ACTION_LABELS_ZH
    return labels.get(action, action)

# Output directory for generated files
EXPORT_DIR = (Path(__file__).resolve().parent.parent.parent / "data" / "exports")
EXPORT_DIR.mkdir(parents=True, exist_ok=True)


def _action_label(action: str, lang: str = "zh-TW") -> str:
    """Convert action code to display label."""
    return _action_label_lang(action, lang)


def _format_details(details: dict) -> str:
    """Format details dict into a readable string."""
    if not details:
        return ""
    parts = []
    for k, v in details.items():
        if v is None or v == "":
            continue
        parts.append(f"{k}: {v}")
    return "; ".join(parts)


def format_audit_table_markdown(records: list, lang: str = "zh-TW") -> str:
    """Format audit records as a Markdown table for chat display."""
    ui = _ui(lang)
    if not records:
        return f"📋 {ui['no_records']}"

    _zh = lang.startswith("zh")
    title = "文件更動紀錄" if _zh else "Document Change Records"
    count_label = "筆" if _zh else "records"

    lines = [
        f"📋 **{title}** ({len(records)} {count_label})\n",
    ]
    if _zh:
        lines.append("| # | 時間 | 操作 | 文件編號 | 操作者 | 詳情 |")
    else:
        lines.append("| # | Time | Action | Document ID | Operator | Details |")
    lines.append("|---|------|------|---------|--------|------|")

    for i, r in enumerate(records, 1):
        ts = r.get("timestamp", "")
        # Format timestamp to shorter form
        try:
            dt = datetime.fromisoformat(ts)
            ts_short = dt.strftime("%Y-%m-%d %H:%M")
        except (ValueError, TypeError):
            ts_short = ts[:16] if ts else ""

        action = _action_label(r.get("action", ""), lang)
        doc_id = r.get("document_id", "")
        user = r.get("user_id", "")
        details = _format_details(r.get("details", {}))
        # Truncate details for table display
        if len(details) > 50:
            details = details[:47] + "..."

        lines.append(f"| {i} | {ts_short} | {action} | {doc_id} | {user} | {details} |")

    lines.append(f"\n✅ {ui['integrity']}: {ui['hash_protected']}")
    return "\n".join(lines)


def export_to_word(records: list, download_stats: dict | None = None, lang: str = "zh-TW") -> str:
    """
    Export audit records to a Word (.docx) file.

    Returns:
        Path to the generated .docx file.
    """
    ui = _ui(lang)
    doc = Document()

    # Title
    title = doc.add_heading(ui["title"], level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = meta.add_run(f"{ui['export_time']}: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run2 = meta2.add_run(f"{ui['total_records']}: {len(records)}")
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()  # spacer

    dl_stats = download_stats or {}

    if not records:
        doc.add_paragraph(ui["no_records"])
    else:
        # Create table
        table = doc.add_table(rows=1, cols=8)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        headers = ui["headers"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)

        # Data rows
        for idx, r in enumerate(records, 1):
            row = table.add_row()
            ts = r.get("timestamp", "")
            try:
                dt = datetime.fromisoformat(ts)
                ts_short = dt.strftime("%Y-%m-%d %H:%M:%S")
            except (ValueError, TypeError):
                ts_short = ts

            doc_id = r.get("document_id", "")
            stat = dl_stats.get(doc_id, {})
            values = [
                str(idx),
                ts_short,
                _action_label(r.get("action", ""), lang),
                doc_id,
                r.get("user_id", ""),
                _format_details(r.get("details", {})),
                str(stat.get("count", 0)),
                stat.get("recent", ""),
            ]
            for i, val in enumerate(values):
                cell = row.cells[i]
                cell.text = val
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

        # Set column widths
        widths = [Cm(0.8), Cm(3), Cm(2.2), Cm(2.5), Cm(2), Cm(4), Cm(1.5), Cm(3)]
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = width

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    run = footer.add_run(
        "本報告由 AI-QMS 品質管理系統自動產生。文件更動紀錄受 SHA-256 雜湊鏈保護，確保資料完整性與不可竄改。"
    )
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Save
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"doc_change_records_{timestamp}.docx"
    filepath = EXPORT_DIR / filename
    doc.save(str(filepath))
    return str(filepath)


def export_to_excel(records: list, download_stats: dict | None = None, lang: str = "zh-TW") -> str:
    """
    Export audit records to an Excel (.xlsx) file.

    Returns:
        Path to the generated .xlsx file.
    """
    ui = _ui(lang)
    wb = Workbook()
    ws = wb.active
    ws.title = ui["sheet_name"]

    # Styles
    header_font = Font(name="Microsoft JhengHei", bold=True, size=10, color="FFFFFF")
    header_fill = PatternFill(
        start_color="2563EB", end_color="2563EB", fill_type="solid"
    )
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    cell_font = Font(name="Microsoft JhengHei", size=9)
    cell_alignment = Alignment(vertical="center", wrap_text=True)
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )

    dl_stats = download_stats or {}

    # Title row
    ws.merge_cells("A1:I1")
    title_cell = ws["A1"]
    title_cell.value = ui["title"]
    title_cell.font = Font(name="Microsoft JhengHei", bold=True, size=14)
    title_cell.alignment = Alignment(horizontal="center")

    # Metadata row
    ws.merge_cells("A2:I2")
    meta_cell = ws["A2"]
    meta_cell.value = (
        f"{ui['export_time']}: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | "
        f"{ui['total_records']}: {len(records)}"
    )
    meta_cell.font = Font(name="Microsoft JhengHei", size=9, color="808080")
    meta_cell.alignment = Alignment(horizontal="right")

    # Headers (row 4) — 9 cols (prepend record-id to the 8-col list)
    base_headers = ui["headers"]
    record_id_label = "Record ID" if lang == "en-US" else ("レコードID" if lang == "ja-JP" else "紀錄 ID")
    headers = [base_headers[0], record_id_label] + base_headers[1:]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=4, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
        cell.border = thin_border

    # Data rows
    for idx, r in enumerate(records, 1):
        row_num = idx + 4
        ts = r.get("timestamp", "")
        try:
            dt = datetime.fromisoformat(ts)
            ts_str = dt.strftime("%Y-%m-%d %H:%M:%S")
        except (ValueError, TypeError):
            ts_str = ts

        doc_id = r.get("document_id", "")
        stat = dl_stats.get(doc_id, {})
        values = [
            idx,
            r.get("record_id", ""),
            ts_str,
            _action_label(r.get("action", ""), lang),
            doc_id,
            r.get("user_id", ""),
            _format_details(r.get("details", {})),
            stat.get("count", 0),
            stat.get("recent", ""),
        ]
        for col, val in enumerate(values, 1):
            cell = ws.cell(row=row_num, column=col, value=val)
            cell.font = cell_font
            cell.alignment = cell_alignment
            cell.border = thin_border

    # Column widths
    col_widths = [5, 25, 20, 12, 15, 12, 35, 10, 30]
    for i, width in enumerate(col_widths, 1):
        ws.column_dimensions[chr(64 + i)].width = width

    # Freeze header row
    ws.freeze_panes = "A5"

    # Hash chain integrity note
    note_row = len(records) + 6
    ws.merge_cells(f"A{note_row}:I{note_row}")
    note_cell = ws.cell(row=note_row, column=1)
    note_cell.value = (
        "本報告由 AI-QMS 品質管理系統自動產生。"
        "文件更動紀錄受 SHA-256 雜湊鏈保護，確保資料完整性與不可竄改。"
    )
    note_cell.font = Font(
        name="Microsoft JhengHei", size=8, italic=True, color="808080"
    )

    # Save
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"doc_change_records_{timestamp}.xlsx"
    filepath = EXPORT_DIR / filename
    wb.save(str(filepath))
    return str(filepath)
