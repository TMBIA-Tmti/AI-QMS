"""
AI-QMS Phase 1 - 法規清單與引用清單匯出模組
Export regulatory standards list and document reference list to Word/Excel formats.
"""

import json
import os

from datetime import datetime
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side


# ── i18n helpers ──


import re as _re


def _region_display(region_key: str, lang: str) -> str:
    """Extract English name for non-zh languages: '美國 (USA)' → 'USA'"""
    if lang.startswith("zh"):
        return region_key
    m = _re.search(r'\(([^)]+)\)', region_key)
    return m.group(1) if m else region_key


def _t(key: str, lang: str = "zh-TW", **kwargs) -> str:
    """Translate a key using locale JSON files."""
    _cache = getattr(_t, "_cache", {})
    if lang not in _cache:
        locale_path = os.path.join(
            os.path.dirname(__file__), "..", "chainlit_app", "locales", f"{lang}.json"
        )
        try:
            with open(locale_path, "r", encoding="utf-8") as f:
                _cache[lang] = json.load(f)
        except Exception:
            _cache[lang] = {}
        _t._cache = _cache
    text = _cache.get(lang, {}).get(key, key)
    if kwargs:
        try:
            text = text.format(**kwargs)
        except (KeyError, IndexError):
            pass
    return text


def _tl(key: str, lang: str = "zh-TW") -> list:
    """Translate a key that returns a list (e.g. table headers)."""
    _cache = getattr(_t, "_cache", {})
    if lang not in _cache:
        _t(key, lang)  # populate cache
        _cache = getattr(_t, "_cache", {})
    val = _cache.get(lang, {}).get(key)
    return val if isinstance(val, list) else [key]


# Output directory for generated files
EXPORT_DIR = (Path(__file__).resolve().parent.parent.parent / "data" / "exports")
EXPORT_DIR.mkdir(parents=True, exist_ok=True)

# Shared styles
HEADER_FILL = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
HEADER_FONT = Font(name="Microsoft JhengHei", size=10, bold=True, color="FFFFFF")
CELL_FONT = Font(name="Microsoft JhengHei", size=9)
THIN_BORDER = Border(
    left=Side(style="thin"),
    right=Side(style="thin"),
    top=Side(style="thin"),
    bottom=Side(style="thin"),
)

# Status colors for Excel
SUCCESS_FILL = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
FAIL_FILL = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")


# ============================================================
# 法規清單 (Regulatory Standards List)
# ============================================================


def format_regulatory_table_markdown(
    scan_result: dict,
    assessment: Optional[str] = None,
    lang: str = "zh-TW",
) -> str:
    """Format regulatory scan result as Markdown for chat display."""
    by_doc = scan_result.get("by_document", [])
    aggregate = scan_result.get("aggregate", [])

    if not aggregate:
        return _t("regulatory_export.no_refs", lang)

    lines = [
        _t(
            "regulatory_export.aggregate_title",
            lang,
            std_count=len(aggregate),
            doc_count=len(by_doc),
        )
        + "\n",
        f"### {_t('regulatory_export.std_summary', lang)}\n",
        f"| {_tl('regulatory_export.std_headers', lang)[0]} | {_tl('regulatory_export.std_headers', lang)[1]} |",
        "|------|-----------|",
    ]

    for entry in aggregate:
        std = entry["standard"]
        refs = entry["referenced_by"]
        lines.append(f"| {std} | {len(refs)} |")

    # Assessment section
    if assessment:
        lines.append("\n---\n")
        lines.append(f"### {_t('regulatory_export.assessment_report', lang)}\n")
        lines.append(assessment)

    return "\n".join(lines)


def _render_assessment_to_word(doc, assessment: str, lang: str = "zh-TW"):
    """Render assessment markdown text into Word document paragraphs."""
    doc.add_heading(_t("regulatory_export.assessment_heading", lang), level=2)
    for para_text in assessment.split("\n"):
        stripped = para_text.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        if stripped.startswith("###"):
            doc.add_heading(stripped.lstrip("#").strip(), level=4)
        elif stripped.startswith("##"):
            doc.add_heading(stripped.lstrip("#").strip(), level=3)
        elif stripped.startswith("#"):
            doc.add_heading(stripped.lstrip("#").strip(), level=2)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(stripped[2:], style="List Bullet")
            for run in p.runs:
                run.font.size = Pt(9)
        elif stripped.startswith(tuple(f"{i}." for i in range(1, 20))):
            p = doc.add_paragraph(stripped, style="List Number")
            for run in p.runs:
                run.font.size = Pt(9)
        else:
            p = doc.add_paragraph(stripped)
            for run in p.runs:
                run.font.size = Pt(9)


def _render_assessment_to_excel(wb, assessment: str, lang: str = "zh-TW"):
    """Render assessment text into a new Excel sheet."""
    ws = wb.create_sheet(_t("regulatory_export.assessment_sheet", lang))

    ws.merge_cells("A1:B1")
    title_cell = ws.cell(row=1, column=1)
    title_cell.value = _t("regulatory_export.assessment_report", lang)
    title_cell.font = Font(name="Microsoft JhengHei", size=14, bold=True)
    title_cell.alignment = Alignment(horizontal="center")

    assessment_lines = assessment.split("\n")
    for row_idx, line in enumerate(assessment_lines, 3):
        cell = ws.cell(row=row_idx, column=1)
        cell.value = line
        if line.strip().startswith("#"):
            cell.font = Font(name="Microsoft JhengHei", size=11, bold=True)
        else:
            cell.font = CELL_FONT

    ws.column_dimensions["A"].width = 100


def _render_verification_to_word(doc, verification_report: dict, lang: str = "zh-TW"):
    """Render verification report section into Word document."""
    doc.add_paragraph()
    doc.add_heading(_t("regulatory_export.verification_heading", lang), level=2)
    ver_passed = verification_report.get("passed_count", 0)
    ver_warn = verification_report.get("warning_count", 0)
    ver_fail = verification_report.get("failed_count", 0)
    ver_total = verification_report.get("total_documents", 0)
    p = doc.add_paragraph(
        _t(
            "regulatory_export.verification_meta",
            lang,
            time=verification_report.get("verified_at", ""),
            passed=ver_passed,
            warning=ver_warn,
            failed=ver_fail,
            total=ver_total,
        )
    )
    for r in p.runs:
        r.font.size = Pt(9)

    # Cross checks
    cross_checks = verification_report.get("cross_checks", [])
    if cross_checks:
        doc.add_heading(_t("regulatory_export.cross_check_heading", lang), level=3)
        for cc in cross_checks:
            icon = "✓" if cc.get("passed") else "✗"
            p = doc.add_paragraph(
                f"{icon} {cc.get('check_name', '')}: {cc.get('message', '')}",
                style="List Bullet",
            )
            for r in p.runs:
                r.font.size = Pt(8)

    # Per-document verification table
    ver_docs = verification_report.get("documents", [])
    if ver_docs:
        doc.add_heading(_t("regulatory_export.per_doc_verification", lang), level=3)
        vtable = doc.add_table(rows=1, cols=5)
        vtable.style = "Table Grid"
        vtable.alignment = WD_TABLE_ALIGNMENT.CENTER
        vheaders = _tl("regulatory_export.verification_headers", lang)
        for i, h in enumerate(vheaders):
            cell = vtable.rows[0].cells[i]
            cell.text = h
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in paragraph.runs:
                    r.font.bold = True
                    r.font.size = Pt(9)
        for vd in ver_docs:
            vrow = vtable.add_row()
            checks = vd.get("checks", [])
            passed_n = sum(1 for c in checks if c.get("passed"))
            issues = "; ".join(
                c.get("message", "") for c in checks if not c.get("passed")
            )
            status_map = {
                "pass": _t("regulatory_export.status_pass", lang),
                "warning": _t("regulatory_export.status_warning", lang),
                "fail": _t("regulatory_export.status_fail", lang),
            }
            vals = [
                _region_display(vd.get("region", ""), lang),
                vd.get("agency", ""),
                status_map.get(vd.get("overall_status"), vd.get("overall_status", "")),
                f"{passed_n}/{len(checks)}",
                issues or _t("regulatory_export.no_issues", lang),
            ]
            for i, val in enumerate(vals):
                cell = vrow.cells[i]
                cell.text = val
                for paragraph in cell.paragraphs:
                    for r in paragraph.runs:
                        r.font.size = Pt(8)
        vwidths = [Cm(2), Cm(2), Cm(1.5), Cm(2), Cm(8)]
        for row in vtable.rows:
            for i, w in enumerate(vwidths):
                row.cells[i].width = w
        doc.add_paragraph()


def _render_verification_to_excel(wb, verification_report: dict, lang: str = "zh-TW"):
    """Render verification report into a new Excel sheet."""
    ws_ver = wb.create_sheet(_t("regulatory_export.verification_sheet", lang))

    ws_ver.merge_cells("A1:E1")
    vtitle = ws_ver.cell(row=1, column=1)
    vtitle.value = _t("regulatory_export.verification_heading", lang)
    vtitle.font = Font(name="Microsoft JhengHei", size=14, bold=True)
    vtitle.alignment = Alignment(horizontal="center")

    ver_passed = verification_report.get("passed_count", 0)
    ver_warn = verification_report.get("warning_count", 0)
    ver_fail = verification_report.get("failed_count", 0)
    ver_total = verification_report.get("total_documents", 0)
    ws_ver.merge_cells("A2:E2")
    vmeta = ws_ver.cell(row=2, column=1)
    vmeta.value = _t(
        "regulatory_export.verification_meta",
        lang,
        time=verification_report.get("verified_at", ""),
        passed=ver_passed,
        warning=ver_warn,
        failed=ver_fail,
        total=ver_total,
    )
    vmeta.font = Font(name="Microsoft JhengHei", size=9, italic=True, color="808080")

    vheaders = _tl("regulatory_export.verification_headers", lang)
    for col, h in enumerate(vheaders, 1):
        cell = ws_ver.cell(row=4, column=col)
        cell.value = h
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = THIN_BORDER

    status_map = {
        "pass": _t("regulatory_export.status_pass", lang),
        "warning": _t("regulatory_export.status_warning", lang),
        "fail": _t("regulatory_export.status_fail", lang),
    }
    ver_docs = verification_report.get("documents", [])
    for row_idx, vd in enumerate(ver_docs, 5):
        checks = vd.get("checks", [])
        passed_n = sum(1 for c in checks if c.get("passed"))
        issues = "; ".join(c.get("message", "") for c in checks if not c.get("passed"))
        vals = [
            _region_display(vd.get("region", ""), lang),
            vd.get("agency", ""),
            status_map.get(vd.get("overall_status"), vd.get("overall_status", "")),
            f"{passed_n}/{len(checks)}",
            issues or _t("regulatory_export.no_issues", lang),
        ]
        for col, val in enumerate(vals, 1):
            cell = ws_ver.cell(row=row_idx, column=col)
            cell.value = val
            cell.font = CELL_FONT
            cell.border = THIN_BORDER
        st_cell = ws_ver.cell(row=row_idx, column=3)
        if vd.get("overall_status") == "pass":
            st_cell.fill = SUCCESS_FILL
        elif vd.get("overall_status") == "fail":
            st_cell.fill = FAIL_FILL

    ws_ver.column_dimensions["A"].width = 15
    ws_ver.column_dimensions["B"].width = 15
    ws_ver.column_dimensions["C"].width = 10
    ws_ver.column_dimensions["D"].width = 12
    ws_ver.column_dimensions["E"].width = 60
    ws_ver.freeze_panes = "A5"


def _source_label(source_command: str, lang: str = "zh-TW") -> str:
    labels = {
        "regulatory_list": _t("source_label.regulatory_list", lang),
        "regulatory_update": _t("source_label.regulatory_update", lang),
    }
    return labels.get(source_command, source_command)


def export_regulatory_to_word(
    scan_result: dict,
    assessment: Optional[str] = None,
    verification_report: Optional[dict] = None,
    lang: str = "zh-TW",
    source_command: str = "regulatory_list",
) -> str:
    """
    Export regulatory standards list to Word (.docx).

    Returns:
        Path to the generated .docx file.
    """
    by_doc = scan_result.get("by_document", [])
    aggregate = scan_result.get("aggregate", [])

    src_label = _source_label(source_command, lang)
    doc = Document()

    # Title
    title = doc.add_heading(
        f"{_t('regulatory_export.title_regulatory', lang)}（{src_label}）", level=1
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = meta.add_run(
        f"{_t('regulatory_export.export_time', lang)}: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | "
        f"{_t('source_label.source', lang)}: {src_label}"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run2 = meta2.add_run(
        f"{_t('regulatory_export.std_count', lang, count=len(aggregate))} | "
        f"{_t('regulatory_export.std_coverage', lang, count=len(by_doc))}"
    )
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    # ── Abbreviation Legend (language-aware) ──
    if lang.startswith("zh"):
        _abbrev_heading = "縮寫說明"
        _abbrev_content = (
            "ISO 13485  — 醫療器材品質管理系統國際標準\n"
            "EU MDR     — 歐盟醫療器材法規 2017/745\n"
            "EU IVDR    — 歐盟體外診斷法規 2017/746\n"
            "TFDA       — 臺灣食品藥物管理署\n"
            "21 CFR 820 — 美國 FDA 品質系統法規（QSR/QMSR）\n"
            "MDSAP      — 醫療器材單一稽核計畫（美國/加拿大/巴西/澳洲/日本）\n"
            "QMS        — 品質管理系統\n"
            "RA         — 法規事務；報告中標記為需 RA 審查的項目\n"
            "GMP        — 良好製造規範\n"
            "IVD        — 體外診斷醫療器材"
        )
        _guide_heading = "報告欄位說明"
        _guide_content = (
            "【法規彙總清單】\n"
            "  標準名稱：文件中引用的法規或標準名稱\n"
            "  引用文件數：引用該法規的文件總數\n"
            "  引用文件清單：列出所有引用該法規的文件 ID\n\n"
            "【逐文件詳情】\n"
            "  文件 ID / 標題 / 文件類型：來自 QMS 文件庫的基本資訊\n"
            "  版本：文件目前版次\n"
            "  引用法規：該文件所引用的全部法規/標準清單\n\n"
            "【評估報告】：AI 對法規引用完整性與合規性的整體評估意見\n"
            "【驗證報告】：對評估結論的二次核查結果"
        )
    elif lang.startswith("ja"):
        _abbrev_heading = "略語一覧"
        _abbrev_content = (
            "ISO 13485  — 医療機器の品質マネジメントシステム国際規格\n"
            "EU MDR     — EU 医療機器規則 2017/745\n"
            "EU IVDR    — EU 体外診断医療機器規則 2017/746\n"
            "TFDA       — 台湾食品薬物管理署\n"
            "21 CFR 820 — 米国 FDA 品質システム規制（QSR/QMSR）\n"
            "MDSAP      — 医療機器単一審査プログラム（米国/カナダ/ブラジル/オーストラリア/日本）\n"
            "QMS        — 品質マネジメントシステム\n"
            "RA         — 薬事；レポートで RA 審査が必要とマークされた項目\n"
            "GMP        — 適正製造基準\n"
            "IVD        — 体外診断医療機器"
        )
        _guide_heading = "レポートフィールドガイド"
        _guide_content = (
            "【法規集計一覧】\n"
            "  標準名称：文書内で引用された法規または標準の名称\n"
            "  引用文書数：当該法規を引用する文書の総数\n"
            "  引用文書一覧：当該法規を引用するすべての文書 ID\n\n"
            "【文書別詳細】\n"
            "  文書 ID / タイトル / 文書種別：QMS 文書ライブラリの基本情報\n"
            "  バージョン：文書の現行版\n"
            "  引用法規：当該文書が引用するすべての法規/標準の一覧\n\n"
            "【評価レポート】：法規引用の完全性とコンプライアンスに関する AI の総合評価\n"
            "【検証レポート】：評価結論の二次確認結果"
        )
    elif lang.startswith("ko"):
        _abbrev_heading = "약어 목록"
        _abbrev_content = (
            "ISO 13485  — 의료기기 품질경영시스템 국제표준\n"
            "EU MDR     — EU 의료기기 규정 2017/745\n"
            "EU IVDR    — EU 체외진단기기 규정 2017/746\n"
            "TFDA       — 대만 식품의약품안전처\n"
            "21 CFR 820 — 미국 FDA 품질시스템 규정(QSR/QMSR)\n"
            "MDSAP      — 의료기기 단일심사 프로그램(미국/캐나다/브라질/호주/일본)\n"
            "QMS        — 품질경영시스템\n"
            "RA         — 규제 업무; 보고서에서 RA 검토가 필요한 항목\n"
            "GMP        — 우수제조기준\n"
            "IVD        — 체외진단 의료기기"
        )
        _guide_heading = "보고서 필드 안내"
        _guide_content = (
            "[규제 종합 목록]\n"
            "  표준명: 문서에서 인용된 법규 또는 표준 이름\n"
            "  인용 문서 수: 해당 법규를 인용한 문서 총수\n"
            "  인용 문서 목록: 해당 법규를 인용한 모든 문서 ID\n\n"
            "[문서별 상세]\n"
            "  문서 ID / 제목 / 문서 유형: QMS 문서 라이브러리 기본 정보\n"
            "  버전: 현재 문서 버전\n"
            "  인용 규제: 해당 문서에서 인용한 모든 법규/표준 목록\n\n"
            "[평가 보고서]: 법규 인용 완전성 및 컴플라이언스에 대한 AI 종합 평가\n"
            "[검증 보고서]: 평가 결론의 이중 검증 결과"
        )
    elif lang.startswith("de"):
        _abbrev_heading = "Abkürzungsverzeichnis"
        _abbrev_content = (
            "ISO 13485  — Internationale Norm für Qualitätsmanagementsysteme für Medizinprodukte\n"
            "EU MDR     — EU-Medizinprodukteverordnung 2017/745\n"
            "EU IVDR    — EU-In-vitro-Diagnostika-Verordnung 2017/746\n"
            "TFDA       — Taiwanesische Lebens- und Arzneimittelbehörde\n"
            "21 CFR 820 — US-amerikanische FDA-Qualitätssystemverordnung (QSR/QMSR)\n"
            "MDSAP      — Single-Audit-Programm für Medizinprodukte (USA/Kanada/Brasilien/Australien/Japan)\n"
            "QMS        — Qualitätsmanagementsystem\n"
            "RA         — Regulatorische Angelegenheiten; im Bericht als RA-prüfungspflichtig markiert\n"
            "GMP        — Gute Herstellungspraxis\n"
            "IVD        — In-vitro-Diagnostika"
        )
        _guide_heading = "Berichtsfeld-Leitfaden"
        _guide_content = (
            "[Aggregate Normen]\n"
            "  Normname: Name der in Dokumenten zitierten Vorschrift/Norm\n"
            "  Anzahl Dokumente: Gesamtanzahl der Dokumente, die diese Norm zitieren\n"
            "  Dokumentenliste: Alle Dokument-IDs, die diese Norm zitieren\n\n"
            "[Dokumentendetails]\n"
            "  Dokument-ID / Titel / Typ: Grundinformationen aus der QMS-Dokumentenbibliothek\n"
            "  Version: Aktuelle Dokumentversion\n"
            "  Zitierte Normen: Alle Vorschriften/Normen, die dieses Dokument zitiert\n\n"
            "[Bewertungsbericht]: KI-Gesamtbewertung der Vollständigkeit und Compliance von Normzitierungen\n"
            "[Verifizierungsbericht]: Zweite Überprüfung der Bewertungsschlussfolgerungen"
        )
    elif lang.startswith("fr"):
        _abbrev_heading = "Liste des abréviations"
        _abbrev_content = (
            "ISO 13485  — Norme internationale pour les systèmes de management de la qualité des dispositifs médicaux\n"
            "EU MDR     — Règlement UE sur les dispositifs médicaux 2017/745\n"
            "EU IVDR    — Règlement UE sur les dispositifs de diagnostic in vitro 2017/746\n"
            "TFDA       — Administration taïwanaise des aliments et des médicaments\n"
            "21 CFR 820 — Réglementation du système qualité de la FDA américaine (QSR/QMSR)\n"
            "MDSAP      — Programme d'audit unique des dispositifs médicaux (USA/Canada/Brésil/Australie/Japon)\n"
            "QMS        — Système de management de la qualité\n"
            "RA         — Affaires réglementaires; éléments marqués comme nécessitant une révision RA\n"
            "GMP        — Bonnes pratiques de fabrication\n"
            "IVD        — Diagnostic in vitro"
        )
        _guide_heading = "Guide des champs du rapport"
        _guide_content = (
            "[Normes agrégées]\n"
            "  Nom de la norme: Nom de la réglementation/norme citée dans les documents\n"
            "  Nombre de documents: Nombre total de documents citant cette norme\n"
            "  Liste de documents: Tous les identifiants de documents citant cette norme\n\n"
            "[Détail par document]\n"
            "  ID / Titre / Type: Informations de base de la bibliothèque QMS\n"
            "  Version: Version actuelle du document\n"
            "  Normes citées: Toutes les réglementations/normes citées dans ce document\n\n"
            "[Rapport d'évaluation]: Évaluation globale de l'IA sur la conformité des citations réglementaires\n"
            "[Rapport de vérification]: Deuxième validation des conclusions d'évaluation"
        )
    else:
        _abbrev_heading = "Abbreviation Legend"
        _abbrev_content = (
            "ISO 13485  — International Standard for Quality Management Systems for Medical Devices\n"
            "EU MDR     — European Union Medical Device Regulation 2017/745\n"
            "EU IVDR    — European Union In Vitro Diagnostic Regulation 2017/746\n"
            "TFDA       — Taiwan Food and Drug Administration\n"
            "21 CFR 820 — US FDA Quality System Regulation (QSR/QMSR)\n"
            "MDSAP      — Medical Device Single Audit Program (US/Canada/Brazil/Australia/Japan)\n"
            "QMS        — Quality Management System\n"
            "RA         — Regulatory Affairs; items marked RA require RA review in report\n"
            "GMP        — Good Manufacturing Practice\n"
            "IVD        — In Vitro Diagnostics"
        )
        _guide_heading = "Report Field Guide"
        _guide_content = (
            "[Aggregate Standards]\n"
            "  Standard Name: Name of regulation/standard referenced in documents\n"
            "  Referenced By (Count): Total number of documents referencing this standard\n"
            "  Reference List: All document IDs referencing this standard\n\n"
            "[Per-Document Detail]\n"
            "  Document ID / Title / Type: Basic information from QMS document library\n"
            "  Version: Current document version\n"
            "  Referenced Standards: All regulatory standards referenced in this document\n\n"
            "[Assessment]: AI's overall assessment of regulatory reference completeness and compliance\n"
            "[Verification]: Second-level validation of assessment conclusions"
        )
    doc.add_heading(_abbrev_heading, level=2)
    doc.add_paragraph(_abbrev_content)

    doc.add_heading(_guide_heading, level=2)
    doc.add_paragraph(_guide_content)

    if not aggregate:
        doc.add_paragraph(_t("regulatory_export.no_refs", lang))
    else:
        # Section 1: Aggregate standards
        doc.add_heading(_t("regulatory_export.std_summary_heading", lang), level=2)

        table1 = doc.add_table(rows=1, cols=3)
        table1.style = "Table Grid"
        table1.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = _tl("regulatory_export.std_headers", lang)
        for i, header in enumerate(headers):
            cell = table1.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)

        for entry in aggregate:
            row = table1.add_row()
            values = [
                entry["standard"],
                str(len(entry["referenced_by"])),
                ", ".join(entry["referenced_by"]),
            ]
            for i, val in enumerate(values):
                cell = row.cells[i]
                cell.text = val
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

        widths = [Cm(5), Cm(2.5), Cm(10)]
        for row in table1.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = width

        doc.add_paragraph()

        # Section 2: Per-document detail
        doc.add_heading(_t("regulatory_export.detail_heading", lang), level=2)

        table2 = doc.add_table(rows=1, cols=5)
        table2.style = "Table Grid"
        table2.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers2 = _tl("regulatory_export.doc_headers", lang)
        for i, header in enumerate(headers2):
            cell = table2.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)

        for d in by_doc:
            row = table2.add_row()
            values = [
                d["doc_id"],
                d["title"],
                d["doc_type"],
                f"v{d['current_version']}",
                ", ".join(d["standards"]),
            ]
            for i, val in enumerate(values):
                cell = row.cells[i]
                cell.text = val
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

        widths2 = [Cm(2.5), Cm(5), Cm(1.5), Cm(1.5), Cm(7)]
        for row in table2.rows:
            for i, width in enumerate(widths2):
                row.cells[i].width = width

    # Section 3: Assessment Report
    if assessment:
        doc.add_paragraph()
        _render_assessment_to_word(doc, assessment, lang)

    # Section: Verification Report (if provided)
    if verification_report and verification_report.get("has_data"):
        _render_verification_to_word(doc, verification_report, lang)

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    run = footer.add_run(
        f"{_t('regulatory_export.footer_regulatory', lang)} | {_t('source_label.source', lang)}: {src_label}"
    )
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    try:
        from src.utils.crossexam_export import _append_crawl_status_word, _load_crawl_results
        _append_crawl_status_word(doc, _load_crawl_results(), lang)
    except Exception:
        pass
    try:
        from src.utils.crossref_export import append_crossref_table_word
        append_crossref_table_word(doc, lang=lang)
    except Exception:
        pass

    # Save
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    cmd_tag = "list" if source_command == "regulatory_list" else "update"
    filename = f"regulatory_standards_{cmd_tag}_{timestamp}.docx"
    filepath = EXPORT_DIR / filename
    doc.save(str(filepath))
    return str(filepath)


def export_regulatory_to_excel(
    scan_result: dict,
    assessment: Optional[str] = None,
    verification_report: Optional[dict] = None,
    lang: str = "zh-TW",
    source_command: str = "regulatory_list",
) -> str:
    """
    Export regulatory standards list to Excel (.xlsx).

    Returns:
        Path to the generated .xlsx file.
    """
    by_doc = scan_result.get("by_document", [])
    aggregate = scan_result.get("aggregate", [])

    src_label = _source_label(source_command, lang)
    wb = Workbook()

    # Sheet 1: Aggregate
    ws1 = wb.active
    ws1.title = _t("regulatory_export.std_summary", lang)

    # Title
    ws1.merge_cells("A1:C1")
    title_cell = ws1.cell(row=1, column=1)
    title_cell.value = (
        f"{_t('regulatory_export.title_regulatory', lang)}（{src_label}）"
    )
    title_cell.font = Font(name="Microsoft JhengHei", size=14, bold=True)
    title_cell.alignment = Alignment(horizontal="center")

    # Metadata
    ws1.merge_cells("A2:C2")
    meta_cell = ws1.cell(row=2, column=1)
    meta_cell.value = (
        f"{_t('regulatory_export.export_time', lang)}: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | "
        f"{_t('source_label.source', lang)}: {src_label} | "
        f"{_t('regulatory_export.std_count', lang, count=len(aggregate))} | "
        f"{_t('regulatory_export.std_coverage', lang, count=len(by_doc))}"
    )
    meta_cell.font = Font(
        name="Microsoft JhengHei", size=9, italic=True, color="808080"
    )

    # Headers
    headers = _tl("regulatory_export.std_headers", lang)
    for col, header in enumerate(headers, 1):
        cell = ws1.cell(row=4, column=col)
        cell.value = header
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = THIN_BORDER

    # Data
    for row_idx, entry in enumerate(aggregate, 5):
        values = [
            entry["standard"],
            len(entry["referenced_by"]),
            ", ".join(entry["referenced_by"]),
        ]
        for col, val in enumerate(values, 1):
            cell = ws1.cell(row=row_idx, column=col)
            cell.value = val
            cell.font = CELL_FONT
            cell.border = THIN_BORDER

    ws1.column_dimensions["A"].width = 30
    ws1.column_dimensions["B"].width = 15
    ws1.column_dimensions["C"].width = 50
    ws1.freeze_panes = "A5"

    # Sheet 2: Per-document detail
    ws2 = wb.create_sheet(_t("regulatory_export.detail_sheet", lang))

    # Title
    ws2.merge_cells("A1:E1")
    title_cell2 = ws2.cell(row=1, column=1)
    title_cell2.value = _t("regulatory_export.detail_title", lang)
    title_cell2.font = Font(name="Microsoft JhengHei", size=14, bold=True)
    title_cell2.alignment = Alignment(horizontal="center")

    # Headers
    headers2 = _tl("regulatory_export.doc_headers", lang)
    for col, header in enumerate(headers2, 1):
        cell = ws2.cell(row=3, column=col)
        cell.value = header
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = THIN_BORDER

    for row_idx, d in enumerate(by_doc, 4):
        values = [
            d["doc_id"],
            d["title"],
            d["doc_type"],
            f"v{d['current_version']}",
            ", ".join(d["standards"]),
        ]
        for col, val in enumerate(values, 1):
            cell = ws2.cell(row=row_idx, column=col)
            cell.value = val
            cell.font = CELL_FONT
            cell.border = THIN_BORDER

    ws2.column_dimensions["A"].width = 15
    ws2.column_dimensions["B"].width = 35
    ws2.column_dimensions["C"].width = 10
    ws2.column_dimensions["D"].width = 10
    ws2.column_dimensions["E"].width = 60
    ws2.freeze_panes = "A4"

    # Footer note
    note_row = len(by_doc) + 5
    ws2.merge_cells(f"A{note_row}:E{note_row}")
    note_cell = ws2.cell(row=note_row, column=1)
    note_cell.value = (
        f"{_t('regulatory_export.footer_regulatory', lang)} | "
        f"{_t('source_label.source', lang)}: {src_label}"
    )
    note_cell.font = Font(
        name="Microsoft JhengHei", size=8, italic=True, color="808080"
    )

    # Sheet 3: Assessment Report (if provided)
    if assessment:
        _render_assessment_to_excel(wb, assessment, lang)

    # Sheet: Verification Report (if provided)
    if verification_report and verification_report.get("has_data"):
        _render_verification_to_excel(wb, verification_report, lang)

    try:
        from src.utils.crossexam_export import _append_crawl_status_excel, _load_crawl_results
        _append_crawl_status_excel(wb, _load_crawl_results(), lang)
    except Exception:
        pass
    try:
        from src.utils.crossref_export import append_crossref_table_excel
        append_crossref_table_excel(wb, lang=lang)
    except Exception:
        pass

    # Save
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    cmd_tag = "list" if source_command == "regulatory_list" else "update"
    filename = f"regulatory_standards_{cmd_tag}_{timestamp}.xlsx"
    filepath = EXPORT_DIR / filename
    wb.save(str(filepath))
    return str(filepath)


# ============================================================
# 引用清單 (Document Reference List after version update)
# ============================================================


def format_reference_table_markdown(
    doc_id: str,
    ref_docs: List[dict],
    lang: str = "zh-TW",
) -> str:
    """Format document reference list as Markdown for chat display."""
    if not ref_docs:
        return _t("regulatory_export.no_refs_doc", lang, doc_id=doc_id)

    ref_headers = _tl("regulatory_export.ref_table_headers", lang)
    lines = [
        _t(
            "regulatory_export.ref_table_title",
            lang,
            doc_id=doc_id,
            count=len(ref_docs),
        )
        + "\n",
        f"| {' | '.join(ref_headers)} |",
        f"|{'|'.join(['------' for _ in ref_headers])}|",
    ]

    for r in ref_docs:
        ref_type = (
            _t("regulatory_export.explicit_ref", lang)
            if r.get("reference_type") == "explicit"
            else _t("regulatory_export.implicit_ref", lang)
        )
        lines.append(
            f"| {r['doc_id']} | {r['title'][:30]} | {r['doc_type']} | v{r['current_version']} | {ref_type} |"
        )

    return "\n".join(lines)


def export_reference_to_word(
    doc_id: str,
    ref_docs: List[dict],
    lang: str = "zh-TW",
) -> str:
    """
    Export document reference list to Word (.docx).

    Returns:
        Path to the generated .docx file.
    """
    doc = Document()

    # Title
    title = doc.add_heading(
        _t("regulatory_export.title_reference", lang, doc_id=doc_id), level=1
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = meta.add_run(
        f"{_t('regulatory_export.export_time', lang)}: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run2 = meta2.add_run(
        _t("regulatory_export.ref_count", lang, doc_id=doc_id, count=len(ref_docs))
    )
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    if not ref_docs:
        doc.add_paragraph(_t("regulatory_export.no_refs_doc_word", lang, doc_id=doc_id))
    else:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = _tl("regulatory_export.ref_headers", lang)
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)

        for idx, r in enumerate(ref_docs, 1):
            row = table.add_row()
            ref_type = (
                _t("regulatory_export.explicit_ref", lang)
                if r.get("reference_type") == "explicit"
                else _t("regulatory_export.implicit_ref", lang)
            )
            values = [
                str(idx),
                r["doc_id"],
                r["title"],
                f"v{r['current_version']}",
                ref_type,
            ]
            for i, val in enumerate(values):
                cell = row.cells[i]
                cell.text = val
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

        widths = [Cm(1), Cm(3), Cm(6), Cm(2), Cm(2.5)]
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = width

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    run = footer.add_run(_t("regulatory_export.footer_reference", lang, doc_id=doc_id))
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    try:
        from src.utils.crossexam_export import _append_crawl_status_word, _load_crawl_results
        _append_crawl_status_word(doc, _load_crawl_results(), lang)
    except Exception:
        pass
    try:
        from src.utils.crossref_export import append_crossref_table_word
        append_crossref_table_word(doc, lang=lang)
    except Exception:
        pass

    # Save
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"reference_list_{doc_id}_{timestamp}.docx"
    filepath = EXPORT_DIR / filename
    doc.save(str(filepath))
    return str(filepath)


def export_reference_to_excel(
    doc_id: str,
    ref_docs: List[dict],
    lang: str = "zh-TW",
) -> str:
    """
    Export document reference list to Excel (.xlsx).

    Returns:
        Path to the generated .xlsx file.
    """
    wb = Workbook()
    ws = wb.active
    ws.title = _t("regulatory_export.ref_sheet", lang)

    # Title
    ws.merge_cells("A1:E1")
    title_cell = ws.cell(row=1, column=1)
    title_cell.value = _t("regulatory_export.title_reference", lang, doc_id=doc_id)
    title_cell.font = Font(name="Microsoft JhengHei", size=14, bold=True)
    title_cell.alignment = Alignment(horizontal="center")

    # Metadata
    ws.merge_cells("A2:E2")
    meta_cell = ws.cell(row=2, column=1)
    meta_cell.value = (
        f"{_t('regulatory_export.export_time', lang)}: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | "
        f"{_t('regulatory_export.ref_count', lang, doc_id=doc_id, count=len(ref_docs))}"
    )
    meta_cell.font = Font(
        name="Microsoft JhengHei", size=9, italic=True, color="808080"
    )

    # Headers
    headers = _tl("regulatory_export.ref_headers", lang)
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=4, column=col)
        cell.value = header
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = THIN_BORDER

    # Data
    for row_idx, r in enumerate(ref_docs, 5):
        ref_type = (
            _t("regulatory_export.explicit_ref", lang)
            if r.get("reference_type") == "explicit"
            else _t("regulatory_export.implicit_ref", lang)
        )
        values = [
            row_idx - 4,
            r["doc_id"],
            r["title"],
            f"v{r['current_version']}",
            ref_type,
        ]
        for col, val in enumerate(values, 1):
            cell = ws.cell(row=row_idx, column=col)
            cell.value = val
            cell.font = CELL_FONT
            cell.border = THIN_BORDER

    ws.column_dimensions["A"].width = 5
    ws.column_dimensions["B"].width = 15
    ws.column_dimensions["C"].width = 40
    ws.column_dimensions["D"].width = 10
    ws.column_dimensions["E"].width = 12
    ws.freeze_panes = "A5"

    # Footer
    note_row = len(ref_docs) + 6
    ws.merge_cells(f"A{note_row}:E{note_row}")
    note_cell = ws.cell(row=note_row, column=1)
    note_cell.value = _t("regulatory_export.footer_reference", lang, doc_id=doc_id)
    note_cell.font = Font(
        name="Microsoft JhengHei", size=8, italic=True, color="808080"
    )

    try:
        from src.utils.crossexam_export import _append_crawl_status_excel, _load_crawl_results
        _append_crawl_status_excel(wb, _load_crawl_results(), lang)
    except Exception:
        pass
    try:
        from src.utils.crossref_export import append_crossref_table_excel
        append_crossref_table_excel(wb, lang=lang)
    except Exception:
        pass

    # Save
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"reference_list_{doc_id}_{timestamp}.xlsx"
    filepath = EXPORT_DIR / filename
    wb.save(str(filepath))
    return str(filepath)
