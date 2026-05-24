"""
AI-QMS — Cross-Examination & Deep Report Export
=================================================

Export utilities for:
1. Individual cross-exam records (Word/Excel)
2. Deep analysis report with ALL LLM interactions (Word/Excel)

Follows the export pattern from doclist_export.py / regulatory_export.py:
  - Uses python-docx for Word
  - Uses openpyxl for Excel
  - Returns file paths
  - Thread-safe via safe_io
"""

from __future__ import annotations

import json
import logging
import re
import time
from datetime import datetime
from pathlib import Path
from typing import Optional

from src.utils.safe_io import safe_save_binary

logger = logging.getLogger(__name__)

__all__ = [
    "export_crossexam_record_word",
    "export_crossexam_record_excel",
    "export_deep_report_word",
    "export_deep_report_excel",
]

EXPORT_DIR = (Path(__file__).resolve().parent.parent.parent / "data" / "exports")


# ============================================================
# Language helpers (bilingual section headers)
# ============================================================

from src.chainlit_app.lang_config import lang_key as _lang_key  # noqa: E402


def _safe_str(value, max_len: int = 0) -> str:
    """Coerce any value to an Excel/Word-safe string.

    Guards against LLMs returning dicts where strings are expected
    (e.g. regulation_citation as {"primary": "...", "supplementary": {...}}).
    """
    if value is None:
        return ""
    if isinstance(value, bool):
        return "Y" if value else ""
    if isinstance(value, dict):
        # Human-readable flatten for primary/supplementary pattern
        if "primary" in value or "supplementary" in value:
            primary = value.get("primary", "")
            supp = value.get("supplementary", {})
            if isinstance(supp, dict) and supp:
                supp_str = "; ".join(
                    f"{k}: {v}" for k, v in supp.items()
                )
                result = f"{primary} | 補充: {supp_str}" if primary else supp_str
            else:
                result = str(primary)
        else:
            result = "; ".join(f"{k}: {v}" for k, v in value.items())
    elif isinstance(value, (list, tuple)):
        result = "; ".join(str(x) for x in value)
    else:
        result = str(value)
    if max_len and len(result) > max_len:
        return result[:max_len]
    return result


def _build_env_table_rows(run_meta: dict, lk: str) -> list[tuple[str, str]]:
    """Build (label, value) pairs for the Execution Environment table."""
    _lb = {
        "zh": {
            "provider":      "LLM Provider",
            "provider_type": "Provider 類型",
            "model":         "使用模型",
            "api_base":      "API Base URL",
            "max_workers":   "並發設定",
            "gpu":           "GPU",
            "vram":          "VRAM",
            "capability":    "GPU Capability",
            "driver_cuda":   "Driver CUDA",
            "torch":         "PyTorch",
            "torch_cuda":    "PyTorch CUDA",
            "compat":        "相容性狀態",
            "platform":      "作業系統",
        },
        "en": {
            "provider":      "LLM Provider",
            "provider_type": "Provider Type",
            "model":         "Model",
            "api_base":      "API Base URL",
            "max_workers":   "Concurrency",
            "gpu":           "GPU",
            "vram":          "VRAM",
            "capability":    "GPU Capability",
            "driver_cuda":   "Driver CUDA",
            "torch":         "PyTorch",
            "torch_cuda":    "PyTorch CUDA",
            "compat":        "Compatibility",
            "platform":      "OS",
        },
        "ja": {
            "provider":      "LLM プロバイダ",
            "provider_type": "プロバイダ種別",
            "model":         "使用モデル",
            "api_base":      "API ベース URL",
            "max_workers":   "並列設定",
            "gpu":           "GPU",
            "vram":          "VRAM",
            "capability":    "GPU Capability",
            "driver_cuda":   "Driver CUDA",
            "torch":         "PyTorch",
            "torch_cuda":    "PyTorch CUDA",
            "compat":        "互換性ステータス",
            "platform":      "OS",
        },
    }
    lb = _lb.get(lk, _lb["en"])

    _no_torch_label = {
        "zh": "僅 CPU（不支援 CUDA GPU 加速）",
        "ja": "CPU のみ（CUDA GPU 非対応）",
        "en": "CPU only (no CUDA support)",
    }
    _torch_cuda_na_label = {
        "zh": "不適用（僅 CPU 版 PyTorch）",
        "ja": "非対応（CPU 専用 PyTorch）",
        "en": "N/A (CPU-only PyTorch)",
    }

    def _compat_str(status: str, warnings: list) -> str:
        if status == "ok":
            return "✅ OK"
        if status == "no-torch":
            return _no_torch_label.get(lk, _no_torch_label["en"])
        if warnings:
            return f"⚠️ {status}: {warnings[0][:80]}"
        return f"⚠️ {status}"

    def _torch_version_str(ver: str) -> str:
        """Format torch version string, clarifying CPU-only editions."""
        if "+cpu" in ver.lower():
            # Extract base version before the +cpu suffix
            base = ver.lower().split("+cpu")[0].rstrip(".")
            _cpu_label = {
                "zh": f"PyTorch {base}（CPU 專用版本）",
                "ja": f"PyTorch {base}（CPU 専用エディション）",
                "en": f"PyTorch {base} (CPU-only edition)",
            }
            return _cpu_label.get(lk, _cpu_label["en"])
        return ver

    rows: list[tuple[str, str]] = []
    if run_meta.get("provider_name"):
        rows.append((lb["provider"], run_meta["provider_name"]))
    if run_meta.get("provider_type"):
        rows.append((lb["provider_type"], run_meta["provider_type"]))
    if run_meta.get("model"):
        rows.append((lb["model"], run_meta["model"]))
    if run_meta.get("api_base_url") and run_meta.get("is_local"):
        rows.append((lb["api_base"], run_meta["api_base_url"]))
    if "max_workers" in run_meta:
        reason = run_meta.get("workers_reason", "")
        wval = f"max_workers = {run_meta['max_workers']}"
        if reason:
            wval += f"  ({reason})"
        rows.append((lb["max_workers"], wval))
    rows.append(("", ""))  # visual separator
    if run_meta.get("gpu_name"):
        rows.append((lb["gpu"], run_meta["gpu_name"]))
    if run_meta.get("vram_gb") is not None and run_meta.get("vram_gb", 0) > 0:
        rows.append((lb["vram"], f"{run_meta['vram_gb']} GB"))
    if run_meta.get("gpu_capability"):
        rows.append((lb["capability"], run_meta["gpu_capability"]))
    if run_meta.get("driver_cuda"):
        rows.append((lb["driver_cuda"], run_meta["driver_cuda"]))
    if run_meta.get("torch_version"):
        rows.append((lb["torch"], _torch_version_str(run_meta["torch_version"])))
    torch_cuda_val = run_meta.get("torch_cuda")
    torch_ver = run_meta.get("torch_version", "")
    is_cpu_only_torch = torch_ver and "+cpu" in torch_ver.lower()
    if torch_cuda_val:
        rows.append((lb["torch_cuda"], torch_cuda_val))
    elif is_cpu_only_torch:
        rows.append((lb["torch_cuda"], _torch_cuda_na_label.get(lk, _torch_cuda_na_label["en"])))
    compat = run_meta.get("gpu_compat_status", "")
    if compat:
        rows.append((lb["compat"], _compat_str(compat, run_meta.get("gpu_compat_warnings", []))))
    if run_meta.get("platform"):
        rows.append((lb["platform"], run_meta["platform"]))
    return rows


_EXPORT_HEADERS: dict[str, dict[str, str]] = {
    "zh": {
        "title_crossexam": "AI-QMS 交叉詰問記錄",
        "title_deep": "AI-QMS 完整分析報告",
        "summary": "摘要",
        "clause_details": "條款交叉詰問詳情",
        "record_id": "記錄 ID",
        "analysis_id": "分析 ID",
        "time": "時間",
        "regulations": "法規",
        "countries": "國家",
        "clause_count": "條款數",
        "agreed": "同意",
        "flagged_ra": "標記 RA",
        "total_rounds": "總輪次",
        "model": "模型",
        "duration": "耗時",
        "none": "無",
        "doc": "文件",
        "verdict": "判定",
        "gap": "差距",
        "yes": "是",
        "no": "否",
        # Word body labels
        "how_it_works_heading": "作用原理 / How Cross-Examination Works",
        "how_it_works_body": (
            "Phase 5 交叉詰問採用辯論式 AI 稽核架構：\n\n"
            "1. 依 ISO 13485 稽核清單（71 條款）抽取當次問題（以日期為 seed 輪替，各條款至少 7 個版本）\n"
            "2. Analyzer（辯護方）分析各 QMS 文件，輸出：立場、信心度、關鍵證據\n"
            "3. Verifier（質疑方）逐條質疑 Analyzer 論點，挑戰未引用的法規要求或證據漏洞\n"
            "4. 雙方進行最多 3 輪辯論，達成 agree / partial / disagree 結論\n"
            "5. QA Auditor（審查者）對整場辯論獨立評分（0–100）\n"
            "6. 最終判定（verdict）依辯論結論與 gap_severity 由風險矩陣自動計算\n\n"
            "縮寫對照：verdict — compliant / improvement_plan / deadline_correction / immediate_correction\n"
            "         gap_severity — none / minor / major / critical\n"
            "         flagged_for_ra — 需 RA 法規事務人員進一步審查"
        ),
        "audit_question": "稽核問題",
        "expected_evidence": "預期書面證據（靜態清單）/ Expected Evidence (Static)",
        "source_b_label": "🤖 [LLM 動態生成問題 (Side B)]",
        "source_a_label": "📋 [靜態題庫問題 (Side A)]",
        "focus_area_label": "聚焦面向",
        "verifiable_by_label": "驗證方式",
        "analyzer_label": "🔍 分析者（實際看到）",
        "verifier_label": "⚖️ 驗證者（期望看到）",
        "qa_audit_heading": "🔎 第三方稽核",
        "score_label": "分數",
        "question_quality_label": "問題品質",
        "answer_accuracy_label": "回答準確",
        "hallucination_label": "幻覺偵測",
        "hallucination_yes": "⚠️ 是",
        "hallucination_no": "否",
        "ra_flag_yes": "⚠️ 是",
        "ra_flag_no": "否",
        "agreed_yes": "✅ 是",
        "agreed_no": "❌ 否",
        "meta_label": "記錄 ID: {rid}  |  分析 ID: {aid}  |  時間: {ts}",
        "summary_body": "法規: {regs}\n國家: {countries}\n條款數: {clauses}  |  同意: {agreed}  |  標記 RA: {flagged}  |  總輪次: {rounds}\n模型: {model}  |  耗時: {duration:.1f}s",
        # Deep report section headings
        "deep_title": "AI-QMS 深度合規性分析報告",
        "deep_meta": "分析 ID: {run_id}  |  匯出時間: {ts}",
        "deep_s1": "第一章 執行摘要",
        "deep_s1_intro": "本次合規性分析共評估 {total} 個條款-文件對照項目。\n\n判定結果分布:\n",
        "deep_s1_risk": "\n風險等級分布:\n",
        "deep_s1_ra": "\n⚠️ 需 RA 審查: {flagged} 項\n",
        "deep_s2": "第二章 GAP 分析詳情 (Phase 1)",
        "deep_s2_doc": "文件",
        "deep_s2_found": "找到",
        "deep_s2_not_found": "未找到",
        "deep_s2_inadequate": "不足",
        "deep_llm_response": "LLM 回應",
        "deep_token_usage": "Token 用量: {tokens:,}  |  模型: {model}",
        "deep_no_p1": "（本次分析無 Phase 1 LLM 互動記錄）",
        "deep_no_record": "（無 LLM 互動記錄可用）",
        "deep_s3": "第三章 驗證詳情 (Phase 2)",
        "deep_no_p2": "（本次分析無 Phase 2 LLM 互動記錄）",
        "deep_s4": "第四章 改善建議 (Phase 4)",
        "deep_no_p4": "（本次分析無 Phase 4 LLM 互動記錄）",
        "deep_s5": "第五章 交叉詰問 (Phase 5)",
        "deep_no_p5": "（本次分析無 Phase 5 LLM 互動記錄）",
        "deep_no_xexam": "（無交叉詰問記錄可用）",
        "deep_s55": "第五章之二 第三方品質稽核 (Phase 5 Step 2)",
        "deep_qa_score": "整體品質分數: {score}/100\n稽核條款數: {count}\n模型: {model}",
        "deep_qa_summary": "稽核摘要",
        "deep_qa_recs": "稽核建議",
        "deep_qa_clause_results": "逐條稽核結果",
        "deep_qa_tbl_headers": ["條款", "分數", "問題品質", "回答準確", "幻覺偵測", "問題"],
        "deep_doc_label": "文件",
        "deep_analyzer_label": "🔍 分析者",
        "deep_verifier_label": "⚖️ 驗證者",
        "deep_s6": "第六章 合規性分析結果表",
        "deep_s6_headers": ["條款", "文件", "稽核影響", "判定", "風險", "差距", "RA 標記", "Pipeline 狀態"],
        "deep_s7": "第七章 交叉詰問品質分析",
        "deep_s7_findings": "發現事項",
        "deep_s7_recs": "建議",
        "deep_s7_prompt": "Prompt 調整記錄",
        "deep_s7_no_summary": "（無分析摘要）",
        "deep_s8": "附錄 LLM 使用統計",
        "deep_s8_calls": "次呼叫",
        "deep_no_issues": "無",
        "deep_qa_skipped": "（已跳過：{summary}）",
        "deep_no_qa_record": "（無第三方品質稽核記錄）",
        "xl_sheet_summary": "摘要",
        "xl_sheet_compliance": "合規分析",
        "xl_sheet_llm": "LLM 互動記錄",
        "xl_sheet_crossexam": "交叉詰問",
        "xl_sheet_meta": "品質分析",
        "xl_run_id": "分析 ID",
        "xl_export_time": "匯出時間",
        "xl_total_rows": "總分析項目",
        "xl_flagged_ra": "需 RA 審查",
        "xl_verdict_prefix": "判定 - ",
        "xl_risk_prefix": "風險 - ",
        "xl_comp_headers": ["條款 ID","條款名稱","文件 ID","文件標題","稽核影響","稽核問題","判定","風險等級","差距嚴重度","證據 (找到/總計)","RA 標記","RA 覆寫","RA 備註","改善建議 (LLM P4)","法規引用 (LLM P4)","分析者立場 (LLM P5)","驗證者評語 (LLM P5)","Pipeline 狀態"],
        "deep_s0": "第零章 Pipeline 執行進度總覽",
        "deep_s0_phase_col_phase": "階段",
        "deep_s0_phase_col_name_zh": "名稱（中文）",
        "deep_s0_phase_col_name_en": "名稱（英文）",
        "deep_s0_phase_col_rows": "行數",
        "deep_s0_phase_col_status": "狀態",
        "deep_s0_skipped": "已跳過階段",
        "deep_s0_budget_title": "LLM Token 使用統計",
        "deep_s0_budget_total": "已使用 Token",
        "deep_s0_budget_calls": "LLM 呼叫次數",
        "deep_s0_budget_pct": "預算使用率",
        "deep_s0_budget_remaining": "剩餘 Token",
        "deep_s0_dq_title": "資料品質檢查結果 (P0)",
        "deep_s0_sc_title": "來源驗證結果 (P6)",
        "deep_s0_none": "（無資料）",
        "xl_llm_headers": ["Phase","文件 ID","條款 ID","角色","Round","LLM 回應 (摘要)","Token 用量","模型","時間"],
        "xl_xe_headers": ["條款 ID","條款名稱","文件 ID","判定","同意","RA 標記","輪次數","R1 分析者立場","R1 分析者信心","R1 驗證者評估","R1 Agreement","QA 分數","問題品質","回答準確","幻覺偵測"],
        # Section 0.5
        "deep_s05_heading": "第 0.5 章 風險優先項目摘要",
        "deep_s05_body": "以下 {n} 個條款-文件對照項目屬於高風險／嚴重不符合，需優先處理。",
        "deep_s05_headers": ["條款 ID", "條款名稱", "文件 ID", "判定", "風險等級", "差距嚴重度", "改善建議 (摘要)"],
        # Crawl status appendix
        "crawl_appendix_heading": "附錄：法規資料來源爬取狀態",
        "crawl_appendix_headers": ["狀態", "機構", "地區", "URL", "爬取時間", "HTTP 狀態", "錯誤訊息"],
        "crawl_appendix_selected": "本次爬取地區",
        "crawl_appendix_all_ts": "資料更新時間",
        "crawl_xl_sheet": "法規爬取狀態",
        "xl_phase_labels": {
            "gap_scan": "P1 差距掃描",
            "checklist_verify": "P2 查核表驗證",
            "remediation": "P4 改善建議",
            "verification": "P5 交叉詰問",
        },
        "xl_role_labels": {"analyzer": "分析者", "verifier": "驗證者"},
        "xl_batch_label": "—（文件層級）",
        "xl_xe_ext_headers": ["條款 ID","條款名稱","文件 ID","判定","同意","RA 標記","輪次數","R1 分析者立場","R1 分析者信心","R1 關鍵證據","R1 驗證者質疑","R1 Agreement","QA 分數","問題品質","回答準確","幻覺偵測"],
    },
    "en": {
        "title_crossexam": "AI-QMS Cross-Examination Record",
        "title_deep": "AI-QMS Full Analysis Report",
        "summary": "Summary",
        "clause_details": "Clause Cross-Examination Details",
        "record_id": "Record ID",
        "analysis_id": "Run ID",
        "time": "Time",
        "regulations": "Regulations",
        "countries": "Countries",
        "clause_count": "Clause count",
        "agreed": "Agreed",
        "flagged_ra": "Flagged for RA",
        "total_rounds": "Total rounds",
        "model": "Model",
        "duration": "Duration",
        "none": "None",
        "doc": "Document",
        "verdict": "Verdict",
        "gap": "Gap",
        "yes": "Yes",
        "no": "No",
        # Word body labels
        "how_it_works_heading": "How Cross-Examination Works",
        "how_it_works_body": (
            "Phase 5 cross-examination uses a debate-based AI audit architecture:\n\n"
            "1. Questions are drawn from the ISO 13485 audit checklist (71 clauses) with date-seeded rotation — at least 7 versions per clause\n"
            "2. Analyzer (defender) analyzes each QMS document and outputs: position, confidence, key evidence\n"
            "3. Verifier (challenger) challenges the Analyzer's arguments clause by clause, targeting uncited regulatory requirements or evidence gaps\n"
            "4. Up to 3 debate rounds, reaching agree / partial / disagree conclusion\n"
            "5. QA Auditor independently scores the entire debate (0–100)\n"
            "6. Final verdict is auto-calculated from debate conclusion and gap_severity via the risk matrix\n\n"
            "Abbreviations: verdict — compliant / improvement_plan / deadline_correction / immediate_correction\n"
            "               gap_severity — none / minor / major / critical\n"
            "               flagged_for_ra — requires further review by RA regulatory affairs personnel"
        ),
        "audit_question": "Audit Question",
        "expected_evidence": "Expected Evidence (Static Checklist)",
        "source_b_label": "🤖 [AI-Generated Question (Side B)]",
        "source_a_label": "📋 [Static Question Pool (Side A)]",
        "focus_area_label": "Focus area",
        "verifiable_by_label": "Verifiable by",
        "analyzer_label": "🔍 Analyzer (Actual Evidence Found)",
        "verifier_label": "⚖️ Verifier (Expected to See)",
        "qa_audit_heading": "🔎 Third-Party Audit",
        "score_label": "Score",
        "question_quality_label": "Question quality",
        "answer_accuracy_label": "Answer accuracy",
        "hallucination_label": "Hallucination detected",
        "hallucination_yes": "⚠️ Yes",
        "hallucination_no": "No",
        "ra_flag_yes": "⚠️ Yes",
        "ra_flag_no": "No",
        "agreed_yes": "✅ Yes",
        "agreed_no": "❌ No",
        "meta_label": "Record ID: {rid}  |  Run ID: {aid}  |  Time: {ts}",
        "summary_body": "Regulations: {regs}\nCountries: {countries}\nClauses: {clauses}  |  Agreed: {agreed}  |  Flagged RA: {flagged}  |  Total rounds: {rounds}\nModel: {model}  |  Duration: {duration:.1f}s",
        # Deep report section headings
        "deep_title": "AI-QMS Deep Compliance Analysis Report",
        "deep_meta": "Run ID: {run_id}  |  Export time: {ts}",
        "deep_s1": "Chapter 1: Executive Summary",
        "deep_s1_intro": "This compliance analysis evaluated {total} clause-document pairs.\n\nVerdict distribution:\n",
        "deep_s1_risk": "\nRisk level distribution:\n",
        "deep_s1_ra": "\n⚠️ Requires RA review: {flagged} items\n",
        "deep_s2": "Chapter 2: GAP Analysis Details (Phase 1)",
        "deep_s2_doc": "Document",
        "deep_s2_found": "Found",
        "deep_s2_not_found": "Not found",
        "deep_s2_inadequate": "Inadequate",
        "deep_llm_response": "LLM Response",
        "deep_token_usage": "Token usage: {tokens:,}  |  Model: {model}",
        "deep_no_p1": "(No Phase 1 LLM interaction records for this analysis)",
        "deep_no_record": "(No LLM interaction records available)",
        "deep_s3": "Chapter 3: Verification Details (Phase 2)",
        "deep_no_p2": "(No Phase 2 LLM interaction records for this analysis)",
        "deep_s4": "Chapter 4: Improvement Suggestions (Phase 4)",
        "deep_no_p4": "(No Phase 4 LLM interaction records for this analysis)",
        "deep_s5": "Chapter 5: Cross-Examination (Phase 5)",
        "deep_no_p5": "(No Phase 5 LLM interaction records for this analysis)",
        "deep_no_xexam": "(No cross-examination records available)",
        "deep_s55": "Chapter 5b: Third-Party QA Audit (Phase 5 Step 2)",
        "deep_qa_score": "Overall quality score: {score}/100\nClauses audited: {count}\nModel: {model}",
        "deep_qa_summary": "Audit Summary",
        "deep_qa_recs": "Audit Recommendations",
        "deep_qa_clause_results": "Per-Clause Audit Results",
        "deep_qa_tbl_headers": ["Clause", "Score", "Question Quality", "Answer Accuracy", "Hallucination", "Issues"],
        "deep_doc_label": "Document",
        "deep_analyzer_label": "🔍 Analyzer",
        "deep_verifier_label": "⚖️ Verifier",
        "deep_s6": "Chapter 6: Compliance Analysis Results",
        "deep_s6_headers": ["Clause", "Document", "Audit Impact", "Verdict", "Risk", "Gap", "RA Flag", "Pipeline Status"],
        "deep_s7": "Chapter 7: Cross-Examination Quality Analysis",
        "deep_s7_findings": "Findings",
        "deep_s7_recs": "Recommendations",
        "deep_s7_prompt": "Prompt Tuning Log",
        "deep_s7_no_summary": "(No analysis summary)",
        "deep_s8": "Appendix: LLM Usage Statistics",
        "deep_s8_calls": "calls",
        "deep_no_issues": "None",
        "deep_qa_skipped": "(Skipped: {summary})",
        "deep_no_qa_record": "(No third-party QA audit records)",
        "xl_sheet_summary": "Summary",
        "xl_sheet_compliance": "Compliance",
        "xl_sheet_llm": "LLM Interactions",
        "xl_sheet_crossexam": "Cross-Exam",
        "xl_sheet_meta": "Quality Analysis",
        "xl_run_id": "Run ID",
        "xl_export_time": "Export Time",
        "xl_total_rows": "Total Items",
        "xl_flagged_ra": "Flagged for RA",
        "xl_verdict_prefix": "Verdict - ",
        "xl_risk_prefix": "Risk - ",
        "xl_comp_headers": ["Clause ID","Clause Title","Doc ID","Doc Title","Audit Impact","Audit Question","Verdict","Risk Level","Gap Severity","Evidence (Found/Total)","RA Flag","RA Override","RA Notes","Improvement Suggestion (P4)","Regulation Cite (P4)","Analyzer Position (P5)","Verifier Assessment (P5)","Pipeline Status"],
        "deep_s0": "Chapter 0: Pipeline Execution Progress",
        "deep_s0_phase_col_phase": "Phase",
        "deep_s0_phase_col_name_zh": "Name (ZH)",
        "deep_s0_phase_col_name_en": "Name (EN)",
        "deep_s0_phase_col_rows": "Rows",
        "deep_s0_phase_col_status": "Status",
        "deep_s0_skipped": "Skipped Phases",
        "deep_s0_budget_title": "LLM Token Usage",
        "deep_s0_budget_total": "Tokens Used",
        "deep_s0_budget_calls": "LLM Calls",
        "deep_s0_budget_pct": "Budget Used %",
        "deep_s0_budget_remaining": "Remaining Tokens",
        "deep_s0_dq_title": "Data Quality Check Results (P0)",
        "deep_s0_sc_title": "Source Verification Results (P6)",
        "deep_s0_none": "(No data)",
        "xl_llm_headers": ["Phase","Doc ID","Clause ID","Role","Round","LLM Response (excerpt)","Token Usage","Model","Timestamp"],
        "xl_xe_headers": ["Clause ID","Clause Title","Doc ID","Verdict","Agreed","RA Flag","Rounds","R1 Analyzer Position","R1 Analyzer Confidence","R1 Verifier Assessment","R1 Agreement","QA Score","Question Quality","Answer Accuracy","Hallucination"],
        # Section 0.5
        "deep_s05_heading": "Chapter 0.5: Risk Priority Summary",
        "deep_s05_body": "The following {n} clause-document pairs are high/critical risk or non-compliant — require priority action.",
        "deep_s05_headers": ["Clause ID", "Clause Title", "Document ID", "Verdict", "Risk Level", "Gap Severity", "Remediation (Summary)"],
        # Crawl status appendix
        "crawl_appendix_heading": "Appendix: Regulatory Source Crawl Status",
        "crawl_appendix_headers": ["Status", "Agency", "Region", "URL", "Crawled At", "HTTP Status", "Error"],
        "crawl_appendix_selected": "Regions crawled this session",
        "crawl_appendix_all_ts": "Data last updated",
        "crawl_xl_sheet": "Crawl Status",
        "xl_phase_labels": {
            "gap_scan": "P1 Gap Scan",
            "checklist_verify": "P2 Checklist Verify",
            "remediation": "P4 Remediation",
            "verification": "P5 Cross-Exam",
        },
        "xl_role_labels": {"analyzer": "Analyzer", "verifier": "Verifier"},
        "xl_batch_label": "— (Batch/Doc level)",
        "xl_xe_ext_headers": ["Clause ID","Clause Title","Doc ID","Verdict","Agreed","RA Flag","Rounds","R1 Analyzer Position","R1 Analyzer Confidence","R1 Key Evidence","R1 Verifier Challenges","R1 Agreement","QA Score","Question Quality","Answer Accuracy","Hallucination"],
    },
    "ja": {
        "title_crossexam": "AI-QMS 相互尋問記録",
        "title_deep": "AI-QMS 完全分析レポート",
        "summary": "サマリー",
        "clause_details": "条項相互尋問詳細",
        "record_id": "記録ID",
        "analysis_id": "実行ID",
        "time": "時刻",
        "regulations": "規制",
        "countries": "国",
        "clause_count": "条項数",
        "agreed": "同意",
        "flagged_ra": "RA要確認",
        "total_rounds": "総ラウンド数",
        "model": "モデル",
        "duration": "所要時間",
        "none": "なし",
        "doc": "文書",
        "verdict": "判定",
        "gap": "ギャップ",
        "yes": "はい",
        "no": "いいえ",
        # Word body labels
        "how_it_works_heading": "相互尋問の仕組み / How Cross-Examination Works",
        "how_it_works_body": (
            "フェーズ5の相互尋問は、ディベート型AI監査アーキテクチャを使用します：\n\n"
            "1. ISO 13485監査チェックリスト（71条項）から日付シードによる輪替での質問抽出 — 各条項最低7バージョン\n"
            "2. 分析者（弁護側）が各QMS文書を分析し出力：立場、信頼度、主要証拠\n"
            "3. 検証者（質疑側）が分析者の論点を条項ごとに質疑し、未引用の規制要件または証拠の欠落を指摘\n"
            "4. 最大3ラウンドのディベートを経てagree / partial / disagreeの結論に達する\n"
            "5. QA監査員が議論全体を独立採点（0〜100）\n"
            "6. 最終判定はリスクマトリクスにより議論結論とgap_severityから自動算出\n\n"
            "略語：verdict — compliant / improvement_plan / deadline_correction / immediate_correction\n"
            "      gap_severity — none / minor / major / critical\n"
            "      flagged_for_ra — RA規制担当者による追加審査が必要"
        ),
        "audit_question": "監査質問",
        "expected_evidence": "期待される書面証拠（静的チェックリスト）",
        "source_b_label": "🤖 [AI生成質問 (Side B)]",
        "source_a_label": "📋 [静的質問プール (Side A)]",
        "focus_area_label": "焦点領域",
        "verifiable_by_label": "検証方法",
        "analyzer_label": "🔍 分析者（実際に確認した証拠）",
        "verifier_label": "⚖️ 検証者（期待される証拠）",
        "qa_audit_heading": "🔎 第三者監査",
        "score_label": "スコア",
        "question_quality_label": "質問品質",
        "answer_accuracy_label": "回答精度",
        "hallucination_label": "ハルシネーション検出",
        "hallucination_yes": "⚠️ あり",
        "hallucination_no": "なし",
        "ra_flag_yes": "⚠️ あり",
        "ra_flag_no": "なし",
        "agreed_yes": "✅ はい",
        "agreed_no": "❌ いいえ",
        "meta_label": "記録ID: {rid}  |  実行ID: {aid}  |  時刻: {ts}",
        "summary_body": "規制: {regs}\n国: {countries}\n条項数: {clauses}  |  同意: {agreed}  |  RAフラグ: {flagged}  |  総ラウンド: {rounds}\nモデル: {model}  |  所要時間: {duration:.1f}s",
        # Deep report section headings
        "deep_title": "AI-QMS 詳細コンプライアンス分析レポート",
        "deep_meta": "実行ID: {run_id}  |  エクスポート時刻: {ts}",
        "deep_s1": "第1章 エグゼクティブサマリー",
        "deep_s1_intro": "このコンプライアンス分析では {total} 件の条項-文書ペアを評価しました。\n\n判定結果の分布:\n",
        "deep_s1_risk": "\nリスクレベルの分布:\n",
        "deep_s1_ra": "\n⚠️ RA審査が必要: {flagged} 件\n",
        "deep_s2": "第2章 GAP分析詳細 (Phase 1)",
        "deep_s2_doc": "文書",
        "deep_s2_found": "確認済み",
        "deep_s2_not_found": "未確認",
        "deep_s2_inadequate": "不十分",
        "deep_llm_response": "LLM回答",
        "deep_token_usage": "トークン使用量: {tokens:,}  |  モデル: {model}",
        "deep_no_p1": "（この分析ではPhase 1のLLMインタラクション記録はありません）",
        "deep_no_record": "（利用可能なLLMインタラクション記録はありません）",
        "deep_s3": "第3章 検証詳細 (Phase 2)",
        "deep_no_p2": "（この分析ではPhase 2のLLMインタラクション記録はありません）",
        "deep_s4": "第4章 改善提案 (Phase 4)",
        "deep_no_p4": "（この分析ではPhase 4のLLMインタラクション記録はありません）",
        "deep_s5": "第5章 相互尋問 (Phase 5)",
        "deep_no_p5": "（この分析ではPhase 5のLLMインタラクション記録はありません）",
        "deep_no_xexam": "（利用可能な相互尋問記録はありません）",
        "deep_s55": "第5章b 第三者QA監査 (Phase 5 Step 2)",
        "deep_qa_score": "全体品質スコア: {score}/100\n監査条項数: {count}\nモデル: {model}",
        "deep_qa_summary": "監査サマリー",
        "deep_qa_recs": "監査推奨事項",
        "deep_qa_clause_results": "条項別監査結果",
        "deep_qa_tbl_headers": ["条項", "スコア", "質問品質", "回答精度", "ハルシネーション", "問題点"],
        "deep_doc_label": "文書",
        "deep_analyzer_label": "🔍 分析者",
        "deep_verifier_label": "⚖️ 検証者",
        "deep_s6": "第6章 コンプライアンス分析結果表",
        "deep_s6_headers": ["条項", "文書", "監査影響", "判定", "リスク", "ギャップ", "RAフラグ", "パイプライン状態"],
        "deep_s7": "第7章 相互尋問品質分析",
        "deep_s7_findings": "発見事項",
        "deep_s7_recs": "推奨事項",
        "deep_s7_prompt": "プロンプト調整記録",
        "deep_s7_no_summary": "（分析サマリーなし）",
        "deep_s8": "付録 LLM使用統計",
        "deep_s8_calls": "回の呼び出し",
        "deep_no_issues": "なし",
        "deep_qa_skipped": "（スキップ：{summary}）",
        "deep_no_qa_record": "（第三者QA監査記録なし）",
        "xl_sheet_summary": "サマリー",
        "xl_sheet_compliance": "コンプライアンス",
        "xl_sheet_llm": "LLMインタラクション",
        "xl_sheet_crossexam": "相互尋問",
        "xl_sheet_meta": "品質分析",
        "xl_run_id": "実行ID",
        "xl_export_time": "エクスポート時刻",
        "xl_total_rows": "総分析項目",
        "xl_flagged_ra": "RA要確認",
        "xl_verdict_prefix": "判定 - ",
        "xl_risk_prefix": "リスク - ",
        "xl_comp_headers": ["条項ID","条項名","文書ID","文書タイトル","監査影響","監査質問","判定","リスクレベル","ギャップ重大度","証拠（確認/合計）","RAフラグ","RAオーバーライド","RAメモ","改善提案 (P4)","規制引用 (P4)","分析者立場 (P5)","検証者評価 (P5)","パイプライン状態"],
        "deep_s0": "フェーズ0：パイプライン実行進捗",
        "deep_s0_phase_col_phase": "フェーズ",
        "deep_s0_phase_col_name_zh": "名前（中国語）",
        "deep_s0_phase_col_name_en": "名前（英語）",
        "deep_s0_phase_col_rows": "行数",
        "deep_s0_phase_col_status": "状態",
        "deep_s0_skipped": "スキップされたフェーズ",
        "deep_s0_budget_title": "LLMトークン使用統計",
        "deep_s0_budget_total": "使用トークン数",
        "deep_s0_budget_calls": "LLM呼び出し回数",
        "deep_s0_budget_pct": "予算使用率",
        "deep_s0_budget_remaining": "残りトークン",
        "deep_s0_dq_title": "データ品質チェック結果 (P0)",
        "deep_s0_sc_title": "ソース検証結果 (P6)",
        "deep_s0_none": "（データなし）",
        "xl_llm_headers": ["Phase","文書ID","条項ID","役割","Round","LLM回答（抜粋）","トークン使用量","モデル","タイムスタンプ"],
        "xl_xe_headers": ["条項ID","条項名","文書ID","判定","同意","RAフラグ","ラウンド数","R1分析者立場","R1分析者信頼度","R1検証者評価","R1 Agreement","QAスコア","質問品質","回答精度","ハルシネーション"],
        # Section 0.5
        "deep_s05_heading": "第0.5章 リスク優先事項サマリー",
        "deep_s05_body": "以下の {n} 件の条項-文書ペアは高リスク／重大な不適合であり、優先的に対処が必要です。",
        "deep_s05_headers": ["条項 ID", "条項名", "文書 ID", "判定", "リスクレベル", "ギャップ重大度", "是正提案 (概要)"],
        # Crawl status appendix
        "crawl_appendix_heading": "付録：規制データソースクロール状態",
        "crawl_appendix_headers": ["状態", "機関", "地域", "URL", "クロール日時", "HTTPステータス", "エラー"],
        "crawl_appendix_selected": "今回クロールした地域",
        "crawl_appendix_all_ts": "データ最終更新",
        "crawl_xl_sheet": "クロール状態",
        "xl_phase_labels": {
            "gap_scan": "P1 差距スキャン",
            "checklist_verify": "P2 チェックリスト検証",
            "remediation": "P4 改善提案",
            "verification": "P5 相互尋問",
        },
        "xl_role_labels": {"analyzer": "分析者", "verifier": "検証者"},
        "xl_batch_label": "—（バッチ）",
        "xl_xe_ext_headers": ["条項ID","条項名","文書ID","判定","同意","RAフラグ","ラウンド数","R1分析者立場","R1分析者信頼度","R1主要証拠","R1検証者指摘","R1 Agreement","QAスコア","質問品質","回答精度","ハルシネーション"],
    },
}


# ============================================================
# Shared: Word visual helpers
# ============================================================


def _shade_cell(cell, fill_hex: str) -> None:
    """Apply background fill to a Word table cell (tcPr level, not pPr)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _bold_cell(cell, text: str, fill_hex: str | None = None, font_color: str | None = None,
               font_size: int = 9) -> None:
    """Set cell text bold, optionally with background fill and font color."""
    from docx.shared import Pt, RGBColor
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)
    if font_color:
        r, g, b = int(font_color[0:2], 16), int(font_color[2:4], 16), int(font_color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)
    if fill_hex:
        _shade_cell(cell, fill_hex)


# ============================================================
# Shared: AI Roles Legend
# ============================================================


def _add_ai_roles_legend(doc, lang: str = "zh-TW") -> None:
    """Insert the AI roles and scoring legend into any Word document."""
    _lk = _lang_key(lang)

    if _lk == "en":
        doc.add_heading("AI Role Definitions & Scoring", level=2)
        doc.add_paragraph(
            "This system uses a three-role debate architecture for QMS compliance review."
        )
        p = doc.add_paragraph()
        p.add_run("🔍 Analyzer / Defender").bold = True
        doc.add_paragraph(
            "  Role: Analyzes QMS documents against each regulatory clause, takes a clear position\n"
            "        (compliant / non-compliant) and defends it with evidence.\n"
            "  Output fields:\n"
            "    • position: compliant / non-compliant / partially_compliant\n"
            "    • confidence: high / medium / low\n"
            "    • key_evidence: quoted passages supporting the position\n"
            "    • regulatory_references: cited clause numbers and text\n"
            "  Scoring: QA Auditor evaluates citation accuracy → Dim A (0–100)\n"
            "    90–100 Precise | 70–89 Minor gaps | 50–69 Partial | 30–49 Superficial | 0–29 Inadequate"
        )
        p = doc.add_paragraph()
        p.add_run("⚖️ Verifier / Reviewer").bold = True
        doc.add_paragraph(
            "  Role: Challenges the Analyzer's argument as devil's advocate, identifies uncited\n"
            "        regulatory requirements, contradictions, or evidence gaps.\n"
            "  Output fields:\n"
            "    • agreement_level: agree / partial / disagree\n"
            "    • challenges: specific weaknesses or gaps in the Analyzer's argument\n"
            "    • overall_assessment: summary evaluation of the debate quality\n"
            "    • remaining_concerns: unresolved issues after debate\n"
            "  Scoring: QA Auditor evaluates challenge quality and depth → Dim B (0–100)\n"
            "    90–100 Deep & balanced | 70–89 Minor gaps | 50–69 Notable gaps | 30–49 Superficial | 0–29 Inadequate"
        )
        p = doc.add_paragraph()
        p.add_run("🔎 QA Auditor").bold = True
        doc.add_paragraph(
            "  Role: Independent third-party auditor — does not debate, only evaluates debate quality.\n"
            "  Output fields:\n"
            "    • overall_score: 0–100 composite debate quality score\n"
            "    • score_rationale: explanation of the score band and specific reasons\n"
            "    • question_quality: good / acceptable / poor\n"
            "    • answer_accuracy: accurate / partially_accurate / inaccurate\n"
            "    • logic_consistency: consistent / minor_issues / inconsistent\n"
            "    • hallucination_detected: true / false — whether AI cited non-existent regulations\n"
            "    • issues: list of specific quality problems found in the debate\n"
            "  Review steps:\n"
            "    1. Assess whether audit questions target the core requirement (question_quality)\n"
            "    2. Verify that cited regulations exist and are accurately quoted (answer_accuracy)\n"
            "    3. Evaluate whether Verifier challenges are grounded and constructive (Dim B)\n"
            "    4. Check that both reasoning chains are complete and internally consistent (logic_consistency)\n"
            "    5. Detect fabricated clauses or erroneous citations (hallucination_detected)"
        )

    elif _lk == "ja":
        doc.add_heading("AIロール定義とスコアリング", level=2)
        doc.add_paragraph(
            "本システムはQMSコンプライアンス審査のために三役割討論アーキテクチャを採用しています。"
        )
        p = doc.add_paragraph()
        p.add_run("🔍 分析者 / 弁護方（Analyzer / Defender）").bold = True
        doc.add_paragraph(
            "  役割：各規制条項に対してQMS文書を分析し、明確な立場（適合／非適合）を取り、証拠で弁護します。\n"
            "  出力フィールド：\n"
            "    • position（立場）：compliant / non-compliant / partially_compliant\n"
            "    • confidence（確信度）：high / medium / low\n"
            "    • key_evidence（主要証拠）：立場を支持する文書引用箇所\n"
            "    • regulatory_references（規制引用）：引用した条項番号と条文\n"
            "  採点方法：QA Auditor が規制引用の正確性を評価 → Dim A（0–100）\n"
            "    90–100 引用精確 | 70–89 軽微な漏れ | 50–69 部分的 | 30–49 表面的 | 0–29 不十分"
        )
        p = doc.add_paragraph()
        p.add_run("⚖️ 検証者 / 質問方（Verifier / Reviewer）").bold = True
        doc.add_paragraph(
            "  役割：悪魔の代弁者として Analyzer の論拠に異議を唱え、未引用の規制要件・矛盾・証拠の欠落を指摘します。\n"
            "  出力フィールド：\n"
            "    • agreement_level（同意度）：agree / partial / disagree\n"
            "    • challenges（質問点）：Analyzer の論拠の弱点や欠落を具体的に指摘\n"
            "    • overall_assessment（総合評価）：討論品質の文章サマリー\n"
            "    • remaining_concerns（未解決懸念）：討論後も残る争点\n"
            "  採点方法：QA Auditor が質問の質と深さを評価 → Dim B（0–100）\n"
            "    90–100 深く均衡的 | 70–89 軽微な欠落 | 50–69 顕著な欠落 | 30–49 形式的 | 0–29 不十分"
        )
        p = doc.add_paragraph()
        p.add_run("🔎 品質監査員（QA Auditor）").bold = True
        doc.add_paragraph(
            "  役割：独立した第三者監査員として討論には参加せず、討論品質のみを客観的に評価します。\n"
            "  出力フィールド：\n"
            "    • overall_score（総合スコア）：0–100、討論品質の総合評点\n"
            "    • score_rationale（採点根拠）：スコア帯と具体的理由の説明\n"
            "    • question_quality（質問品質）：good / acceptable / poor\n"
            "    • answer_accuracy（回答精度）：accurate / partially_accurate / inaccurate\n"
            "    • logic_consistency（論理一貫性）：consistent / minor_issues / inconsistent\n"
            "    • hallucination_detected（幻覚検出）：true / false — AIが存在しない規制を引用したか\n"
            "    • issues（問題リスト）：討論中に発見された品質上の問題を具体的に列挙\n"
            "  審査手順：\n"
            "    1. 監査質問が当該条項の核心要件を対象としているか評価（question_quality）\n"
            "    2. Analyzer が引用した規制条文が存在し正確かを確認（answer_accuracy）\n"
            "    3. Verifier の質問が根拠のある建設的なものかを評価（Dim B）\n"
            "    4. 双方の推論チェーンが完整で内部矛盾がないか確認（logic_consistency）\n"
            "    5. 架空条項や誤引用などの幻覚現象を検出（hallucination_detected）"
        )

    else:
        doc.add_heading("AI 角色說明 / AI Role Definitions & Scoring", level=2)
        doc.add_paragraph(
            "本系統採用三角色辩論架構進行 QMS 合規性審查。\n"
            "This system uses a three-role debate architecture for QMS compliance review."
        )
        p = doc.add_paragraph()
        p.add_run("🔍 分析者 / 辩護方（Analyzer / Defender）").bold = True
        doc.add_paragraph(
            "  角色定位：針對每個法規条款，分析 QMS 文件是否提供充分書面證據，採取明確立場（符合/不符合）並加以辩護。\n"
            "  Role: Analyzes QMS documents against each regulatory clause, takes a clear position\n"
            "        (compliant / non-compliant) and defends it with evidence.\n"
            "  輸出欄位：\n"
            "    • position（立場）：compliant / non-compliant / partially_compliant\n"
            "    • confidence（信心度）：high / medium / low\n"
            "    • key_evidence（關鍵證據）：文件中支持立場的引用段落\n"
            "    • regulatory_references（法規引用）：引用的条款編號與条文\n"
            "  評分方式：QA Auditor 評估其法規引用準確性，計入 Dim A（0–100）\n"
            "    90–100 引用精確，完全符合 | 70–89 輕微遗漏 | 50–69 部分符合 | 30–49 表面符合 | 0–29 完全不符"
        )
        p = doc.add_paragraph()
        p.add_run("⚖️ 驗證者 / 質疑方（Verifier / Reviewer）").bold = True
        doc.add_paragraph(
            "  角色定位：以魔鬼代言人角色質疑 Analyzer 的論點，指出未引用的法規要求、矛盾或證據漏洞，\n"
            "            迫使 Analyzer 進行更深入論證，最終給出同意程度結論。\n"
            "  Role: Challenges the Analyzer's argument as devil's advocate, identifies uncited\n"
            "        regulatory requirements, contradictions, or evidence gaps.\n"
            "  輸出欄位：\n"
            "    • agreement_level（同意程度）：agree / partial / disagree\n"
            "    • challenges（質疑點）：具體指出 Analyzer 論點的弱點或漏洞\n"
            "    • overall_assessment（整體評語）：對整場辩論品質的文字總結\n"
            "    • remaining_concerns（未解疑慮）：最終仍存在的爭議點\n"
            "  評分方式：QA Auditor 評估其詰問品質與深度，計入 Dim B（0–100）\n"
            "    90–100 深度均衡可操作 | 70–89 輕微缺失 | 50–69 明顯缺失 | 30–49 流於形式 | 0–29 嚴重不足"
        )
        p = doc.add_paragraph()
        p.add_run("🔎 品質稽核員 / 審查者（QA Auditor）").bold = True
        doc.add_paragraph(
            "  角色定位：模擬獨立第三方稽核員，不參與辩論，對整場 Analyzer↔Verifier 辩論進行客觀品質評核。\n"
            "  Role: Independent third-party auditor — does not debate, only evaluates debate quality.\n"
            "  輸出欄位：\n"
            "    • overall_score（整體分數）：0–100，綜合辩論品質評分\n"
            "    • score_rationale（評分依據）：說明落在哪個分數區間及具體原因\n"
            "    • question_quality（問題品質）：good / acceptable / poor — 評估稽核問題是否聰焦可查\n"
            "    • answer_accuracy（回答準確性）：accurate / partially_accurate / inaccurate\n"
            "    • logic_consistency（邏輯一致性）：consistent / minor_issues / inconsistent\n"
            "    • hallucination_detected（幻覺偉測）：true / false — AI 是否引用了不存在的法規內容\n"
            "    • issues（問題清單）：具體列出辩論中發現的品質問題\n"
            "  審核流程：\n"
            "    1. 評估稽核問題是否針對該条款的核心要求（question_quality）\n"
            "    2. 核查 Analyzer 引用的法規条文是否存在且準確（answer_accuracy）\n"
            "    3. 評估 Verifier 的質疑是否有根據且具建設性（Dim B）\n"
            "    4. 檢查雙方推理鏈是否完整、無內部矛盾（logic_consistency）\n"
            "    5. 偵測是否存在虛構条款、錯誤引用等幻覺現象（hallucination_detected）"
        )


# ============================================================
# Crawl Status Helpers (Word + Excel)
# ============================================================


def _load_crawl_results() -> Optional[dict]:
    """Load most recent crawl results from storage. Returns None if unavailable."""
    try:
        from src.storage.regulatory_storage import get_regulatory_store
        return get_regulatory_store().load_last_results()
    except Exception:
        return None


def _append_crawl_status_word(doc, crawl_results: Optional[dict], lang: str = "zh-TW") -> None:
    """Append a crawl status appendix section to a Word document."""
    from docx.shared import Pt
    lk = _lang_key(lang)
    dh = _EXPORT_HEADERS[lk]
    doc.add_heading(dh["crawl_appendix_heading"], level=2)

    if not crawl_results or not crawl_results.get("results"):
        _no_crawl_note = {
            "zh": (
                "（尚無法規爬蟲資料）\n\n"
                "說明：法規更新連線狀態資料需透過「法規清單更新」功能執行爬蟲後方可顯示。\n"
                "請在主介面選擇「法規清單更新 / Regulatory Update」並執行爬蟲，重新匯出報告後此區塊將顯示各機構連線狀態及爬蟲網址。"
            ),
            "en": (
                "(No regulatory crawl data available)\n\n"
                "Note: Crawl status data is generated by running the Regulatory Update crawler.\n"
                "Select 'Regulatory Update' from the main interface and run the crawler. "
                "Re-export the report to populate this section with agency connection status and crawl URLs."
            ),
            "ja": (
                "（規制クロールデータなし）\n\n"
                "備考：規制更新接続状態データは「規制更新 / Regulatory Update」機能でクローラを実行することで生成されます。\n"
                "メイン画面から「規制更新」を選択してクローラを実行し、レポートを再エクスポートするとこのセクションにデータが表示されます。"
            ),
        }
        p = doc.add_paragraph(_no_crawl_note.get(lk, _no_crawl_note["zh"]))
        if p.runs:
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = None  # use default color
        return

    results = crawl_results.get("results", [])

    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    overall_ts = (crawl_results.get("crawl_timestamp") or "")[:19].replace("T", " ")
    selected = crawl_results.get("selected_regions") or []
    meta_lines = []
    if overall_ts:
        meta_lines.append(f"{dh['crawl_appendix_all_ts']}: {overall_ts}")
    if selected:
        meta_lines.append(f"{dh['crawl_appendix_selected']}: {', '.join(selected)}")
    if meta_lines:
        p = doc.add_paragraph("\n".join(meta_lines))
        p.runs[0].font.size = Pt(8)

    headers = dh["crawl_appendix_headers"]
    tbl = doc.add_table(rows=1 + len(results), cols=7)
    tbl.style = "Table Grid"
    for ci, hdr in enumerate(headers):
        _bold_cell(tbl.rows[0].cells[ci], hdr, fill_hex="1F3864", font_color="FFFFFF", font_size=8)

    for ri, r in enumerate(results, 1):
        status_ok = r.get("crawl_status") == "success"
        cells = tbl.rows[ri].cells
        r_ts = (r.get("crawl_timestamp") or "")[:16].replace("T", " ")
        vals = [
            "✓" if status_ok else "✗",
            r.get("agency", ""),
            r.get("region", ""),
            r.get("url") or "",
            r_ts,
            str(r.get("http_status") or ""),
            r.get("failure_reason") or "",
        ]
        bg = "E8F5E9" if status_ok else "FFE0E0"
        for ci, val in enumerate(vals):
            cells[ci].text = val
            _shade_cell(cells[ci], bg)
            for run in cells[ci].paragraphs[0].runs:
                run.font.size = Pt(7)


def _append_crawl_status_excel(wb, crawl_results: Optional[dict], lang: str = "zh-TW") -> None:
    """Append a crawl status sheet to an openpyxl workbook."""
    lk = _lang_key(lang)
    dh = _EXPORT_HEADERS[lk]

    try:
        from openpyxl.styles import PatternFill, Font, Alignment
        ws = wb.create_sheet(title=dh["crawl_xl_sheet"][:31])

        if not crawl_results or not crawl_results.get("results"):
            _no_crawl_msgs = {
                "zh": "（尚無法規爬蟲資料）— 請執行「法規清單更新 / Regulatory Update」爬蟲後重新匯出報告。",
                "en": "(No regulatory crawl data) — Run 'Regulatory Update' crawler and re-export the report.",
                "ja": "（規制クロールデータなし）— 「規制更新 / Regulatory Update」クローラを実行してレポートを再エクスポートしてください。",
            }
            ws.append([_no_crawl_msgs.get(lk, _no_crawl_msgs["zh"])])
            ws.cell(row=1, column=1).font = Font(italic=True, color="666666", size=9)
            ws.column_dimensions["A"].width = 80
            return

        results = crawl_results.get("results", [])

        # Meta rows: timestamp and selected regions
        overall_ts = (crawl_results.get("crawl_timestamp") or "")[:19].replace("T", " ")
        selected = crawl_results.get("selected_regions") or []
        if overall_ts:
            ws.append([dh["crawl_appendix_all_ts"], overall_ts])
            ws.cell(row=ws.max_row, column=1).font = Font(bold=True, size=9)
        if selected:
            ws.append([dh["crawl_appendix_selected"], ", ".join(selected)])
            ws.cell(row=ws.max_row, column=1).font = Font(bold=True, size=9)
        if overall_ts or selected:
            ws.append([])  # blank spacer row

        headers = dh["crawl_appendix_headers"]
        ws.append(headers)
        hdr_fill = PatternFill("solid", fgColor="1F3864")
        hdr_font = Font(bold=True, color="FFFFFF", size=9)
        for cell in ws[ws.max_row]:
            cell.fill = hdr_fill
            cell.font = hdr_font

        ok_fill = PatternFill("solid", fgColor="E8F5E9")
        err_fill = PatternFill("solid", fgColor="FFE0E0")
        for r in results:
            status_ok = r.get("crawl_status") == "success"
            r_ts = (r.get("crawl_timestamp") or "")[:16].replace("T", " ")
            row = [
                "✓" if status_ok else "✗",
                r.get("agency", ""),
                r.get("region", ""),
                r.get("url") or "",
                r_ts,
                r.get("http_status") or "",
                r.get("failure_reason") or "",
            ]
            ws.append(row)
            row_fill = ok_fill if status_ok else err_fill
            for cell in ws[ws.max_row]:
                cell.fill = row_fill
                cell.font = Font(size=8)
                cell.alignment = Alignment(wrap_text=True)

        ws.column_dimensions["A"].width = 6
        ws.column_dimensions["B"].width = 16
        ws.column_dimensions["C"].width = 16
        ws.column_dimensions["D"].width = 46
        ws.column_dimensions["E"].width = 17
        ws.column_dimensions["F"].width = 12
        ws.column_dimensions["G"].width = 38
    except Exception:
        pass


# ============================================================
# Cross-Exam Record Export (individual records)
# ============================================================


def export_crossexam_record_word(record_dict: dict, lang: str = "zh-TW") -> Path:
    """Export a single cross-exam record as a Word document.

    Args:
        record_dict: CrossExamRecord.to_dict() output
        lang: UI language code (e.g., 'zh-TW', 'en', 'ja') — controls section headers

    Returns:
        Path to the generated .docx file
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    lk = _lang_key(lang)
    h = _EXPORT_HEADERS[lk]

    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    filepath = EXPORT_DIR / f"crossexam_{record_dict.get('record_id', 'unknown')}.docx"

    doc = Document()
    title = doc.add_heading(h["title_crossexam"], level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = meta.add_run(
        h["meta_label"].format(
            rid=record_dict.get("record_id", ""),
            aid=record_dict.get("run_id", ""),
            ts=record_dict.get("timestamp", ""),
        )
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # ── AI roles legend ──
    _add_ai_roles_legend(doc, lang)

    doc.add_heading(h["how_it_works_heading"], level=2)
    doc.add_paragraph(h["how_it_works_body"])

    # Summary
    doc.add_heading(h["summary"], level=2)
    regs = ", ".join(record_dict.get("selected_regulations", [])) or h["none"]
    countries = ", ".join(record_dict.get("countries", [])) or h["none"]
    doc.add_paragraph(
        h["summary_body"].format(
            regs=regs,
            countries=countries,
            clauses=record_dict.get("total_clauses", 0),
            agreed=record_dict.get("total_agreed", 0),
            flagged=record_dict.get("total_flagged", 0),
            rounds=record_dict.get("total_rounds", 0),
            model=record_dict.get("llm_model", ""),
            duration=record_dict.get("duration_seconds", 0),
        )
    )

    # Load ISO checklist and cross-exam question generator for country sections
    try:
        from src.analysis.compliance_rules import ISO_13485_CHECKLIST as _ISO_CL, generate_cross_exam_questions as _gen_cxq
    except Exception:
        _ISO_CL = {}
        _gen_cxq = None

    _selected_regs = record_dict.get("selected_regulations", [])

    # Problem E fix: Country-Specific Questions section (before clause-by-clause details)
    if _selected_regs and _gen_cxq:
        _cxq_heading = {"zh": "各國特有稽核問題彙整", "en": "Country-Specific Audit Questions Summary", "ja": "各国固有の監査質問まとめ"}
        doc.add_heading(_cxq_heading.get(lk, _cxq_heading["en"]), level=2)
        _cxq_intro = {
            "zh": "以下各節列出每個選定市場中超越 ISO 13485 標準的特有稽核問題，LLM 驗證者將在交叉詰問時重點檢查這些要求。",
            "en": "The following sections list country-specific audit questions that exceed ISO 13485 requirements for each selected market. The LLM Verifier will focus on these during cross-examination.",
            "ja": "以下の各節では、選択した各市場においてISO 13485の要件を超える固有の監査質問をリストします。LLM検証者は交差検証時にこれらに重点を置きます。",
        }
        doc.add_paragraph(_cxq_intro.get(lk, _cxq_intro["en"]))
        # Get all clause IDs from this run
        _all_clause_ids = list({c.get("clause_id", "") for c in record_dict.get("clauses", []) if c.get("clause_id")})
        # Build per-regulation country-specific questions
        _by_reg: dict[str, list] = {}
        for _cid in _all_clause_ids:
            try:
                _qs = _gen_cxq(doc_id="", doc_title="", baseline_clause=_cid, selected_regulations=_selected_regs)
            except Exception:
                continue
            for _q in _qs:
                if _q.get("question_type") not in ("delta",):
                    continue
                _rid = _q.get("regulation_id", "")
                if _rid not in _by_reg:
                    _by_reg[_rid] = []
                _by_reg[_rid].append(_q)
        # Render per-country sections
        _country_heading = {"zh": "特有要求", "en": "Specific Requirements", "ja": "固有の要件"}
        _q_label = {"zh": "問題", "en": "Question", "ja": "質問"}
        _ref_label = {"zh": "法規依據", "en": "Regulatory Basis", "ja": "法規根拠"}
        _impact_label = {"zh": "影響等級", "en": "Impact", "ja": "影響レベル"}
        for _rid, _qlist in _by_reg.items():
            if not _qlist:
                continue
            _country = _qlist[0].get("country", _rid)
            _reg_name = _qlist[0].get("regulation_name", _rid)
            doc.add_heading(f"{_country} — {_reg_name}", level=3)
            for _q in _qlist:
                _q_text = (_q.get("question_zh") if lk == "zh" else (_q.get("question_en") or _q.get("question_zh", "")))
                _title_text = (_q.get("title_zh") if lk == "zh" else (_q.get("title_en") or _q.get("title_zh", "")))
                doc.add_paragraph(f"[{_q.get('audit_impact', '').upper()}] {_title_text}")
                doc.add_paragraph(f"  {_q_label[lk]}: {_q_text}", style="Quote")
                doc.add_paragraph(f"  {_ref_label[lk]}: {_q.get('method', '')}  {_impact_label[lk]}: {_q.get('audit_impact', '')}")

    # Clause details
    doc.add_heading(h["clause_details"], level=2)
    for clause in record_dict.get("clauses", []):
        cid = clause.get("clause_id", "")
        doc.add_heading(
            f"{cid} — {clause.get('clause_title', '')}",
            level=3,
        )
        doc.add_paragraph(
            f"{h['doc']}: {clause.get('doc_id', '')} {clause.get('doc_title', '')}\n"
            f"{h['verdict']}: {clause.get('verdict', '')}  |  {h['gap']}: {clause.get('gap_severity', '')}\n"
            f"{h['agreed']}: {h['agreed_yes'] if clause.get('agreed') else h['agreed_no']}  |  "
            f"{h['flagged_ra']}: {h['ra_flag_yes'] if clause.get('flagged_for_ra') else h['ra_flag_no']}"
        )

        # Audit question source label (A/B hybrid)
        _cl_def = _ISO_CL.get(cid, {})
        _q_source = clause.get("question_source", "")
        if _q_source == "B":
            doc.add_paragraph(h["source_b_label"], style="Quote")
            _focus = clause.get("focus_area", "")
            if _focus:
                doc.add_paragraph(f"  {h['focus_area_label']}: {_focus}")
            _verifiable = clause.get("verifiable_by", "")
            if _verifiable:
                doc.add_paragraph(f"  {h['verifiable_by_label']}: {_verifiable}")
        elif _q_source == "A":
            doc.add_paragraph(h["source_a_label"], style="Quote")

        # Audit question
        _aq = clause.get("audit_question") or _cl_def.get("audit_question", "")
        if _aq:
            doc.add_paragraph(f"{h['audit_question']}: {_aq}")

        # Expected evidence (static checklist)
        _exp_ev = _cl_def.get("expected_evidence", [])
        if _exp_ev:
            doc.add_paragraph(
                f"{h['expected_evidence']}:\n"
                + "\n".join(f"  • {e}" for e in _exp_ev),
                style="Quote",
            )

        for rd in clause.get("rounds", []):
            doc.add_heading(f"Round {rd.get('round', '?')}", level=4)

            analyzer = rd.get("analyzer", {})
            p = doc.add_paragraph()
            r = p.add_run(f"{h['analyzer_label']}: ")
            r.bold = True
            _render_role_content(p, analyzer)

            verifier = rd.get("verifier", {})
            p = doc.add_paragraph()
            r = p.add_run(f"{h['verifier_label']}: ")
            r.bold = True
            _render_role_content(p, verifier)

            agreement = verifier.get("agreement_level", "")
            doc.add_paragraph(f"Agreement: {agreement}")

        qa = clause.get("qa_audit", {})
        if qa:
            doc.add_heading(h["qa_audit_heading"], level=4)
            doc.add_paragraph(
                f"{h['score_label']}: {qa.get('score', 0)}/100  |  "
                f"{h['question_quality_label']}: {qa.get('question_quality', '')}  |  "
                f"{h['answer_accuracy_label']}: {qa.get('answer_accuracy', '')}\n"
                f"{h['hallucination_label']}: "
                f"{h['hallucination_yes'] if qa.get('hallucination_detected') else h['hallucination_no']}"
            )
            issues = qa.get("issues", [])
            if issues:
                for iss in issues:
                    doc.add_paragraph(f"  • {iss}")

    try:
        _append_crawl_status_word(doc, _load_crawl_results(), lang)
    except Exception:
        pass
    try:
        from src.utils.crossref_export import append_crossref_table_word
        append_crossref_table_word(doc, lang=lang)
    except Exception:
        pass

    safe_save_binary(filepath, doc.save)
    return filepath


def export_crossexam_record_excel(record_dict: dict, lang: str = "zh-TW") -> Path:
    """Export a single cross-exam record as an Excel file.

    Args:
        record_dict: CrossExamRecord.to_dict() output
        lang: UI language code (e.g., 'zh-TW', 'en', 'ja')

    Returns:
        Path to the generated .xlsx file
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill

    # lang reserved for future localization of sheet/section names
    _lk = _lang_key(lang)

    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    filepath = EXPORT_DIR / f"crossexam_{record_dict.get('record_id', 'unknown')}.xlsx"

    wb = Workbook()

    eh = _EXPORT_HEADERS[_lk]
    _duration_label = {"zh": "耗時 (秒)", "en": "Duration (s)", "ja": "所要時間 (秒)"}[_lk]
    _agreed_count_label = {"zh": "同意數", "en": "Agreed count", "ja": "同意数"}[_lk]

    # Sheet 1: Summary
    ws_sum = wb.active
    ws_sum.title = eh["summary"]
    summary_data = [
        (eh["record_id"], record_dict.get("record_id", "")),
        (eh["analysis_id"], record_dict.get("run_id", "")),
        (eh["time"], record_dict.get("timestamp", "")),
        (eh["regulations"], ", ".join(record_dict.get("selected_regulations", []))),
        (eh["countries"], ", ".join(record_dict.get("countries", []))),
        (eh["clause_count"], record_dict.get("total_clauses", 0)),
        (_agreed_count_label, record_dict.get("total_agreed", 0)),
        (eh["flagged_ra"], record_dict.get("total_flagged", 0)),
        (eh["total_rounds"], record_dict.get("total_rounds", 0)),
        (eh["model"], record_dict.get("llm_model", "")),
        (_duration_label, record_dict.get("duration_seconds", 0)),
    ]
    header_fill = PatternFill(
        start_color="4472C4", end_color="4472C4", fill_type="solid"
    )
    header_font = Font(bold=True, color="FFFFFF")
    for ri, (label, value) in enumerate(summary_data, 1):
        c1 = ws_sum.cell(row=ri, column=1, value=label)
        c1.fill = header_fill
        c1.font = header_font
        ws_sum.cell(row=ri, column=2, value=str(value))
    ws_sum.column_dimensions["A"].width = 15
    ws_sum.column_dimensions["B"].width = 50

    _sheet2_headers: dict[str, list] = {
        "zh": [
            "條款 ID", "條款名稱", "文件 ID", "判定", "差距", "同意", "RA 標記",
            "問題來源(A/B)", "輪次數", "R1 分析者立場", "R1 分析者信心",
            "R1 實際看到(證據)", "R1 驗證者評估", "R1 Agreement",
            "R1 期望看到(證據)", "QA 分數", "問題品質", "回答準確", "幻覺偵測",
        ],
        "en": [
            "Clause ID", "Clause Name", "Doc ID", "Verdict", "Gap", "Agreed", "RA Flag",
            "Question Source(A/B)", "Rounds", "R1 Analyzer Position", "R1 Analyzer Confidence",
            "R1 Actual Evidence", "R1 Verifier Assessment", "R1 Agreement",
            "R1 Expected Evidence", "QA Score", "Question Quality", "Answer Accuracy", "Hallucination",
        ],
        "ja": [
            "条項ID", "条項名", "文書ID", "判定", "ギャップ", "同意", "RAフラグ",
            "質問ソース(A/B)", "ラウンド数", "R1 分析者立場", "R1 分析者信頼度",
            "R1 実際の証拠", "R1 検証者評価", "R1 Agreement",
            "R1 期待される証拠", "QAスコア", "質問品質", "回答精度", "ハルシネーション",
        ],
    }
    _sheet2_title: dict[str, str] = {
        "zh": "條款詳情", "en": "Clause Details", "ja": "条項詳細"
    }
    ws_detail = wb.create_sheet(_sheet2_title[_lk])
    headers = _sheet2_headers[_lk]
    for ci, h in enumerate(headers, 1):
        c = ws_detail.cell(row=1, column=ci, value=h)
        c.fill = header_fill
        c.font = header_font
        c.alignment = Alignment(horizontal="center")

    for ri, clause in enumerate(record_dict.get("clauses", []), 2):
        ws_detail.cell(row=ri, column=1, value=clause.get("clause_id", ""))
        ws_detail.cell(row=ri, column=2, value=clause.get("clause_title", ""))
        ws_detail.cell(row=ri, column=3, value=clause.get("doc_id", ""))
        ws_detail.cell(row=ri, column=4, value=clause.get("verdict", ""))
        ws_detail.cell(row=ri, column=5, value=clause.get("gap_severity", ""))
        ws_detail.cell(row=ri, column=6, value="Y" if clause.get("agreed") else "N")
        ws_detail.cell(
            row=ri, column=7, value="Y" if clause.get("flagged_for_ra") else ""
        )
        ws_detail.cell(row=ri, column=8, value=clause.get("question_source", ""))
        rounds = clause.get("rounds", [])
        ws_detail.cell(row=ri, column=9, value=len(rounds))
        if rounds:
            r1 = rounds[0]
            a = r1.get("analyzer", {})
            v = r1.get("verifier", {})
            ws_detail.cell(
                row=ri,
                column=10,
                value=_flatten_role_text(a, "position")[:500],
            )
            ws_detail.cell(
                row=ri,
                column=11,
                value=str(a.get("confidence", a.get("confidence_score", ""))),
            )
            # R1 實際看到：Analyzer key_evidence
            _actual_ev = a.get("key_evidence", [])
            if isinstance(_actual_ev, list):
                _actual_ev = "; ".join(str(x) for x in _actual_ev)
            ws_detail.cell(row=ri, column=12, value=str(_actual_ev)[:500])
            ws_detail.cell(
                row=ri,
                column=13,
                value=_flatten_role_text(v, "assessment")[:500],
            )
            ws_detail.cell(
                row=ri,
                column=14,
                value=v.get("agreement_level", ""),
            )
            # R1 期望看到：Verifier challenges[*].expected_evidence
            _challenges = v.get("challenges", [])
            _exp_parts = []
            for _ch in (_challenges if isinstance(_challenges, list) else []):
                if isinstance(_ch, dict) and _ch.get("expected_evidence"):
                    _exp_parts.append(_ch["expected_evidence"])
            ws_detail.cell(row=ri, column=15, value="; ".join(_exp_parts)[:500])
        qa = clause.get("qa_audit", {})
        ws_detail.cell(row=ri, column=16, value=qa.get("score", "") if qa else "")
        ws_detail.cell(
            row=ri, column=17, value=qa.get("question_quality", "") if qa else ""
        )
        ws_detail.cell(
            row=ri, column=18, value=qa.get("answer_accuracy", "") if qa else ""
        )
        ws_detail.cell(
            row=ri,
            column=19,
            value="Yes"
            if qa and qa.get("hallucination_detected")
            else ("No" if qa else ""),
        )

    # Auto-width
    for col in ws_detail.columns:
        max_len = max((len(str(c.value or "")) for c in col), default=8)
        ws_detail.column_dimensions[col[0].column_letter].width = min(max_len + 2, 60)

    # Problem E fix: Country-Specific Questions sheet in Excel
    _ex_sel_regs = record_dict.get("selected_regulations", [])
    try:
        from src.analysis.compliance_rules import generate_cross_exam_questions as _excxq
        if _ex_sel_regs and _excxq:
            _cxq_sheet_title = {"zh": "各國特殊問題", "en": "Country Questions", "ja": "各国固有質問"}
            ws_cxq = wb.create_sheet(_cxq_sheet_title.get(_lk, "Country Questions"))
            _cxq_hdrs = {
                "zh": ["條款ID", "國家", "Profile ID", "問題標題(zh)", "問題標題(en)", "影響等級", "問題(zh)", "問題(en)", "法規依據", "Delta類型"],
                "en": ["Clause ID", "Country", "Profile ID", "Title (zh)", "Title (en)", "Impact", "Question (zh)", "Question (en)", "Regulation Ref", "Delta Type"],
                "ja": ["条項ID", "国", "Profile ID", "タイトル(zh)", "タイトル(en)", "影響", "質問(zh)", "質問(en)", "法規根拠", "Deltaタイプ"],
            }
            for ci, hdr in enumerate(_cxq_hdrs.get(_lk, _cxq_hdrs["en"]), 1):
                c = ws_cxq.cell(row=1, column=ci, value=hdr)
                c.fill = header_fill
                c.font = header_font
            _all_cids = list({cl.get("clause_id", "") for cl in record_dict.get("clauses", []) if cl.get("clause_id")})
            _ri = 2
            for _cid in sorted(_all_cids):
                try:
                    _qs = _excxq(doc_id="", doc_title="", baseline_clause=_cid, selected_regulations=_ex_sel_regs)
                except Exception:
                    continue
                for _q in _qs:
                    if _q.get("question_type") not in ("delta",):
                        continue
                    ws_cxq.cell(row=_ri, column=1, value=_cid)
                    ws_cxq.cell(row=_ri, column=2, value=_q.get("country", ""))
                    ws_cxq.cell(row=_ri, column=3, value=_q.get("regulation_id", ""))
                    ws_cxq.cell(row=_ri, column=4, value=_q.get("title_zh", ""))
                    ws_cxq.cell(row=_ri, column=5, value=_q.get("title_en", ""))
                    ws_cxq.cell(row=_ri, column=6, value=_q.get("audit_impact", ""))
                    ws_cxq.cell(row=_ri, column=7, value=_q.get("question_zh", ""))
                    ws_cxq.cell(row=_ri, column=8, value=_q.get("question_en", ""))
                    ws_cxq.cell(row=_ri, column=9, value=_q.get("method", ""))
                    ws_cxq.cell(row=_ri, column=10, value=_q.get("regulation_name", ""))
                    _ri += 1
            for col in ws_cxq.columns:
                max_len = max((len(str(c.value or "")) for c in col), default=8)
                ws_cxq.column_dimensions[col[0].column_letter].width = min(max_len + 2, 60)
    except Exception:
        pass

    try:
        _append_crawl_status_excel(wb, _load_crawl_results(), lang)
    except Exception:
        pass
    try:
        from src.utils.crossref_export import append_crossref_table_excel
        append_crossref_table_excel(wb, lang=lang)
    except Exception:
        pass

    safe_save_binary(filepath, wb.save)
    return filepath


# ============================================================
# Deep Analysis Report Export
# ============================================================


def export_deep_report_word(
    run_id: str,
    flat_rows: list[dict],
    summary: dict,
    interactions: list[dict] | None = None,
    crossexam_record: dict | None = None,
    meta_analysis: dict | None = None,
    qa_audit_summary: dict | None = None,
    lang: str = "zh-TW",
    progress: dict | None = None,
    data_quality: dict | None = None,
    source_check: dict | None = None,
    skipped_phases: list[str] | None = None,
    output_path: "Path | None" = None,
) -> Path:
    """Export a deep analysis report as Word document.

    Includes ALL LLM interactions, GAP analysis, verification, remediation.

    Args:
        run_id: Pipeline run ID
        flat_rows: ComparisonTable.to_flat_rows() output
        summary: ComparisonTable.summary() output
        interactions: InteractionLog interactions list
        crossexam_record: CrossExamRecord.to_dict() (optional)
        meta_analysis: Meta-analysis QA results (optional)
        qa_audit_summary: QA-audit summary (optional)
        lang: UI language code (e.g., 'zh-TW', 'en', 'ja')

    Returns:
        Path to generated .docx
    """
    _lk = _lang_key(lang)
    dh = _EXPORT_HEADERS[_lk]
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    filepath = output_path if output_path is not None else (EXPORT_DIR / f"deep_report_{run_id}.docx")

    doc = Document()

    # ── Title ──
    title = doc.add_heading(dh["deep_title"], level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = meta_p.add_run(
        dh["deep_meta"].format(
            run_id=run_id,
            ts=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        )
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # ── AI Roles Legend ──
    _add_ai_roles_legend(doc, lang)

    # ── Section 0: Pipeline Execution Progress ──
    doc.add_heading(dh["deep_s0"], level=2)
    # Phase progress table
    _prog = progress or {}
    _budget = _prog.get("llm_budget", {})
    _phase_dist = _prog.get("phase_distribution", {})
    _skipped = skipped_phases or []
    phase_tbl = doc.add_table(rows=1 + len(_PHASE_ORDER), cols=4)
    phase_tbl.style = "Table Grid"
    _ph_headers = [dh["deep_s0_phase_col_phase"], dh["deep_s0_phase_col_name_zh"], dh["deep_s0_phase_col_name_en"], dh["deep_s0_phase_col_rows"]]
    for ci, hdr in enumerate(_ph_headers):
        phase_tbl.rows[0].cells[ci].text = hdr
    for ri, ph in enumerate(_PHASE_ORDER, 1):
        ph_label = _PHASE_LABELS[ph]
        is_skipped = ph in _skipped
        row_count = _phase_dist.get(ph, 0)
        phase_tbl.rows[ri].cells[0].text = ph_label + (" (↷)" if is_skipped else "")
        phase_tbl.rows[ri].cells[1].text = _PHASE_NAMES_ZH.get(ph, ph)
        phase_tbl.rows[ri].cells[2].text = _PHASE_NAMES_EN.get(ph, ph)
        phase_tbl.rows[ri].cells[3].text = str(row_count) if not is_skipped else "↷ 跳過"
    # LLM budget
    doc.add_heading(dh["deep_s0_budget_title"], level=3)
    doc.add_paragraph(
        f"{dh['deep_s0_budget_total']}: {_budget.get('total_tokens_used', 0):,}\n"
        f"{dh['deep_s0_budget_calls']}: {_budget.get('calls_made', 0)}\n"
        f"{dh['deep_s0_budget_pct']}: {_budget.get('usage_percent', 0)}%\n"
        f"{dh['deep_s0_budget_remaining']}: {_budget.get('remaining', 0):,}"
    )
    # Skipped phases summary
    if _skipped:
        doc.add_heading(dh["deep_s0_skipped"], level=3)
        for sp in _skipped:
            doc.add_paragraph(f"  • {_PHASE_LABELS.get(sp, sp)} — {_PHASE_NAMES_ZH.get(sp, sp)} / {_PHASE_NAMES_EN.get(sp, sp)}")
    # ── Execution Environment ───────────────────────────────────
    _run_meta = _prog.get("run_metadata", {})
    if _run_meta:
        _env_head = {"zh": "執行環境", "en": "Execution Environment", "ja": "実行環境"}
        doc.add_heading(_env_head.get(_lk, "Execution Environment"), level=3)
        _env_rows = _build_env_table_rows(_run_meta, _lk)
        _eh_labels = {"zh": ["項目", "值"], "en": ["Item", "Value"], "ja": ["項目", "値"]}
        _eh = _eh_labels.get(_lk, ["Item", "Value"])
        env_tbl = doc.add_table(rows=1 + len(_env_rows), cols=2)
        env_tbl.style = "Table Grid"
        for ci, h in enumerate(_eh):
            cell = env_tbl.rows[0].cells[ci]
            cell.text = h
            runs = cell.paragraphs[0].runs
            (runs[0] if runs else cell.paragraphs[0].add_run(h)).bold = True
        for ri, (label, value) in enumerate(_env_rows, 1):
            env_tbl.rows[ri].cells[0].text = label
            env_tbl.rows[ri].cells[1].text = value
    # P0 Data Quality
    if data_quality:
        doc.add_heading(dh["deep_s0_dq_title"], level=3)
        # Scalar summary stats
        for k, v in data_quality.items():
            if k == "documents_checked":
                continue
            if isinstance(v, list):
                doc.add_paragraph(f"  {k}: {len(v)} 項")
            elif isinstance(v, dict):
                doc.add_paragraph(f"  {k}:")
                for kk, vv in v.items():
                    doc.add_paragraph(f"    • {kk}: {vv}")
            else:
                doc.add_paragraph(f"  {k}: {v}")
        # documents_checked as Word table
        docs = data_quality.get("documents_checked")
        if docs and isinstance(docs, dict):
            doc.add_paragraph(f"  檢核文件清單（共 {len(docs)} 份）：")
            dq_tbl = doc.add_table(rows=1 + len(docs), cols=5)
            dq_tbl.style = "Table Grid"
            dq_hdr = dq_tbl.rows[0].cells
            for ci, h in enumerate(["文件 ID", "存在", "有內容", "字元數", "廢止"]):
                dq_hdr[ci].text = h
                r = dq_hdr[ci].paragraphs[0].runs
                (r[0] if r else dq_hdr[ci].paragraphs[0].add_run(h)).bold = True
            for di, (doc_id, info) in enumerate(docs.items(), 1):
                if isinstance(info, dict):
                    dq_tbl.rows[di].cells[0].text = doc_id
                    dq_tbl.rows[di].cells[1].text = "✓" if info.get("exists") else "✗"
                    dq_tbl.rows[di].cells[2].text = "✓" if info.get("has_content") else "✗"
                    dq_tbl.rows[di].cells[3].text = f"{info.get('content_length', 0):,}"
                    dq_tbl.rows[di].cells[4].text = "是" if info.get("is_obsolete") else "否"
                else:
                    dq_tbl.rows[di].cells[0].text = doc_id
                    dq_tbl.rows[di].cells[1].text = str(info)
            doc.add_paragraph("")
    # P6 Source Check
    if source_check:
        doc.add_heading(dh["deep_s0_sc_title"], level=3)
        for k, v in source_check.items():
            if k == "verification_results" and isinstance(v, list):
                # Summary line
                total_u = len(v)
                acc = sum(1 for i in v if isinstance(i, dict) and i.get("accessible"))
                broken = total_u - acc
                doc.add_paragraph(f"  共 {total_u} 筆 URL：可存取 {acc}，無法存取 {broken}")
                # Word table: 狀態 | 機構 | 地區 | URL | 變更 | 錯誤
                tbl = doc.add_table(rows=1, cols=6)
                tbl.style = "Table Grid"
                hdr = tbl.rows[0].cells
                for ci, h in enumerate(["狀態", "機構", "地區", "URL", "內容變更", "錯誤訊息"]):
                    hdr[ci].text = h
                    run = hdr[ci].paragraphs[0].runs[0]
                    run.bold = True
                from docx.shared import Pt
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                for item in v:
                    if not isinstance(item, dict):
                        continue
                    row_cells = tbl.add_row().cells
                    row_cells[0].text = "✓" if item.get("accessible") else "✗"
                    row_cells[1].text = item.get("agency", "")
                    row_cells[2].text = item.get("region", "")
                    url_val = item.get("url", "")
                    row_cells[3].text = url_val[:80] + ("…" if len(url_val) > 80 else "")
                    row_cells[4].text = "是" if item.get("content_changed") else ("—" if item.get("accessible") else "")
                    row_cells[5].text = (item.get("error") or "")[:80]
                doc.add_paragraph("")  # spacer
            elif isinstance(v, list):
                doc.add_paragraph(f"  {k}: {len(v)} 項")
            elif isinstance(v, dict):
                doc.add_paragraph(f"  {k}:")
                for kk, vv in v.items():
                    doc.add_paragraph(f"    • {kk}: {vv}")
            else:
                doc.add_paragraph(f"  {k}: {v}")

    # ── Section 0.5: Risk Priority Summary Table ──
    _risk_rows = [r for r in flat_rows if r.get("risk_level", "") in ("immediate_correction", "deadline_correction")
                  or r.get("verdict", "") in ("non_compliance", "partial_compliance")]
    if _risk_rows:
        doc.add_heading(dh["deep_s05_heading"], level=2)
        doc.add_paragraph(dh["deep_s05_body"].format(n=len(_risk_rows)))
        _risk_sort_order = {"immediate_correction": 0, "deadline_correction": 1, "improvement_plan": 2, "compliant": 9}
        _risk_rows_sorted = sorted(_risk_rows, key=lambda r: (_risk_sort_order.get(r.get("risk_level",""), 9), r.get("clause_id","")))
        risk_tbl = doc.add_table(rows=1 + len(_risk_rows_sorted), cols=7)
        risk_tbl.style = "Table Grid"
        _risk_hdr_labels = dh["deep_s05_headers"]
        for ci, h in enumerate(_risk_hdr_labels):
            cell = risk_tbl.rows[0].cells[ci]
            _bold_cell(cell, h, fill_hex="1F3864", font_color="FFFFFF", font_size=9)
        for ti, r in enumerate(_risk_rows_sorted, 1):
            risk_level = r.get("risk_level", "").lower()
            verdict = r.get("verdict", "")
            row_bg = "FFE0E0" if risk_level == "immediate_correction" else ("FFF2CC" if risk_level == "deadline_correction" else "E8F5E9")
            cells = risk_tbl.rows[ti].cells
            ri_icon = "🔴" if risk_level == "immediate_correction" else ("🟡" if risk_level == "deadline_correction" else "🟢")
            vals = [
                r.get("clause_id", ""),
                r.get("clause_title", "")[:60],
                r.get("doc_id", ""),
                r.get("verdict_label", verdict),
                f"{ri_icon} {r.get('risk_label', risk_level)}",
                r.get("gap_severity", "") or "—",
                "" if "phase_4" in _skipped else (r.get("remediation_suggestion") or "")[:150],
            ]
            for ci, val in enumerate(vals):
                cells[ci].text = val
                _shade_cell(cells[ci], row_bg)
                for run in cells[ci].paragraphs[0].runs:
                    run.font.size = Pt(8)
                    if risk_level == "immediate_correction":
                        run.bold = True
        doc.add_paragraph("")

    # ── Section 1: Executive Summary ──
    doc.add_heading(dh["deep_s1"], level=2)
    verdict_dist = summary.get("verdict_distribution", {})
    risk_dist = summary.get("risk_distribution", {})
    total = summary.get("total_rows", len(flat_rows))
    flagged = summary.get("flagged_for_ra", 0)

    from src.analysis.risk_matrix import VERDICT_DISPLAY, RISK_LEVEL_DISPLAY
    _label_key = "label_en" if _lk == "en" else "label_ja" if _lk == "ja" else "label_zh"
    summary_text = dh["deep_s1_intro"].format(total=total)
    for v, count in verdict_dist.items():
        _vlabel = VERDICT_DISPLAY.get(v, {}).get(_label_key, v)
        summary_text += f"  • {_vlabel}: {count}\n"
    summary_text += dh["deep_s1_risk"]
    for r, count in risk_dist.items():
        _rlabel = RISK_LEVEL_DISPLAY.get(r, {}).get(_label_key, r)
        summary_text += f"  • {_rlabel}: {count}\n"
    if flagged:
        summary_text += dh["deep_s1_ra"].format(flagged=flagged)

    doc.add_paragraph(summary_text)

    # ── Section 2: GAP Analysis Details ──
    doc.add_heading(dh["deep_s2"], level=2)
    if interactions:
        gap_interactions = [i for i in interactions if i.get("phase") == "gap_scan"]
        if gap_interactions:
            for gi in gap_interactions:
                doc.add_heading(
                    f"{dh['deep_s2_doc']}: {gi.get('doc_id', '')} — {gi.get('doc_title', '')}",
                    level=3,
                )
                extra = gi.get("extra", {})
                ev_sum = extra.get("evidence_summary", {})
                if ev_sum:
                    doc.add_paragraph(
                        f"{dh['deep_s2_found']}: {ev_sum.get('found', 0)}  |  "
                        f"{dh['deep_s2_not_found']}: {ev_sum.get('not_found', 0)}  |  "
                        f"{dh['deep_s2_inadequate']}: {ev_sum.get('inadequate', 0)}"
                    )
                resp = gi.get("llm_response", "")
                if resp:
                    doc.add_heading(dh["deep_llm_response"], level=4)
                    _render_gap_json_response(doc, resp, lang=lang)
                usage = gi.get("usage", {})
                if usage:
                    doc.add_paragraph(
                        dh["deep_token_usage"].format(
                            tokens=usage.get("total_tokens", 0),
                            model=gi.get("model", ""),
                        )
                    )
        else:
            doc.add_paragraph(dh["deep_no_p1"])
    else:
        doc.add_paragraph(dh["deep_no_record"])

    # ── Section 3: Verification Details ──
    doc.add_heading(dh["deep_s3"], level=2)
    if interactions:
        verify_interactions = [
            i for i in interactions if i.get("phase") == "checklist_verify"
        ]
        if verify_interactions:
            for vi in verify_interactions:
                doc.add_heading(
                    f"{dh['deep_s2_doc']}: {vi.get('doc_id', '')} — {vi.get('doc_title', '')}",
                    level=3,
                )
                resp = vi.get("llm_response", "")
                if resp:
                    _render_gap_json_response(doc, resp, lang=lang)
                usage = vi.get("usage", {})
                if usage:
                    doc.add_paragraph(
                        dh["deep_token_usage"].format(
                            tokens=usage.get("total_tokens", 0),
                            model=vi.get("model", ""),
                        )
                    )
        else:
            doc.add_paragraph(dh["deep_no_p2"])
    else:
        doc.add_paragraph(dh["deep_no_record"])

    # ── Section 4: Remediation ──
    doc.add_heading(dh["deep_s4"], level=2)
    if interactions:
        remed_interactions = [
            i for i in interactions if i.get("phase") == "remediation"
        ]
        if remed_interactions:
            for ri in remed_interactions:
                doc.add_heading(
                    f"{dh['deep_s2_doc']}: {ri.get('doc_id', '')} — {ri.get('doc_title', '')}",
                    level=3,
                )
                resp = ri.get("llm_response", "")
                if resp:
                    _render_gap_json_response(doc, resp, lang=lang)
        else:
            doc.add_paragraph(dh["deep_no_p4"])
    else:
        doc.add_paragraph(dh["deep_no_record"])

    # ── Section 5: Cross-Examination ──
    doc.add_heading(dh["deep_s5"], level=2)
    if interactions:
        xexam_interactions = [
            i for i in interactions if i.get("phase") == "verification"
        ]
        if xexam_interactions:
            clause_groups: dict[str, list[dict]] = {}
            for xi in xexam_interactions:
                cid = xi.get("clause_id", "unknown")
                clause_groups.setdefault(cid, []).append(xi)

            for cid, group in clause_groups.items():
                clause_title = group[0].get("clause_title", "")
                doc.add_heading(f"{cid} — {clause_title}", level=3)
                doc_id = group[0].get("doc_id", "")
                if doc_id:
                    doc.add_paragraph(f"{dh['deep_doc_label']}: {doc_id}")

                group.sort(key=lambda x: (x.get("round_number") or 0, x.get("role") or ""))

                for xi in group:
                    # P5 llm_response is a JSON array of rounds
                    raw = xi.get("llm_response", "")
                    rounds_data = []
                    if raw:
                        try:
                            parsed_rounds = json.loads(raw)
                            if isinstance(parsed_rounds, list):
                                rounds_data = parsed_rounds
                            elif isinstance(parsed_rounds, dict):
                                rounds_data = [parsed_rounds]
                        except Exception:
                            pass
                    if not rounds_data:
                        # Fallback: treat as a single-role interaction
                        parsed = xi.get("parsed_response")
                        role = xi.get("role", "")
                        role_label = dh["deep_analyzer_label"] if role == "analyzer" else dh["deep_verifier_label"]
                        p = doc.add_paragraph()
                        p.add_run(f"{role_label}: ").bold = True
                        if parsed and isinstance(parsed, dict):
                            _render_role_content(p, parsed)
                        else:
                            p.add_run((raw or "")[:2000])
                        continue

                    for rd in rounds_data:
                        rd_num = rd.get("round", "?")
                        doc.add_heading(f"── Round {rd_num} ──", level=4)
                        # Analyzer block
                        analyzer = rd.get("analyzer") or {}
                        if analyzer:
                            _render_p5_role_block(doc, dh["deep_analyzer_label"], analyzer, "DDEEFF")
                        # Verifier block
                        verifier = rd.get("verifier") or {}
                        if verifier:
                            agreement = verifier.get("agreement_level", "")
                            _render_p5_role_block(doc, dh["deep_verifier_label"], verifier, "FCE4D6")
                            if agreement:
                                ag_p = doc.add_paragraph()
                                ag_p.add_run("Agreement: ").bold = True
                                ag_color = "C00000" if "disagree" in agreement.lower() else ("FF8000" if "partial" in agreement.lower() else "00A000")
                                from docx.shared import RGBColor as _RGB
                                r_run = ag_p.add_run(agreement)
                                r_run.bold = True
                                r_run.font.color.rgb = _RGB(*bytes.fromhex(ag_color.ljust(6, "0")))
        else:
            doc.add_paragraph(dh["deep_no_p5"])
    elif crossexam_record:
        for clause in crossexam_record.get("clauses", []):
            doc.add_heading(
                f"{clause.get('clause_id', '')} — {clause.get('clause_title', '')}",
                level=3,
            )
            for rd in clause.get("rounds", []):
                doc.add_heading(f"Round {rd.get('round', '?')}", level=4)
                p = doc.add_paragraph()
                r = p.add_run(f"{dh['deep_analyzer_label']}: ")
                r.bold = True
                _render_role_content(p, rd.get("analyzer", {}))

                p = doc.add_paragraph()
                r = p.add_run(f"{dh['deep_verifier_label']}: ")
                r.bold = True
                _render_role_content(p, rd.get("verifier", {}))

                agreement = rd.get("verifier", {}).get("agreement_level", "")
                if agreement:
                    doc.add_paragraph(f"Agreement: {agreement}")

            qa = clause.get("qa_audit", {})
            if qa:
                doc.add_heading(dh["qa_audit_heading"], level=4)
                doc.add_paragraph(
                    f"{dh['score_label']}: {qa.get('score', 0)}/100  |  "
                    f"{dh['question_quality_label']}: {qa.get('question_quality', '')}  |  "
                    f"{dh['answer_accuracy_label']}: {qa.get('answer_accuracy', '')}\n"
                    f"{dh['hallucination_label']}: "
                    f"{dh['hallucination_yes'] if qa.get('hallucination_detected') else dh['hallucination_no']}"
                )
                issues = qa.get("issues", [])
                if issues:
                    for iss in issues:
                        doc.add_paragraph(f"  • {iss}")
    else:
        doc.add_paragraph(dh["deep_no_xexam"])

    # ── Section 5.5: Third-Party QA Audit ──
    doc.add_heading(dh["deep_s55"], level=2)
    _qa_sum = qa_audit_summary
    if not _qa_sum and crossexam_record:
        _qa_sum = crossexam_record.get("qa_audit_summary")
    if _qa_sum and not _qa_sum.get("skipped"):
        score = _qa_sum.get("overall_score", 0)
        doc.add_paragraph(
            dh["deep_qa_score"].format(
                score=score,
                count=_qa_sum.get("clause_count", 0),
                model=_qa_sum.get("llm_model", ""),
            )
        )
        qa_summary_text = _qa_sum.get("summary", "")
        if qa_summary_text:
            doc.add_heading(dh["deep_qa_summary"], level=3)
            for chunk in _split_text(qa_summary_text, 3000):
                doc.add_paragraph(chunk)
        qa_recs = _qa_sum.get("recommendations", [])
        if qa_recs:
            doc.add_heading(dh["deep_qa_recs"], level=3)
            for rec in qa_recs:
                doc.add_paragraph(f"• {rec}")
        clause_audits = _qa_sum.get("clause_audits", [])
        if clause_audits:
            doc.add_heading(dh["deep_qa_clause_results"], level=3)
            qa_tbl = doc.add_table(rows=1 + len(clause_audits), cols=6)
            qa_tbl.style = "Table Grid"
            qa_headers = dh["deep_qa_tbl_headers"]
            for i, h in enumerate(qa_headers):
                qa_tbl.rows[0].cells[i].text = h
            for qi, ca in enumerate(clause_audits, 1):
                qa_tbl.rows[qi].cells[0].text = ca.get("clause_id", "")
                qa_tbl.rows[qi].cells[1].text = str(ca.get("score", 0))
                qa_tbl.rows[qi].cells[2].text = ca.get("question_quality", "")
                qa_tbl.rows[qi].cells[3].text = ca.get("answer_accuracy", "")
                qa_tbl.rows[qi].cells[4].text = (
                    dh["hallucination_yes"] if ca.get("hallucination_detected") else dh["hallucination_no"]
                )
                issues = ca.get("issues", [])
                qa_tbl.rows[qi].cells[5].text = "; ".join(issues) if issues else dh["deep_no_issues"]
    elif _qa_sum and _qa_sum.get("skipped"):
        doc.add_paragraph(dh["deep_qa_skipped"].format(summary=_qa_sum.get("summary", "")))
    else:
        doc.add_paragraph(dh["deep_no_qa_record"])
        _qa_word_note = {
            "zh": (
                "※ 說明：第三方 QA 稽核資料需執行「每日稽核 / Daily Audit」功能後才會產生。"
                "請在主介面選擇「每日稽核 / Daily Audit」並完成至少一次稽核流程後，重新匯出報告即可看到此欄位資料。"
            ),
            "en": (
                "Note: Third-party QA audit data is generated by running the \"Daily Audit\" function. "
                "Please select \"Daily Audit\" from the main interface and complete at least one audit cycle, "
                "then re-export the report to see this data."
            ),
            "ja": (
                "※ 備考：第三者 QA 監査データは「毎日監査 / Daily Audit」機能を実行することで生成されます。"
                "メインインターフェイスで「毎日監査」を選択し、少なくとも 1 回の監査を完了してから、レポートを再エクスポートしてください。"
            ),
        }
        _qa_lk = "ja" if str(lang or "").startswith("ja") else "en" if str(lang or "").startswith("en") else "zh"
        _note_p = doc.add_paragraph(_qa_word_note[_qa_lk])
        _note_p.runs[0].italic = True
        _note_p.runs[0].font.color.rgb = None  # inherit theme color (gray-ish)

    # ── Section 6: Compliance Table ──
    doc.add_heading(dh["deep_s6"], level=2)
    if flat_rows:
        headers = dh["deep_s6_headers"]
        tbl = doc.add_table(rows=1 + len(flat_rows), cols=len(headers))
        tbl.style = "Table Grid"
        _S6_HDR_FILL = "1F3864"
        for i, h in enumerate(headers):
            _bold_cell(tbl.rows[0].cells[i], h, fill_hex=_S6_HDR_FILL, font_color="FFFFFF", font_size=9)
        _S6_RISK_FILL = {
            "immediate_correction": "FFE0E0",
            "deadline_correction":  "FFF2CC",
            "improvement_plan":     "E8F5E9",
            "compliant":            "F2F2F2",
        }
        for ri, row in enumerate(flat_rows, 1):
            rl = row.get("risk_level", "")
            row_bg = _S6_RISK_FILL.get(rl, "FFFFFF")
            row_cells = tbl.rows[ri].cells
            row_cells[0].text = f"{row.get('clause_id', '')} {row.get('clause_title', '')}"
            row_cells[1].text = f"{row.get('doc_id', '')}"
            row_cells[2].text = row.get("audit_impact", "")
            row_cells[3].text = f"{row.get('verdict_icon', '')} {row.get('verdict_label', '')}"
            row_cells[4].text = f"{row.get('risk_icon', '')} {row.get('risk_label', '')}"
            row_cells[5].text = row.get("gap_severity", "") or ""
            row_cells[6].text = "⚠️" if row.get("flagged_for_ra") else ""
            row_cells[7].text = _pipeline_status_str(row.get("phase_status_summary", {}))
            for ci in range(8):
                _shade_cell(row_cells[ci], row_bg)
                for run in row_cells[ci].paragraphs[0].runs:
                    run.font.size = Pt(8)
                    if rl == "immediate_correction":
                        run.bold = True

    # ── Section 7: Meta-Analysis (if available) ──
    if meta_analysis:
        doc.add_heading(dh["deep_s7"], level=2)
        doc.add_paragraph(meta_analysis.get("summary", dh["deep_s7_no_summary"]))

        findings = meta_analysis.get("findings", [])
        if findings:
            doc.add_heading(dh["deep_s7_findings"], level=3)
            for f in findings:
                doc.add_paragraph(
                    f"• [{f.get('severity', '')}] {f.get('description', '')}",
                )

        recommendations = meta_analysis.get("recommendations", [])
        if recommendations:
            doc.add_heading(dh["deep_s7_recs"], level=3)
            for rec in recommendations:
                doc.add_paragraph(f"• {rec}")

        tuning = meta_analysis.get("prompt_tuning", {})
        if tuning:
            doc.add_heading(dh["deep_s7_prompt"], level=3)
            for key, val in tuning.items():
                doc.add_paragraph(f"• {key}: {val}")

    # ── Section 8: LLM Usage Statistics ──
    doc.add_heading(dh["deep_s8"], level=2)
    if interactions:
        phase_counts: dict[str, dict] = {}
        for i in interactions:
            phase = i.get("phase_label", i.get("phase", "unknown"))
            if phase not in phase_counts:
                phase_counts[phase] = {"count": 0, "tokens": 0}
            phase_counts[phase]["count"] += 1
            phase_counts[phase]["tokens"] += i.get("usage", {}).get("total_tokens", 0)

        for phase, stats in phase_counts.items():
            doc.add_paragraph(
                f"• {phase}: {stats['count']} {dh['deep_s8_calls']}, {stats['tokens']:,} tokens"
            )

    try:
        _append_crawl_status_word(doc, _load_crawl_results(), lang)
    except Exception:
        pass
    try:
        from src.utils.crossref_export import append_crossref_table_word
        append_crossref_table_word(doc, lang=lang)
    except Exception:
        pass

    safe_save_binary(filepath, doc.save)
    return filepath


def export_deep_report_excel(
    run_id: str,
    flat_rows: list[dict],
    summary: dict,
    interactions: list[dict] | None = None,
    crossexam_record: dict | None = None,
    meta_analysis: dict | None = None,
    qa_audit_summary: dict | None = None,
    lang: str = "zh-TW",
    progress: dict | None = None,
    data_quality: dict | None = None,
    source_check: dict | None = None,
    skipped_phases: list[str] | None = None,
    output_path: "Path | None" = None,
) -> Path:
    """Export a deep analysis report as Excel workbook.

    Multiple sheets: Summary, Compliance Table, LLM Interactions, Cross-Exam, Meta-Analysis.

    Args:
        lang: UI language code (e.g., 'zh-TW', 'en', 'ja') — reserved for
            future localization of sheet/section names.
        output_path: Optional custom output path. If None, writes to EXPORT_DIR.
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill
    from src.analysis.risk_matrix import VERDICT_DISPLAY, RISK_LEVEL_DISPLAY

    _lk = _lang_key(lang)
    dh = _EXPORT_HEADERS[_lk]
    _label_key = "label_en" if _lk == "en" else "label_ja" if _lk == "ja" else "label_zh"
    _skipped_xl = skipped_phases or []

    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    filepath = output_path if output_path is not None else (EXPORT_DIR / f"deep_report_{run_id}.xlsx")

    wb = Workbook()
    header_fill = PatternFill(
        start_color="4472C4", end_color="4472C4", fill_type="solid"
    )
    header_font = Font(bold=True, color="FFFFFF", size=10)

    # Risk-level fill colors reused across sheets
    _RISK_XE_FILL = {
        "immediate_correction": PatternFill(start_color="FFE0E0", end_color="FFE0E0", fill_type="solid"),
        "deadline_correction":  PatternFill(start_color="FFF2CC", end_color="FFF2CC", fill_type="solid"),
        "improvement_plan":     PatternFill(start_color="E8F5E9", end_color="E8F5E9", fill_type="solid"),
        "compliant":            PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid"),
    }

    # ── Sheet 1: Summary ──
    ws_sum = wb.active
    ws_sum.title = dh["xl_sheet_summary"]
    summary_data = [
        (dh["xl_run_id"], run_id),
        (dh["xl_export_time"], datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        (dh["xl_total_rows"], summary.get("total_rows", len(flat_rows))),
        (dh["xl_flagged_ra"], summary.get("flagged_for_ra", 0)),
    ]
    verdict_dist = summary.get("verdict_distribution", {})
    for v, count in verdict_dist.items():
        _vlabel = VERDICT_DISPLAY.get(v, {}).get(_label_key, v)
        summary_data.append((f"{dh['xl_verdict_prefix']}{_vlabel}", count))
    risk_dist = summary.get("risk_distribution", {})
    for r, count in risk_dist.items():
        _rlabel = RISK_LEVEL_DISPLAY.get(r, {}).get(_label_key, r)
        summary_data.append((f"{dh['xl_risk_prefix']}{_rlabel}", count))

    for ri, (label, value) in enumerate(summary_data, 1):
        c1 = ws_sum.cell(row=ri, column=1, value=label)
        c1.fill = header_fill
        c1.font = header_font
        ws_sum.cell(row=ri, column=2, value=str(value))
    ws_sum.column_dimensions["A"].width = 20
    ws_sum.column_dimensions["B"].width = 40

    # ── Sheet 2: Compliance Table ──
    ws_comp = wb.create_sheet(dh["xl_sheet_compliance"])
    comp_headers = dh["xl_comp_headers"]
    _skip_hdr_fill = PatternFill(start_color="BFBFBF", end_color="BFBFBF", fill_type="solid")
    _skip_hdr_font = Font(bold=True, color="666666", size=10)
    for ci, h in enumerate(comp_headers, 1):
        c = ws_comp.cell(row=1, column=ci, value=h)
        # Grey out P4 header columns (14-15) and P5 header columns (16-17) when skipped
        if ci in (14, 15) and "phase_4" in _skipped_xl:
            c.fill = _skip_hdr_fill
            c.font = _skip_hdr_font
        elif ci in (16, 17) and "phase_5" in _skipped_xl:
            c.fill = _skip_hdr_fill
            c.font = _skip_hdr_font
        else:
            c.fill = header_fill
            c.font = header_font
    for ri, row in enumerate(flat_rows, 2):
        ws_comp.cell(row=ri, column=1, value=row.get("clause_id", ""))
        ws_comp.cell(row=ri, column=2, value=row.get("clause_title", ""))
        ws_comp.cell(row=ri, column=3, value=row.get("doc_id", ""))
        ws_comp.cell(row=ri, column=4, value=row.get("doc_title", ""))
        ws_comp.cell(row=ri, column=5, value=row.get("audit_impact", ""))
        ws_comp.cell(row=ri, column=6, value=row.get("audit_question", ""))
        ws_comp.cell(
            row=ri,
            column=7,
            value=f"{row.get('verdict_icon', '')} {row.get('verdict_label', '')}",
        )
        ws_comp.cell(
            row=ri,
            column=8,
            value=f"{row.get('risk_icon', '')} {row.get('risk_label', '')}",
        )
        ws_comp.cell(row=ri, column=9, value=row.get("gap_severity", "") or "")
        ws_comp.cell(
            row=ri,
            column=10,
            value=f"{row.get('evidence_found', 0)}/{row.get('evidence_total', 0)}",
        )
        ws_comp.cell(row=ri, column=11, value="Y" if row.get("flagged_for_ra") else "")
        override = row.get("ra_override")
        ws_comp.cell(
            row=ri,
            column=12,
            value=override.get("reason", "") if isinstance(override, dict) else "",
        )
        ws_comp.cell(row=ri, column=13, value=_safe_str(row.get("ra_notes"), 500))
        # P4 columns: only write when phase_4 not skipped
        if "phase_4" not in _skipped_xl:
            ws_comp.cell(row=ri, column=14, value=_safe_str(row.get("remediation_suggestion"), 500))
            ws_comp.cell(row=ri, column=15, value=_safe_str(row.get("remediation_regulation_cite"), 500))
        # P5 columns: only write when phase_5 not skipped
        if "phase_5" not in _skipped_xl:
            ws_comp.cell(row=ri, column=16, value=_safe_str(row.get("analyzer_position"), 500))
            ws_comp.cell(row=ri, column=17, value=_safe_str(row.get("verifier_position"), 500))
        ws_comp.cell(row=ri, column=18, value=_pipeline_status_str(row.get("phase_status_summary", {})))

    # ── Sheet: Phase Progress (unified single table) ──
    _prog = progress or {}
    _budget_xl = _prog.get("llm_budget", {})
    _phase_dist_xl = _prog.get("phase_distribution", {})
    _xl_safe = lambda s: re.sub(r'[:\\/*?\[\]]', '', s)[:31]
    ws_prog = wb.create_sheet(_xl_safe(dh.get("deep_s0", "Pipeline Progress")))

    _sec_fill = PatternFill(start_color="1F3864", end_color="1F3864", fill_type="solid")
    _sec_font = Font(bold=True, color="FFFFFF", size=11)
    _sub_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
    _sub_font = Font(bold=True, size=10)

    def _prog_sec_hdr(ws, row, label, ncols=7):
        c = ws.cell(row=row, column=1, value=label)
        c.fill = _sec_fill
        c.font = _sec_font
        for col in range(2, ncols + 1):
            ws.cell(row=row, column=col).fill = _sec_fill

    def _prog_col_hdr(ws, row, labels):
        for ci, lbl in enumerate(labels, 1):
            c = ws.cell(row=row, column=ci, value=lbl)
            c.fill = _sub_fill
            c.font = _sub_font

    cur = 1  # current row pointer

    # ── Section A: Execution Phases ──
    _prog_sec_hdr(ws_prog, cur, "▌ 執行階段狀態 / Execution Phase Status")
    cur += 1
    _prog_col_hdr(ws_prog, cur, [
        dh.get("deep_s0_phase_col_phase", "階段"),
        dh.get("deep_s0_phase_col_name_zh", "名稱（中）"),
        dh.get("deep_s0_phase_col_name_en", "Name (EN)"),
        dh.get("deep_s0_phase_col_rows", "Row 數"),
        dh.get("deep_s0_phase_col_status", "狀態"),
    ])
    cur += 1
    for ph in _PHASE_ORDER:
        is_skipped = ph in _skipped_xl
        ws_prog.cell(row=cur, column=1, value=_PHASE_LABELS[ph])
        ws_prog.cell(row=cur, column=2, value=_PHASE_NAMES_ZH.get(ph, ph))
        ws_prog.cell(row=cur, column=3, value=_PHASE_NAMES_EN.get(ph, ph))
        ws_prog.cell(row=cur, column=4, value=_phase_dist_xl.get(ph, 0))
        ws_prog.cell(row=cur, column=5, value="↷ Skipped" if is_skipped else "✓")
        cur += 1

    # ── Section B: LLM Budget ──
    _prog_sec_hdr(ws_prog, cur, "▌ LLM 預算 / LLM Budget")
    cur += 1
    _prog_col_hdr(ws_prog, cur, ["項目", "數值"])
    cur += 1
    ws_prog.cell(row=cur, column=1, value=dh.get("deep_s0_budget_total", "Token 總用量"))
    ws_prog.cell(row=cur, column=2, value=_budget_xl.get("total_tokens_used", 0))
    cur += 1
    ws_prog.cell(row=cur, column=1, value=dh.get("deep_s0_budget_calls", "呼叫次數"))
    ws_prog.cell(row=cur, column=2, value=_budget_xl.get("calls_made", 0))
    cur += 1
    ws_prog.cell(row=cur, column=1, value=dh.get("deep_s0_budget_pct", "預算使用%"))
    ws_prog.cell(row=cur, column=2, value=f"{_budget_xl.get('usage_percent', 0)}%")
    cur += 1

    # ── Section C: Data Quality P0 ──
    _DQ_KEY_LABEL = {
        "total_rows":                  "總分析筆數 / Total Rows / 総分析行数",
        "rows_with_doc_content":       "有文件內容筆數 / Rows w/ Doc Content / 文書内容あり",
        "rows_without_doc_content":    "無文件內容筆數 / Rows w/o Doc Content / 文書内容なし",
        "rows_with_regulatory_data":   "有法規資料筆數 / Rows w/ Reg. Data / 規制データあり",
        "rows_without_regulatory_data":"無法規資料筆數 / Rows w/o Reg. Data / 規制データなし",
        "regulatory_data_available":   "法規資料可用 / Reg. Data Available / 規制データ利用可能",
        "overall_pass":                "整體通過 / Overall Pass / 全体合格",
        "issues":                      "問題數量 / Issues / 問題数",
    }
    _SC_KEY_LABEL = {
        "total_urls":        "URL 總數 / Total URLs / URL総数",
        "accessible":        "可存取數 / Accessible / アクセス可能",
        "inaccessible":      "無法存取 / Inaccessible / アクセス不可",
        "content_changed":   "內容已變更 / Content Changed / 内容変更あり",
        "errors":            "錯誤數 / Errors / エラー数",
        "overall_pass":      "整體通過 / Overall Pass / 全体合格",
    }
    if data_quality:
        _prog_sec_hdr(ws_prog, cur, "▌ 資料品質 P0 / Data Quality (Phase 0)")
        cur += 1
        _prog_col_hdr(ws_prog, cur, ["項目 (中/EN/日)", "數值"])
        cur += 1
        # Summary stats (scalar fields)
        for k, v in data_quality.items():
            if k == "documents_checked":
                continue
            if not isinstance(v, (dict, list)):
                ws_prog.cell(row=cur, column=1, value=_DQ_KEY_LABEL.get(k, k))
                ws_prog.cell(row=cur, column=2, value=v)
                cur += 1
        # documents_checked table
        if "documents_checked" in data_quality and isinstance(data_quality["documents_checked"], dict):
            docs = data_quality["documents_checked"]
            ws_prog.cell(row=cur, column=1, value=f"檢核文件清單 ({len(docs)} 份)").font = _sub_font
            cur += 1
            _prog_col_hdr(ws_prog, cur, ["文件 ID", "存在", "有內容", "字元數", "廢止"])
            cur += 1
            for doc_id, info in docs.items():
                if isinstance(info, dict):
                    ws_prog.cell(row=cur, column=1, value=doc_id)
                    ws_prog.cell(row=cur, column=2, value="✓" if info.get("exists") else "✗")
                    ws_prog.cell(row=cur, column=3, value="✓" if info.get("has_content") else "✗")
                    ws_prog.cell(row=cur, column=4, value=info.get("content_length", 0))
                    ws_prog.cell(row=cur, column=5, value="是" if info.get("is_obsolete") else "否")
                    cur += 1

    # ── Section D: Source Check P6 ──
    if source_check:
        _prog_sec_hdr(ws_prog, cur, "▌ 法規來源驗證 P6 / Source Check (Phase 6)")
        cur += 1
        _prog_col_hdr(ws_prog, cur, ["項目 (中/EN/日)", "數值"])
        cur += 1
        for k, v in source_check.items():
            if k == "verification_results":
                continue
            if not isinstance(v, (dict, list)):
                ws_prog.cell(row=cur, column=1, value=_SC_KEY_LABEL.get(k, k))
                ws_prog.cell(row=cur, column=2, value=v)
                cur += 1
        vr = source_check.get("verification_results", [])
        if vr:
            ws_prog.cell(row=cur, column=1, value=f"URL 驗證結果 ({len(vr)} 筆)").font = _sub_font
            cur += 1
            _prog_col_hdr(ws_prog, cur, ["狀態", "機構", "地區", "URL", "內容變更", "HTTP 狀態", "錯誤訊息"])
            cur += 1
            for item in vr:
                if not isinstance(item, dict):
                    continue
                ws_prog.cell(row=cur, column=1, value="✓" if item.get("accessible") else "✗")
                ws_prog.cell(row=cur, column=2, value=item.get("agency", ""))
                ws_prog.cell(row=cur, column=3, value=item.get("region", ""))
                ws_prog.cell(row=cur, column=4, value=(item.get("url", ""))[:150])
                ws_prog.cell(row=cur, column=5, value="是" if item.get("content_changed") else ("—" if item.get("accessible") else ""))
                ws_prog.cell(row=cur, column=6, value=item.get("status_code") or "")
                ws_prog.cell(row=cur, column=7, value=(item.get("error") or "")[:120])
                cur += 1

    # Auto-width for progress sheet
    for col in ws_prog.columns:
        max_len = max((len(str(c.value or "")) for c in col), default=8)
        ws_prog.column_dimensions[col[0].column_letter].width = min(max_len + 2, 60)

    # track dq_row for backward compat (not used below, just kept)
    dq_row = cur

    # ── Sheet 3: LLM Interactions ──
    _PHASE_LABEL_MAP = dh["xl_phase_labels"]
    _ROLE_LABEL_MAP = dh["xl_role_labels"]
    _batch_label = dh["xl_batch_label"]
    # Phase-based alternating fill colors for LLM sheet
    _LLM_PHASE_FILL = {
        "gap_scan":         PatternFill(start_color="DDEEFF", end_color="DDEEFF", fill_type="solid"),
        "checklist_verify": PatternFill(start_color="E8F5E9", end_color="E8F5E9", fill_type="solid"),
        "remediation":      PatternFill(start_color="FFF9C4", end_color="FFF9C4", fill_type="solid"),
        "verification":     PatternFill(start_color="FCE4D6", end_color="FCE4D6", fill_type="solid"),
    }
    # Map phase key → interaction phase name for filtering
    _SKIP_PHASE_MAP = {"phase_4": "remediation", "phase_5": "verification"}
    _filtered_interactions = [
        i for i in (interactions or [])
        if i.get("phase") not in {_SKIP_PHASE_MAP[p] for p in _skipped_xl if p in _SKIP_PHASE_MAP}
    ] if interactions else None
    if _filtered_interactions:
        ws_llm = wb.create_sheet(dh["xl_sheet_llm"])
        llm_headers = dh["xl_llm_headers"]
        for ci, h in enumerate(llm_headers, 1):
            c = ws_llm.cell(row=1, column=ci, value=h)
            c.fill = header_fill
            c.font = header_font
        for ri, interaction in enumerate(_filtered_interactions, 2):
            phase = interaction.get("phase", "")
            phase_label = interaction.get("phase_label") or _PHASE_LABEL_MAP.get(phase, phase)
            clause_id = interaction.get("clause_id") or ""
            role = interaction.get("role") or ""
            role_label = _ROLE_LABEL_MAP.get(role, role) if role else _batch_label
            # For non-P5 phases (gap_scan/checklist_verify/remediation), clause_id is per-doc
            clause_display = clause_id if clause_id else (_batch_label if phase != "verification" else "")
            row_fill = _LLM_PHASE_FILL.get(phase)
            ws_llm.cell(row=ri, column=1, value=_safe_str(phase_label))
            ws_llm.cell(row=ri, column=2, value=_safe_str(interaction.get("doc_id", "")))
            ws_llm.cell(row=ri, column=3, value=_safe_str(clause_display))
            ws_llm.cell(row=ri, column=4, value=_safe_str(role_label))
            _rn = interaction.get("round_number")
            ws_llm.cell(row=ri, column=5, value=_rn if isinstance(_rn, (int, float)) else (_safe_str(_rn) or "—"))
            resp_text = _format_llm_response_for_excel(
                interaction.get("llm_response", "") or "",
                phase,
                interaction.get("parsed_response"),
                lang_key=_lk,
            )
            ws_llm.cell(row=ri, column=6, value=resp_text[:1200])
            _usage = interaction.get("usage") or {}
            ws_llm.cell(row=ri, column=7, value=_usage.get("total_tokens", 0) if isinstance(_usage, dict) else 0)
            ws_llm.cell(row=ri, column=8, value=_safe_str(interaction.get("model", "")))
            ws_llm.cell(row=ri, column=9, value=_safe_str(interaction.get("timestamp", "—")))
            if row_fill:
                for ci in range(1, 10):
                    ws_llm.cell(row=ri, column=ci).fill = row_fill

        # Freeze header row
        ws_llm.freeze_panes = "A2"
        for col in ws_llm.columns:
            max_len = max((len(str(c.value or "")) for c in col), default=8)
            ws_llm.column_dimensions[col[0].column_letter].width = min(max_len + 2, 60)
        # Set fixed widths for key columns
        ws_llm.column_dimensions["A"].width = 28  # Phase
        ws_llm.column_dimensions["B"].width = 16  # Doc ID
        ws_llm.column_dimensions["C"].width = 20  # Clause ID
        ws_llm.column_dimensions["D"].width = 22  # Role
        ws_llm.column_dimensions["F"].width = 60  # LLM Response

    # ── Sheet 4: Cross-Exam Details ──
    # Build cross-exam rows from crossexam_record (preferred) or flat_rows (fallback)
    ws_xe = wb.create_sheet(dh["xl_sheet_crossexam"])
    _xe_ext_headers = dh["xl_xe_ext_headers"]
    for ci, h in enumerate(_xe_ext_headers, 1):
        c = ws_xe.cell(row=1, column=ci, value=h)
        c.fill = header_fill
        c.font = header_font
    # Fill rows from interactions (most data-complete source for this run)
    _xe_rows_written = 0
    _p5_interactions = [i for i in (interactions or []) if i.get("phase") == "verification"]
    if _p5_interactions:
        _p5_seen: dict[str, dict] = {}
        for ix in _p5_interactions:
            cid = ix.get("clause_id", "")
            did = ix.get("doc_id", "")
            key = f"{cid}||{did}"
            if key not in _p5_seen:
                _p5_seen[key] = {"clause_id": cid, "doc_id": did, "clause_title": ix.get("clause_title",""), "rounds": []}
            raw = ix.get("llm_response", "") or ""
            try:
                rds = json.loads(raw)
                if isinstance(rds, list):
                    _p5_seen[key]["rounds"].extend(rds)
            except Exception:
                pass
        for ri, entry in enumerate(sorted(_p5_seen.values(), key=lambda x: x["clause_id"]), 2):
            rds = entry.get("rounds", [])
            # Merge first round data
            r1_a = (rds[0].get("analyzer") or {}) if rds else {}
            r1_v = (rds[0].get("verifier") or {}) if rds else {}
            # Find matching flat_row for verdict/agreed/qa
            fr = next((r for r in flat_rows if r.get("clause_id") == entry["clause_id"] and r.get("doc_id") == entry["doc_id"]), {})
            qa = fr.get("qa_audit") or {}
            row_verdict = fr.get("verdict", "")
            row_rl = fr.get("risk_level", "")
            row_fill_xe = _RISK_XE_FILL.get(row_rl)
            a_ev_list = r1_a.get("key_evidence") or []
            a_ev_str = "; ".join(str(e)[:100] for e in a_ev_list[:3])
            v_chal_list = r1_v.get("challenges") or []
            v_chal_str = "; ".join(
                str(c.get("point", c) if isinstance(c, dict) else c)[:100]
                for c in v_chal_list[:3]
            )
            row_vals = [
                entry["clause_id"],
                entry["clause_title"][:80],
                entry["doc_id"],
                fr.get("verdict_label", row_verdict),
                "Y" if fr.get("verification_agreed") else "N",
                "⚠️" if fr.get("flagged_for_ra") else "",
                len(rds),
                str(r1_a.get("position") or "")[:500],
                str(r1_a.get("confidence") or r1_a.get("confidence_score") or ""),
                a_ev_str[:400],
                v_chal_str[:400],
                r1_v.get("agreement_level", ""),
                qa.get("score", "") if qa else "",
                qa.get("question_quality", "") if qa else "",
                qa.get("answer_accuracy", "") if qa else "",
                "是" if qa and qa.get("hallucination_detected") else ("否" if qa else ""),
            ]
            for ci2, v in enumerate(row_vals, 1):
                c = ws_xe.cell(row=ri, column=ci2, value=v)
                c.font = Font(size=9)
                if row_fill_xe:
                    c.fill = row_fill_xe
            _xe_rows_written += 1
    elif crossexam_record:
        for ri, clause in enumerate(crossexam_record.get("clauses", []), 2):
            rounds = clause.get("rounds", [])
            r1_a = rounds[0].get("analyzer", {}) if rounds else {}
            r1_v = rounds[0].get("verifier", {}) if rounds else {}
            qa = clause.get("qa_audit", {}) or {}
            a_ev_list = r1_a.get("key_evidence") or []
            v_chal_list = r1_v.get("challenges") or []
            row_vals = [
                clause.get("clause_id", ""), clause.get("clause_title", ""), clause.get("doc_id", ""),
                clause.get("verdict", ""),
                "Y" if clause.get("agreed") else "N",
                "Y" if clause.get("flagged_for_ra") else "",
                len(rounds),
                _flatten_role_text(r1_a, "position")[:500],
                str(r1_a.get("confidence", r1_a.get("confidence_score", ""))),
                "; ".join(str(e)[:100] for e in a_ev_list[:3]),
                "; ".join(str(c.get("point",c) if isinstance(c,dict) else c)[:100] for c in v_chal_list[:3]),
                r1_v.get("agreement_level", ""),
                qa.get("score", ""), qa.get("question_quality", ""), qa.get("answer_accuracy", ""),
                dh.get("yes", "是") if qa.get("hallucination_detected") else dh.get("no", "否"),
            ]
            for ci2, v in enumerate(row_vals, 1):
                ws_xe.cell(row=ri, column=ci2, value=v).font = Font(size=9)
            _xe_rows_written += 1
    elif flat_rows and "phase_5" not in _skipped_xl:
        for ri, row in enumerate(flat_rows, 2):
            qa = row.get("qa_audit") or {}
            row_vals = [
                row.get("clause_id",""), row.get("clause_title","")[:80], row.get("doc_id",""),
                row.get("verdict_label", row.get("verdict","")),
                "Y" if row.get("verification_agreed") else "N",
                "⚠️" if row.get("flagged_for_ra") else "",
                str(row.get("verification_rounds", "")),
                (row.get("r1_analyzer_position") or row.get("analyzer_position") or "")[:500],
                (row.get("r1_analyzer_confidence") or ""),
                (row.get("r1_key_evidence") or "")[:400],
                (row.get("r1_verifier_challenges") or "")[:400],
                (row.get("r1_agreement_level") or ""),
                qa.get("score","") if qa else "",
                qa.get("question_quality","") if qa else "",
                qa.get("answer_accuracy","") if qa else "",
                (dh.get("yes", "是") if qa.get("hallucination_detected") else dh.get("no", "否")) if qa else "",
            ]
            for ci2, v in enumerate(row_vals, 1):
                ws_xe.cell(row=ri, column=ci2, value=v).font = Font(size=9)
            _xe_rows_written += 1
    if not _xe_rows_written:
        _skip_note = "（P5 交叉詰問已跳過 — 此報告類型不包含交叉詰問資料）" if "phase_5" in _skipped_xl else "（無交叉詰問資料）"
        ws_xe.cell(row=2, column=1, value=_skip_note)
    ws_xe.freeze_panes = "A2"
    for col in ws_xe.columns:
        max_len = max((len(str(c.value or "")) for c in col), default=8)
        ws_xe.column_dimensions[col[0].column_letter].width = min(max_len + 2, 50)
    ws_xe.column_dimensions["B"].width = 40
    ws_xe.column_dimensions["H"].width = 50
    ws_xe.column_dimensions["J"].width = 40
    ws_xe.column_dimensions["K"].width = 40

    # ── Sheet: 第三方稽核 — use per-row qa_audit data from flat_rows ──
    _qa_rows = [r for r in (flat_rows or []) if r.get("qa_audit")]
    _qa_sheet_name = {"zh": "第三方稽核", "en": "Third-Party QA", "ja": "第三者QA"}[_lk]
    ws_qa = wb.create_sheet(_qa_sheet_name)
    _qa_col_headers = {
        "zh": ["條款 ID", "文件 ID", "分數", "評分說明", "問題品質", "回答準確性", "幻覺偵測", "問題清單"],
        "en": ["Clause ID", "Doc ID", "Score", "Score Rationale", "Question Quality", "Answer Accuracy", "Hallucination", "Issue List"],
        "ja": ["条項ID", "文書ID", "スコア", "採点根拠", "質問品質", "回答精度", "幻覚検出", "問題一覧"],
    }
    _qa_row_headers = _qa_col_headers.get(_lk, _qa_col_headers["zh"])
    for ci, h in enumerate(_qa_row_headers, 1):
        c = ws_qa.cell(row=1, column=ci, value=h)
        c.fill = header_fill
        c.font = header_font
    if _qa_rows:
        for qi, row in enumerate(_qa_rows, 2):
            qa = row.get("qa_audit", {}) or {}
            ws_qa.cell(row=qi, column=1, value=row.get("clause_id", ""))
            ws_qa.cell(row=qi, column=2, value=row.get("doc_id", ""))
            ws_qa.cell(row=qi, column=3, value=qa.get("score", ""))
            ws_qa.cell(row=qi, column=4, value=(qa.get("score_rationale", "") or "")[:300])
            ws_qa.cell(row=qi, column=5, value=qa.get("question_quality", ""))
            ws_qa.cell(row=qi, column=6, value=qa.get("answer_accuracy", ""))
            ws_qa.cell(row=qi, column=7, value=dh.get("yes", "是") if qa.get("hallucination_detected") else dh.get("no", "否"))
            issues = qa.get("issues", []) or []
            ws_qa.cell(row=qi, column=8, value="; ".join(str(i) for i in issues) if issues else "")
    else:
        _qa_no_data_note = {
            "zh": (
                "（無第三方稽核資料）\n\n"
                "※ 說明：第三方稽核（Third-Party QA）資料需透過「每日稽核」功能執行後方可產生。\n"
                "請在主介面選擇「每日稽核 / Daily Audit」，對已完成分析的文件執行稽核後，重新匯出報告即可看到此欄位資料。"
            ),
            "en": (
                "(No third-party QA data available)\n\n"
                "Note: Third-Party QA data is generated by running the Daily Audit function.\n"
                "Please select 'Daily Audit' from the main interface and run the audit on analyzed documents, "
                "then re-export the report to populate this sheet."
            ),
            "ja": (
                "（第三者QAデータなし）\n\n"
                "※ 備考：第三者QAデータは「日次監査」機能を実行することで生成されます。\n"
                "メイン画面から「日次監査 / Daily Audit」を選択し、分析済み文書に対して監査を実行後、"
                "レポートを再エクスポートするとこのシートにデータが表示されます。"
            ),
        }
        note_cell = ws_qa.cell(row=2, column=1, value=_qa_no_data_note.get(_lk, _qa_no_data_note["zh"]))
        note_cell.alignment = Alignment(wrap_text=True, vertical="top")
        note_cell.font = Font(size=10, color="666666", italic=True)
        ws_qa.row_dimensions[2].height = 80
        ws_qa.merge_cells(start_row=2, start_column=1, end_row=2, end_column=len(_qa_row_headers))
    for col in ws_qa.columns:
        max_len = max((len(str(c.value or "")) for c in col), default=8)
        ws_qa.column_dimensions[col[0].column_letter].width = min(max_len + 2, 50)

    # ── Sheet 5: Meta-Analysis ──
    if meta_analysis:
        _meta_title = {"zh": "交叉詰問品質分析結果", "en": "Cross-Examination Quality Analysis", "ja": "相互尋問品質分析結果"}[_lk]
        ws_ma = wb.create_sheet(dh["xl_sheet_meta"])
        ws_ma.cell(row=1, column=1, value=_meta_title).font = Font(
            bold=True, size=14
        )
        ws_ma.cell(row=3, column=1, value=dh["summary"]).font = Font(bold=True)
        ws_ma.cell(row=4, column=1, value=meta_analysis.get("summary", ""))

        findings = meta_analysis.get("findings", [])
        if findings:
            row_offset = 6
            ws_ma.cell(row=row_offset, column=1, value=dh["deep_s7_findings"]).font = Font(
                bold=True
            )
            for fi, f in enumerate(findings, row_offset + 1):
                ws_ma.cell(row=fi, column=1, value=f.get("severity", ""))
                ws_ma.cell(row=fi, column=2, value=f.get("description", ""))

    # Auto-width for compliance sheet
    for col in ws_comp.columns:
        max_len = max((len(str(c.value or "")) for c in col), default=8)
        ws_comp.column_dimensions[col[0].column_letter].width = min(max_len + 2, 50)
    # ── Sheet: 風險分析 / Risk Analysis ──
    _risk_sheet_name = {"zh": "風險分析", "en": "Risk Analysis", "ja": "リスク分析"}[_lk]
    ws_risk = wb.create_sheet(_risk_sheet_name)
    _risk_prio = {"immediate_correction": 0, "deadline_correction": 1, "improvement_plan": 2, "compliant": 9}
    _risk_rows_xl = sorted(flat_rows, key=lambda r: (_risk_prio.get(r.get("risk_level",""), 9), r.get("clause_id","")))
    _risk_hdrs = ["優先度", "條款 ID", "條款名稱", "文件 ID", "判定", "判定圖示", "風險等級", "差距嚴重度", "RA 標記", "改善建議", "法規引用"]
    for ci, h in enumerate(_risk_hdrs, 1):
        c = ws_risk.cell(row=1, column=ci, value=h)
        c.fill = header_fill; c.font = header_font
    _RISK_FILL = {
        "immediate_correction": PatternFill(start_color="FFE0E0", end_color="FFE0E0", fill_type="solid"),
        "deadline_correction":  PatternFill(start_color="FFF2CC", end_color="FFF2CC", fill_type="solid"),
        "improvement_plan":     PatternFill(start_color="E8F5E9", end_color="E8F5E9", fill_type="solid"),
        "compliant":            PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid"),
    }
    _PRIO_LABEL = {"immediate_correction": "🔴 立即改正", "deadline_correction": "🟡 期限改正", "improvement_plan": "🟢 改進計劃", "compliant": "✅ 符合"}
    for ri2, row in enumerate(_risk_rows_xl, 2):
        rl = row.get("risk_level", "")
        row_fill = _RISK_FILL.get(rl)
        vals = [
            _PRIO_LABEL.get(rl, rl),
            row.get("clause_id",""), row.get("clause_title","")[:60],
            row.get("doc_id",""),
            row.get("verdict_label", row.get("verdict","")),
            row.get("verdict_icon",""),
            row.get("risk_label", rl),
            row.get("gap_severity","") or "—",
            "⚠️" if row.get("flagged_for_ra") else "",
            _safe_str(row.get("remediation_suggestion"), 300),
            _safe_str(row.get("remediation_regulation_cite"), 200),
        ]
        for ci2, v in enumerate(vals, 1):
            c = ws_risk.cell(row=ri2, column=ci2, value=v)
            if row_fill: c.fill = row_fill
            c.font = Font(size=9)
    ws_risk.freeze_panes = "A2"
    ws_risk.column_dimensions["A"].width = 16
    ws_risk.column_dimensions["B"].width = 12
    ws_risk.column_dimensions["C"].width = 40
    ws_risk.column_dimensions["D"].width = 14
    ws_risk.column_dimensions["E"].width = 14
    ws_risk.column_dimensions["J"].width = 50
    ws_risk.column_dimensions["K"].width = 30

    try:
        _append_crawl_status_excel(wb, _load_crawl_results(), lang)
    except Exception:
        pass
    try:
        from src.utils.crossref_export import append_crossref_table_excel
        append_crossref_table_excel(wb, lang=lang)
    except Exception:
        pass

    # ── Sheet: Execution Environment ───────────────────────────
    _run_meta_xl = _prog.get("run_metadata", {}) if _prog else {}
    if _run_meta_xl:
        _env_sheet_name = {"zh": "執行環境", "en": "Exec Environment", "ja": "実行環境"}
        ws_env = wb.create_sheet(_xl_safe(_env_sheet_name.get(_lk, "Exec Environment")))
        _env_rows_xl = _build_env_table_rows(_run_meta_xl, _lk)
        _eh_xl = {"zh": ["項目", "值"], "en": ["Item", "Value"], "ja": ["項目", "値"]}
        _hdr_xl = _eh_xl.get(_lk, ["Item", "Value"])
        _env_hdr_fill = PatternFill(start_color="1F3864", end_color="1F3864", fill_type="solid")
        for ci, h in enumerate(_hdr_xl, 1):
            c = ws_env.cell(row=1, column=ci, value=h)
            c.font = Font(bold=True, color="FFFFFF")
            c.fill = _env_hdr_fill
            c.alignment = Alignment(horizontal="center")
        for ri, (label, value) in enumerate(_env_rows_xl, 2):
            ws_env.cell(row=ri, column=1, value=label)
            ws_env.cell(row=ri, column=2, value=value)

        # ── Actual models used per phase (from interaction log) ──
        if interactions:
            _phase_models: dict[str, set] = {}
            for _ix in interactions:
                _ph = _ix.get("phase") or _ix.get("phase_label", "")
                _m = _ix.get("model", "")
                if _ph and _m and not _m.startswith("["):
                    _phase_models.setdefault(_ph, set()).add(_m)

            if _phase_models:
                _actual_model_label = {"zh": "實際使用模型（依階段）", "en": "Models Used by Phase", "ja": "フェーズ別使用モデル"}
                cur_ri = ws_env.max_row + 2
                ws_env.cell(row=cur_ri, column=1, value=_actual_model_label.get(_lk, "Models Used by Phase")).font = Font(bold=True, size=10)
                cur_ri += 1
                _phase_name_map = {"gap_scan": "P1 Gap Scan", "checklist_verify": "P2 Checklist", "remediation": "P4 Remediation", "verification": "P5 Cross-Exam"}
                for _ph, _models in sorted(_phase_models.items()):
                    _ph_label = _phase_name_map.get(_ph, _ph)
                    ws_env.cell(row=cur_ri, column=1, value=_ph_label)
                    ws_env.cell(row=cur_ri, column=2, value="; ".join(sorted(_models)))
                    cur_ri += 1

        ws_env.column_dimensions["A"].width = 28
        ws_env.column_dimensions["B"].width = 65

    safe_save_binary(filepath, wb.save)
    return filepath


# ============================================================
# Helpers


def _strip_md_fence(s) -> str:
    """Remove ```json ... ``` or ``` ... ``` fences. Accepts any type — coerces to str first."""
    if not isinstance(s, str):
        s = _safe_str(s)
    import re
    s = s.strip()
    m = re.match(r"^```(?:json)?\s*(.*?)\s*```$", s, re.DOTALL)
    return m.group(1) if m else s


def _format_llm_response_for_excel(raw_resp: str, phase: str, parsed_response=None, lang_key: str = "zh") -> str:
    """Convert LLM response to human-readable single-line summary for Excel cell.

    Each phase has a different JSON structure:
      gap_scan        → clause_results[cid].evidence_results[]
      checklist_verify→ clause_results[cid].verification_results[]
      remediation     → clause_results[cid].remediation{}
      verification    → [{round, analyzer{position,key_evidence}, verifier{challenges,agreement_level}}]
    """
    _LABELS: dict[str, dict[str, str]] = {
        "zh": {
            "analyzer": "分析者", "evidence": "證據", "challenge": "質疑",
            "found": "找到 {found}/{total}", "missing": "缺",
            "full": "完整", "partial": "部分", "priority": "優先",
        },
        "en": {
            "analyzer": "Analyzer", "evidence": "Evidence", "challenge": "Challenge",
            "found": "{found}/{total} found", "missing": "Missing",
            "full": "Full", "partial": "Partial", "priority": "Priority",
        },
        "ja": {
            "analyzer": "分析者", "evidence": "証拠", "challenge": "指摘",
            "found": "{found}/{total}確認", "missing": "欠落",
            "full": "完全", "partial": "部分", "priority": "優先",
        },
    }
    lb = _LABELS.get(lang_key, _LABELS["en"])

    if not raw_resp:
        return ""
    if not isinstance(raw_resp, str):
        raw_resp = _safe_str(raw_resp)

    # ── Phase 5 (verification): JSON array of rounds ──
    if phase == "verification":
        try:
            rds = json.loads(_strip_md_fence(raw_resp))
            if isinstance(rds, list):
                parts = []
                for rd in rds[:3]:
                    rn = rd.get("round", "?")
                    a = rd.get("analyzer") or {}
                    v = rd.get("verifier") or {}
                    a_pos = str(a.get("position") or "")[:100]
                    a_ev_list = a.get("key_evidence") or []
                    a_ev = "; ".join(str(e)[:50] for e in a_ev_list[:2])
                    v_chals = "; ".join(
                        str(c.get("point", c) if isinstance(c, dict) else c)[:60]
                        for c in (v.get("challenges") or [])[:2]
                    )
                    v_ag = v.get("agreement_level", "")
                    summary = f"[R{rn}] {lb['analyzer']}: {a_pos}"
                    if a_ev: summary += f" ▸{lb['evidence']}: {a_ev}"
                    if v_chals: summary += f" ⚠{lb['challenge']}: {v_chals}"
                    if v_ag: summary += f" → {v_ag}"
                    parts.append(summary)
                return " ‖ ".join(parts)
        except Exception:
            pass
        return (raw_resp[:800]).replace("\n", " ")

    # ── Try to parse JSON (may have ```json fence or LLM quirks) ──
    try:
        parsed = json.loads(_strip_md_fence(raw_resp))
    except Exception:
        try:
            parsed = json.loads(_sanitize_llm_json(_strip_md_fence(raw_resp)))
        except Exception:
            return raw_resp[:800].replace("\n", " ")

    if not isinstance(parsed, dict):
        return str(parsed)[:800]

    clause_results = parsed.get("clause_results", {})
    if not clause_results:
        return _format_flat_dict_summary(parsed)

    out_parts = []
    for cid, data in clause_results.items():
        if not isinstance(data, dict):
            continue

        # ── gap_scan ──
        ev_results = data.get("evidence_results", [])
        if ev_results:
            found_count = sum(1 for e in ev_results if e.get("found"))
            total = len(ev_results)
            missing = [e.get("evidence_name", "") for e in ev_results if not e.get("found")]
            found_str = lb["found"].format(found=found_count, total=total)
            s = f"[{cid}] {found_str}"
            if missing:
                s += f" ✗{lb['missing']}: {'; '.join(missing[:3])}"
            out_parts.append(s)
            continue

        # ── checklist_verify ──
        ver_results = data.get("verification_results", [])
        if ver_results:
            full = sum(1 for e in ver_results if e.get("adequacy") == "full")
            partial = sum(1 for e in ver_results if e.get("adequacy") == "partial")
            missing_v = [e.get("evidence_name","") for e in ver_results if e.get("adequacy") in ("missing","none","")]
            s = f"[{cid}] {lb['full']}:{full} {lb['partial']}:{partial}"
            if missing_v:
                s += f" ✗{lb['missing']}: {'; '.join(missing_v[:3])}"
            out_parts.append(s)
            continue

        # ── remediation ──
        remed = data.get("remediation", {})
        if remed and isinstance(remed, dict):
            summary = str(remed.get("summary") or "")[:150]
            priority = remed.get("priority", "")
            suggestions = remed.get("suggestions", []) or []
            actions = "; ".join(
                str(sg.get("action", sg) if isinstance(sg, dict) else sg)[:80]
                for sg in suggestions[:2]
            )
            s = f"[{cid}]"
            if priority: s += f" {lb['priority']}:{priority}"
            if summary: s += f" {summary}"
            if actions: s += f" ▶{actions}"
            out_parts.append(s)
            continue

        out_parts.append(f"[{cid}] {str(data)[:120]}")

    return " | ".join(out_parts) if out_parts else raw_resp[:400].replace("\n", " ")


def _format_flat_dict_summary(d: dict) -> str:
    """Format a flat dict response as a brief summary string."""
    parts = []
    for key in ("summary", "position", "assessment", "overall_assessment", "conclusion"):
        val = d.get(key)
        if val and isinstance(val, str):
            parts.append(f"{key}: {val[:200]}")
            break
    return " | ".join(parts) if parts else str(d)[:300]
# ============================================================


_PHASE_NAMES_ZH = {
    "phase_0": "資料品質檢查", "phase_0_5": "法規參照對應",
    "phase_1": "差距掃描", "phase_2": "查核表驗證",
    "phase_3": "風險評估", "phase_4": "改善建議",
    "phase_5": "獨立驗證", "phase_6": "來源驗證",
}
_PHASE_NAMES_EN = {
    "phase_0": "Data Quality Gate", "phase_0_5": "Reference Mapping",
    "phase_1": "Gap Scan", "phase_2": "Checklist Verification",
    "phase_3": "Risk Assessment", "phase_4": "Remediation",
    "phase_5": "Independent Verification", "phase_6": "Source Verification",
}
_PHASE_ORDER = ["phase_0", "phase_0_5", "phase_1", "phase_2", "phase_3", "phase_4", "phase_5", "phase_6"]
_PHASE_LABELS = {"phase_0": "P0", "phase_0_5": "P0.5", "phase_1": "P1", "phase_2": "P2",
                 "phase_3": "P3", "phase_4": "P4", "phase_5": "P5", "phase_6": "P6"}
_STATUS_ICONS = {"completed": "✓", "skipped": "↷", "failed": "✗", "running": "⟳", "pending": "○", "paused": "⏸"}


def _pipeline_status_str(phase_status_summary: dict) -> str:
    parts = []
    for p in _PHASE_ORDER:
        st = phase_status_summary.get(p, "pending")
        icon = _STATUS_ICONS.get(st, "?")
        parts.append(f"{_PHASE_LABELS[p]}{icon}")
    return " ".join(parts)


def _split_text(text: str, max_len: int = 3000) -> list[str]:
    if len(text) <= max_len:
        return [text]
    chunks = []
    while text:
        chunks.append(text[:max_len])
        text = text[max_len:]
    return chunks


def _sanitize_llm_json(s: str) -> str:
    """Fix common LLM JSON mistakes so json.loads() can parse the result.

    Handles:
    - Bare N/A → null  (e.g. "source_section": N/A)
    - Python None/True/False → null/true/false
    - Trailing commas before } or ]
    """
    import re
    # N/A bare value (with or without trailing comma/newline)
    s = re.sub(r':\s*N/A\s*([,\n\r\}])', lambda m: ': null' + m.group(1), s)
    s = re.sub(r':\s*N/A\s*$', ': null', s, flags=re.MULTILINE)
    # Python-style literals
    s = re.sub(r'\bNone\b', 'null', s)
    s = re.sub(r'\bTrue\b', 'true', s)
    s = re.sub(r'\bFalse\b', 'false', s)
    # Trailing commas before closing bracket
    s = re.sub(r',\s*([\}\]])', r'\1', s)
    return s


def _render_gap_json_response(doc, resp: str, lang: str = "zh-TW") -> None:
    """Parse and render a phase JSON llm_response as readable Word content.

    Handles responses wrapped in ```json ... ``` code fences.
    Falls back to raw paragraph dump if JSON parsing fails.

    Supported structures:
      gap_scan:        {"clause_results": {"<cid>": {"evidence_results": [...]}}}
      checklist_verify:{"clause_results": {"<cid>": {"verification_results": [...]}}}
                        or flat {"verification_results": [...]}
      remediation:     {"remediation": {"summary", "priority", "suggestions": [...],
                                        "regulation_citation"}}
    """
    if not resp:
        return
    if not isinstance(resp, str):
        resp = _safe_str(resp)

    # Strip ```json / ``` fences
    stripped = resp.strip()
    if stripped.startswith("```"):
        lines = stripped.splitlines()
        inner_lines = lines[1:]
        if inner_lines and inner_lines[-1].strip() == "```":
            inner_lines = inner_lines[:-1]
        stripped = "\n".join(inner_lines).strip()

    try:
        data = json.loads(stripped)
    except (json.JSONDecodeError, ValueError):
        try:
            data = json.loads(_sanitize_llm_json(stripped))
        except (json.JSONDecodeError, ValueError):
            for chunk in _split_text(resp, 3000):
                doc.add_paragraph(chunk)
            return

    # ── Remediation: flat {"remediation": {...}} ──
    if "remediation" in data and isinstance(data.get("remediation"), dict):
        _render_remediation_block(doc, data["remediation"], lang=lang)
        return

    # ── Flat checklist_verify (older format): {"verification_results": [...]} ──
    if "verification_results" in data and isinstance(data.get("verification_results"), list):
        _render_verification_results_table(doc, data["verification_results"], lang=lang)
        return

    clause_results = data.get("clause_results", {})
    if not clause_results:
        for chunk in _split_text(resp, 3000):
            doc.add_paragraph(chunk)
        return

    for clause_id, clause_data in clause_results.items():
        # ── gap_scan ──
        evidence_list = clause_data.get("evidence_results", [])
        if evidence_list:
            ev_tbl = doc.add_table(rows=1 + len(evidence_list), cols=5)
            ev_tbl.style = "Table Grid"
            hdr_cells = ev_tbl.rows[0].cells
            _ev_hdrs = {
                "zh": ["狀態", "證據名稱", "原文位置", "原文引用", "分析說明"],
                "en": ["Status", "Evidence Name", "Source Location", "Source Quote", "Analysis"],
                "ja": ["状態", "証拠名称", "出典位置", "原文引用", "分析説明"],
            }
            _ev_lk = "ja" if str(lang or "").startswith("ja") else "en" if str(lang or "").startswith("en") else "zh"
            for ci, h in enumerate(_ev_hdrs[_ev_lk]):
                hdr_cells[ci].text = h
                run = hdr_cells[ci].paragraphs[0].runs[0] if hdr_cells[ci].paragraphs[0].runs else hdr_cells[ci].paragraphs[0].add_run(h)
                run.bold = True
            for ei, ev in enumerate(evidence_list, 1):
                found = ev.get("found", False)
                inadequate = ev.get("is_inadequate", False)
                outdated = ev.get("is_outdated", False)
                flags = []
                if inadequate:
                    flags.append("不足")
                if outdated:
                    flags.append("過時")
                status_str = ("✓" if found else "✗") + (f" [{','.join(flags)}]" if flags else "")
                score = ev.get("relevance_score", "")
                name_str = ev.get("evidence_name", "")
                if score != "":
                    name_str += f" [相關度:{score}]"
                source_quote = ev.get("source_quote", "") or ""
                q_str = source_quote[:300] + ("…" if len(source_quote) > 300 else "")
                reasoning = ev.get("reasoning", "") or ""
                r_str = reasoning[:400] + ("…" if len(reasoning) > 400 else "")
                ev_tbl.rows[ei].cells[0].text = status_str
                ev_tbl.rows[ei].cells[1].text = name_str
                ev_tbl.rows[ei].cells[2].text = (ev.get("source_section") or "")[:200]
                ev_tbl.rows[ei].cells[3].text = q_str
                ev_tbl.rows[ei].cells[4].text = r_str
            continue

        # ── checklist_verify (clause_results-wrapped) ──
        ver_list = clause_data.get("verification_results", [])
        if ver_list:
            _render_verification_results_table(doc, ver_list, lang=lang)
            continue

        # ── remediation (clause_results-wrapped, unlikely but handled) ──
        remed = clause_data.get("remediation", {})
        if remed and isinstance(remed, dict):
            _render_remediation_block(doc, remed, lang=lang)


def _render_verification_results_table(doc, ver_list: list, lang: str = "zh-TW") -> None:
    """Render checklist_verify verification_results[] as a Word table."""
    _ADEQUACY_ICON = {"full": "✓", "partial": "△", "irrelevant": "—", "not_found": "✗"}
    tbl = doc.add_table(rows=1 + len(ver_list), cols=4)
    tbl.style = "Table Grid"
    hdr_cells = tbl.rows[0].cells
    _vr_hdrs = {
        "zh": ["充分性", "證據名稱", "相關度", "說明"],
        "en": ["Adequacy", "Evidence Name", "Relevance", "Explanation"],
        "ja": ["充足性", "証拠名称", "関連性", "説明"],
    }
    _vr_lk = "ja" if str(lang or "").startswith("ja") else "en" if str(lang or "").startswith("en") else "zh"
    for ci, h in enumerate(_vr_hdrs[_vr_lk]):
        hdr_cells[ci].text = h
        run = hdr_cells[ci].paragraphs[0].runs[0] if hdr_cells[ci].paragraphs[0].runs else hdr_cells[ci].paragraphs[0].add_run(h)
        run.bold = True
    for ei, vr in enumerate(ver_list, 1):
        adequacy = vr.get("adequacy", "")
        icon = _ADEQUACY_ICON.get(adequacy, adequacy)
        score = vr.get("semantic_score", "")
        score_str = f"{score:.2f}" if isinstance(score, float) else (str(score) if score != "" else "")
        explanation = vr.get("explanation", "") or ""
        tbl.rows[ei].cells[0].text = icon
        tbl.rows[ei].cells[1].text = vr.get("evidence_name", "")
        tbl.rows[ei].cells[2].text = score_str
        tbl.rows[ei].cells[3].text = explanation[:500] + ("…" if len(explanation) > 500 else "")


def _render_remediation_block(doc, remed: dict, lang: str = "zh-TW") -> None:
    """Render a remediation dict as structured Word content."""
    summary = _safe_str(remed.get("summary")).strip()
    priority = _safe_str(remed.get("priority")).strip()
    regulation_citation = _safe_str(remed.get("regulation_citation")).strip()
    suggestions = remed.get("suggestions") or []

    if summary or priority:
        p = doc.add_paragraph()
        if priority:
            p.add_run(f"優先級 / Priority: {priority}　").bold = True
        if summary:
            p.add_run(summary)

    for idx, sg in enumerate(suggestions, 1):
        if not isinstance(sg, dict):
            doc.add_paragraph(str(sg)[:500])
            continue
        doc.add_paragraph(f"建議 {idx} / Suggestion {idx}", style="Heading 4" if "Heading 4" in [s.name for s in doc.styles] else None)
        tbl = doc.add_table(rows=4, cols=2)
        tbl.style = "Table Grid"
        _rem_lbls = {
            "zh": ["具體動作", "目標章節", "法規依據", "範例內容"],
            "en": ["Action", "Target Section", "Regulation Basis", "Example Content"],
            "ja": ["具体的なアクション", "対象セクション", "法規依拠", "例示内容"],
        }
        _rem_lk = "ja" if str(lang or "").startswith("ja") else "en" if str(lang or "").startswith("en") else "zh"
        labels = _rem_lbls[_rem_lk]
        keys = ["action", "target_section", "regulation_basis", "example_content"]
        for ri2, (lbl, key) in enumerate(zip(labels, keys)):
            val = _safe_str(sg.get(key)).strip()
            tbl.rows[ri2].cells[0].text = lbl
            run0 = tbl.rows[ri2].cells[0].paragraphs[0].runs[0] if tbl.rows[ri2].cells[0].paragraphs[0].runs else tbl.rows[ri2].cells[0].paragraphs[0].add_run(lbl)
            run0.bold = True
            tbl.rows[ri2].cells[1].text = val

    if regulation_citation:
        p = doc.add_paragraph()
        p.add_run("法規引用 / Regulation Citation: ").bold = True
        p.add_run(regulation_citation[:800] + ("…" if len(regulation_citation) > 800 else ""))


def _render_p5_role_block(doc, label: str, data: dict, fill_hex: str = "FFFFFF") -> None:
    """Render a P5 Analyzer or Verifier block as a single-row Word table for visual distinction."""
    from docx.shared import Pt, RGBColor
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    _shade_cell(cell, fill_hex)
    p = cell.paragraphs[0]
    label_run = p.add_run(f"{label}\n")
    label_run.bold = True
    label_run.font.size = Pt(10)

    parts = []
    position = (data.get("position") or data.get("overall_assessment") or "")
    if position:
        parts.append(("立場 / Position", str(position)[:800]))
    key_evidence = data.get("key_evidence") or []
    if key_evidence:
        ev_lines = "\n".join(f"  • {str(e)[:200]}" for e in key_evidence[:5])
        parts.append(("▸ 關鍵證據 / Key Evidence", ev_lines))
    confidence = data.get("confidence") or data.get("confidence_score") or ""
    if confidence:
        parts.append(("信心度 / Confidence", str(confidence)))
    challenges = data.get("challenges") or []
    if challenges:
        ch_lines = []
        for ch in challenges[:3]:
            if isinstance(ch, dict):
                ch_lines.append(f"  ⚠ {str(ch.get('point', ch))[:200]}")
            else:
                ch_lines.append(f"  ⚠ {str(ch)[:200]}")
        parts.append(("⚠ 質疑 / Challenges", "\n".join(ch_lines)))
    weaknesses = data.get("acknowledged_weaknesses") or []
    if weaknesses:
        wk_lines = "\n".join(f"  • {str(w)[:200]}" for w in weaknesses[:3])
        parts.append(("已承認弱點 / Weaknesses", wk_lines))
    agreement = data.get("agreement_level") or ""
    if agreement:
        parts.append(("Agreement", agreement))

    for field_label, val in parts:
        cell.add_paragraph("")
        sub_p = cell.add_paragraph()
        sub_p.add_run(f"{field_label}: ").bold = True
        sub_p.add_run(val)

    doc.add_paragraph("")


def _render_role_content(paragraph, data: dict) -> None:
    """Render Analyzer/Verifier structured data as readable text in a Word paragraph.

    Includes structured display of Verifier challenges with:
      - 質疑要點 (point)
      - 法規依據 (regulation_basis)
      - ▸ 期望看到 (expected_evidence)   ← "expected" side
      - ⚠ 風險影響 (worst_case_impact)
    Analyzer key_evidence renders as "▸ 實際看到" for the "actual" side.
    """
    if not data or not isinstance(data, dict):
        paragraph.add_run("（無資料）")
        return

    parts = []
    for key in (
        "position",
        "assessment",
        "confidence",
        "confidence_score",
        "evidence",
        "evidence_cited",
        "key_evidence",
        "reasoning",
        "analysis",
        "agreement_level",
        "concerns",
        "recommendation",
        "acknowledged_weaknesses",
        "response",
    ):
        val = data.get(key)
        if val is None:
            continue
        if isinstance(val, list):
            label = "▸ 實際看到" if key == "key_evidence" else key
            parts.append(f"{label}:")
            for item in val:
                parts.append(f"    • {item}")
        elif key == "response" and isinstance(val, dict):
            # Structured analyzer response: arbitrary numbered section keys
            for sec_title, sec_data in val.items():
                parts.append(f"  【{sec_title}】")
                if isinstance(sec_data, dict):
                    content = (sec_data.get("content") or sec_data.get("analysis")
                               or sec_data.get("回應內容") or sec_data.get("position") or "")
                    if content:
                        parts.append(f"    {str(content)[:600]}{'…' if len(str(content)) > 600 else ''}")
                    reg_ev = sec_data.get("regulation_evidence") or sec_data.get("引用法規") or []
                    if isinstance(reg_ev, list) and reg_ev:
                        parts.append("    ▸ 引用法規:")
                        for re_item in reg_ev[:4]:
                            if isinstance(re_item, dict):
                                parts.append(f"      • [{re_item.get('source','')}] {str(re_item.get('text',''))[:100]}")
                    doc_ev = sec_data.get("document_evidence") or sec_data.get("引用文件") or []
                    if isinstance(doc_ev, list) and doc_ev:
                        parts.append("    ▸ 引用文件:")
                        for de_item in doc_ev[:4]:
                            if isinstance(de_item, dict):
                                parts.append(f"      • {de_item.get('name','')} §{de_item.get('section','')} v{de_item.get('version','')}")
                elif isinstance(sec_data, str):
                    parts.append(f"    {sec_data[:400]}")
        elif key == "response" and isinstance(val, str):
            # Try to parse as JSON, otherwise show plain
            try:
                parsed_resp = json.loads(val.strip().lstrip("```json").lstrip("```").rstrip("```").strip())
                if isinstance(parsed_resp, dict):
                    for sec_title, sec_data in parsed_resp.items():
                        parts.append(f"  【{sec_title}】")
                        if isinstance(sec_data, dict):
                            content = (sec_data.get("content") or sec_data.get("analysis")
                                       or sec_data.get("回應內容") or "")
                            if content:
                                parts.append(f"    {str(content)[:600]}{'…' if len(str(content)) > 600 else ''}")
                        elif isinstance(sec_data, str):
                            parts.append(f"    {sec_data[:400]}")
                else:
                    parts.append(f"response: {val[:2000]}")
            except (json.JSONDecodeError, ValueError):
                parts.append(f"response: {val[:2000]}")
        elif isinstance(val, dict):
            parts.append(f"{key}: {json.dumps(val, ensure_ascii=False)[:500]}")
        else:
            parts.append(f"{key}: {val}")

    # Render Verifier challenges in structured "期望 vs 實際" format
    challenges = data.get("challenges", [])
    if challenges and isinstance(challenges, list):
        parts.append("─── 驗證者質疑 / Verifier Challenges ───")
        for i, ch in enumerate(challenges, 1):
            if isinstance(ch, dict):
                parts.append(f"  [{i}] {ch.get('point', '')}")
                reg_basis = ch.get("regulation_basis", "")
                if reg_basis:
                    parts.append(f"      法規依據: {reg_basis}")
                exp_ev = ch.get("expected_evidence", "")
                if exp_ev:
                    if isinstance(exp_ev, list):
                        parts.append("      ▸ 期望看到:")
                        for ei in exp_ev[:4]:
                            parts.append(f"        • {str(ei)[:200]}")
                    else:
                        parts.append(f"      ▸ 期望看到: {str(exp_ev)[:300]}")
                wci = ch.get("worst_case_impact", "")
                if wci:
                    if isinstance(wci, dict):
                        scenario = wci.get("scenario", "") or wci.get("worst_case", "")
                        sev = wci.get("severity", "")
                        conseq = wci.get("regulatory_consequences", [])
                        if scenario:
                            parts.append(f"      ⚠ 情境: {str(scenario)[:200]}")
                        if sev:
                            parts.append(f"        嚴重度: {sev}")
                        if isinstance(conseq, list) and conseq:
                            parts.append("        法規後果:")
                            for c in conseq[:3]:
                                parts.append(f"          • {str(c)[:150]}")
                    else:
                        parts.append(f"      ⚠ 風險影響: {str(wci)[:200]}")
            else:
                parts.append(f"  [{i}] {ch}")

    remaining = data.get("remaining_concerns", [])
    if remaining and isinstance(remaining, list):
        parts.append("─── 未解疑慮 / Remaining Concerns ───")
        for c in remaining:
            if isinstance(c, dict):
                point = c.get("point", "")
                reg_basis = c.get("regulation_basis", "")
                exp_ev = c.get("expected_evidence", "")
                wci = c.get("worst_case_impact", "")
                if point:
                    parts.append(f"  • {str(point)[:500]}{'…' if len(str(point)) > 500 else ''}")
                if reg_basis:
                    parts.append(f"    法規依據: {str(reg_basis)[:200]}")
                if exp_ev:
                    if isinstance(exp_ev, list):
                        parts.append("    ▸ 期望證據:")
                        for ei in exp_ev[:4]:
                            parts.append(f"      • {str(ei)[:200]}")
                    else:
                        parts.append(f"    ▸ 期望證據: {str(exp_ev)[:300]}{'…' if len(str(exp_ev)) > 300 else ''}")
                if wci:
                    if isinstance(wci, dict):
                        scenario = wci.get("scenario", "") or wci.get("worst_case", "")
                        sev = wci.get("severity", "")
                        conseq = wci.get("regulatory_consequences", [])
                        if scenario:
                            parts.append(f"    ⚠ 情境: {str(scenario)[:200]}")
                        if sev:
                            parts.append(f"      嚴重度: {sev}")
                        if isinstance(conseq, list) and conseq:
                            parts.append("      法規後果:")
                            for c in conseq[:3]:
                                parts.append(f"        • {str(c)[:150]}")
                    else:
                        parts.append(f"    ⚠ 風險影響: {str(wci)[:200]}")
            else:
                parts.append(f"  • {c}")

    if parts:
        paragraph.add_run("\n".join(parts)[:4000])
    else:
        # Fallback: render as readable key-value pairs
        kv_lines = [f"{k}: {v}" for k, v in data.items() if v is not None]
        paragraph.add_run("\n".join(kv_lines)[:4000] if kv_lines else "（無內容）")


def _flatten_role_text(data: dict, primary_key: str) -> str:
    """Extract primary field from Analyzer/Verifier data as a readable string."""
    if not data or not isinstance(data, dict):
        return ""
    val = data.get(primary_key, "")
    if isinstance(val, list):
        return "; ".join(str(v) for v in val)
    if isinstance(val, dict):
        return "; ".join(f"{k}: {v}" for k, v in val.items())
    return str(val) if val else "; ".join(f"{k}: {v}" for k, v in data.items() if v is not None)
