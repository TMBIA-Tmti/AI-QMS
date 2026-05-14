"""
AI-QMS — Regulatory Update Report Export Module
================================================

Export regulatory crawl update results to Word (.docx) and Excel (.xlsx) formats.
Follows the same patterns as regulatory_export.py for consistency.
"""

import json
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side


from src.utils.regulatory_export import _t, _tl, _source_label as _source_label
import re as _re


def _region_display(region_key: str, lang: str) -> str:
    """Extract English name for non-zh languages: '美國 (USA)' → 'USA'"""
    if lang.startswith("zh"):
        return region_key
    m = _re.search(r'\(([^)]+)\)', region_key)
    return m.group(1) if m else region_key


# Output directory for generated files — use absolute path so cl.File always resolves correctly
EXPORT_DIR = (Path(__file__).resolve().parent.parent.parent / "data" / "exports")
EXPORT_DIR.mkdir(parents=True, exist_ok=True)

# Shared styles (same as regulatory_export.py)
HEADER_FILL = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
HEADER_FONT = Font(name="Microsoft JhengHei", size=10, bold=True, color="FFFFFF")
CELL_FONT = Font(name="Microsoft JhengHei", size=9)
THIN_BORDER = Border(
    left=Side(style="thin"),
    right=Side(style="thin"),
    top=Side(style="thin"),
    bottom=Side(style="thin"),
)

# Status colors for Excel
SUCCESS_FILL = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
FAIL_FILL = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")


# ============================================================
# Helper Functions
# ============================================================


def _strip_markdown(text: str) -> str:
    """Strip markdown formatting for plain text preview."""
    text = re.sub(r"#{1,6}\s+", "", text)  # headers
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)  # bold
    text = re.sub(r"\*([^*]+)\*", r"\1", text)  # italic
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)  # links
    text = re.sub(r"[-*+]\s+", "", text)  # list markers
    text = re.sub(r"\n+", " ", text)  # newlines
    return text.strip()


def _get_qms_mapping(agency: str, region: str = "") -> str:
    """Map agency/region to related QMS standards/regulations."""
    mapping = {
        "FDA": "21 CFR 820 (QSR/QMSR), FDA Guidance",
        "eCFR": "21 CFR Part 820, Part 11",
        "TFDA": "醫療器材管理法, GMP",
        "PMDA": "QMS省令 (MHLW Ordinance 169), JIS T 2304",
        "NMPA": "醫療器械生產質量管理規範",
        "TGA": "Australian Medical Device Standards",
        "MFDS": "KGMP (Korean GMP)",
        "HSA": "Singapore Medical Device GMP",
        "CDSCO": "India MDR 2017",
        "ANVISA": "Brazil RDC 665/2022",
        "Health_Canada": "CMDR (SOR/98-282), ISO 13485",
        "Swissmedic": "MepV (Swiss Medical Devices Ordinance)",
        "COFEPRIS": "Mexico NOM Standards",
        "EMA": "EU MDR 2017/745, EU IVDR 2017/746",
        "MDCG": "EU MDR Guidance Documents",
        "BSI": "EU MDR, UKCA Marking",
        "ISO": "ISO 13485, ISO 14971, IEC 62304",
        "EUR-Lex": "EU MDR/IVDR Regulations",
    }
    for key, val in mapping.items():
        if key.upper() in agency.upper():
            return val
    # Generic fallback based on region
    if any(kw in region.upper() for kw in ["EU", "歐盟", "欧盟", "EUROPE"]):
        return "EU MDR 2017/745"
    return "ISO 13485"


# ============================================================
# Markdown Format (for Chainlit chat display)
# ============================================================


def format_regulatory_update_markdown(
    crawl_results: dict,
    assessment: Optional[str] = None,
    lang: str = "zh-TW",
) -> str:
    """Format crawl results as Markdown for Chainlit chat display."""
    results = crawl_results.get("results", [])
    summary = crawl_results.get("summary", {})

    total = summary.get("total_sites", 0)
    success = summary.get("success_count", 0)

    duration = summary.get("crawl_duration_seconds", 0)
    regions = summary.get("regions_covered", [])

    md_headers = _tl("regulatory_update_export.table_headers_md", lang)

    lines = [
        _t(
            "regulatory_update_export.result_title",
            lang,
            success=success,
            total=total,
            duration=f"{duration:.1f}",
        )
        + "\n",
        _t("regulatory_update_export.regions_covered", lang, regions=", ".join(_region_display(r, lang) for r in regions))
        + "\n",
        f"### {_t('regulatory_update_export.summary_table', lang)}\n",
        f"| {' | '.join(md_headers)} |",
        f"|{'|'.join(['------' for _ in md_headers])}|",
    ]

    for r in results:
        region = r.get("region", "")
        agency = r.get("agency", "")
        status = "✅" if r.get("crawl_status") == "success" else "❌"
        qms = _get_qms_mapping(agency, region)

        if r.get("crawl_status") == "success":
            content = r.get("content_markdown", "")
            if content:
                # Strip markdown and skip the first line if it's a region header
                stripped = _strip_markdown(content)
                lines = stripped.splitlines()
                # Skip lines that look like "RegionName — AgencyName" headers
                import re as _re2
                _header_pat = _re2.compile(r'^[\w\s\u4e00-\u9fff\uff08\uff09（）()]+\s*[—–-]\s*\w', _re2.UNICODE)
                while lines and _header_pat.match(lines[0].strip()):
                    lines.pop(0)
                preview = " ".join(lines)[:50] if lines else ""
            else:
                preview = ""
        elif r.get("crawl_status") == "failed":
            reason = r.get(
                "failure_reason", _t("regulatory_update_export.unknown_reason", lang)
            )
            preview = reason[:50] + "..." if len(reason) > 50 else reason
        else:
            preview = r.get("note", "")[:50]

        lines.append(f"| {_region_display(region, lang)} | {agency} | {status} | {preview} | {qms[:30]} |")

    # Failed sites section
    failed_results = [r for r in results if r.get("crawl_status") == "failed"]
    if failed_results:
        lines.append(f"\n### {_t('regulatory_update_export.failed_sites', lang)}\n")
        for r in failed_results:
            reason = r.get(
                "failure_reason", _t("regulatory_update_export.unknown_reason", lang)
            )
            lines.append(f"- **{_region_display(r['region'], lang)} — {r['agency']}**: {reason}")
        lines.append("")

    # Assessment section
    if assessment:
        lines.append("\n---\n")
        lines.append(f"### {_t('regulatory_export.assessment_report', lang)}\n")
        lines.append(assessment)

    return "\n".join(lines)


# ============================================================
# Word Export
# ============================================================


def export_regulatory_update_to_word(
    crawl_results: dict,
    assessment: Optional[str] = None,
    verification_report: Optional[dict] = None,
    lang: str = "zh-TW",
    source_command: str = "regulatory_update",
) -> str:
    """Export regulatory update results to Word (.docx).

    Returns:
        Path to the generated .docx file.
    """
    results = crawl_results.get("results", [])
    summary = crawl_results.get("summary", {})

    total = summary.get("total_sites", 0)
    success = summary.get("success_count", 0)
    failed = summary.get("failed_count", 0)
    duration = summary.get("crawl_duration_seconds", 0)

    src_label = _source_label(source_command, lang)
    doc = Document()

    # Title
    title = doc.add_heading(
        f"{_t('regulatory_update_export.title', lang)}（{src_label}）", level=1
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = meta.add_run(
        f"{_t('regulatory_export.export_time', lang)}: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | "
        f"{_t('source_label.source', lang)}: {src_label}"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run2 = meta2.add_run(
        _t(
            "regulatory_update_export.meta_stats",
            lang,
            total=total,
            success=success,
            failed=failed,
            duration=f"{duration:.1f}",
        )
    )
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    # ── Abbreviation Legend (language-aware) ──
    if lang.startswith("zh"):
        abbrev_heading = "縮寫說明"
        abbrev_body = (
            "FDA / eCFR  — 美國食品藥物管理局 / 聯邦法規電子資料庫（21 CFR 820 QSR/QMSR）\n"
            "TFDA        — 臺灣食品藥物管理署（醫療器材管理法, GMP）\n"
            "EMA / MDCG  — 歐洲藥品管理局 / 醫療器材協調小組（EU MDR 2017/745, EU IVDR 2017/746）\n"
            "PMDA        — 日本藥品與醫療器材局（QMS省令, JIS T 2304）\n"
            "NMPA        — 中國國家藥品監督管理局（醫療器械生產質量管理規範）\n"
            "TGA         — 澳洲治療用品管理局（Australian Medical Device Standards）\n"
            "MFDS        — 韓國食品藥品安全處（KGMP）\n"
            "HSA         — 新加坡衛生科學局（Singapore Medical Device GMP）\n"
            "CDSCO       — 印度中央藥品標準控制組織（India MDR 2017）\n"
            "ANVISA      — 巴西國家衛生監督局（Brazil RDC 665/2022）\n"
            "Health Canada — 加拿大衛生部（CMDR SOR/98-282, ISO 13485）\n"
            "MDSAP       — 醫療器材單一稽核計畫（美國/加拿大/巴西/澳洲/日本）\n"
            "QMS         — 品質管理系統\n"
            "RA          — 法規事務"
        )
        guide_heading = "報告欄位說明"
        guide_body = (
            "【彙總更新清單欄位說明】\n"
            "  地區：法規機構所在地區，如 US / EU / TW / JP / CN\n"
            "  機構：法規發布機構代碼，如 FDA / TFDA / EMA\n"
            "  爬取狀態：success 表示成功取得更新；failed 表示無法存取\n"
            "  新增項目數：本次爬取中偵測到的新發布文件或指引數量\n"
            "  更新類型：regulation（法規本文）/ guidance（指引）/ notice（公告）\n"
            "  相關性：AI 評估與本 QMS 系統的相關程度（high / medium / low）\n"
            "  對應 QMS 標準：該機構發布的內容所對應的主要 QMS 標準\n\n"
            "【爬取結果作用原理】\n"
            "  系統定期從各國法規機構官網爬取最新公告，比對既有紀錄偵測新增內容，\n"
            "  並由 AI 評估每筆更新與現行 QMS 文件的相關性，協助 RA 人員快速識別需回應的法規變動。\n\n"
            "【內容摘要】：爬取內容的前 100 字預覽\n"
            "【儲存路徑】：成功爬取的原始文件儲存於 regulatory_markdown_storage 的路徑"
        )
    elif lang.startswith("ja"):
        abbrev_heading = "略語一覧"
        abbrev_body = (
            "FDA / eCFR  — 米国食品医薬品局 / 連邦規則電子データベース（21 CFR 820 QSR/QMSR）\n"
            "TFDA        — 台湾食品薬物管理署（医療機器管理法, GMP）\n"
            "EMA / MDCG  — 欧州医薬品庁 / 医療機器調整グループ（EU MDR 2017/745, EU IVDR 2017/746）\n"
            "PMDA        — 医薬品医療機器総合機構（QMS省令, JIS T 2304）\n"
            "NMPA        — 中国国家薬品監督管理局（医療機器生産品質管理基準）\n"
            "TGA         — オーストラリア医薬品管理局（Australian Medical Device Standards）\n"
            "MFDS        — 韓国食品医薬品安全処（KGMP）\n"
            "HSA         — シンガポール保健科学局（Singapore Medical Device GMP）\n"
            "CDSCO       — インド中央医薬品標準管理機構（India MDR 2017）\n"
            "ANVISA      — ブラジル国家衛生監督局（Brazil RDC 665/2022）\n"
            "Health Canada — カナダ保健省（CMDR SOR/98-282, ISO 13485）\n"
            "MDSAP       — 医療機器単一審査プログラム（米国/カナダ/ブラジル/オーストラリア/日本）\n"
            "QMS         — 品質マネジメントシステム\n"
            "RA          — 薬事"
        )
        guide_heading = "レポートフィールドガイド"
        guide_body = (
            "【集計更新一覧フィールド説明】\n"
            "  地域：法規機関の所在地域（US / EU / TW / JP / CN など）\n"
            "  機関：法規発行機関コード（FDA / TFDA / EMA など）\n"
            "  クロール状態：success = 更新取得成功、failed = アクセス不可\n"
            "  新規項目数：今回のクロールで検出された新規文書・ガイダンスの数\n"
            "  更新種別：regulation（法規本文）/ guidance（ガイダンス）/ notice（公告）\n"
            "  関連性：本 QMS システムとの関連度の AI 評価（high / medium / low）\n"
            "  対応 QMS 標準：当該機関が発行する内容に対応する主要 QMS 標準\n\n"
            "【クロール結果の仕組み】\n"
            "  システムは各国規制機関の公式サイトから最新公告を定期クロールし、既存記録と比較して新着内容を検出します。\n"
            "  AI が各更新と現行 QMS 文書との関連性を評価し、RA 担当者が対応すべき規制変更を迅速に特定できるよう支援します。\n\n"
            "【内容サマリー】：クロール内容の最初の 100 文字プレビュー\n"
            "【保存パス】：クロール成功した原文の regulatory_markdown_storage 内保存パス"
        )
    else:
        abbrev_heading = "Abbreviation Legend"
        abbrev_body = (
            "FDA / eCFR  — U.S. Food & Drug Administration / Electronic Code of Federal Regulations (21 CFR 820 QSR/QMSR)\n"
            "TFDA        — Taiwan Food and Drug Administration (Medical Devices Act, GMP)\n"
            "EMA / MDCG  — European Medicines Agency / Medical Device Coordination Group (EU MDR 2017/745, EU IVDR 2017/746)\n"
            "PMDA        — Japan Pharmaceuticals and Medical Devices Agency (QMS Ordinance, JIS T 2304)\n"
            "NMPA        — China National Medical Products Administration (Medical Device GMP)\n"
            "TGA         — Australia Therapeutic Goods Administration (Australian Medical Device Standards)\n"
            "MFDS        — Korea Ministry of Food and Drug Safety (KGMP)\n"
            "HSA         — Singapore Health Sciences Authority (Singapore Medical Device GMP)\n"
            "CDSCO       — India Central Drugs Standard Control Organisation (India MDR 2017)\n"
            "ANVISA      — Brazil National Health Surveillance Agency (Brazil RDC 665/2022)\n"
            "Health Canada — Canada (CMDR SOR/98-282, ISO 13485)\n"
            "MDSAP       — Medical Device Single Audit Program (US/Canada/Brazil/Australia/Japan)\n"
            "QMS         — Quality Management System\n"
            "RA          — Regulatory Affairs"
        )
        guide_heading = "Report Field Guide"
        guide_body = (
            "[Summary Table Fields]\n"
            "  Region: Regulatory agency jurisdiction (US / EU / TW / JP / CN etc.)\n"
            "  Agency: Regulatory body code (FDA / TFDA / EMA etc.)\n"
            "  Crawl Status: success = update retrieved; failed = access failed\n"
            "  New Items: Number of newly published documents/guidance detected this crawl\n"
            "  Update Type: regulation / guidance / notice\n"
            "  Relevance: AI-assessed relevance to this QMS (high / medium / low)\n"
            "  QMS Mapping: Primary QMS standards corresponding to the agency's content\n\n"
            "[How Crawl Results Work]\n"
            "  The system periodically crawls official regulatory agency websites, detects new content\n"
            "  by comparing against existing records, and uses AI to assess each update's relevance\n"
            "  to the current QMS documents — helping RA teams quickly identify regulatory changes.\n\n"
            "[Content Summary]: First 100-character preview of crawled content\n"
            "[Storage Path]: Path within regulatory_markdown_storage where crawled content is saved"
        )
    doc.add_heading(abbrev_heading, level=2)
    doc.add_paragraph(abbrev_body)

    doc.add_heading(guide_heading, level=2)
    doc.add_paragraph(guide_body)

    # Section 1: Summary Table (enhanced with structured columns)
    doc.add_heading(_t("regulatory_update_export.summary_heading", lang), level=2)

    table1 = doc.add_table(rows=1, cols=7)
    table1.style = "Table Grid"
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = _tl("regulatory_update_export.table_headers", lang)
    for i, header in enumerate(headers):
        cell = table1.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)

    for r in results:
        row = table1.add_row()
        region = r.get("region", "")
        agency = r.get("agency", "")
        status_text = (
            _t("regulatory_update_export.crawl_status_success", lang)
            if r.get("crawl_status") == "success"
            else _t("regulatory_update_export.crawl_status_fail", lang)
        )

        # Content summary
        content = r.get("content_markdown", "")
        content_summary = (
            _strip_markdown(content)[:100]
            if content
            else (r.get("failure_reason", "")[:100])
        )

        # Storage path
        storage_path = (
            f"regulatory_markdown_storage/documents/{region}/{agency}_*.md"
            if r.get("crawl_status") == "success"
            else ""
        )

        # QMS mapping
        qms = _get_qms_mapping(agency, region)

        values = [
            _region_display(region, lang),
            agency,
            r.get("url", "")[:80],
            status_text,
            content_summary,
            storage_path,
            qms,
        ]
        for i, val in enumerate(values):
            cell = row.cells[i]
            cell.text = val
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    widths = [Cm(2), Cm(2), Cm(3.5), Cm(1.2), Cm(3.5), Cm(2.5), Cm(3)]
    for row in table1.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = width

    doc.add_paragraph()

    # Section 2: Content Details
    success_results = [r for r in results if r.get("crawl_status") == "success"]
    if success_results:
        doc.add_heading(
            _t("regulatory_update_export.content_detail_heading", lang), level=2
        )

        for r in success_results:
            doc.add_heading(
                f"{_region_display(r.get('region', ''), lang)} — {r.get('agency', '')}",
                level=3,
            )
            content = r.get("content_markdown") or ""
            for line in content.split("\n"):
                stripped_line = line.strip()
                if not stripped_line:
                    doc.add_paragraph()
                    continue
                p = doc.add_paragraph(stripped_line)
                for run in p.runs:
                    run.font.size = Pt(8)
            doc.add_paragraph()

    # Section 3: Assessment Report
    if assessment:
        doc.add_heading(_t("regulatory_export.assessment_heading", lang), level=2)
        # Split assessment into paragraphs
        for para_text in assessment.split("\n"):
            stripped = para_text.strip()
            if not stripped:
                doc.add_paragraph()
                continue
            if stripped.startswith("###"):
                doc.add_heading(stripped.lstrip("#").strip(), level=4)
            elif stripped.startswith("##"):
                doc.add_heading(stripped.lstrip("#").strip(), level=3)
            elif stripped.startswith("#"):
                doc.add_heading(stripped.lstrip("#").strip(), level=2)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                p = doc.add_paragraph(stripped[2:], style="List Bullet")
                for run in p.runs:
                    run.font.size = Pt(9)
            elif stripped.startswith(tuple(f"{i}." for i in range(1, 20))):
                p = doc.add_paragraph(stripped, style="List Number")
                for run in p.runs:
                    run.font.size = Pt(9)
            else:
                p = doc.add_paragraph(stripped)
                for run in p.runs:
                    run.font.size = Pt(9)

    # Section: Verification Report (if provided)
    if verification_report and verification_report.get("has_data"):
        # Reuse the shared verification renderer from regulatory_export
        from src.utils.regulatory_export import _render_verification_to_word

        _render_verification_to_word(doc, verification_report, lang)

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    run = footer.add_run(
        f"{_t('regulatory_update_export.footer', lang)} | "
        f"{_t('source_label.source', lang)}: {src_label}"
    )
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Save
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    cmd_tag = "list" if source_command == "regulatory_list" else "update"
    filename = f"regulatory_update_{cmd_tag}_{timestamp}.docx"
    filepath = EXPORT_DIR / filename
    doc.save(str(filepath))
    return str(filepath)


# ============================================================
# Excel Export
# ============================================================


def export_regulatory_update_to_excel(
    crawl_results: dict,
    assessment: Optional[str] = None,
    verification_report: Optional[dict] = None,
    lang: str = "zh-TW",
    source_command: str = "regulatory_update",
) -> str:
    """Export regulatory update results to Excel (.xlsx).

    Returns:
        Path to the generated .xlsx file.
    """
    results = crawl_results.get("results", [])
    summary = crawl_results.get("summary", {})

    total = summary.get("total_sites", 0)
    success = summary.get("success_count", 0)
    failed = summary.get("failed_count", 0)
    duration = summary.get("crawl_duration_seconds", 0)

    src_label = _source_label(source_command, lang)
    wb = Workbook()

    # Sheet 1: Summary (enhanced with structured columns)
    ws1 = wb.active
    ws1.title = _t("regulatory_update_export.summary_sheet", lang)

    # Title
    ws1.merge_cells("A1:G1")
    title_cell = ws1.cell(row=1, column=1)
    title_cell.value = f"{_t('regulatory_update_export.title', lang)}（{src_label}）"
    title_cell.font = Font(name="Microsoft JhengHei", size=14, bold=True)
    title_cell.alignment = Alignment(horizontal="center")

    # Metadata
    ws1.merge_cells("A2:G2")
    meta_cell = ws1.cell(row=2, column=1)
    meta_cell.value = (
        f"{_t('regulatory_export.export_time', lang)}: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | "
        f"{_t('source_label.source', lang)}: {src_label} | "
        + _t(
            "regulatory_update_export.meta_stats",
            lang,
            total=total,
            success=success,
            failed=failed,
            duration=f"{duration:.1f}",
        )
    )
    meta_cell.font = Font(
        name="Microsoft JhengHei", size=9, italic=True, color="808080"
    )

    # Headers (enhanced)
    headers = _tl("regulatory_update_export.table_headers", lang)
    for col, header in enumerate(headers, 1):
        cell = ws1.cell(row=4, column=col)
        cell.value = header
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = THIN_BORDER

    # Data
    for row_idx, r in enumerate(results, 5):
        region = r.get("region", "")
        agency = r.get("agency", "")
        status_text = (
            _t("regulatory_update_export.crawl_status_success", lang)
            if r.get("crawl_status") == "success"
            else _t("regulatory_update_export.crawl_status_fail", lang)
        )

        # Content summary
        content = r.get("content_markdown", "")
        content_summary = (
            _strip_markdown(content)
            if content
            else r.get("failure_reason", "")
        )

        # Storage path
        storage_path = (
            f"regulatory_markdown_storage/documents/{region}/{agency}_*.md"
            if r.get("crawl_status") == "success"
            else ""
        )

        # QMS mapping
        qms = _get_qms_mapping(agency, region)

        values = [
            _region_display(region, lang),
            agency,
            r.get("url", ""),
            status_text,
            content_summary,
            storage_path,
            qms,
        ]
        for col, val in enumerate(values, 1):
            cell = ws1.cell(row=row_idx, column=col)
            cell.value = val
            cell.font = CELL_FONT
            cell.border = THIN_BORDER

        # Color status cell
        status_cell = ws1.cell(row=row_idx, column=4)
        if status_text == _t("regulatory_update_export.crawl_status_success", lang):
            status_cell.fill = SUCCESS_FILL
        else:
            status_cell.fill = FAIL_FILL

    ws1.column_dimensions["A"].width = 15
    ws1.column_dimensions["B"].width = 15
    ws1.column_dimensions["C"].width = 40
    ws1.column_dimensions["D"].width = 8
    ws1.column_dimensions["E"].width = 40
    ws1.column_dimensions["F"].width = 25
    ws1.column_dimensions["G"].width = 30
    ws1.freeze_panes = "A5"

    # Sheet 2: Assessment Report (if provided)
    if assessment:
        from src.utils.regulatory_export import _render_assessment_to_excel

        _render_assessment_to_excel(wb, assessment, lang)

    # Sheet: Verification Report (if provided)
    if verification_report and verification_report.get("has_data"):
        from src.utils.regulatory_export import _render_verification_to_excel

        _render_verification_to_excel(wb, verification_report, lang)

    # Footer note
    note_row = len(results) + 6
    ws1.merge_cells(f"A{note_row}:G{note_row}")
    note_cell = ws1.cell(row=note_row, column=1)
    note_cell.value = (
        f"{_t('regulatory_update_export.footer', lang)} | "
        f"{_t('source_label.source', lang)}: {src_label}"
    )
    note_cell.font = Font(
        name="Microsoft JhengHei", size=8, italic=True, color="808080"
    )

    # Save
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    cmd_tag = "list" if source_command == "regulatory_list" else "update"
    filename = f"regulatory_update_{cmd_tag}_{timestamp}.xlsx"
    filepath = EXPORT_DIR / filename
    wb.save(str(filepath))
    return str(filepath)
