"""
AI-QMS — Cross-Reference Table Export Helpers
==============================================

Appends the N-country x ISO 13485 cross-reference table as a section
to an existing python-docx Document or openpyxl Workbook.

Two public functions:
  - append_crossref_table_word(doc, lang, show_original_text)
  - append_crossref_table_excel(wb, lang, show_original_text)

Both functions support zh-TW / ja-JP / en-US. Cell fill color encodes status
(green=full, orange=partial, blue=exceeds, grey=na). Unique requirements sheet
includes original-language regulatory text and flags within-clause deltas.
"""

from __future__ import annotations

import json
import logging
import os
from typing import Optional

logger = logging.getLogger(__name__)

__all__ = [
    "append_crossref_table_word",
    "append_crossref_table_excel",
]


# ── i18n helper (same pattern as regulatory_export.py) ──


def _t(key: str, lang: str = "zh-TW", **kwargs) -> str:
    _cache = getattr(_t, "_cache", {})
    if lang not in _cache:
        locale_path = os.path.join(
            os.path.dirname(__file__), "..", "chainlit_app", "locales", f"{lang}.json"
        )
        try:
            with open(locale_path, "r", encoding="utf-8") as f:
                _cache[lang] = json.load(f)
        except Exception:
            _cache[lang] = {}
        _t._cache = _cache
    text = _cache.get(lang, {}).get(key, key)
    if kwargs:
        try:
            text = text.format(**kwargs)
        except (KeyError, IndexError):
            pass
    return text


# ── Shared helpers ──


def _get_crossref_data() -> Optional[dict]:
    """Call crossref report generator and return data dict, or None on failure."""
    try:
        from src.analysis.crossref_report import generate_crossref_validation_report
        from src.analysis.compliance_rules import PREDEFINED_REGULATIONS

        reg_ids = [
            rid for rid, p in PREDEFINED_REGULATIONS.items() if p.iso_mapped
        ]
        if not reg_ids:
            return None
        return generate_crossref_validation_report(reg_ids)
    except Exception as exc:
        logger.debug("crossref data unavailable: %s", exc)
        return None


def _lang_key(lang: str) -> str:
    """Normalize lang to 'zh', 'ja', or 'en'."""
    if lang.startswith("zh"):
        return "zh"
    if lang.startswith("ja"):
        return "ja"
    return "en"


def _status_label(status: str, lang: str) -> str:
    """Localized + emoji label for a mapping status."""
    lk = _lang_key(lang)
    # Use \U codepoints (not surrogate pairs) so lxml can encode them
    _labels = {
        "full": {"zh": "\u2705 \u5b8c\u5168\u5c0d\u61c9", "ja": "\u2705 \u5b8c\u5168\u5bfe\u5fdc", "en": "\u2705 Full"},
        "partial": {"zh": "\U0001f536 \u90e8\u5206\u5c0d\u61c9", "ja": "\U0001f536 \u90e8\u5206\u5bfe\u5fdc", "en": "\U0001f536 Partial"},
        "exceeds": {"zh": "\u2b06 \u8d85\u51faISO", "ja": "\u2b06 ISO\u8d85\u904e", "en": "\u2b06 Exceeds"},
        "not_applicable": {"zh": "\u2b1c \u4e0d\u9069\u7528", "ja": "\u2b1c \u5bfe\u8c61\u5916", "en": "\u2b1c N/A"},
        "na": {"zh": "\u2b1c \u4e0d\u9069\u7528", "ja": "\u2b1c \u5bfe\u8c61\u5916", "en": "\u2b1c N/A"},
        "not_mapped": {"zh": "\u2753 \u672a\u6620\u5c04", "ja": "\u2753 \u672a\u30de\u30c3\u30d4\u30f3\u30b0", "en": "\u2753 Not Mapped"},
    }
    entry = _labels.get(status, _labels["not_mapped"])
    return entry.get(lk, entry["en"])


def _impact_emoji(impact: str) -> str:
    # Use \U codepoints (not surrogate pairs) so lxml can encode them
    if impact == "critical":
        return "\U0001f534"
    if impact == "major":
        return "\U0001f7e1"
    if impact == "minor":
        return "\U0001f7e2"
    return "\u26aa"


# ── Status colors ──

_STATUS_COLORS_HEX = {
    "full": "E8F5E9",
    "partial": "FFF3E0",
    "exceeds": "E3F2FD",
    "not_applicable": "F5F5F5",
    "na": "F5F5F5",
    "not_mapped": "FFFFFF",
}

_STATUS_COLORS_EXCEL = {
    "full": "C8E6C9",
    "partial": "FFE0B2",
    "exceeds": "BBDEFB",
    "not_applicable": "F5F5F5",
    "na": "F5F5F5",
    "not_mapped": "FFFFFF",
}

_IMPACT_FILL_EXCEL = {
    "critical": "FFCDD2",
    "major": "FFF9C4",
    "minor": "F1F8E9",
}

# Status merge priority: higher = better coverage
_STATUS_PRIORITY = {
    "exceeds": 5,
    "full": 4,
    "partial": 3,
    "not_applicable": 2,
    "na": 2,
    "not_mapped": 1,
}


def _merge_status(statuses: list) -> str:
    """Return the best status from a list (highest coverage wins)."""
    if not statuses:
        return "not_mapped"
    return max(statuses, key=lambda s: _STATUS_PRIORITY.get(s, 0))


def _build_country_groups(reg_ids: list, country_stats: dict) -> list:
    """Group regulation IDs by country code, merging same-country regs.

    Returns a list of group dicts in original order:
      {country, reg_ids: [...], name_zh, name_en}
    """
    try:
        from src.analysis.compliance_rules import PREDEFINED_REGULATIONS
    except Exception:
        PREDEFINED_REGULATIONS = {}

    seen: dict = {}  # country_key → group
    order: list = []

    for rid in reg_ids:
        profile = PREDEFINED_REGULATIONS.get(rid)
        # Use country code from profile; fall back to reg_id so it gets its own column
        country_key = (profile.country if profile and profile.country else rid)

        if country_key not in seen:
            stats = country_stats.get(rid, {})
            seen[country_key] = {
                "country": country_key,
                "reg_ids": [rid],
                "name_zh": stats.get("country_name_zh") or stats.get("name_zh") or rid,
                "name_en": stats.get("country_name_en") or stats.get("name_en") or rid,
            }
            order.append(country_key)
        else:
            seen[country_key]["reg_ids"].append(rid)

    return [seen[c] for c in order]


def _group_display(group: dict, lang: str) -> str:
    """Display name for a country group."""
    lk = _lang_key(lang)
    if lk == "zh":
        return group.get("name_zh", group.get("country", ""))
    return group.get("name_en", group.get("country", ""))


def _merge_clause_data(clause_entry: dict, reg_ids: list) -> dict:
    """Merge clause data from multiple regulation IDs into one cell's data."""
    all_statuses = []
    all_wcds = []
    total_conf = 0.0
    conf_count = 0

    for rid in reg_ids:
        cdata = clause_entry.get("countries", {}).get(rid, {})
        status = cdata.get("status", "not_mapped")
        all_statuses.append(status)
        conf = cdata.get("confidence", 0.0)
        if conf:
            total_conf += conf
            conf_count += 1
        all_wcds.extend(cdata.get("within_clause_deltas", []))

    return {
        "status": _merge_status(all_statuses),
        "confidence": total_conf / conf_count if conf_count else 0.0,
        "within_clause_deltas": all_wcds,
    }


def _merge_delta_items(delta_items: dict, groups: list) -> list:
    """Merge delta_items by country group.

    Returns a list of (group, reqs) tuples preserving group order.
    """
    result = []
    for group in groups:
        merged_reqs = []
        seen_req_ids: set = set()
        for rid in group["reg_ids"]:
            for req in delta_items.get(rid, []):
                req_id = req.get("req_id", "")
                if req_id and req_id in seen_req_ids:
                    continue
                if req_id:
                    seen_req_ids.add(req_id)
                merged_reqs.append(req)
        if merged_reqs:
            result.append((group, merged_reqs))
    return result


def _clause_title(clause_info: dict, lang: str) -> str:
    """Get localized clause title."""
    lk = _lang_key(lang)
    if lk == "en":
        return clause_info.get("title_en", clause_info.get("title", ""))
    if lk == "ja":
        return clause_info.get("title_ja", clause_info.get("title", ""))
    return clause_info.get("title", "")


def _country_display(stats: dict, lang: str) -> str:
    """Display name for a country."""
    lk = _lang_key(lang)
    if lk == "zh":
        return stats.get("name_zh", stats.get("regulation_id", ""))
    return stats.get("name_en", stats.get("regulation_id", ""))


def _req_title(req: dict, lang: str) -> str:
    lk = _lang_key(lang)
    if lk == "zh":
        return req.get("title_zh", req.get("title_en", ""))
    if lk == "ja":
        return req.get("title_ja", req.get("title_en", req.get("title_zh", "")))
    return req.get("title_en", req.get("title_zh", ""))


def _req_text(req: dict, lang: str) -> str:
    lk = _lang_key(lang)
    if lk == "zh":
        return req.get("requirement_zh", req.get("requirement_en", ""))
    if lk == "ja":
        return req.get("requirement_ja", req.get("requirement_en", req.get("requirement_zh", "")))
    return req.get("requirement_en", req.get("requirement_zh", ""))


# ============================================================
# Word Export
# ============================================================


def append_crossref_table_word(
    doc,
    lang: str = "zh-TW",
    show_original_text: bool = True,
) -> None:
    """Append cross-reference section to an open python-docx Document.

    Args:
        doc: python-docx Document instance
        lang: UI language code ('zh-TW', 'ja-JP', 'en-US', etc.)
        show_original_text: if True, include original-language regulatory text
    """
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    lk = _lang_key(lang)

    # Page break to start on new page
    doc.add_page_break()

    # Section heading
    _headings = {
        "zh": "\u5404\u570b\u6cd5\u898f \u00d7 ISO 13485 \u4ea4\u53c9\u6bd4\u5c0d\u8868",
        "ja": "\u5404\u56fd\u898f\u5236 \u00d7 ISO 13485 \u4ea4\u53c9\u6bd4\u8f03\u8868",
        "en": "Multi-Country \u00d7 ISO 13485 Cross-Reference Table",
    }
    doc.add_heading(_headings.get(lk, _headings["en"]), level=2)

    # Get data
    data = _get_crossref_data()
    if data is None:
        _no_data = {
            "zh": "\uff08\u66ab\u7121\u6cd5\u898f Profile \u8cc7\u6599\uff09",
            "ja": "\uff08\u898f\u5236\u30d7\u30ed\u30d5\u30a1\u30a4\u30eb\u30c7\u30fc\u30bf\u306a\u3057\uff09",
            "en": "(No regulation profile data available)",
        }
        doc.add_paragraph(_no_data.get(lk, _no_data["en"]))
        return

    reg_ids = data["metadata"]["regulation_ids"]
    country_stats = data["summary"]["country_stats"]
    clauses = data["clauses"]
    delta_items = data.get("delta_items", {})

    # Build country groups (merges same-country regulations into one column)
    country_groups = _build_country_groups(reg_ids, country_stats)

    # Country list
    country_names = [_group_display(g, lang) for g in country_groups]
    _country_prefix = {
        "zh": "\u6bd4\u5c0d\u570b\u5bb6 / Countries",
        "ja": "\u6bd4\u8f03\u5bfe\u8c61\u56fd / Countries",
        "en": "Countries",
    }
    doc.add_paragraph(f"{_country_prefix.get(lk, _country_prefix['en'])}: {', '.join(country_names)}")

    # ── Helper to set cell background ──
    def _set_cell_bg(cell, hex_color: str):
        """Set background color of a Word table cell."""
        try:
            shading_elm = OxmlElement("w:shd")
            shading_elm.set(qn("w:fill"), hex_color)
            shading_elm.set(qn("w:val"), "clear")
            cell._tc.get_or_add_tcPr().append(shading_elm)
        except Exception:
            pass

    # ── Load ISO checklist for titles ──
    try:
        from src.analysis.compliance_rules import ISO_13485_CHECKLIST
    except ImportError:
        ISO_13485_CHECKLIST = {}

    # ── Main clause table ──
    n_cols = 2 + len(country_groups)  # clause_id, title, per-country-group

    _col0_label = {"zh": "ISO \u689d\u6b3e", "ja": "ISO \u6761\u9805", "en": "ISO Clause"}
    _col1_label = {"zh": "\u689d\u6b3e\u540d\u7a31", "ja": "\u6761\u9805\u540d", "en": "Title"}

    tbl = doc.add_table(rows=1 + len(clauses), cols=n_cols)
    tbl.style = "Table Grid"

    # Header row
    hdr_row = tbl.rows[0]
    hdr_row.cells[0].text = _col0_label.get(lk, _col0_label["en"])
    hdr_row.cells[1].text = _col1_label.get(lk, _col1_label["en"])
    for ci, group in enumerate(country_groups):
        hdr_row.cells[2 + ci].text = _group_display(group, lang)

    # Bold + small font for headers
    for cell in hdr_row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)

    # Data rows
    for ri, clause_entry in enumerate(clauses, 1):
        clause_id = clause_entry["clause_id"]
        row = tbl.rows[ri]

        # Col 0: clause id
        row.cells[0].text = clause_id
        for p in row.cells[0].paragraphs:
            for r in p.runs:
                r.font.size = Pt(7)

        # Col 1: title
        clause_info = ISO_13485_CHECKLIST.get(clause_id, {})
        row.cells[1].text = _clause_title(clause_info, lang) if clause_info else clause_entry.get("title", "")
        for p in row.cells[1].paragraphs:
            for r in p.runs:
                r.font.size = Pt(7)

        # Per-country-group columns
        for ci, group in enumerate(country_groups):
            cell = row.cells[2 + ci]
            merged = _merge_clause_data(clause_entry, group["reg_ids"])
            status = merged["status"]
            confidence = merged["confidence"]

            label = _status_label(status, lang)
            conf_str = f" ({confidence:.0%})" if confidence else ""
            cell_text = f"{label}{conf_str}"

            # Within-clause deltas
            wcds = merged["within_clause_deltas"]
            if wcds and status == "exceeds":
                for idx_d, wcd in enumerate(wcds, 1):
                    d_title = wcd.get("title_zh", wcd.get("title_en", "")) if lk == "zh" else (
                        wcd.get("title_ja", wcd.get("title_en", "")) if lk == "ja" else wcd.get("title_en", "")
                    )
                    d_impact = _impact_emoji(wcd.get("audit_impact", ""))
                    iso_base = wcd.get("iso_baseline_zh", wcd.get("iso_baseline_en", "")) if lk == "zh" else (
                        wcd.get("iso_baseline_ja", wcd.get("iso_baseline_en", "")) if lk == "ja" else wcd.get("iso_baseline_en", "")
                    )
                    c_spec = wcd.get("country_specific_zh", wcd.get("country_specific_en", "")) if lk == "zh" else (
                        wcd.get("country_specific_ja", wcd.get("country_specific_en", "")) if lk == "ja" else wcd.get("country_specific_en", "")
                    )
                    cell_text += f"\n\u21b3 \u2460{idx_d} [{d_impact}] {d_title}"
                    if iso_base and c_spec:
                        cell_text += f"\n   ISO: {iso_base[:60]}\n   \u2192 {c_spec[:60]}"

            cell.text = cell_text
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(6)

            # Cell background color
            color = _STATUS_COLORS_HEX.get(status, "FFFFFF")
            _set_cell_bg(cell, color)

    # ── Unique requirements section ──
    _uniq_heading = {
        "zh": "\u7368\u6709\u8981\u6c42\u6e05\u55ae / Unique Requirements",
        "ja": "\u56fa\u6709\u8981\u4ef6\u30ea\u30b9\u30c8 / Unique Requirements",
        "en": "Unique Requirements",
    }
    doc.add_heading(_uniq_heading.get(lk, _uniq_heading["en"]), level=3)

    merged_delta = _merge_delta_items(delta_items, country_groups)
    if not merged_delta:
        _no_uniq = {"zh": "\uff08\u7121\u7368\u6709\u8981\u6c42\uff09", "ja": "\uff08\u56fa\u6709\u8981\u4ef6\u306a\u3057\uff09", "en": "(No unique requirements)"}
        doc.add_paragraph(_no_uniq.get(lk, _no_uniq["en"]))
        return

    # Headers for unique req table
    _urq_headers = {
        "zh": ["#", "ID", "\u689d\u6b3e", "\u6a19\u984c", "\u5f71\u97ff", "\u8981\u6c42", "\u7a3d\u6838\u554f\u984c", "\u671f\u671b\u8b49\u64da"],
        "ja": ["#", "ID", "\u6761\u9805", "\u30bf\u30a4\u30c8\u30eb", "\u5f71\u97ff", "\u8981\u4ef6", "\u76e3\u67fb\u8cea\u554f", "\u671f\u5f85\u8a3c\u62e0"],
        "en": ["#", "ID", "Clause", "Title", "Impact", "Requirement", "Audit Q", "Evidence"],
    }
    headers = _urq_headers.get(lk, _urq_headers["en"])

    for group, reqs in merged_delta:
        if not reqs:
            continue
        country_name = _group_display(group, lang)
        doc.add_heading(f"{country_name}", level=4)

        urq_tbl = doc.add_table(rows=1, cols=len(headers))
        urq_tbl.style = "Table Grid"
        hdr = urq_tbl.rows[0]
        for hi, h in enumerate(headers):
            hdr.cells[hi].text = h
            for p in hdr.cells[hi].paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(7)

        for idx_r, req in enumerate(reqs, 1):
            row = urq_tbl.add_row()
            row.cells[0].text = str(idx_r)
            row.cells[1].text = req.get("req_id", "")
            row.cells[2].text = ", ".join(req.get("related_iso_clauses", []))
            row.cells[3].text = _req_title(req, lang)

            impact = req.get("audit_impact", "")
            row.cells[4].text = f"{_impact_emoji(impact)} {impact}"

            # Requirement text + within-clause delta note
            req_text = _req_text(req, lang)
            is_wcd = req.get("is_within_clause_delta", False)
            wcd_note = req.get("within_clause_delta_vs_iso", "")
            if is_wcd and wcd_note:
                clauses_str = ", ".join(req.get("related_iso_clauses", []))
                req_text += f"\n\u26a1 Within-clause delta vs ISO {clauses_str}: {wcd_note}"
            row.cells[5].text = req_text

            # Audit question
            if lk == "zh":
                row.cells[6].text = req.get("audit_question_zh", req.get("audit_question_en", ""))
            elif lk == "ja":
                row.cells[6].text = req.get("audit_question_ja", req.get("audit_question_en", req.get("audit_question_zh", "")))
            else:
                row.cells[6].text = req.get("audit_question_en", req.get("audit_question_zh", ""))

            # Expected evidence
            evidence = req.get("expected_evidence", [])
            row.cells[7].text = "; ".join(evidence) if evidence else ""

            # Small font for all cells
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(6)

            # Original text row
            if show_original_text and req.get("original_text"):
                orig_row = urq_tbl.add_row()
                # Merge all cells
                orig_row.cells[0].merge(orig_row.cells[-1])
                orig_text = f"\u539f\u6587 ({req.get('original_lang', '')}): {req['original_text']}"
                eng_trans = req.get("english_translation", "")
                if eng_trans:
                    orig_text += f"\nEnglish: {eng_trans}"
                orig_row.cells[0].text = orig_text
                for p in orig_row.cells[0].paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(6)
                        r.italic = True


# ============================================================
# Excel Export
# ============================================================


def append_crossref_table_excel(
    wb,
    lang: str = "zh-TW",
    show_original_text: bool = True,
) -> None:
    """Append cross-reference sheets to an open openpyxl Workbook.

    Creates two sheets:
      - CrossRef / (sheet name localized)
      - UniqueReqs / (sheet name localized)

    Args:
        wb: openpyxl Workbook instance
        lang: UI language code
        show_original_text: include original-language text columns
    """
    from openpyxl.styles import Font, PatternFill, Alignment

    lk = _lang_key(lang)

    data = _get_crossref_data()
    if data is None:
        return

    reg_ids = data["metadata"]["regulation_ids"]
    country_stats = data["summary"]["country_stats"]
    clauses = data["clauses"]
    delta_items = data.get("delta_items", {})

    # Build country groups (merges same-country regulations into one column)
    country_groups = _build_country_groups(reg_ids, country_stats)

    # Load checklist for titles
    try:
        from src.analysis.compliance_rules import ISO_13485_CHECKLIST
    except ImportError:
        ISO_13485_CHECKLIST = {}

    # ── Sheet 1: CrossRef ──
    _sheet1_name = {
        "zh": "\u4ea4\u53c9\u6bd4\u5c0d",
        "ja": "\u4ea4\u53c9\u6bd4\u8f03",
        "en": "CrossRef",
    }
    ws = wb.create_sheet(_sheet1_name.get(lk, _sheet1_name["en"]))

    header_fill = PatternFill(start_color="1565C0", end_color="1565C0", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=9)

    # Header row
    headers = ["ISO Clause", "Title"]
    for group in country_groups:
        headers.append(_group_display(group, lang))

    for ci, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=ci, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", wrap_text=True)

    # Data rows
    for ri, clause_entry in enumerate(clauses, 2):
        clause_id = clause_entry["clause_id"]
        clause_info = ISO_13485_CHECKLIST.get(clause_id, {})
        title = _clause_title(clause_info, lang) if clause_info else clause_entry.get("title", "")

        ws.cell(row=ri, column=1, value=clause_id)
        ws.cell(row=ri, column=2, value=title).alignment = Alignment(wrap_text=True)

        for ci, group in enumerate(country_groups):
            merged = _merge_clause_data(clause_entry, group["reg_ids"])
            status = merged["status"]
            confidence = merged["confidence"]

            label = _status_label(status, lang)
            conf_str = f" ({confidence:.0%})" if confidence else ""
            cell_text = f"{label}{conf_str}"

            # Within-clause deltas
            wcds = merged["within_clause_deltas"]
            if wcds and status == "exceeds":
                delta_titles = []
                for wcd in wcds:
                    d_title = wcd.get("title_zh", wcd.get("title_en", "")) if lk == "zh" else (
                        wcd.get("title_ja", wcd.get("title_en", "")) if lk == "ja" else wcd.get("title_en", "")
                    )
                    delta_titles.append(d_title)
                if delta_titles:
                    cell_text += "\n" + "; ".join(delta_titles)

            cell = ws.cell(row=ri, column=3 + ci, value=cell_text)
            cell.alignment = Alignment(wrap_text=True, vertical="top")

            # Fill color
            hex_c = _STATUS_COLORS_EXCEL.get(status, "FFFFFF")
            cell.fill = PatternFill(start_color=hex_c, end_color=hex_c, fill_type="solid")

        ws.row_dimensions[ri].height = 45

    # Freeze header
    ws.freeze_panes = "A2"

    # Column widths
    ws.column_dimensions["A"].width = 10
    ws.column_dimensions["B"].width = 35
    from openpyxl.utils import get_column_letter
    for ci in range(len(country_groups)):
        ws.column_dimensions[get_column_letter(3 + ci)].width = 28

    # ── Sheet 2: UniqueReqs ──
    _sheet2_name = {
        "zh": "\u7368\u6709\u8981\u6c42",
        "ja": "\u56fa\u6709\u8981\u4ef6",
        "en": "UniqueReqs",
    }
    ws2 = wb.create_sheet(_sheet2_name.get(lk, _sheet2_name["en"]))

    _uq_header_map = {
        "zh": ["國家", "法規 ID", "ISO 條款", "標題", "影響等級", "要求說明", "稽核問題", "預期佐證", "類型", "法規原文", "語言", "英文翻譯"],
        "ja": ["国", "規制 ID", "ISO 条項", "タイトル", "影響レベル", "要件説明", "監査質問", "期待される証拠", "タイプ", "原文", "言語", "英語翻訳"],
        "en": ["Country", "Reg ID", "ISO Clauses", "Title", "Impact", "Requirement", "Audit Question", "Expected Evidence", "Type", "Original Text", "Language", "English Translation"],
    }
    uq_headers = _uq_header_map.get(lk, _uq_header_map["en"])
    for ci, h in enumerate(uq_headers, 1):
        cell = ws2.cell(row=1, column=ci, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", wrap_text=True)

    row_idx = 2
    merged_delta2 = _merge_delta_items(delta_items, country_groups)
    for group, reqs in merged_delta2:
        if not reqs:
            continue
        country_name = _group_display(group, lang)
        for req in reqs:
            is_wcd = req.get("is_within_clause_delta", False)
            impact = req.get("audit_impact", "")

            ws2.cell(row=row_idx, column=1, value=country_name)
            ws2.cell(row=row_idx, column=2, value=req.get("req_id", ""))
            ws2.cell(row=row_idx, column=3, value=", ".join(req.get("related_iso_clauses", [])))
            ws2.cell(row=row_idx, column=4, value=_req_title(req, lang))
            ws2.cell(row=row_idx, column=5, value=f"{_impact_emoji(impact)} {impact}")

            req_text = _req_text(req, lang)
            wcd_note = req.get("within_clause_delta_vs_iso", "")
            if is_wcd and wcd_note:
                clauses_str = ", ".join(req.get("related_iso_clauses", []))
                req_text += f"\n\u26a1 Within-clause delta vs ISO {clauses_str}: {wcd_note}"
            ws2.cell(row=row_idx, column=6, value=req_text).alignment = Alignment(wrap_text=True)

            # Audit question
            if lk == "zh":
                ws2.cell(row=row_idx, column=7, value=req.get("audit_question_zh", req.get("audit_question_en", "")))
            elif lk == "ja":
                ws2.cell(row=row_idx, column=7, value=req.get("audit_question_ja", req.get("audit_question_en", req.get("audit_question_zh", ""))))
            else:
                ws2.cell(row=row_idx, column=7, value=req.get("audit_question_en", req.get("audit_question_zh", "")))

            evidence = req.get("expected_evidence", [])
            ws2.cell(row=row_idx, column=8, value="; ".join(evidence) if evidence else "")

            # Type
            _type_wcd = {"zh": "\u689d\u6b3e\u5167\u5dee\u7570 \u26a1", "ja": "\u6761\u9805\u5185\u5dee\u7570 \u26a1", "en": "Within-Clause Delta \u26a1"}
            _type_uniq = {"zh": "\u7368\u6709\u8981\u6c42", "ja": "\u56fa\u6709\u8981\u4ef6", "en": "Unique Requirement"}
            type_label = _type_wcd.get(lk, _type_wcd["en"]) if is_wcd else _type_uniq.get(lk, _type_uniq["en"])
            ws2.cell(row=row_idx, column=9, value=type_label)

            # Original text
            if show_original_text:
                ws2.cell(row=row_idx, column=10, value=req.get("original_text", "")).alignment = Alignment(wrap_text=True)
                ws2.cell(row=row_idx, column=11, value=req.get("original_lang", ""))
                ws2.cell(row=row_idx, column=12, value=req.get("english_translation", "")).alignment = Alignment(wrap_text=True)

            # Row fill by impact
            fill_hex = _IMPACT_FILL_EXCEL.get(impact)
            if fill_hex:
                fill = PatternFill(start_color=fill_hex, end_color=fill_hex, fill_type="solid")
                for c in range(1, 13):
                    ws2.cell(row=row_idx, column=c).fill = fill

            row_idx += 1

    # Freeze header
    ws2.freeze_panes = "A2"

    # Column widths
    _uq_widths = [14, 14, 14, 25, 10, 45, 35, 30, 18, 45, 8, 45]
    for ci, w in enumerate(_uq_widths, 1):
        ws2.column_dimensions[get_column_letter(ci)].width = w
